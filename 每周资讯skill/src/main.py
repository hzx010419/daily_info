#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
每周资讯Skill主程序（升级版 v2）
核心思路：每日资讯已做好摘要/分类，每周只需 去重+整合
升级点：
  1. 所有AI调用改为8线程批量并发
  2. 移植每日资讯skill的42条噪音过滤规则（纯规则，无需API）
  3. TF-IDF+余弦相似度去重
  4. 去重后大簇AI多视角整合、小簇选优（同每日资讯逻辑）
  5. 批量并发MECE分类
  6. 使用auth模块管理API Key
"""

import warnings
import os
import re
import json
import logging
import threading
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass, field
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import sys

warnings.filterwarnings('ignore')

from ai_client import (
    DeepSeekClient, API_MAX_CONCURRENCY,
    HIGH_SIM_THRESHOLD, MEDIUM_SIM_THRESHOLD, LOW_SIM_THRESHOLD,
)

# ====================== 日志 ======================
current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
logs_dir = os.path.join(current_dir, 'logs')
os.makedirs(logs_dir, exist_ok=True)

file_handler = logging.FileHandler(os.path.join(logs_dir, 'app.log'), encoding='utf-8')
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
console_handler = logging.StreamHandler()
console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

logging.basicConfig(level=logging.INFO, handlers=[file_handler, console_handler], force=True)
logging.getLogger('urllib3').setLevel(logging.WARNING)
logging.getLogger('requests').setLevel(logging.WARNING)
logger = logging.getLogger(__name__)


def _p(msg: str) -> None:
    try:
        print(msg, flush=True)
    except Exception:
        pass


# ====================== Article ======================
@dataclass
class Article:
    source_name: str
    title: str
    link: str
    pub_date: datetime
    ai_summary: str = ""
    category_tag: str = ""
    mece_category: str = ""
    original_file: str = ""
    is_merged: bool = False
    merged_data: dict = None
    is_advertisement: bool = False
    rejection_reason: str = ""

    def __post_init__(self):
        if self.merged_data is None:
            self.merged_data = {}

    MECE_CATEGORIES = {
        "1": "国际局势与地缘政治", "1.1": "地区冲突与战争", "1.2": "国际关系与外交", "1.3": "全球安全与反恐",
        "2": "宏观经济与金融", "2.1": "国内宏观经济", "2.2": "宏观消费与零售", "2.3": "资本市场", "2.4": "房地产市场", "2.5": "全球金融",
        "3": "文化娱乐内容", "3.1": "文化产业", "3.2": "体育休闲", "3.3": "体验消费", "3.4": "网络与数字文化",
        "4": "科技产业动态与创新", "4.1": "人工智能", "4.2": "信息技术", "4.3": "智能制造与机器人", "4.4": "新能源与绿色科技",
        "5": "医疗健康与生命科学", "5.1": "医疗改革与政策", "5.2": "医学研究与突破", "5.3": "公共卫生",
        "6": "教育发展与人才培养", "6.1": "教育政策与改革", "6.2": "高等教育", "6.3": "教育科技", "6.4": "就业与人才",
        "7": "社会民生与消费", "7.1": "民生消费", "7.2": "就业与职场", "7.3": "人口与家庭", "7.4": "社会治理",
        "8": "企业与商业", "8.1": "企业经营", "8.2": "商业动态", "8.3": "电商与新零售", "8.4": "创业与创新",
        "9": "政策法规与监管", "9.1": "国家政策", "9.2": "行业监管", "9.3": "地方治理", "9.4": "反腐与廉政",
        "10": "能源与资源", "10.1": "传统能源", "10.2": "新能源", "10.3": "资源与环境",
        "11": "社会热点与舆论动态", "11.1": "热点事件", "11.2": "舆情分析", "11.3": "跨界热点",
    }

    HIGH_QUALITY_SOURCES = [
        'New Economist', '新经济学家', '财新', '财经', '经济学人',
        '三联生活周刊', '南方周末', '新京报', '第一财经', '界面',
        '中国新闻周刊', '瞭望', '半月谈',
        '麦可思研究', '中美聚焦', '智谷趋势', '秦朔朋友圈', '吴晓波频道',
        '财经十一人', '虎嗅APP', '钛媒体', '36氪',
    ]

    CROSS_DOMAIN_INDICATORS = [
        '助推', '推动', '影响', '冲击', '带动', '拉动', '倒逼',
        '重塑', '改写', '牵动', '波及', '传导',
        '背景', '背后', '启示', '机遇', '挑战', '博弈', '关系',
        '战略', '格局', '趋势', '走向', '变局', '演变',
        '如何影响', '意味着', '背后逻辑', '深层',
    ]

    def is_noise_content(self) -> tuple:
        """42条规则噪音过滤（纯规则，无需API）"""
        text = self.title + " " + (self.ai_summary[:500] if self.ai_summary else "")

        if not self.ai_summary or self.ai_summary == "摘要生成失败":
            return True, "摘要缺失"

        for p in [r'原文内容在此中断', r'原文未提供完整信息', r'摘要仅能涵盖已给出的部分', r'无法补充未提及的细节']:
            if re.search(p, self.ai_summary):
                return True, "摘要不完整"

        is_hq = any(s in self.source_name for s in Article.HIGH_QUALITY_SOURCES)

        # 1. 鸡汤文
        chicken_soup = [r'高敏感', r'过度努力', r'心理内耗', r'职场困境',
            r'情绪.{0,4}(失控|崩溃|低落|内耗|勒索|化|垃圾)',
            r'负面情绪', r'(管理|控制|调节).{0,4}情绪',
            r'感悟', r'人生哲理', r'励志', r'成功学', r'心得', r'自我反思', r'个人成长',
            r'学会.*相处', r'心态', r'幸福感', r'心理.*建议']
        fin_emotions = ['市场情绪', '投资者情绪', '情绪指数', '情绪面', '恐慌情绪', '风险情绪', '消费情绪']
        for p in chicken_soup:
            if re.search(p, text):
                if '情绪' in p and any(c in text for c in fin_emotions): continue
                if not any(k in text for k in ['数据','政策','分析','研究','报告','统计','市场','股市','经济','金融']): 
                    return True, f"鸡汤文: {p}"

        # 2. 医学案例
        for p in [r'\d+岁.{0,10}(患者|女士|先生|男子|女孩|男孩)', r'罕见.{0,10}(病例|疾病|肿瘤)',
                  r'(双子宫|心脏骤停|癌症晚期)', r'手术.{0,10}(成功|顺利)', r'经历.{0,5}(手术|化疗)']:
            if re.search(p, text):
                if not any(k in text for k in ['疫情','公共卫生','疫苗','医保','政策','突破','基因编辑']):
                    return True, f"医学案例: {p}"

        # 4. 餐饮业
        for p in [r'餐饮店', r'餐饮业', r'餐饮行业', r'人均.{0,5}(消费|餐饮)', r'外卖.{0,5}(平台|骑手)', r'翻台率', r'坪效']:
            if re.search(p, text):
                if not any(k in text for k in ['政策','监管','处罚','重大','食品安全']):
                    return True, f"餐饮业: {p}"

        # 5. 院校动态
        for p in [r'技师学院', r'职业技术学院', r'学院.{0,5}(并入|合并|组建|揭牌)',
                  r'(学院|大学).{0,10}(升格|升本|更名)', r'升格.{0,5}(本科|大学)']:
            if re.search(p, text):
                if not any(k in text for k in ['违规','丑闻','犯罪','舆情','贪污','性侵','学术不端']):
                    return True, f"院校动态: {p}"

        # 7. 职场盘点
        for p in [r'职场.{0,10}(减分项|加分项|禁忌|雷区|建议|盘点)',
                  r'(医生|护士|教师|公务员).{0,10}(过得好不好|辛苦|忙碌|不容易)',
                  r'(医生|护士|教师).{0,10}(一眼.{0,5}(看出|看穿)|真实状态|日常)',
                  r'(忙起来|忙到).{0,15}(顾不上|没时间|来不及)']:
            if re.search(p, text):
                if not any(k in text for k in ['数据','报告','研究','政策','改革','行业分析','趋势']):
                    return True, f"职场盘点: {p}"

        # 9. 聚合快讯
        for p in [r'科股快报', r'财经早知道', r'今日要闻', r'新闻速递', r'早间快讯',
                  r'涨停板', r'涨跌停', r'【今日导读】', r'【早餐内参】',
                  r'影响.{0,5}(下周|本周|市场).{0,5}(十大|几大)?(消息|资讯|事件)',
                  r'(十大|几大|八大).{0,5}(消息|资讯|事件|信号)']:
            if re.search(p, self.title):
                if len(self.title) < 30 or '、' in self.title:
                    return True, f"聚合快讯: {p}"
                break

        # 9.1 内容杂糅（保护跨领域分析）
        is_cross = any(ind in text for ind in Article.CROSS_DOMAIN_INDICATORS)
        if not is_cross:
            domain_map = {
                '国际军事': ['军事','战争','武器','军演','导弹','制裁'],
                '金融': ['股市','A股','美股','基金','债券','利率','汇率','央行'],
                '科技': ['AI','人工智能','芯片','半导体','5G','大模型'],
                '医疗': ['医疗','医药','医院','药品','疫苗','医保'],
                '教育': ['教育','高考','大学','学校','招生','考研'],
                '房地产': ['房地产','楼市','房价','限购'],
                '能源': ['石油','天然气','煤炭','核电','风电','光伏','新能源'],
            }
            matched = [d for d, kws in domain_map.items() if any(k in text for k in kws)]
            if len(matched) >= 3:
                connectives = ['与此同时','此外','另外','除此之外','而在']
                if sum(1 for w in connectives if w in text) >= 1:
                    return True, f"内容杂糅: {len(matched)}个领域"

        # 更多噪音规则
        noise_rules = {
            r'推荐.{0,5}(书单|必读)': "书单", r'(十大|百强|TOP\d+).{0,10}(城市|企业|榜单)': "榜单",
            r'(震撼|泪目|看哭).{0,10}(全国|全网)': "标题党", r'个人.{0,5}(成长|感悟|分享)': "个人感悟",
            r'(涨停|跌停|暴雷).{0,5}(名单|汇总)': "个股名单",
            r'(最美|最佳|最好看).{0,10}(电影|景点)': "主观审美",
            r'(多吃|少吃|不能吃).{0,10}(这些|这种)': "饮食科普",
            r'(今日|本周).{0,5}热搜.{0,5}(排行|盘点)': "热搜盘点",
            r'(明星|网红).{0,10}(恋情|离婚|出轨)': "娱乐八卦",
            r'(回顾|揭秘).{0,10}(历史|古代|朝代)': "历史回顾",
            r'(大赛|竞赛|挑战赛).{0,10}(启动|报名|开幕)': "比赛宣传",
            r'(降温|暴雨|台风).{0,10}(预警|预报|提醒)': "天气预报",
            r'(省|市|区).{0,10}(政府|人大).{0,10}(任免|任命)': "人事任免",
            r'(街道|社区|小区).{0,10}(通知|公告|处理)': "微小事件",
            r'(结构方程|中介效应|问卷调查).{0,10}(研究|分析)': "学术论文",
            r'(退学|处分|开除).{0,10}(学生|研究生)': "学生事件",
            r'(维生素|矿物质|蛋白质).{0,10}(补充|缺乏)': "营养科普",
        }
        for pattern, reason in noise_rules.items():
            if re.search(pattern, text):
                if not any(k in text for k in ['政策','重大','国务院','全国','突破','监管','处罚','犯罪','反腐']):
                    return True, f"{reason}: {pattern}"

        return False, ""


# ====================== 摘要清理 ======================
def _clean_summaries(articles: List[Article]):
    meta_pats = [r'[，。]?\s*文章由[^。]+发布[，，][^。]+', r'[，。]?\s*作者[：:][^，。]+[，，]编辑[：:][^，。]+',
                 r'[，。]?\s*本文由[^。]+发布', r'[，。]?\s*转载自[^，。]+']
    noinfo_pats = [r'[，。]?\s*原文未[进一步]{0,3}(说明|提及|公布|披露)[^。]*。',
                  r'[，。]?\s*(具体|详细|后续)(处理|处罚|结果|措施)(未|暂未)[^。]*。',
                  r'[，。]?\s*(未|暂未)(公布|披露|说明|提及)(具体|详细|后续)[^。]*。']
    cnt = 0
    for a in articles:
        if a.is_advertisement or not a.ai_summary: continue
        orig = a.ai_summary
        c = orig
        for p in meta_pats: c = re.sub(p, '', c)
        for p in noinfo_pats: c = re.sub(p, '', c)
        c = re.sub(r'\s{2,}', ' ', c)
        c = re.sub(r'[，。]{2,}', '，', c).strip()
        if c != orig: a.ai_summary = c; cnt += 1
    if cnt: logger.info(f"摘要清理: {cnt}篇")


# ====================== 字体设置 ======================
def _set_font(paragraph, font_name, size):
    from docx.oxml.ns import qn
    from docx.shared import Pt
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)


# ====================== DocumentGenerator ======================
class DocumentGenerator:
    MECE_MAIN_CATEGORIES = {
        "1": "国际局势与地缘政治", "2": "宏观经济与金融", "3": "文化娱乐内容",
        "4": "科技产业动态与创新", "5": "医疗健康与生命科学", "6": "教育发展与人才培养",
        "7": "社会民生与消费", "8": "企业与商业", "9": "政策法规与监管",
        "10": "能源与资源", "11": "社会热点与舆论动态",
    }

    @staticmethod
    def create_document(articles, output_path, week_start, week_end) -> bool:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            sorted_articles = DocumentGenerator._sort_by_mece_category(articles)
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            style.font.size = Pt(10.5)

            title = doc.add_heading(level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = title.add_run('每周资讯汇总 - '); r1.font.name = '微软雅黑'; r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); r1.font.size = Pt(16); r1.font.bold = True
            r2 = title.add_run(f'{week_start.strftime("%Y年%m月%d日")}至{week_end.strftime("%Y年%m月%d日")}'); r2.font.name = '微软雅黑'; r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); r2.font.size = Pt(14); r2.font.bold = True

            sp = doc.add_paragraph(); sp.add_run(f"本周共收录 {len(sorted_articles)} 篇资讯"); _set_font(sp, '微软雅黑', 10.5); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp.paragraph_format.space_after = Pt(15)
            sep = doc.add_paragraph(); r = sep.add_run("═"*71); r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(128,128,128); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER; sep.paragraph_format.space_after = Pt(15)

            items = []; cur_cat = None; anum = 0
            for art in sorted_articles:
                cat = art.mece_category or "11.3"; mc = cat.split('.')[0] if '.' in cat else cat
                if mc != cur_cat:
                    cur_cat = mc; anum = 0
                    items.append({'type':'category','name': DocumentGenerator.MECE_MAIN_CATEGORIES.get(mc,"其他")})
                anum += 1; items.append({'type':'article','art':art,'num':anum})

            cn = ['一','二','三','四','五','六','七','八','九','十',
                  '十一','十二','十三','十四','十五','十六','十七','十八','十九','二十',
                  '二十一','二十二','二十三','二十四','二十五','二十六','二十七','二十八','二十九','三十',
                  '三十一','三十二','三十三','三十四','三十五','三十六','三十七','三十八','三十九','四十',
                  '四十一','四十二','四十三','四十四','四十五','四十六','四十七','四十八','四十九','五十',
                  '五十一','五十二','五十三','五十四','五十五','五十六','五十七','五十八','五十九','六十',
                  '六十一','六十二','六十三','六十四','六十五','六十六','六十七','六十八','六十九','七十',
                  '七十一','七十二','七十三','七十四','七十五','七十六','七十七','七十八','七十九','八十',
                  '八十一','八十二','八十三','八十四','八十五','八十六','八十七','八十八','八十九','九十',
                  '九十一','九十二','九十三','九十四','九十五','九十六','九十七','九十八','九十九','一百']

            for i, it in enumerate(items):
                if it['type'] == 'category':
                    p = doc.add_paragraph(); r = p.add_run(f"\n【{it['name']}】"); r.font.name='微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.bold=True; r.font.size=Pt(16); p.paragraph_format.space_before=Pt(15); p.paragraph_format.space_after=Pt(10)
                else:
                    art = it['art']; ns = cn[it['num']-1] if it['num']-1 < len(cn) else str(it['num'])
                    dt = art.category_tag or art.title
                    tp = doc.add_paragraph(); r = tp.add_run(f'{ns}、{dt}'); r.font.name='微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.bold=True; r.font.size=Pt(12); tp.paragraph_format.space_after=Pt(6)
                    sp2 = doc.add_paragraph(f'         来源：{art.source_name}'); sp2.paragraph_format.space_after=Pt(6); _set_font(sp2,'微软雅黑',10.5)
                    tp2 = doc.add_paragraph(f'         发布时间：{art.pub_date.strftime("%Y-%m-%d %H:%M:%S")}'); tp2.paragraph_format.space_after=Pt(6); _set_font(tp2,'微软雅黑',10.5)
                    cp = doc.add_paragraph(f'         内容摘要：{art.ai_summary}'); cp.paragraph_format.space_after=Pt(6); cp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; _set_font(cp,'微软雅黑',10.5)
                    if art.is_merged and art.merged_data.get('links'):
                        lp = doc.add_paragraph('         文章链接：'); _set_font(lp,'微软雅黑',10.5)
                        links = art.merged_data.get('links',[]); sources = art.merged_data.get('sources',[])
                        for idx2, lk in enumerate(links):
                            src = sources[idx2] if idx2 < len(sources) else ""
                            lp2 = doc.add_paragraph(); r = lp2.add_run(f"         • {src}: {lk}"); r.font.name='微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.size=Pt(10.5); r.font.color.rgb=RGBColor(0,0,255); lp2.paragraph_format.space_after=Pt(3)
                        us = list(set(sources))
                        if len(us) > 1:
                            ss = doc.add_paragraph(); ss.add_run(f"         （融合来源：{'、'.join(us[:5])}{'等' if len(us)>5 else ''}）"); _set_font(ss,'微软雅黑',10.5); ss.paragraph_format.space_after=Pt(12)
                    else:
                        lp = doc.add_paragraph('         文章链接：'); _set_font(lp,'微软雅黑',10.5); r = lp.add_run(art.link); r.font.color.rgb=RGBColor(0,0,255); lp.paragraph_format.space_after=Pt(12)
                    if i < len(items)-1 and items[i+1]['type'] != 'category':
                        sp3 = doc.add_paragraph(); r = sp3.add_run("——————————————————————————————"); r.font.name='微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); r.font.bold=True; r.font.size=Pt(9); sp3.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp3.paragraph_format.space_before=Pt(10); sp3.paragraph_format.space_after=Pt(6)

            doc.save(output_path); logger.info(f"文档已生成: {output_path}"); return True
        except Exception as e:
            logger.error(f"生成文档失败: {e}"); import traceback; traceback.print_exc(); return False

    @staticmethod
    def _sort_by_mece_category(articles):
        def key(a):
            cat = a.mece_category or "11.3"; parts = cat.split('.')
            try: mc = int(parts[0]); sc = int(parts[1]) if len(parts)>1 else 0
            except: mc = 11; sc = 3
            return (mc, sc, -a.pub_date.timestamp() if a.pub_date else 0)
        return sorted(articles, key=key)


# ====================== WeeklyNewsProcessor ======================
class WeeklyNewsProcessor:

    def __init__(self, daily_news_dir, output_dir, api_key):
        self.daily_news_dir = daily_news_dir
        self.output_dir = output_dir
        self.ai = DeepSeekClient(api_key)
        self.articles = []

    def get_week_range(self, target_date=None):
        if target_date is None: target_date = datetime.now()
        d = target_date.weekday()
        ws = target_date - timedelta(days=d)
        ws = ws.replace(hour=0, minute=0, second=0, microsecond=0)
        we = ws + timedelta(days=6, hours=23, minutes=59, seconds=59)
        return ws, we

    # ---------- 解析 ----------
    def parse_daily_news_file(self, file_path):
        articles = []
        try:
            from docx import Document
            doc = Document(file_path)
            fname = os.path.basename(file_path)
            dm = re.search(r'(\d{4})-(\d{2})-(\d{2})', fname)
            fdate = datetime(int(dm.group(1)),int(dm.group(2)),int(dm.group(3))) if dm else datetime.now()
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            cur = {}
            mece_names = set(Article.MECE_CATEGORIES.values())

            for para in paras:
                if para.startswith('每日资讯汇总') or para.startswith('【今日导读】') or para.startswith('本周共收录'): continue
                if para.startswith('═') or para.startswith('—'): continue
                if para.startswith('【') and para.endswith('】') and '、' not in para: continue
                if para in ('目  录','目录','目 录'): continue

                tm = re.match(r'^[一二三四五六七八九十百千]+、(.+)', para)
                if tm:
                    mt = tm.group(1).strip()
                    if mt in mece_names: continue  # 目录区分类标题
                    if cur and 'title' in cur:
                        articles.append(Article(
                            source_name=cur.get('source',''), title=cur.get('title',''),
                            link=cur.get('link',''), pub_date=cur.get('pub_date',fdate),
                            ai_summary=cur.get('summary',''), category_tag=cur.get('category_tag',''),
                            mece_category=cur.get('mece_category','11.3'), original_file=fname))
                    cur = {'title': mt, 'category_tag': mt if mt.startswith('【') else f'【{mt}】'}
                elif para.startswith('来源：') or para.startswith('         来源：'):
                    cur['source'] = para.replace('来源：','').replace('         来源：','').strip()
                elif '发布时间：' in para:
                    try: cur['pub_date'] = datetime.strptime(para.split('发布时间：')[-1].strip(), '%Y-%m-%d %H:%M:%S')
                    except: cur['pub_date'] = fdate
                elif '内容摘要：' in para:
                    cur['summary'] = para.split('内容摘要：')[-1].strip()
                elif para.startswith('文章链接：') or para.startswith('         文章链接：'):
                    cur['link'] = para.replace('文章链接：','').replace('         文章链接：','').strip()

            if cur and 'title' in cur:
                articles.append(Article(
                    source_name=cur.get('source',''), title=cur.get('title',''),
                    link=cur.get('link',''), pub_date=cur.get('pub_date',fdate),
                    ai_summary=cur.get('summary',''), category_tag=cur.get('category_tag',''),
                    mece_category=cur.get('mece_category','11.3'), original_file=fname))
            logger.info(f"从 {fname} 解析到 {len(articles)} 篇文章")
        except Exception as e:
            logger.error(f"解析文件失败 {file_path}: {e}")
        return articles

    def collect_weekly_articles(self, ws, we):
        if not os.path.exists(self.daily_news_dir):
            logger.error(f"目录不存在: {self.daily_news_dir}"); return []
        all_art = []
        for fn in os.listdir(self.daily_news_dir):
            if not fn.endswith('.docx'): continue
            dm = re.search(r'(\d{4})-(\d{2})-(\d{2})', fn)
            if dm:
                fd = datetime(int(dm.group(1)),int(dm.group(2)),int(dm.group(3)))
                if ws <= fd <= we:
                    for a in self.parse_daily_news_file(os.path.join(self.daily_news_dir, fn)):
                        if ws <= a.pub_date <= we: all_art.append(a)
        logger.info(f"共收集 {len(all_art)} 篇")
        return all_art

    # ---------- 步骤2: 规则噪音过滤（纯规则，极快） ----------
    def filter_noise_rules(self, articles):
        _p("  规则噪音过滤...")
        filtered = []; noise = 0; reasons = defaultdict(int)
        for a in articles:
            is_n, r = a.is_noise_content()
            if is_n: a.is_advertisement = True; a.rejection_reason = r; noise += 1; reasons[r.split(':')[0]] += 1
            else: filtered.append(a)
        _p(f"    移除 {noise} 篇，保留 {len(filtered)} 篇")
        return filtered

    # ---------- 步骤3: 批量AI广告过滤（8线程并发） ----------
    def batch_filter_ai_advertisement(self, articles):
        _p("  批量AI广告过滤...")
        data = [{'title': a.category_tag or a.title, 'content': a.ai_summary or ""} for a in articles]
        results = self.ai.batch_classify_advertisement(data)
        filtered = []
        for i, (a, r) in enumerate(zip(articles, results)):
            if r.get('is_ad', False):
                a.is_advertisement = True; a.rejection_reason = f"AI: {r.get('reason','')}"
            else: filtered.append(a)
        _p(f"    移除 {len(articles)-len(filtered)} 篇，保留 {len(filtered)} 篇")
        return filtered

    # ---------- 步骤4: 批量AI聚合/评论/杂糅过滤（8线程并发） ----------
    def batch_filter_combined(self, articles):
        """一次性并发判断 聚合新闻+社会评论+杂糅，三合一，减少轮次"""
        _p("  批量AI聚合/评论/杂糅过滤...")
        n = len(articles)
        results = [{'is_aggregated': False, 'is_commentary': False, 'is_mess': False}] * n
        lock = threading.Lock()

        def _judge(idx, art):
            title = art.category_tag or art.title
            summary = art.ai_summary or ""
            res = {'is_aggregated': False, 'is_commentary': False, 'is_mess': False}
            try:
                # 合并为一次AI调用判断3种过滤类型
                prompt = f"""请判断以下文章是否属于以下3种需要过滤的类型：

文章标题：{title}
文章摘要：{summary[:800] if summary else "无"}

判断类型：
1. **聚合新闻** = 一条资讯像新闻简报一样汇总多个完全不相关主题
2. **社会评论** = 揭示"外界对某职业/群体的误解"，用"XX不是XX"句式，大量感慨性语言
3. **主题杂糅** = 用"与此同时""此外"等连接词把3个以上彼此无因果的独立新闻硬拼到一起，缺乏统一主题

以下情况**不是**上述类型：
- 围绕单一主题的深度分析
- 跨领域因果/影响分析（如"战争助推能源战略"）
- 客观新闻报道、政策解读

请严格按照以下JSON格式回复（只回复JSON）：
{{"is_aggregated": true/false, "is_commentary": true/false, "is_mess": true/false, "reason": "简要理由"}}"""

                resp = self.ai._call_api(prompt)
                parsed = json.loads(resp)
                res = {
                    'is_aggregated': bool(parsed.get('is_aggregated', False)),
                    'is_commentary': bool(parsed.get('is_commentary', False)),
                    'is_mess': bool(parsed.get('is_mess', False)),
                    'reason': parsed.get('reason', ''),
                }
            except Exception as e:
                logger.error(f"综合判断失败: {e}")
            with lock:
                results[idx] = res

        with ThreadPoolExecutor(max_workers=API_MAX_CONCURRENCY) as exe:
            futures = [exe.submit(_judge, i, a) for i, a in enumerate(articles)]
            for f in as_completed(futures):
                try: f.result()
                except: pass

        filtered = []
        for a, r in zip(articles, results):
            if r['is_aggregated']:
                a.is_advertisement = True; a.rejection_reason = f"聚合新闻: {r.get('reason','')}"
            elif r['is_commentary']:
                a.is_advertisement = True; a.rejection_reason = f"社会评论: {r.get('reason','')}"
            elif r['is_mess']:
                a.is_advertisement = True; a.rejection_reason = f"主题杂糅: {r.get('reason','')}"
            else:
                filtered.append(a)
        removed = len(articles) - len(filtered)
        _p(f"    移除 {removed} 篇（聚合/评论/杂糅），保留 {len(filtered)} 篇")
        return filtered

    # ---------- 步骤5: TF-IDF去重 + 核心实体匹配 + AI整合 ----------
    def deduplicate_and_merge(self, articles):
        """TF-IDF + 核心实体匹配 → 连通分量 → AI整合（所有簇均融合）"""
        if not articles: return articles
        _p("  TF-IDF+实体去重+整合...")

        sorted_a = sorted(articles, key=lambda x: x.pub_date, reverse=True)
        n = len(sorted_a)

        # 1) 精确去重
        exact = self._find_exact_duplicates(sorted_a)
        _p(f"    精确重复: {len(exact)}对")

        # 2) TF-IDF相似
        cluster_data = [{'title': a.category_tag or a.title, 'ai_summary': a.ai_summary} for a in sorted_a]
        sim_pairs = self.ai.find_similar_groups_v2(cluster_data)
        _p(f"    TF-IDF相似: {len(sim_pairs)}对")

        # 3) 人名/事件匹配
        person_pairs = self._find_same_person_pairs(sorted_a)
        event_pairs = self._find_same_event_pairs(sorted_a)

        # 4) 核心实体匹配（补充TF-IDF对同一大事件不同角度文章的漏检）
        entity_pairs = self._find_core_entity_pairs(sorted_a)
        _p(f"    核心实体匹配: {len(entity_pairs)}对")

        # 合并
        all_pairs = set(exact)
        for p in sim_pairs:
            if len(p) >= 2: all_pairs.add((p[0], p[1]))
        for p in person_pairs: all_pairs.add(p)
        for p in event_pairs: all_pairs.add(p)
        for p in entity_pairs: all_pairs.add(p)

        if not all_pairs:
            _p("    无相似文章，跳过融合"); return sorted_a

        # Union-Find
        parent = list(range(n))
        def find(x):
            while parent[x] != x: parent[x] = parent[parent[x]]; x = parent[x]
            return x
        def union(x, y):
            px, py = find(x), find(y)
            if px != py: parent[px] = py

        for i, j in all_pairs:
            if i < n and j < n: union(i, j)

        groups = defaultdict(list)
        for i in range(n): groups[find(i)].append(i)

        result = []; merged_cnt = 0; skipped_cnt = 0

        for root, indices in groups.items():
            grp = [sorted_a[i] for i in indices]
            if len(indices) == 1:
                result.append(grp[0])
            else:
                # AI验证
                val_data = [{'title': a.category_tag or a.title, 'summary': (a.ai_summary or "")[:500]} for a in grp]
                if not self.ai.verify_same_topic(val_data):
                    result.extend(grp); skipped_cnt += 1; continue

                # 所有簇均走AI融合（每周整合需保留全部信息）
                _p(f"    簇({len(grp)}篇) → AI融合...")
                mr = self._ai_merge_group(grp)
                result.append(mr); merged_cnt += 1

        _p(f"    整合完成: 融合{merged_cnt}组, 跳过{skipped_cnt}组, 产出{len(result)}篇")
        return result

    def _ai_merge_group(self, grp):
        """大簇AI多视角整合"""
        merge_data = [{'title': a.category_tag or a.title, 'summary': a.ai_summary,
                       'link': a.link, 'source_name': a.source_name,
                       'pub_date': a.pub_date.strftime('%Y-%m-%d %H:%M:%S') if a.pub_date else ''} for a in grp]
        mr = self.ai.merge_articles(merge_data)
        return Article(
            source_name=" | ".join(list(set(a.source_name for a in grp))),
            title=mr.get('merged_title', grp[0].title),
            link="(见下方链接汇总)",
            pub_date=min(a.pub_date for a in grp),
            ai_summary=mr.get('merged_content', grp[0].ai_summary),
            category_tag=mr.get('merged_title', grp[0].title),
            mece_category=grp[0].mece_category,
            original_file=grp[0].original_file,
            is_merged=True,
            merged_data={'links': mr.get('links',[a.link for a in grp]),
                         'sources': mr.get('sources',[a.source_name for a in grp]),
                         'timeline': mr.get('timeline',[]),
                         'viewpoints': mr.get('viewpoints',[]),
                         'original_titles': [a.title for a in grp]})

    def _ai_select_best(self, grp):
        """小簇AI选优"""
        dicts = [{'title': a.category_tag or a.title, 'content': a.ai_summary, 'source_name': a.source_name} for a in grp]
        best_idx = 0
        for i in range(1, len(grp)):
            r = self.ai.compare_articles(dicts[best_idx], dicts[i])
            if r == -1: best_idx = i
        best = grp[best_idx]
        # 保留多来源信息
        if len(grp) > 1:
            best.is_merged = True
            best.merged_data = {
                'links': [a.link for a in grp],
                'sources': [a.source_name for a in grp],
                'original_titles': [a.title for a in grp],
            }
        return best

    # ---------- 辅助去重方法 ----------
    def _find_exact_duplicates(self, articles):
        pairs = []; n = len(articles)
        for i in range(n):
            for j in range(i+1, n):
                a1, a2 = articles[i], articles[j]
                if a1.source_name != a2.source_name: continue
                if self._sim(a1.title.lower(), a2.title.lower()) > 0.8:
                    pairs.append((i,j)); continue
                s1, s2 = (a1.ai_summary or "")[:200], (a2.ai_summary or "")[:200]
                if s1 and s2 and self._sim(s1, s2) > 0.9: pairs.append((i,j))
        return pairs

    def _sim(self, t1, t2):
        if not t1 or not t2: return 0.0
        s1, s2 = set(t1), set(t2); u = len(s1|s2)
        return len(s1&s2)/u if u else 0.0

    def _find_same_person_pairs(self, articles):
        pairs = []; n = len(articles)
        gen = {'官员','干部','企业家','商人','老板','明星','网红','公众人物','某人'}
        pat = re.compile(r'([\u4e00-\u9fa5]{2,4})(认罪|悔罪|被捕|被查|被诉|判刑|起诉|立案|调查|开庭|审理|贪污|受贿)')
        for i in range(n):
            for j in range(i+1, n):
                t1 = (articles[i].category_tag or articles[i].title)+" "+(articles[i].ai_summary or "")
                t2 = (articles[j].category_tag or articles[j].title)+" "+(articles[j].ai_summary or "")
                n1 = [p for p,_ in pat.findall(t1) if p not in gen]
                n2 = [p for p,_ in pat.findall(t2) if p not in gen]
                if n1 and n2 and set(n1)&set(n2):
                    pairs.append((i,j)); continue
                # 标题人名
                tn1 = [w for w in re.findall(r'[\u4e00-\u9fa5]{2,3}', articles[i].category_tag or articles[i].title) if w not in gen]
                tn2 = [w for w in re.findall(r'[\u4e00-\u9fa5]{2,3}', articles[j].category_tag or articles[j].title) if w not in gen]
                cn = set(tn1)&set(tn2)
                if len(cn)==1:
                    name = list(cn)[0]
                    if name in t1 and name in t2:
                        if any(v in t1 for v in ['认罪','被捕','被查','判刑','开庭','贪污','受贿']):
                            if any(v in t2 for v in ['认罪','被捕','被查','判刑','开庭','贪污','受贿']):
                                pairs.append((i,j))
        return pairs

    def _find_same_event_pairs(self, articles):
        pairs = []; n = len(articles)
        feats = []
        for a in articles:
            t = (a.category_tag or a.title)+" "+(a.ai_summary or "")
            feats.append({'numbers': set(re.findall(r'\d+', t)), 'text': t})
        for i in range(n):
            for j in range(i+1, n):
                f1, f2 = feats[i], feats[j]
                has_big_num = any(len(x)>=4 for x in f1['numbers']&f2['numbers'])
                p1 = set(re.findall(r'[\u4e00-\u9fa5]{6,15}', f1['text']))
                p2 = set(re.findall(r'[\u4e00-\u9fa5]{6,15}', f2['text']))
                sp = p1&p2
                if has_big_num and len(sp)>=3: pairs.append((i,j)); continue
                kp1 = set(p for p in p1 if len(p)>=7); kp2 = set(p for p in p2 if len(p)>=7)
                if len(kp1&kp2)>=4: pairs.append((i,j))
        return pairs

    def _find_core_entity_pairs(self, articles):
        """基于核心实体匹配，识别讨论同一大主题的文章对（补充TF-IDF对同主题不同角度文章的漏检）

        匹配规则：
        1. 国际冲突规则：通过国家别名映射识别同一冲突（如"美军"→美国、"伊以"→伊朗+以色列）
        2. 组织/产品规则：共享英文组织/产品名 + 任何中文关键词重叠
        3. 特定事件规则：全文关键词重叠度达到阈值
        """
        pairs = []
        n = len(articles)

        # 国家别名映射：别名 → 标准国家名（一个别名可映射到多个国家，如"美以"→美国+以色列）
        COUNTRY_ALIAS_MAP = {
            '美国': ['美国'], '美军': ['美国'], '美方': ['美国'], '白宫': ['美国'],
            '五角大楼': ['美国'], '华盛顿': ['美国'], '特朗普': ['美国'], '拜登': ['美国'],
            '美以': ['美国', '以色列'], '美伊': ['美国', '伊朗'], '美俄': ['美国', '俄罗斯'],
            '美中': ['美国', '中国'], '美朝': ['美国', '朝鲜'],
            '伊朗': ['伊朗'], '伊方': ['伊朗'], '伊军': ['伊朗'], '德黑兰': ['伊朗'],
            '革命卫队': ['伊朗'], '伊以': ['伊朗', '以色列'], '以伊': ['以色列', '伊朗'],
            '以色列': ['以色列'], '以方': ['以色列'], '以军': ['以色列'],
            '内塔尼亚胡': ['以色列'], '特拉维夫': ['以色列'],
            '俄罗斯': ['俄罗斯'], '俄方': ['俄罗斯'], '俄军': ['俄罗斯'],
            '普京': ['俄罗斯'], '莫斯科': ['俄罗斯'], '俄乌': ['俄罗斯', '乌克兰'],
            '乌克兰': ['乌克兰'], '乌方': ['乌克兰'], '乌军': ['乌克兰'], '基辅': ['乌克兰'],
            '朝鲜': ['朝鲜'], '朝方': ['朝鲜'], '平壤': ['朝鲜'], '金正恩': ['朝鲜'],
            '韩国': ['韩国'], '韩方': ['韩国'], '首尔': ['韩国'],
            '日本': ['日本'], '日方': ['日本'], '东京': ['日本'],
            '中国': ['中国'], '中方': ['中国'], '北京': ['中国'],
            '英国': ['英国'], '英方': ['英国'], '伦敦': ['英国'],
            '法国': ['法国'], '法方': ['法国'], '巴黎': ['法国'],
            '德国': ['德国'], '德方': ['德国'], '柏林': ['德国'],
            '印度': ['印度'], '印方': ['印度'], '新德里': ['印度'],
            '沙特': ['沙特'], '利雅得': ['沙特'],
            '土耳其': ['土耳其'], '安卡拉': ['土耳其'],
            '叙利亚': ['叙利亚'], '伊拉克': ['伊拉克'], '阿富汗': ['阿富汗'],
            '巴勒斯坦': ['巴勒斯坦'], '巴基斯坦': ['巴基斯坦'],
            '菲律宾': ['菲律宾'], '越南': ['越南'], '缅甸': ['缅甸'],
            '索马里': ['索马里'], '苏丹': ['苏丹'], '也门': ['也门'], '黎巴嫩': ['黎巴嫩'],
            '埃及': ['埃及'], '利比亚': ['利比亚'],
        }
        MILITARY_TERMS = {
            '战争', '冲突', '打击', '导弹', '无人机', '军事', '交火', '轰炸', '袭击',
            '入侵', '进攻', '反击', '自卫', '停火', '撤军', '部署', '军演',
            '革命卫队', '空袭', '炮击', '坦克', '战斗机', '军舰', '核武',
            '直升机', '坠毁', '击落', '发射', '拦截', '报复', '制裁',
        }
        # 全文关键词提取的停用词
        KEYWORD_STOPWORDS = {
            '的是', '在一', '不了', '如何', '什么', '怎么', '这个', '那个',
            '可以', '已经', '可能', '需要', '我们', '他们', '自己', '因为',
            '所以', '但是', '然而', '虽然', '如果', '就是', '不是', '没有',
            '这样', '那样', '之后', '之前', '之间', '以及', '进行', '通过',
            '关于', '目前', '表示', '认为', '成为', '引发', '关注', '情况',
            '问题', '原因', '影响', '记者', '消息', '报道', '公开', '相关',
            '事件', '回应', '处理', '调查', '官方', '说明', '全国', '各地',
            '社会', '网络', '网友', '不断', '持续', '发布', '宣布', '重要',
            '关键', '核心', '主要', '基本', '显示', '指出', '强调', '提出',
            '分析', '预测', '预计', '未来', '今年', '去年', '本月', '上周',
            '近日', '当天', '此前', '此时', '其中', '为何', '哪些', '怎样',
            '一个', '这些', '那些', '一些', '这种', '那种', '来自', '称为',
            '名为', '属于', '位于', '获得', '受到', '出现', '发生', '开始',
            '结束', '继续', '导致', '造成', '更是', '还有', '而已', '不过',
        }

        # 预提取每篇文章的实体
        article_info = []
        for a in articles:
            t = (a.category_tag or a.title) + " " + (a.ai_summary or "")
            # 通过别名映射提取标准国家名
            countries = set()
            for alias, country_list in COUNTRY_ALIAS_MAP.items():
                if alias in t:
                    countries.update(country_list)
            # 军事术语
            military = {m for m in MILITARY_TERMS if m in t}
            # 英文组织/产品名
            eng_names = set(re.findall(r'[A-Z][a-zA-Z]{2,}', t))
            # 全文中文关键词（2字词，重叠匹配，去停用词）
            all_cn_words = set(re.findall(r'(?=([\u4e00-\u9fa5]{2}))', t)) - KEYWORD_STOPWORDS
            # 标题中文关键词（2字词，重叠匹配，去停用词，权重更高）
            title_cn_words = set(re.findall(r'(?=([\u4e00-\u9fa5]{2}))', a.category_tag or a.title)) - KEYWORD_STOPWORDS
            article_info.append({
                'countries': countries,
                'military': military,
                'eng_names': eng_names,
                'all_cn_words': all_cn_words,
                'title_cn_words': title_cn_words,
            })

        for i in range(n):
            for j in range(i + 1, n):
                e1, e2 = article_info[i], article_info[j]

                # 规则1：国际冲突/事件匹配
                # 1a: 共享≥2个标准国家名 + 任一方有军事术语
                shared_countries = e1['countries'] & e2['countries']
                if len(shared_countries) >= 2 and (e1['military'] or e2['military']):
                    pairs.append((i, j))
                    continue
                # 1b: 共享1个标准国家名 + 双方都有军事术语 + 有中文关键词重叠
                if len(shared_countries) >= 1 and e1['military'] and e2['military']:
                    shared_kw = e1['title_cn_words'] & e2['title_cn_words']
                    if len(shared_kw) >= 1:
                        pairs.append((i, j))
                        continue
                # 1c: 共享≥2个标准国家名 + 有标题关键词重叠（非军事类国际事件，如拒签等）
                if len(shared_countries) >= 2:
                    shared_kw = e1['title_cn_words'] & e2['title_cn_words']
                    if len(shared_kw) >= 2:
                        pairs.append((i, j))
                        continue

                # 规则2：组织/产品事件匹配
                shared_eng = e1['eng_names'] & e2['eng_names']
                if len(shared_eng) >= 2:
                    # 2个以上共享英文名直接匹配
                    pairs.append((i, j))
                    continue
                if len(shared_eng) >= 1:
                    # 1个共享英文名 + 有中文关键词重叠
                    shared_kw = e1['all_cn_words'] & e2['all_cn_words']
                    if len(shared_kw) >= 2:
                        pairs.append((i, j))
                        continue

                # 规则3：全文关键词重叠（非泛化词）
                shared_all_kw = e1['all_cn_words'] & e2['all_cn_words']
                if len(shared_all_kw) >= 8:
                    pairs.append((i, j))
                    continue

        return pairs

    # ---------- 步骤6: 小热点过滤（纯规则） ----------
    def filter_minor_hotspots(self, articles):
        filtered = []
        for a in articles:
            if a.mece_category.startswith('11'):
                t = (a.category_tag or a.title)+" "+(a.ai_summary or "")
                minor = False
                for p in [r'某[省市县区镇].{2,6}(医院|学校|药店)', r'[^\x00-\x7F]{2,4}[市县][^\x00-\x7F]{2,6}(中学|小学|医院)',
                          r'某[地校院].{2,8}(改制|合并|改革)', r'[^\x00-\x7F]{2,4}区[^\x00-\x7F]{2,6}(小学|中学|幼儿园)']:
                    if re.search(p, t):
                        if not any(k in t for k in ['国务院','教育部','国家卫健委','发改委','全国','中央']): minor = True; break
                if minor: continue
            filtered.append(a)
        return filtered

    # ---------- 主流程 ----------
    def process_weekly_news(self, target_date=None):
        try:
            import json as _json  # 确保 json 可用
            ws, we = self.get_week_range(target_date)
            _p(f"\n{'='*70}")
            _p(f"每周资讯: {ws.strftime('%Y-%m-%d')} ~ {we.strftime('%Y-%m-%d')}")
            _p(f"{'='*70}")

            # 1. 收集
            _p("\n[1/6] 收集文章...")
            self.articles = self.collect_weekly_articles(ws, we)
            if not self.articles: logger.warning("无文章"); return False
            _p(f"  共 {len(self.articles)} 篇")

            # 2. 规则过滤（极快，纯规则）
            _p("\n[2/6] 规则噪音过滤...")
            self.articles = self.filter_noise_rules(self.articles)
            if not self.articles: logger.warning("规则过滤后无文章"); return False

            # 3. 批量AI广告过滤（8线程）
            _p("\n[3/6] 批量AI广告过滤...")
            self.articles = self.batch_filter_ai_advertisement(self.articles)

            # 4. 批量AI聚合/评论/杂糅过滤（8线程，三合一）
            _p("\n[4/6] 批量AI聚合/评论/杂糅过滤...")
            self.articles = self.batch_filter_combined(self.articles)

            # 5. TF-IDF去重+整合
            _p("\n[5/6] TF-IDF去重+AI整合...")
            self.articles = self.deduplicate_and_merge(self.articles)

            # 6. 批量MECE分类 + 小热点过滤 + 清理 + 生成
            _p("\n[6/6] 批量MECE分类 + 过滤 + 生成文档...")
            # 批量分类
            cls_data = [{'title': a.category_tag or a.title, 'summary': a.ai_summary} for a in self.articles]
            categories = self.ai.batch_classify_mece(cls_data)
            for i, a in enumerate(self.articles):
                if i < len(categories): a.mece_category = categories[i]
            # 小热点过滤
            self.articles = self.filter_minor_hotspots(self.articles)
            # 摘要清理
            _clean_summaries(self.articles)
            # 生成文档
            os.makedirs(self.output_dir, exist_ok=True)
            out_fn = f"每周资讯_{ws.strftime('%Y%m%d')}_{we.strftime('%Y%m%d')}.docx"
            out_path = os.path.join(self.output_dir, out_fn)
            ok = DocumentGenerator.create_document(self.articles, out_path, ws, we)
            if ok: _p(f"\n生成成功: {out_path}\n共 {len(self.articles)} 篇")
            return ok
        except Exception as e:
            logger.error(f"处理出错: {e}"); import traceback; traceback.print_exc(); return False


# ====================== main ======================
def main(target_date=None, custom_daily_dir=None, custom_output_dir=None):
    try:
        import json as _json
        key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
        if not key:
            try:
                sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
                from auth import get_api_key
                key = get_api_key() or ""
            except Exception as e:
                logger.error(f"加载授权模块失败: {e}")
            if not key:
                print("[错误] 未获取到API密钥"); return False

        cur = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        par = os.path.dirname(cur)
        daily_dir = custom_daily_dir or os.path.join(par, "每日资讯")
        output_dir = custom_output_dir or os.path.join(par, "每周资讯")

        _p("="*70); _p("每周资讯Skill v2 - 并行去重整合版"); _p(f"开始: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"); _p("="*70)

        proc = WeeklyNewsProcessor(daily_dir, output_dir, key)
        ok = proc.process_weekly_news(target_date)
        if ok:
            _p("\n"+"="*70); _p("[完成] 任务成功！"); _p(f"收录 {len(proc.articles)} 篇"); _p("="*70)
        else:
            _p("\n[错误] 任务失败")
        return ok
    except Exception as e:
        logger.error(f"程序出错: {e}", exc_info=True); _p(f"\n[错误] {e}"); return False


if __name__ == "__main__":
    exit(0 if main() else 1)

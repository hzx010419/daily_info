#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
每日资讯skill主程序
使用DeepSeek AI进行广告判断和摘要生成
"""

import warnings
import pandas as pd
import requests
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
import time
import os
import sys
import re
import json
import logging
from typing import List, Optional, Tuple
from dataclasses import dataclass

warnings.filterwarnings('ignore')

# 导入AI客户端
from ai_client import DeepSeekClient, SIMILARITY_THRESHOLD, HIGH_SIM_THRESHOLD, MEDIUM_SIM_THRESHOLD, LOW_SIM_THRESHOLD

# 并行摘要生成器
from concurrent.futures import ThreadPoolExecutor, as_completed

class ParallelSummaryGenerator:
    """并行摘要生成器"""

    def __init__(self, ai_client, max_workers: int = 8):
        self.ai_client = ai_client
        self.max_workers = max_workers

    def generate_summaries(self, articles: List) -> None:
        """并行生成摘要"""
        articles_to_process = [
            a for a in articles
            if not a.is_advertisement and not getattr(a, '_from_cache', False)
        ]
        if not articles_to_process:
            logger.info("没有需要生成摘要的文章")
            return

        total = len(articles_to_process)
        logger.info(f"并行生成摘要开始，共 {total} 篇，最大并行数 {self.max_workers}")

        def generate_single(article):
            try:
                summary, category_tag = self.ai_client.generate_summary_with_retry(
                    article.title, article.full_content
                )
                return article, summary, category_tag, None
            except Exception as e:
                logger.error(f"摘要生成失败: {article.title[:30]}... 错误: {e}")
                return article, "摘要生成失败", "", str(e)

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = {executor.submit(generate_single, a): a for a in articles_to_process}
            completed = 0
            for future in as_completed(futures):
                article, summary, category_tag, error = future.result()
                article.ai_summary = summary
                if category_tag:
                    article.category_tag = category_tag
                completed += 1
                if error:
                    logger.warning(f"[{completed}/{total}] {article.title[:30]}... 失败")
                else:
                    logger.info(f"[{completed}/{total}] {article.title[:30]}... -> {len(summary)}字")

        logger.info(f"并行生成摘要完成，共 {completed} 篇")

# 设置日志
import os
current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
logs_dir = os.path.join(current_dir, 'logs')
os.makedirs(logs_dir, exist_ok=True)

log_file = os.path.join(logs_dir, 'app.log')


class _FlushStreamHandler(logging.StreamHandler):
    """每次写入后立刻 flush，前台终端能实时看到日志。"""
    def emit(self, record):
        super().emit(record)
        try:
            self.flush()
        except Exception:
            pass


# 控制台 handler：写到 stdout（和 print 同一流，顺序不错乱）并实时 flush
_console_handler = _FlushStreamHandler(stream=sys.stdout)
_console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

_file_handler = logging.FileHandler(log_file, encoding='utf-8')
_file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

logging.basicConfig(
    level=logging.INFO,
    handlers=[_file_handler, _console_handler],
    force=True,  # 覆盖可能已有的 basicConfig
)
logger = logging.getLogger(__name__)

# ============ 重大国际事件关键词常量（去重/整合多处复用，避免重复定义） ============
MAJOR_EVENT_KEYWORD_GROUPS = [
    {
        'name': '美伊战争/冲突/和谈',
        'keywords': ['美伊', '伊朗', '霍尔木兹', '海峡', '鲁比奥', '停火', '和谈', '特朗普', '核问题', '制裁', '解冻',
                     '美伊战争', '空袭', '轰炸', '美军', '打击伊朗', '伊朗反击', '伊朗导弹', '伊朗报复',
                     '德黑兰', '波斯湾', '伊核', '伊朗核', '军事打击', '军事行动', '战争', '开战',
                     '革命卫队', '石油', '导弹', '约旦', '商船', '谈判', '袭击',
                     '摧毁', '毁灭性', '代价', '重建', '援助', '关闭海峡', '通行',
                     '运出', '港口', '基地', '报复性', '反击', '进攻', '撤军',
                     '以伊', '以色列伊朗', '以军', '空防', '防空', '拦截',
                     # 【修复】补充外交/和平/协议相关关键词，解决谅解备忘录、和平协议等文章未整合的问题
                     '谅解', '备忘录', '和平协议', '浓缩铀', '护航', '稀释', '赔偿', '撤出',
                     '伊美', '草案', '圣城旅', '幽灵', '封锁', '突破封锁', '原油',
                     '和平', '协议', '14点', '最终文本', '航运', '经济议题',
                     # 【修复】补充停战/签署/投降相关关键词，解决美伊停战谅解备忘录等和平协议文章未能被匹配的问题
                     '停战', '签署', '签字', '瑞士', '投降', '战争结束', '恢复通行', '伊美协议'],
        # 【关键修复】锚定关键词：只有两篇文章共享至少1个锚定关键词时，才允许通过该事件组建立连接。
        # 锚定关键词是该事件独有的、能唯一标识该事件的词；不含'战争''制裁''导弹'等泛化冲突词，
        # 避免俄乌、日韩等不同冲突事件因共享泛化词而被错误连通到同一簇。
        'anchor_keywords': ['美伊', '伊朗', '霍尔木兹', '鲁比奥', '德黑兰', '波斯湾', '伊核', '伊朗核',
                           '打击伊朗', '伊朗反击', '伊朗导弹', '伊朗报复', '革命卫队', '以伊', '以色列伊朗',
                           # 【修复】补充外交/和平特有锚定词
                           '谅解', '备忘录', '浓缩铀', '护航', '圣城旅', '和平协议', '伊美',
                           # 【修复】补充停战/签署特有锚定词
                           '停战', '瑞士', '投降', '恢复通行', '伊美协议'],
        'min_overlap': 2,
        'is_major_event': True,
    },
    {
        'name': '朝核/朝鲜半岛',
        'keywords': ['朝鲜', '朝核', '金正恩', '六方会谈', '半岛'],
        'anchor_keywords': ['朝鲜', '朝核', '金正恩'],
        'min_overlap': 2,
    },
    {
        'name': '俄乌冲突',
        'keywords': ['俄乌', '乌克兰', '泽连斯基', '克里米亚', '顿巴斯'],
        'anchor_keywords': ['俄乌', '乌克兰', '泽连斯基', '克里米亚', '顿巴斯'],
        'min_overlap': 2,
    },
    {
        'name': '以巴/以黎冲突',
        'keywords': ['以色列', '真主党', '加沙', '哈马斯', '内塔尼亚胡', '以黎'],
        'anchor_keywords': ['以色列', '真主党', '加沙', '哈马斯', '内塔尼亚胡', '以黎'],
        'min_overlap': 2,
    },
]

# 所有重大事件关键词的合集（用于 _merge_large_cluster 实体补充）
# 【关键修复】仅收集锚定关键词（事件特有词），不包含'战争''制裁'等泛化冲突词，
# 避免不同事件的文章因共享泛化词而被误判为同主题
_MAJOR_EVENT_ALL_KEYWORDS = set()
_MAJOR_EVENT_ANCHOR_KEYWORDS = set()  # 仅锚定关键词的合集
for _meg in MAJOR_EVENT_KEYWORD_GROUPS:
    _MAJOR_EVENT_ALL_KEYWORDS.update(_meg['keywords'])
    _anchor_kw = _meg.get('anchor_keywords', None)
    if _anchor_kw:
        _MAJOR_EVENT_ANCHOR_KEYWORDS.update(_anchor_kw)
    else:
        # 兼容：没有定义anchor_keywords的事件组，全部视为锚定关键词
        _MAJOR_EVENT_ANCHOR_KEYWORDS.update(_meg['keywords'])





def _p(msg: str) -> None:
    """前台进度打印：统一的实时进度输出入口，立刻 flush。"""
    try:
        print(msg, flush=True)
    except (TypeError, BlockingIOError, OSError):
        # TypeError: 某些环境 print 不接受 flush kwarg
        # BlockingIOError: 管道缓冲区满，写操作会阻塞
        # OSError: 其他I/O错误
        try:
            sys.stdout.flush()
        except Exception:
            pass

@dataclass
class Article:
    """文章数据结构"""
    source_name: str
    title: str
    link: str
    pub_date: datetime
    full_content: str
    ai_summary: str = ""
    is_advertisement: bool = False
    rejection_reason: str = ""
    category_tag: str = ""  # 摘要开头的【主题标签】，用于去重和显示
    mece_category: str = ""  # MECE分类编号（如"1.1"、"2.3"等）
    is_merged: bool = False  # 是否为多来源整合文章
    merged_sources: list = None  # 整合来源列表

    # MECE分类名称映射
    MECE_CATEGORIES = {
        "1": "国际局势与地缘政治",
        "1.1": "地区冲突与战争",
        "1.2": "国际关系与外交",
        "1.3": "全球安全与反恐",
        "2": "宏观经济与金融",
        "2.1": "国内宏观经济",
        "2.2": "宏观消费与零售",
        "2.3": "资本市场",
        "2.4": "房地产市场",
        "2.5": "全球金融",
        "3": "文体娱乐内容",
        "3.1": "文化产业",
        "3.2": "体育休闲",
        "3.3": "体验消费",
        "3.4": "网络与数字文化",
        "4": "科技产业动态与创新",
        "4.1": "人工智能",
        "4.2": "信息技术",
        "4.3": "智能制造与机器人",
        "4.4": "新能源与绿色科技",
        "5": "医疗健康与生命科学",
        "5.1": "医疗改革与政策",
        "5.2": "医学研究与突破",
        "5.3": "公共卫生",
        "6": "教育发展与人才培养",
        "6.1": "教育政策与改革",
        "6.2": "高等教育",
        "6.3": "教育科技",
        "6.4": "就业与人才",
        "7": "社会民生与消费",
        "7.1": "民生消费",
        "7.2": "就业与职场",
        "7.3": "人口与家庭",
        "7.4": "社会治理",
        "8": "企业与商业",
        "8.1": "企业经营",
        "8.2": "商业动态",
        "8.3": "电商与新零售",
        "8.4": "创业与创新",
        "9": "政策法规与监管",
        "9.1": "国家政策",
        "9.2": "行业监管",
        "9.3": "地方治理",
        "9.4": "反腐与廉政",
        "10": "能源与资源",
        "10.1": "传统能源",
        "10.2": "新能源",
        "10.3": "资源与环境",
        "11": "社会热点与舆论动态",
        "11.1": "热点事件",
        "11.2": "舆情分析",
        "11.3": "跨界热点",
    }

    # 高质量来源白名单 - 这些来源的文章慎重过滤，规则匹配命中时需要更强的证据才过滤
    HIGH_QUALITY_SOURCES = [
        'New Economist', '新经济学家', '财新', '财经', '经济学人',
        '三联生活周刊', '南方周末', '新京报', '第一财经', '界面',
        '中国新闻周刊', '瞭望', '半月谈',
        # 行业垂直/高质量观察类，之前被误杀的来源
        '麦可思研究', '中美聚焦', '智谷趋势', '秦朔朋友圈', '吴晓波频道',
        '财经十一人', '虎嗅APP', '钛媒体', '36氪',
    ]

    # 跨领域分析型标题特征词 —— 标题出现这些词说明文章天然要跨多个领域
    # （如"战争助推能源战略"、"AI冲击就业"），不应判为"主题杂糅"
    CROSS_DOMAIN_ANALYSIS_INDICATORS = [
        '助推', '推动', '影响', '冲击', '带动', '拉动', '倒逼',
        '重塑', '改写', '牵动', '波及', '传导',
        '背景', '背后', '启示', '机遇', '挑战', '博弈', '关系',
        '战略', '格局', '趋势', '走向', '变局', '演变',
        '如何影响', '意味着', '背后逻辑', '深层',
    ]

    def is_noise_content(self) -> tuple[bool, str]:
        """
        判断是否为噪音内容（鸡汤文、医学案例、综合式文章、地方新闻等）
        返回: (是否为噪音, 原因)
        """
        import re
        text = self.title + " " + (self.ai_summary[:500] if self.ai_summary else "")
        
        # 0. 摘要生成失败直接过滤
        if self.ai_summary == "摘要生成失败":
            return True, "摘要生成失败"
        
        # 0.0 原文内容中断/不完整的摘要直接过滤
        incomplete_patterns = [
            r'原文内容在此中断',
            r'原文未提供完整信息',
            r'摘要仅能涵盖已给出的部分',
            r'无法补充未提及的细节',
            r'无法补充未提及的细节或后续发展',
        ]
        for pattern in incomplete_patterns:
            if re.search(pattern, self.ai_summary or ""):
                return True, f"原文内容中断/摘要不完整: 匹配'{pattern}'"
        
        # 0.1 检查是否为高质量来源，如果是则放宽过滤标准
        is_high_quality = any(src in self.source_name for src in Article.HIGH_QUALITY_SOURCES)
        
        # 1. 鸡汤文/个人感悟类
        chicken_soup_patterns = [
            r'高敏感', r'过度努力', r'心理内耗', r'职场困境',
            r'情绪.{0,4}(失控|崩溃|低落|内耗|勒索|化|价值|垃圾|垃圾桶)',  # 鸡汤文语境的"情绪"
            r'负面情绪',  # 明确的个人负面情绪
            r'(管理|控制|调节|照顾|安抚).{0,4}情绪',  # 情绪管理类鸡汤
            r'感悟',
            r'人生哲理', r'励志', r'成功学', r'心得', r'自我反思', r'个人成长',
            r'学会.*相处', r'如何.*面对', r'心态', r'幸福感', r'心理.*建议'
        ]
        # 财经语境下"情绪"的常见搭配（市场情绪、投资者情绪等），出现则不认为是鸡汤
        financial_emotion_contexts = ['市场情绪', '投资者情绪', '投资情绪', '情绪指数', '情绪指标',
                                       '情绪面', '情绪偏弱', '情绪回暖', '情绪低迷', '情绪修复',
                                       '情绪高涨', '情绪乐观', '情绪悲观', '恐慌情绪', '风险情绪',
                                       '消费情绪', '交易情绪', '情绪发酵', '情绪蔓延']
        for pattern in chicken_soup_patterns:
            if re.search(pattern, text):
                # 如果命中情绪相关模式，但文本中存在财经语境的情绪搭配，则跳过
                if '情绪' in pattern and any(ctx in text for ctx in financial_emotion_contexts):
                    continue
                # 检查是否缺乏实质内容（含财经上下文关键词）
                if not any(kw in text for kw in ['数据', '政策', '分析', '研究', '报告', '统计', '调查',
                                                   '市场', '股市', '行情', '投资', '交易', '经济', '金融',
                                                   '股价', '指数', '熔断', '崩跌', '暴跌', '大涨']):
                    return True, f"鸡汤文/个人感悟类: 匹配'{pattern}'"
        
        # 2. 医学案例水文类（具体病例故事）
        medical_case_patterns = [
            r'\d+岁.{0,10}(患者|女士|先生|男子|女孩|男孩)',  # 某岁患者
            r'罕见.{0,10}(病例|畸形|疾病|肿瘤|综合症)',  # 罕见病例
            r'(双子宫|双阴道|心脏骤停|癌症晚期|绝症)',  # 具体病症名
            r'手术.{0,10}(成功|顺利|完成)',  # 手术成功
            r'经历.{0,5}(手术|化疗|治疗)',  # 经历治疗
            r'某某医院.{0,10}(专家|主任|医生)',  # 某某医院医生
        ]
        for pattern in medical_case_patterns:
            if re.search(pattern, text):
                # 排除有公共卫生意义的文章
                if not any(kw in text for kw in ['疫情', '公共卫生', '防控', '疫苗', '药物审批', '医保', '政策', '突破', '治愈', '基因编辑']):
                    return True, f"医学案例水文类: 匹配'{pattern}'"
        
        # 3. 综合式/大杂烩文章（已合并到规则#10和#9.1，此处删除重复逻辑）
        # 原规则#3与规则#10（连接词检测）和规则#9.1（摘要多领域检测）逻辑重复
        # 且容易误杀深度分析文章（如"沃什提名听证会"涉及金融+国际军事，但主题是单一的）
        # 现在统一由规则#9.1和#10处理，此处不再单独判断
        
        # 4. 餐饮业非专业化分析点评类
        # 识别餐饮行业分析、地方小餐饮店、产业观察等缺乏重大新闻价值的内容
        # 注意：只匹配明确是餐饮业的关键词，避免误伤文旅等其他行业
        restaurant_analysis_patterns = [
            r'餐饮店', r'餐饮业', r'餐饮行业', r'餐饮市场', r'餐厅.{0,5}(扩张|关门|倒闭)',
            r'人均.{0,5}(消费|餐饮)', r'快餐店', r'小吃店', r'外卖.{0,5}(平台|骑手|小哥)',
            r'厨师', r'翻台率', r'坪效', r'预制菜.{0,5}(争议|标准|监管)',
        ]
        for pattern in restaurant_analysis_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏重大事件或政策
                if not any(kw in text for kw in ['政策', '监管', '处罚', '重大', '上市', '收购', '融资', '倒闭', '食品安全']):
                    return True, f"餐饮业非专业化分析: 匹配'{pattern}'"
        
        # 5. 一般院校/技师学院/不知名学校动态
        # 识别非重点高校、地方院校、技师学院的普通动态
        # 注意：院校升格/更名/合并等属于常规行政操作，即使提到"教育部批复"也不是重大新闻
        common_school_patterns = [
            r'技师学院', r'职业技术学院', r'专科学校', r'职业学院',
            r'学院.{0,5}(并入|合并|组建|成立|揭牌)',  # 一般院校的机构调整
            r'大学.{0,5}(并入|合并|组建|成立|揭牌|对接)',  # 大学机构调整
            r'(学院|大学|学校).{0,10}(升格|升本|升为本科|更名为|改名)',  # 院校升格/更名
            r'升格.{0,5}(本科|大学)',  # 升格本科/大学
            r'更名.{0,5}(大学|学院|职业)',  # 院校更名
        ]
        for pattern in common_school_patterns:
            if re.search(pattern, text):
                # 院校动态只有在涉及真正重大事件（违规/犯罪/舆情等）时才保留
                # "教育部批复"只是行政流程，不算重大事件
                if not any(kw in text for kw in ['违规', '丑闻', '犯罪', '舆情', '热搜', '重大事故', '贪污', '腐败', '性侵', '学术不端']):
                    return True, f"一般院校动态: 匹配'{pattern}'"
        
        # 5.5 不知名学校的普通福利/慰问/奖学金/补贴事件
        # 识别非985/211/双一流院校的教职工福利、慰问金、带薪假等普通事件
        # 如"江西服装学院为教职工子女高考发5000元慰问金"——普通民办学校的内部福利，无全国新闻价值
        # 知名高校列表：985/211/双一流核心院校，这些学校的动态可能有新闻价值
        well_known_universities = ['清华', '北大', '人民大学', '北航', '北师大', '北理工', '中国农大',
                                    '中央民大', '复旦', '上海交大', '同济', '华东师大', '南开', '天大',
                                    '浙大', '南大', '东南大学', '中科大', '武大', '华科', '中南大学',
                                    '湖南大学', '中山', '华南理工', '川大', '电子科大', '重大', '西交大',
                                    '西北工大', '兰大', '哈工大', '吉大', '大连理工', '东北大学',
                                    '国防科大', '厦大', '山大', '中国海洋', '中央财经', '上海财经',
                                    '对外经贸', '政法大学', '外交学院', '北外', '北邮', '北交大',
                                    '中科院', '社科院', '港大', '港中文', '港科技', '台大']
        is_well_known_school = any(sch in text for sch in well_known_universities)
        if not is_well_known_school:
            school_welfare_patterns = [
                r'(学院|学校|职校|院校).{0,20}(慰问金|带薪假|福利|奖金|补贴|慰问|发放|慰问品)',
                r'(教职工|教师|老师).{0,15}(慰问金|带薪假|福利|奖金|补贴|慰问品)',
                r'(学院|学校).{0,20}(高考|子女).{0,20}(慰问|福利|带薪|补贴|奖金)',
                r'(为|向).{0,10}(教职工|教师).{0,15}(子女|高考).{0,15}(发放|提供|慰问|带薪)',
            ]
            for pattern in school_welfare_patterns:
                if re.search(pattern, text):
                    # 排除：涉及全国性政策/重大改革/舆情事件
                    if any(kw in text for kw in ['教育部', '国务院', '全国', '政策', '改革', '违规', '犯罪', '舆情', '热搜']):
                        continue
                    return True, f"不知名学校普通福利事件: 匹配'{pattern}'"
        
        # 6. 小范围不知名学校/单位的普通事件
        # 识别具体地方学校、不知名单位的一般性事件
        local_school_patterns = [
            r'中学.{0,10}(公开批评|通报|处理|回应)',  # 不知名中学的一般事件
            r'学校.{0,10}(乱扔垃圾|学生.*行为|师德|处分)',  # 学校学生管理类事件
            r'小学.{0,10}(学生|教师|家长)',  # 小学的日常事件
        ]
        for pattern in local_school_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏重大性
                if not any(kw in text for kw in ['教育部', '国务院', '重大', '舆情', '热搜', '全国', '政策', '犯罪', '欺凌']):
                    return True, f"小范围学校事件: 匹配'{pattern}'"
        
        # 7. 职场盘点/总结/指南类水文 + 职业日常/职场生活类
        workplace_water_patterns = [
            r'职场.{0,10}(减分项|加分项|禁忌|雷区|注意事项|建议|盘点)',
            r'(医生|护士|教师|公务员|员工).{0,10}(表现|行为|特征|禁忌|雷区|盘点)',
            r'(职场|工作中).{0,10}(不成熟|成熟|成熟表现|加分|减分)',
            r'如何.{0,10}(表现|避免|应对|处理).{0,20}(职场|工作)',
            r'职场.*(生存|法则|定律|法则|定律)',
            r'盘点.{0,5}(常见|七种|八种|九种|十种)',
            # 职业日常/职场生活类：描述某个职业忙、累、辛苦、好不好等缺乏新闻价值的文章
            r'(医生|护士|教师|公务员|程序员|律师|记者|警察|消防员).{0,10}(过得好不好|好不好|辛苦|忙碌|忙|累|不容易)',
            r'(医生|护士|教师|公务员).{0,10}(一眼.{0,5}(看出|看透|看穿)|脱下.*白大褂|下班后|真实状态|日常|生活状态)',
            r'(忙起来|忙到).{0,15}(顾不上|没时间|来不及|喝不上水|来不及吃饭)',
            r'(诊室|科室|病房|办公室).{0,10}(排.{0,5}队|忙碌|连轴转|加班)',
            r'(房贷|接孩子|伴侣|争吵).{0,10}(日常|生活|代价)',
            r'(被.{0,5}需要|满足感).{0,15}(代价|忙|累|辛苦)',
        ]
        for pattern in workplace_water_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏具体数据和案例
                if not any(kw in text for kw in ['数据', '报告', '研究', '统计', '政策', '案例', '重大', '改革', '制度', '行业分析', '趋势']):
                    # 高质量来源放宽：职业日常文章即使来自高质量来源也不应保留
                    return True, f"职场盘点/职业日常类: 匹配'{pattern}'"
        
        # 8. 地方小新闻类
        # 识别"某市某区/县"类地方新闻
        local_patterns = [
            r'[^\x00-\x7F]{2,4}省[^\x00-\x7F]{2,6}市',  # 某省某市
            r'[^\x00-\x7F]{2,4}市[^\x00-\x7F]{2,6}区',  # 某市某区
            r'[^\x00-\x7F]{2,4}市[^\x00-\x7F]{2,6}县',  # 某市某县
            r'[^\x00-\x7F]{2,6}(学院|大学|学校|医院)',  # 地方院校/医院
        ]
        local_hits = 0
        for pattern in local_patterns:
            if re.search(pattern, text):
                local_hits += 1
        
        # 如果命中多个地方特征，且不包含全国性关键词
        if local_hits >= 2:
            national_keywords = ['国务院', '教育部', '国家卫健委', '发改委', '央行', '证监会', '银保监会', '外交部', '商务部']
            # 【关键修复】重大公共安全/恶性事件豁免：即使是地方报道，只要涉及重大伤亡/事故/恶性
            # 公共事件（如煤矿爆炸、矿难、踩踏、客机事故、严重火灾、官员重大违纪等），都应保留
            major_event_keywords = [
                # 重大事故/灾害
                '煤矿', '矿难', '瓦斯爆炸', '塌方', '坍塌', '矿井', '透水', '尾矿',
                '坠机', '空难', '失事', '客机', '客船', '沉船', '翻船',
                '重大事故', '重大伤亡', '重大火灾', '严重火灾', '化工厂.{0,5}爆炸', '危化品',
                '爆炸', '踩踏', '群体性',
                # 严重伤亡/死亡数字
                r'\d+人.{0,5}(死亡|遇难|失联|被困|罹难)',
                r'(死亡|遇难|失联|罹难).{0,5}\d+人',
                # 重大违法/腐败/反腐
                '副省级', '副部级', '正部级', '中管干部', '严重违纪违法', '中央纪委',
                '一把手.{0,5}被查', '落马',
                # 重大公共卫生
                '聚集性疫情', '传染病', '重大食安', '集体中毒', '公共卫生事件',
            ]
            has_major_event = any(re.search(kw, text) for kw in major_event_keywords)
            if has_major_event:
                pass  # 重大事件豁免，不过滤
            elif not any(kw in text for kw in national_keywords):
                # 检查是否是具体地方高校/医院动态
                if re.search(r'[^\x00-\x7F]{4,20}(学院|大学|医院|职校)', text):
                    return True, f"地方小新闻类: 包含{local_hits}个地方特征"
        
        # 7.5 股市投资建议/技术分析类（针对股市给出操作建议的文章）
        stock_advice_patterns = [
            r'(建议|操作).{0,10}(观望|买入|卖出|减仓|加仓|清仓|持仓|建仓|调仓)',
            r'(观望为主|多看少做|不操作|少操作)',
            r'(市场.*收出.{0,5}(中阴线|大阴线|中阳线|大阳线))',
            r'(分钟.*(低点|高点|结构|顶部|底部))',
            r'(日线.*(顶部|底部|结构|高点|低点))',
            r'(短线操作难度|操作难度较大)',
            r'(下跌.*级别|高点级别|顶部结构)',
            r'(速度.*降下来|急跌.*缓跌)',
            r'(继续.*观望|以观望为主)',
            r'(市场.{0,10}(选择向下|向下运行|运行过程中))',
            r'(A股|股市|大盘|沪深).{0,20}(观望|买入|卖出|减仓|操作建议|仓位)',
            r'(反弹.*减仓|回调.*加仓|突破.*追涨)',
            r'(支撑位|压力位|阻力位|关键位|止损位)',
            r'(MACD|KDJ|RSI|布林带|均线.*金叉|均线.*死叉)',
            r'(缩量.{0,5}(横盘|整理|下跌)|放量.{0,5}(上涨|突破|拉升)|横盘.{0,5}(整理|突破))',
        ]
        for pattern in stock_advice_patterns:
            if re.search(pattern, text):
                # 如果包含宏观经济政策分析（如央行政策、利率变化等），则保留
                if any(kw in text for kw in ['央行', '利率', '货币政策', '降准', '降息', '加息', 'LPR', '政策', '监管', '改革']):
                    continue
                return True, f"股市投资建议类: 匹配'{pattern}'"
        
        # 8.4 缅怀纪念/人物传记类（逝世周年纪念、追忆、回顾等，不是新闻）
        memorial_patterns = [
            r'(逝世|去世|离世).{0,10}(周年|纪念|追忆|缅怀|怀念)',
            r'(缅怀|追忆|怀念|纪念).{0,10}(逝世|去世|离世|伟人|先辈|大师)',
            r'(逝世.{0,5}(五|十|百|一|二|三|四|五|六|七|八|九|十)周年)',
            r'(亲授弟子|关门弟子).{0,10}(追忆|回忆|怀念|缅怀)',
            r'(追忆|回忆|怀念|缅怀).{0,10}(一生|鲜为人知|不为人知|另一面)',
            r'(鲜为人知|不为人知).{0,10}(一面|故事|细节)',
            r'(爱地摊|老小孩|老顽童)',  # 典型的人物轶事描述
            r'(诞辰|冥诞).{0,10}(周年|纪念|缅怀|追忆)',
            r'(永远怀念|深切缅怀|致敬|天堂|安息)',
            r'(人生.{0,5}(故事|传奇|轨迹|经历)|传奇一生)',
            r'(传记|人物志|人物特写|人物故事)',
        ]
        for pattern in memorial_patterns:
            if re.search(pattern, text):
                # 排除：如果文章涉及当前仍在发展的重大事件（不只是纪念）
                if any(kw in text for kw in ['突破', '发现', '政策', '改革', '调查', '起诉', '审判', '逮捕', '制裁', '冲突']):
                    continue
                return True, f"缅怀纪念/人物传记类: 匹配'{pattern}'"
        
        # 8.45 无新闻价值的生活方式/酒店介绍类（仅描述消费场景，无重大事件）
        lifestyle_no_news_patterns = [
            r'(每晚|一晚|房价|房价|房费).{0,10}(万元|万块|万起|千元)',
            r'(套房|房型|豪华套房|皇家套房|总统套房).{0,20}(面积|价格|装修|配备)',
            r'(酒店.{0,10}(开业|房间数|入住|房型|设施|水疗|米其林|餐厅))',
            r'(平方米.*露台|独立电梯.*直达)',
            r'(品牌.{0,10}(扩张|入驻|开业|开业))',
        ]
        for pattern in lifestyle_no_news_patterns:
            if re.search(pattern, text):
                # 排除：如果涉及重大事件（安全事故、重大政策、丑闻等）
                if any(kw in text for kw in ['事故', '安全', '火灾', '坍塌', '死亡', '伤亡', '处罚', '违规', '丑闻', '倒闭', '破产', '政策', '监管', '整治']):
                    continue
                # 排除：高质量来源的深度行业分析
                if is_high_quality and any(kw in text for kw in ['行业', '市场', '趋势', '分析', '数据', '报告']):
                    continue
                return True, f"无新闻价值的生活方式类: 匹配'{pattern}'"
        
        # 8.5 书籍/小说分析评论类（不是新闻，没有信息增量）
        book_review_patterns = [
            r'(小说|长篇小说|短篇|散文|诗集|纪实文学|传记).{0,10}(分析|解读|评论|赏析|读后|推荐|书评)',
            r'(读完|看完).{0,10}(这本|这部).{0,5}(书|小说|作品)',
            r'(作家|作者).{0,10}(创作|写作|笔下)',
            r'(主人公|主角).{0,10}(奋斗|成长|经历|故事)',
            r'(侨批|史诗|文学).{0,10}(描绘|讲述|展现)',
            r'(同根同源|同源异流)',
        ]
        for pattern in book_review_patterns:
            if re.search(pattern, text):
                return True, f"书籍分析评论类: 匹配'{pattern}'"
        
        # 8.6 琐碎/无新闻价值类（读音争议、字词用法等无信息增量的内容）
        trivial_content_patterns = [
            r'(读音|发音|念错|怎么读|怎么念).{0,10}(争议|辩论|讨论|正确|标准)',
            r'(字典|词典).{0,10}(收录|未收录|标准读音)',
            r'(方言|口音).{0,10}(读音|发音|念法|说法)',
            r'(嬷|嫲|嘛).{0,10}(读音|发音|mó|mà)',
            r'(错别字|用字|用词).{0,10}(争议|纠正|规范)',
        ]
        for pattern in trivial_content_patterns:
            if re.search(pattern, text):
                return True, f"琐碎无新闻价值类: 匹配'{pattern}'"
        
        # 8.7 中医药/偏方汇总类（养生偏方汇总，无新闻价值）
        chinese_medicine_patterns = [
            r'(中成药|中药|偏方|秘方).{0,10}(汇总|总结|盘点|大全|清单|推荐)',
            r'(失眠|头痛|胃病|高血压|糖尿病).{0,10}(中成药|中药|偏方|方剂)',
            r'(胶囊|丸|散|汤|膏).{0,10}(汇总|总结|盘点|分类)',
            r'(证型|辨证|分型).{0,10}(用药|治疗|中成药)',
            r'(安神|补气|活血|祛湿|健脾).{0,10}(胶囊|丸|片|汤)',
        ]
        for pattern in chinese_medicine_patterns:
            if re.search(pattern, text):
                return True, f"中医药汇总类: 匹配'{pattern}'"
        
        # 8.8 节气/时令养生建议类（没有信息增量，纯生活建议）
        seasonal_lifestyle_patterns = [
            r'(小满|大满|芒种|夏至|小暑|大暑|立秋|处暑|白露|秋分|寒露|霜降|立冬|小雪|大雪|冬至|小寒|大寒|立春|雨水|惊蛰|春分|清明|谷雨|立夏).{0,15}(养生|饮食|调理|注意|建议|必看|别太)',
            r'(今日|今天是).{0,5}(小满|大满|芒种|夏至|小暑|大暑|立秋|处暑|白露|秋分|寒露|霜降|立冬|小雪|大雪|冬至|小寒|大寒|立春|雨水|惊蛰|春分|清明|谷雨|立夏)',
            r'(节气|时令).{0,10}(养生|饮食|调理|习俗|习俗)',
        ]
        for pattern in seasonal_lifestyle_patterns:
            if re.search(pattern, text):
                # 排除真正有新闻价值的文章（如节气相关的政策/经济数据发布）
                if not any(kw in text for kw in ['政策', '数据', '发布', '经济', '市场', '通胀', 'GDP', '利率', '央行']):
                    return True, f"节气养生建议类: 匹配'{pattern}'"
        
        # 8.9 个人故事/人物特写类（以个人经历为主线，缺乏新闻价值）
        # 如：侨批守了四十年、某人的感人故事等
        personal_story_patterns = [
            r'(守了|等了|熬了|守着).{0,10}(四十年|三十年|二十年|多年|半辈子|一辈子)',
            r'(泛黄的|褪色的|尘封的).{0,10}(信|日记|照片|侨批|书信|档案)',
            r'(一头.{0,5}短发|花白短发|轻声细语).{0,30}(丈夫|妻子|儿女|孩子)',
            r'(过番|下南洋|闯关东).{0,15}(一去不回|再没回来|杳无音信)',
            r'(侨批|家书|家信).{0,15}(情感纽带|寻根|华侨史|跨国汇款)',
            r'(她守|他守|她等|他等).{0,15}(一箱子|一封信|一个人|半辈子)',
            r'(一个人的).{0,10}(坚守|等待|故事|传奇)',
            r'(见证|承载).{0,15}(华侨|历史|岁月|时代).{0,10}(记忆|变迁|印记)',
        ]
        for pattern in personal_story_patterns:
            if re.search(pattern, text):
                # 排除：如果文章有明确新闻事件（不只是人物故事）
                if not any(kw in text for kw in ['政策', '改革', '调查', '起诉', '突破', '发现', '数据', '报告', '发布', '制裁', '冲突', '监管', '处罚']):
                    return True, f"个人故事/人物特写类: 匹配'{pattern}'"
        
        # 8.10 艺术流派/美术史/纯学术研究类（缺乏新闻价值，与每日资讯无关）
        # 如：研究揭示印象派画作中女性人物体重指数变化趋势
        art_history_patterns = [
            r'(印象派|后印象派|野兽派|立体派|超现实主义|抽象表现主义).{0,20}(画作|绘画|作品|画家|风格)',
            r'(画作|绘画|雕塑).{0,15}(体重指数|BMI|体态|纤瘦|丰腴|体形)',
            r'(雷诺阿|德加|莫奈|梵高|塞尚|马蒂斯|毕加索).{0,20}(作品|画风|技法|风格)',
            r'(艺术流派|美术史|绘画史|艺术史).{0,15}(研究|分析|考察|趋势|演变)',
            r'(审美观念|审美标准).{0,15}(转变|变化|演变|变迁)',
            r'(艺术创作).{0,15}(时代|社会|文化).{0,10}(影响|反映|映射)',
        ]
        for pattern in art_history_patterns:
            if re.search(pattern, text):
                # 排除：如果文章涉及当前正在发生的重大文化事件（如博物馆重大发现、文物追索等）
                if not any(kw in text for kw in ['发现', '追索', '回归', '被盗', '拍卖纪录', '天价', '政策', '争议']):
                    return True, f"艺术流派/美术史研究类: 匹配'{pattern}'"
        
        # 8.11 纯学术研究类（非Nature/Science级别的普通学术论文，无新闻价值）
        # 如：某研究分析画作中人物体重指数变化、问卷调查结论等
        academic_research_patterns = [
            r'研究揭示.{0,20}(趋势|变化|关系|差异|特征)',
            r'(一项.{0,5}(最新|新的|近期)研究).{0,20}(分析|发现|揭示|表明)',
            r'(研究人员|科学家|学者).{0,15}(测量|分析|发现).{0,15}(比例|趋势|变化|差异)',
        ]
        for pattern in academic_research_patterns:
            if re.search(pattern, text):
                # 排除：Nature/Science级别突破、涉及政策变化、公共卫生意义的保留
                if not any(kw in text for kw in ['Nature', 'Science', 'Cell', 'Lancet', '政策', '突破', '新药', '疫苗', '基因编辑', '临床试验', 'FDA', '批准']):
                    return True, f"纯学术研究类（非重大突破）: 匹配'{pattern}'"
        
        # 8.12 非时事类理论分析文（没有涉及最新消息，纯理论/立场/历史论述）
        # 如"统一是民族复兴关键"这类没有当日新闻由头、纯理论主张的文章
        non_news_theory_patterns = [
            r'(民族复兴|伟大复兴).{0,20}(关键|必然|核心|根本|前提)',
            r'(统一.{0,5}(是|是.*的).{0,10}(关键|必然|核心|根本|历史))',
            r'(历史.{0,5}(必然|选择|潮流|趋势).{0,10}(统一|复兴))',
            r'(必须.{0,10}(坚定|坚决|坚持).{0,15}(信念|信心|决心))',
            r'(是.{0,5}(当代|我们|这一代).{0,10}(人).{0,10}(历史|使命|责任))',
            r'(克服.{0,10}(困难|障碍|阻力).{0,10}(实现|完成|达成).{0,10}(统一|复兴))',
        ]
        for pattern in non_news_theory_patterns:
            if re.search(pattern, text):
                # 排除：如果文章涉及具体的最新事件/政策/数据（不是纯理论论述）
                if not any(kw in text for kw in ['发布', '公布', '宣布', '通过', '签署', '生效', '批准',
                                                   '数据', '统计', '调查', '报告', '突破', '发现',
                                                   '会议', '决议', '法案', '制裁', '冲突', '停火']):
                    return True, f"非时事类理论分析文: 匹配'{pattern}'"

        # 8.13 古今对比感慨文（以古代故事对比现代，无新闻信息量）
        # 如"从古代驿卒累死六千匹马到现代冷链次日达"这种文章
        ancient_modern_patterns = [
            r'(古代|古时|古人).{0,20}(现代|今天|如今|现在).{0,15}(对比|变化|进步|差异|不同)',
            r'(从.{2,15}到.{2,15}).{0,10}(折射|反映|见证|体现).{0,15}(进步|革命|变化|发展)',
            r'(千年|百年|几百年|几千年).{0,10}(变迁|变化|巨变|演进|进步)',
            r'(从.*到.*).{0,5}(折射|见证).{0,10}(人类|文明|社会).{0,10}(进步|发展|变迁)',
            r'(累死|死.*匹.{0,5}马).{0,20}(现代|如今|今天).{0,15}(冷链|物流|运输)',
        ]
        for pattern in ancient_modern_patterns:
            if re.search(pattern, text):
                # 排除：如果有重大新闻事件（不只是古今对比感慨）
                if not any(kw in text for kw in ['发布', '公布', '政策', '突破', '发现', '制裁', '冲突', '协议']):
                    return True, f"古今对比感慨文: 匹配'{pattern}'"

        # 8.14 生物/植物/生态科普论文综述类（非人类健康相关的自然科学论文报道）
        # 如"绵阳师范学院在《自然》发表植物真菌病害研究"——虽然是Nature，但属于生物科普
        biology_science_patterns = [
            r'(植物|真菌|病原菌|lncRNA|长链非编码).{0,15}(研究|机制|分泌)',
            r'(植物.{0,10}(病害|真菌|病原|致病)).{0,15}(研究|发现|揭示|机制)',
            r'(物种|生态|进化|种群).{0,15}(研究|发现|揭示|机制|演化)',
            r'(昆虫|鸟类|哺乳动物|爬行动物).{0,15}(行为|研究|发现|生态)',
            r'(基因.{0,5}(编辑|修饰|表达)).{0,15}(植物|作物|农作物|真菌)',
            r'(生物.{0,10}(多样性|保护|入侵|灭绝)).{0,15}(研究|发现|报告)',
        ]
        for pattern in biology_science_patterns:
            if re.search(pattern, text):
                # 排除：与人类健康/医药/公共卫生直接相关的保留
                if not any(kw in text for kw in ['人类健康', '新药', '疫苗', '临床试验', '基因治疗',
                                                   '患者', '治愈', '药物', 'FDA', '批准上市', '公共卫生']):
                    # 排除：Nature/Science级别但涉及人类重大突破的保留
                    if any(kw in text for kw in ['Nature', 'Science', 'Cell']):
                        # 即使是Nature级别，植物/真菌/生态类也不属于每日资讯范畴
                        # 除非与粮食安全/人类生存直接相关
                        if any(kw in text for kw in ['粮食安全', '全球饥荒', '作物减产', '农业危机']):
                            continue
                    return True, f"生物科普论文综述类: 匹配'{pattern}'"

        # 8.15 个人宣传/先进事迹类（尤其是不知名小人物的正面宣传）
        # 如"北京民警陈长军微信添加3000余名居民"这类文章
        personal_promotion_patterns = [
            r'(民警|警察|干部|社区.{0,3}(工作者|干部|人员)|村支书).{0,15}(坚持|扎根|深入|服务|守护)',
            r'(微信.{0,5}(添加|好友).{0,10}\d+.{0,5}(居民|群众|村民))',
            r'(手写.{0,5}(感谢信|表扬信|锦旗))',
            r'(群众.{0,5}(找他|信任|称赞|好评|点赞|感动))',
            r'(腿勤|嘴勤|手勤|眼勤).{0,10}(多跑|多问|多干|多看)',
            r'(入党|党员).{0,15}(先锋|模范|示范|带头|志愿)',
            r'(扎根.{0,10}(基层|社区|农村|一线|边疆)).{0,15}(多年|十余年|几十年)',
            r'(最美|先进|模范|优秀).{0,5}(民警|教师|医生|护士|工人|干部|公务员)',
        ]
        for pattern in personal_promotion_patterns:
            if re.search(pattern, text):
                # 排除：涉及重大社会事件/犯罪/腐败/争议的保留
                if not any(kw in text for kw in ['犯罪', '腐败', '违法', '查处', '逮捕', '丑闻',
                                                   '重大', '舆情', '争议', '抗议', '罢工', '事故', '伤亡']):
                    return True, f"个人宣传/先进事迹类: 匹配'{pattern}'"

        # 9. 聚合新闻/快讯类文章过滤
        # 识别标题包含"科股快报"、"财经早知道"等明确是快讯/汇总类的文章
        aggregator_patterns = [
            r'科股快报', r'财经早知道', r'今日要闻', r'新闻速递',
            r'早间快讯', r'盘前机会', r'涨停板', r'涨跌停', r'异动追踪',
            r'股市动态', r'期市动态', r'汇市动态', r'宏观早报',
            r'【今日导读】', r'【早餐内参】', r'【市场早知道】',
            r'【?\d+】',  # 标题以编号开头，如"【1】"、"【今日要点】"
            r'影响.{0,5}(下周|本周|市场|A股|股市).{0,5}(十大|几大|重要)?(消息|资讯|新闻|事件)',  # "影响下周市场的十大消息"
            r'(必看|必读).{0,5}影响.{0,10}(消息|资讯|新闻|事件)',  # "必看！影响下周市场的十大消息"
            r'(十大|几大|八大|六大).{0,5}(消息|资讯|新闻|事件|信号)',  # "十大消息"、"八大信号"
            r'8点1氪',  # 36氪的聚合快讯栏目
            r'财经早餐',  # 财经早餐类聚合快讯
            r'新闻早参考',  # 新闻早参考
            r'(每日|今日|每天).{0,5}(简报|播报|速递|资讯)',  # 每日简报/今日速递
            r'(早间|午间|晚间).{0,5}(要闻|速递|简报|资讯)',  # 早间要闻/晚间速递
        ]
        
        # 不再仅凭"|"或"｜"分隔符判断为聚合新闻，因为有些文章虽有分隔符但内容是单一主题
        # 只有当标题明显是多主题快讯时才过滤
        for pattern in aggregator_patterns:
            if re.search(pattern, self.title):
                # 检查是否是标题的一部分还是全文，且是否包含多个主题分隔
                if len(self.title) < 30 or '、' in self.title:
                    return True, f"聚合快讯类: 匹配'{pattern}'"
                break
        
        # 9.0 标题逗号/分号分隔多主题检测（强化）
        # 识别标题用逗号或分号分隔、每段是不同领域独立话题的聚合文章
        # 如："美伊和谈再陷僵局油价震荡，中国光纤迎AI东风，外卖平台因食安问题被罚36亿"
        # 如："8点1氪：俄罗斯总统普京抵达北京；DeepSeek回应用户"对话泄露"疑虑；全国存款十强城市出炉"
        if '，' in self.title or ',' in self.title or '；' in self.title or ';' in self.title:
            title_parts = re.split(r'[，,；;]', self.title)
            if len(title_parts) >= 3:
                # 定义不同领域的关键词集合
                domain_keywords = {
                    '国际军事': ['美伊', '伊朗', '以色列', '俄乌', '普京', '拜登', '特朗普', '北约', '和谈', '制裁', '战争', '停火', '军事', '海峡', '油价', '原油', '访华', '访美', '抵达', '首脑'],
                    '科技': ['AI', '人工智能', '芯片', '光纤', '5G', '半导体', '量子', '算力', '机器人', 'DeepSeek', '华为', '英伟达', '模型', '对话', '泄露', 'Gemini', 'Google', '谷歌'],
                    '消费商业': ['外卖', '美团', '抖音', '京东', '电商', '平台', '罚单', '罚款', '食品安全'],
                    '金融经济': ['股市', 'A股', '港股', '美股', '基金', '利率', '央行', '降息', '加息', '汇率', '存款', 'GDP', '通胀', '美联储', 'IPO', '上市'],
                    '房地产': ['楼市', '房价', '地产', '房贷', '公积金'],
                    '教育': ['高考', '考研', '招生', '学校', '大学'],
                    '医疗': ['医保', '药品', '医院', '疫苗', '医疗'],
                }
                # 统计标题覆盖了多少个不同领域
                domains_found = set()
                for part in title_parts:
                    for domain, keywords in domain_keywords.items():
                        if any(kw in part for kw in keywords):
                            domains_found.add(domain)
                # 如果标题涵盖3个及以上不同领域，判定为聚合新闻
                if len(domains_found) >= 3:
                    return True, f"聚合快讯类: 标题涵盖{len(domains_found)}个不同领域{list(domains_found)}"
        
        # 特别检查：如果是短标题且包含多个"|"或"｜"分隔开的独立主题
        if len(self.title) < 40:
            pipe_parts = re.split(r'[｜|]', self.title)
            if len(pipe_parts) >= 3:
                # 检查每个部分是否都是独立的主题（包含名词/事件词）
                topic_words = ['股', '市', '指数', '涨', '跌', '要闻', '动态', '快讯', '今日', '市场']
                parts_with_topics = sum(1 for p in pipe_parts if any(tw in p for tw in topic_words))
                if parts_with_topics >= 2:
                    return True, "聚合快讯类: 短标题包含多个分隔主题"
        
        # 9.1 检查摘要中是否包含多个不相关的新闻主题（聚合新闻特征）
        # 关键改进：仅凭"涵盖多领域"不能判断为聚合快讯
        # 必须同时满足：(1)涵盖3+领域 AND (2)标题包含聚合快讯特征词 AND (3)没有明确的单一聚焦主题
        # 这样避免误杀深度分析文章（如"沃什提名听证会"涉及金融+国际军事，但主题是单一的）
        summary_domain_keywords = {
            '国际军事': ['伊朗', '以色列', '俄乌', '普京', '拜登', '特朗普', '北约', '和谈', '制裁', '战争', '停火', '海峡', '访华', '访美', '抵达'],
            '科技': ['AI', '人工智能', '芯片', '光纤', '5G', '半导体', '量子', '算力', '机器人', 'DeepSeek', '模型', 'Gemini', '谷歌'],
            '消费商业': ['外卖', '美团', '抖音', '电商', '平台', '罚单', '幽灵'],
            '金融经济': ['股市', 'A股', '港股', '美股', '基金', '利率', '央行', '降息', '存款', '美联储', 'IPO'],
            '房地产': ['楼市', '房价', '地产', '房贷'],
            '能源': ['油价', '原油', '石油', '天然气', '煤炭', '新能源'],
        }
        summary_domains_found = set()
        for domain, keywords in summary_domain_keywords.items():
            if any(kw in text for kw in keywords):
                summary_domains_found.add(domain)
        
        # 只有同时满足以下条件才判定为聚合快讯：
        # 1. 涵盖3个及以上不同领域
        # 2. 标题包含明确的聚合/快讯特征词（如"科股快报"、"快讯"、"早报"等）
        # 3. 文章不是深度分析/解读/讲稿/全文等类型
        if len(summary_domains_found) >= 3:
            # 检查标题是否包含聚合快讯特征词
            aggregator_title_indicators = ['快报', '快讯', '早报', '早知道', '速递', '要闻', '汇总', '一览', '播报', '日报', '8点1氪', '财经早餐']
            title_is_aggregator = any(ind in self.title for ind in aggregator_title_indicators)
            
            # 检查文章是否是深度分析/解读类型（这类文章即使涉及多领域也不应被判定为聚合快讯）
            deep_analysis_indicators = ['讲稿', '全文', '解读', '深度', '分析', '详解', '演讲', '听证', '专访', '对话', '研讨', '展望', '来龙去脉', '回顾']
            is_deep_analysis = any(ind in self.title for ind in deep_analysis_indicators)

            # 【新增】跨领域分析型文章（标题含"助推/影响/冲击/战略"等）不是聚合快讯
            is_cross_domain_analysis = any(
                ind in self.title for ind in Article.CROSS_DOMAIN_ANALYSIS_INDICATORS
            )
            
            # 高质量来源的文章涉及多领域时，更应倾向保留
            # 关键是"标题含快讯特征词"才是聚合快讯；仅涵盖多领域不是充分条件
            if (
                title_is_aggregator
                and not is_deep_analysis
                and not is_cross_domain_analysis
                and not is_high_quality
            ):
                return True, f"聚合快讯类: 摘要涵盖{len(summary_domains_found)}个不同领域{list(summary_domains_found)}且标题含快讯特征"
        
        # 保留原有的固定主题词检测作为补充
        news_topic_indicators = ['恒大', '电影票房', '国际方面', '金融市场', '国内电影']
        topic_count = sum(1 for indicator in news_topic_indicators if indicator in text)
        if topic_count >= 3:
            return True, f"聚合快讯类: 摘要包含{topic_count}个不同新闻主题"
        
        # 9.2 检查是否同时包含国内外多个不相关话题
        if '国际方面' in text and ('国内' in text or '与此同时' in text):
            if '恒大' in text or '票房' in text or '美股' in text:
                return True, "聚合快讯类: 摘要同时包含国内外多个独立话题"
        
        # 9.3 【新增】检测内容中是否包含多个完全不相关的主题（内容杂糅）
        # 例如：一篇文章同时讲"大学生就业"和"某公司财报"，这是两个完全不相关的主题
        # 这种文章应该被过滤，除非它是深度分析类文章
        unrelated_topics = {
            '教育就业': ['大学生', '就业', '高考', '考研', '招生', '毕业', '技校', '学历', '大学'],
            '能源化工': ['宝丰', '能源', '烯烃', '石油', '原油', '天然气', '煤炭', '炼油', '化工', '布伦特'],
            '宏观经济': ['GDP', '经济', '增长', '放缓', '复苏', '通胀', '央行', '降息', '加息'],
            '国际局势': ['伊朗', '美伊', '以色列', '俄乌', '战争', '停火', '制裁', '和谈'],
            '科技AI': ['AI', '人工智能', '芯片', 'DeepSeek', '大模型', '算力'],
        }
        
        detected_topics = []
        for topic_name, keywords in unrelated_topics.items():
            for kw in keywords:
                if kw in text:
                    detected_topics.append(topic_name)
                    break
        
        # 如果检测到3个及以上完全不相关的主题，且文章不是深度分析类，需要过滤
        if len(set(detected_topics)) >= 3:
            deep_analysis_indicators = ['讲稿', '全文', '解读', '深度', '分析', '详解', '演讲', '听证', '专访', '对话', '研讨', '展望', '来龙去脉', '回顾', '复盘', '全面', '系统']
            is_deep_analysis = any(ind in self.title for ind in deep_analysis_indicators)
            
            # 检查文章是否有明确的单一主题词
            single_topic_indicators = ['分析', '解读', '深度', '评论', '观察', '透视', '复盘', '启示', '趋势']
            has_single_topic = any(ind in self.title for ind in single_topic_indicators)

            # 【新增 1】跨领域分析型文章白名单：标题含"助推/影响/冲击/战略..."等
            # 这类文章天然要跨领域论述（如"战争助推能源战略"），不是主题杂糅
            is_cross_domain_analysis = any(
                ind in self.title for ind in Article.CROSS_DOMAIN_ANALYSIS_INDICATORS
            )

            # 【新增 2】标题-主题一致性：若标题本身就命中"检测出的某一主题"的关键词，
            # 说明文章有明确的聚焦主题，摘要中出现其它主题词只是论述支撑，不是杂糅
            title_hits_topics = set()
            for topic_name, keywords in unrelated_topics.items():
                if any(kw in self.title for kw in keywords):
                    title_hits_topics.add(topic_name)
            has_focused_title = len(title_hits_topics) >= 1

            # 【新增 3】多主题但其中存在"语义自然共现组合"时放行
            # 例如：AI 教育 = 科技AI + 教育就业 是同一主题的两个面；
            # 能源战略 = 能源化工 + 国际局势 在地缘分析中天然共现
            detected_set = set(detected_topics)
            natural_cooccurrence_pairs = [
                {'科技AI', '教育就业'},      # "AI 专业/人才"
                {'能源化工', '国际局势'},    # "地缘冲突 → 能源"
                {'宏观经济', '国际局势'},    # "国际局势 → 宏观经济"
                {'宏观经济', '教育就业'},    # "经济 → 就业"
                {'宏观经济', '能源化工'},    # "经济 → 能源"
                {'科技AI', '宏观经济'},      # "AI → 经济"
            ]
            has_natural_cooccurrence = any(
                pair.issubset(detected_set) for pair in natural_cooccurrence_pairs
            )

            # 如果是杂糅文章且不是深度分析，过滤
            # 新规则：白名单任意一条命中即放行，真正杂糅的是"硬拼三个毫无关联的大主题"
            if (
                not is_deep_analysis
                and not has_single_topic
                and not is_high_quality
                and not is_cross_domain_analysis
                and not has_focused_title
                and not has_natural_cooccurrence
            ):
                return True, f"内容杂糅类: 摘要包含{len(set(detected_topics))}个不相关主题{list(set(detected_topics))}"
        
        # 9.4 【新增】检测摘要是否包含明确的多个独立新闻点（以"1.""2.""3."或"第一""第二""第三"分隔）
        # 这种结构通常表示文章是聚合多则新闻的汇总
        independent_news_patterns = [
            r'1\.\s*\S.{0,30}2\.\s*\S.{0,30}3\.\s*\S',  # 1. xxx 2. xxx 3. xxx 结构
            r'第一[^。]{0,20}第二[^。]{0,20}第三',  # 第一...第二...第三...
            r'①[^。]{0,30}②[^。]{0,30}③',  # ① xxx ② xxx ③ xxx
        ]
        for pattern in independent_news_patterns:
            if re.search(pattern, text):
                # 如果摘要中有明确的多个独立新闻点结构，且标题不是深度分析类，很可能是聚合新闻
                deep_analysis_indicators = ['讲稿', '全文', '解读', '深度', '分析', '详解', '演讲', '听证', '专访', '对话', '研讨', '展望', '来龙去脉', '回顾', '复盘', '全面', '系统']
                is_deep_analysis = any(ind in self.title for ind in deep_analysis_indicators)
                
                if not is_deep_analysis and not is_high_quality:
                    return True, f"聚合新闻类: 摘要包含多个独立新闻点（序号结构）"
        
        # 9.5 【新增】检测摘要是否同时报道多个不相关的"具体事件"
        # 如果摘要中出现"X事件"、"Y事件"、"Z事件"等多个人名/事件名，且它们之间没有关联
        # 可能是聚合新闻
        specific_event_patterns = [
            r'\S{2,4}(事件|风波|争议|丑闻|回应|澄清)',  # 某事件/某风波等
        ]
        events_found = []
        for match in re.finditer(specific_event_patterns[0], text):
            event_name = match.group(0)
            if len(event_name) >= 4:  # 确保不是太短的匹配
                events_found.append(event_name)
        
        # 如果发现3个及以上不同的具体事件名，且它们之间没有明显的因果/关联关系
        if len(events_found) >= 3:
            # 检查这些事件是否在标题中也有体现（如果是，说明文章确实是多事件聚合）
            events_in_title = sum(1 for e in events_found if e in self.title)
            if events_in_title == 0:
                # 事件不在标题中，说明标题可能是聚合的，而摘要把它们都包含进来了
                # 需要进一步检查：这些事件之间是否有逻辑关联
                # 简化判断：如果事件之间没有明显的连接词，且摘要中有多个序号/分隔符，可能是聚合
                if any(conn in text for conn in ['与此同时', '此外', '另外', '值得一提的是']):
                    return True, f"聚合新闻类: 摘要包含多个不相关具体事件"
        
        # 9.6 【新增】预测/前瞻性聚合新闻过滤
        # 识别"下周/本周XX事件密集""X月X日展望""未来一周关注"等对尚未发生事件的预测/预告聚合
        # 这类文章罗列多个未来不同领域的预测判断，不是当日已发生的新闻
        prediction_aggregation_patterns = [
            r'(下周|本周|未来一周|接下来一周).{0,10}(事件|关注|看点|重点|预览|展望|前瞻)',
            r'(下周|本周|未来一周).{0,15}(密集|重要|关注|看点|重点)',
            r'(X月X日|下周).{0,10}(事件|数据|发布|公布)',
            r'(预览|前瞻|展望).{0,5}(下周|本周|未来|一周)',
            r'(下周|本周).{0,30}(PCE|PMI|CPI|GDP|议息|数据)',
            r'(将要|即将|预计).{0,10}(发生|公布|发布|召开|举行).{0,20}(多个|一系列|密集)',
        ]
        for pattern in prediction_aggregation_patterns:
            if re.search(pattern, text):
                # 排除：如果文章主要是对当前已发生事件的深度分析（而非罗列未来事件）
                if not any(kw in text for kw in ['已发生', '已公布', '已发布', '结果出炉', '数据公布']):
                    return True, f"预测/前瞻聚合新闻类: 匹配'{pattern}'"
        
        # 9.7 【新增】宠物寻主/领养/流浪动物救助类琐碎新闻过滤
        # 这类文章除非是重大事件（如大规模虐待动物），否则不具有全国新闻价值
        pet_adoption_patterns = [
            r'(萌宠|宠物).{0,10}(寻人|寻主|领养|认养)',
            r'(流浪.{0,3}(狗|猫|动物)).{0,10}(领养|认养|等待|寻主)',
            r'(领养|认养).{0,10}(流浪.{0,3}(狗|猫|动物)|萌宠)',
            r'(寻人启事|寻主启事).{0,10}(狗|猫|宠物)',
            r'(领养.{0,5}条件|领养.{0,5}要求|领养.{0,5}联系方式)',
        ]
        for pattern in pet_adoption_patterns:
            if re.search(pattern, text):
                # 排除：重大动物保护事件（大规模虐待、野生动物保护政策等）
                if not any(kw in text for kw in ['重大', '虐待', '犯罪', '立法', '政策', '违法', '查获', '走私', '濒危', '保护法']):
                    return True, f"宠物领养/寻主琐碎新闻类: 匹配'{pattern}'"
        
        # 9.8 【新增】多事件主观评论/政府公信力批判类文章过滤
        # 这类文章将多个不相关事件拼凑在一起批判政府/体制，主观性强、不客观
        # 如：将杨梅泡药水和矿难拼在一起说"监管失守""信任危机"
        subjective_commentary_patterns = [
            r'(信任危机|公信力).{0,20}(监管|政府|体制|反复|重演)',
            r'(监管失守|监管缺失|监管漏洞).{0,20}(不止|不仅|从.*到)',
            r'(从.*到.*).{0,10}(都是|无非|归根).{0,10}(人祸|纵容|不作为)',
            r'(劣币驱逐良币).{0,20}(监管|不作为|纵容)',
            r'(举一反三).{0,10}(沦为|空文|形式)',
            r'(悲剧.{0,5}反复.{0,5}重演|重演.{0,5}悲剧)',
            r'(坏人.{0,10}(未受严惩|没有代价)|好人.{0,10}(遭殃|受伤|买单))',
        ]
        for pattern in subjective_commentary_patterns:
            if re.search(pattern, text):
                return True, f"多事件主观评论/政府批判类: 匹配'{pattern}'"
        
        # 10. 检查摘要中是否用多个连接词拼接多个话题（综合式文章）
        # 改进：仅凭连接词数量不能判定，必须同时满足连接词>=2 AND 涵盖3+领域 AND 无单一聚焦主题
        # 避免误杀围绕单一主题展开论述的深度文章
        multi_topic_connectors = ['与此同时', '此外', '另外', '值得一提的是', '值得注意']
        connector_count = sum(1 for conn in multi_topic_connectors if conn in text)
        if connector_count >= 2:
            # 需要同时验证涵盖了3个以上不同领域
            connector_domain_keywords = {
                '国际军事': ['伊朗', '以色列', '俄乌', '普京', '拜登', '特朗普', '北约', '和谈', '制裁', '战争', '停火'],
                '科技': ['AI', '人工智能', '芯片', '光纤', '5G', '半导体', '量子', '算力', '机器人', 'DeepSeek'],
                '消费商业': ['外卖', '美团', '抖音', '电商', '平台', '罚单'],
                '金融': ['股市', 'A股', '港股', '美股', '基金', '利率', '央行', '降息'],
                '房地产': ['楼市', '房价', '地产', '房贷'],
                '能源': ['油价', '原油', '石油', '天然气', '煤炭', '新能源'],
            }
            connector_domains = set()
            for domain, keywords in connector_domain_keywords.items():
                if any(kw in text for kw in keywords):
                    connector_domains.add(domain)
            
            # 只有连接词>=2 且 涵盖3+领域 且 不是深度分析类文章，才判定为综合式
            deep_analysis_indicators = ['讲稿', '全文', '解读', '深度', '分析', '详解', '演讲', '听证', '专访', '对话', '研讨', '展望']
            is_deep_analysis = any(ind in self.title for ind in deep_analysis_indicators)
            
            if len(connector_domains) >= 3 and not is_deep_analysis and not is_high_quality:
                return True, f"综合集成式文章: 包含{connector_count}个话题连接词，覆盖{len(connector_domains)}个领域{list(connector_domains)}"
        
        # 11. 推荐书单/推荐商品/推荐清单类文章过滤
        # 识别"推荐X本书"、"X本必读书单"、"推荐X个商品"等缺乏新闻价值的列表式推荐
        recommendation_list_patterns = [
            r'推荐.{0,5}本书', r'推荐.{0,5}本书单', r'推荐.{0,5}本.*书',
            r'必读.{0,5}本书', r'必看.{0,5}本书', r'收藏.{0,5}本书',
            r'\d+本.{0,10}书单', r'书单', r'推荐.{0,5}清单', r'推荐.{0,5}列表',
            r'推荐.{0,5}商品', r'推荐.{0,5}产品', r'推荐.{0,5}好物',
            r'年度.{0,5}书单', r'暑期.{0,5}书单', r'假期.{0,5}书单',
        ]
        for pattern in recommendation_list_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏新闻事件或重大分析
                if not any(kw in text for kw in ['政策', '重大', '违规', '违法', '犯罪', '舆情', '热搜', '丑闻', '曝光', '调查', '处罚']):
                    return True, f"推荐书单/清单类: 匹配'{pattern}'"
        
        # 12. 虚假故事/情景演绎/虚构场景类文章过滤
        # 识别"虚构"、"假设"、"如果...那么"、"情景模拟"等明显虚构或推理性内容
        fictional_story_patterns = [
            r'虚构.{0,10}(场景|饭局|对话|故事|情节)', r'假设.{0,10}(场景|饭局|对话)',
            r'如果.{0,10}(那么|就|的话)', r'情景.{0,5}(模拟|演绎|再现)',
            r'一场.{0,10}(虚构|假设)的.{0,10}(饭局|对话|场景|故事)',
            r'模拟.{0,10}(家庭|饭局|职场)', r'想象中', r'假如.{0,10}会',
            r'角色扮演', r'演绎.{0,5}(家庭|职场|社会)', r'家庭.{0,5}(互动|聚餐|模拟)',
        ]
        for pattern in fictional_story_patterns:
            if re.search(pattern, text):
                return True, f"虚假故事/情景演绎类: 匹配'{pattern}'"
        
        # 13. 盘点/榜单类水文（除非有重大新闻事件）
        # 识别"盘点"、"榜单"、"排行"、"TOP"等缺乏新闻价值的整理式文章
        # 【重要修复】只匹配标题明确是盘点/榜单的文章，或文章主体就是榜单排行的，
        # 不应仅因文章内容中提及"榜单"一词就过滤（文章可能只是在论述中引用榜单数据）
        ranking_water_patterns = [
            r'盘点.{0,5}(十|九|八|七|六|五|四|三|二|一)?(种|个|类|件)',
            r'排行榜', r'TOP\d+', r'十大', r'八大', r'六大',
            r'排名.{0,5}(出炉|发布|揭晓|公布)', r'年度.{0,5}(榜单|排行|排名)',
            r'十大.{0,5}(事件|新闻|人物|热词|关键词)',
        ]
        # 仅在标题中出现的榜单/盘点模式（标题中出现则更可能是纯榜单水文）
        ranking_title_patterns = [
            r'榜单', r'排行榜', r'TOP\d+', r'盘点.{0,5}(十|九|八|七|六|五|四|三|二|一)?(种|个|类|件)',
        ]
        for pattern in ranking_water_patterns:
            if re.search(pattern, self.title):
                # 标题匹配到盘点/榜单：较大概率是纯榜单文章
                # 排除有重大新闻事件的盘点
                if not any(kw in text for kw in ['重大', '政策', '违规', '违法', '犯罪', '舆情', '热搜', '丑闻', '曝光', '调查', '处罚', '突破', '创新', '历史首次']):
                    return True, f"盘点/榜单水文类: 匹配'{pattern}'"
        # 正文匹配到"榜单"时，需要更严格判断：文章主体是否就是榜单排行
        # 如果文章标题没有明显榜单特征，但正文中出现"榜单"，不轻易过滤
        # 除非文章有多处榜单特征且缺乏实质新闻内容
        for pattern in ranking_water_patterns:
            if re.search(pattern, text):
                # 正文匹配，但标题不含榜单特征 → 可能只是引用榜单数据
                # 需要文章中有多个榜单特征才过滤（说明文章主体就是榜单）
                ranking_count = sum(1 for p in ranking_water_patterns if re.search(p, text))
                if ranking_count >= 3:
                    # 多处榜单特征 + 无重大新闻 → 纯榜单水文
                    if not any(kw in text for kw in ['重大', '政策', '违规', '违法', '犯罪', '舆情', '热搜', '丑闻', '曝光', '调查', '处罚', '突破', '创新', '历史首次']):
                        return True, f"盘点/榜单水文类: 匹配'{pattern}'(正文多处榜单特征)"
        
        # 14. 人物/事件多角度重复报道过滤
        # 检测同一人物或事件的"后续"、"回应"、"分析"等重复报道
        # 如果一篇文章只是在"回顾"、"复盘"、"解读"已报道过的事件，缺乏新信息
        repeat_report_patterns = [
            r'回顾.{0,10}(事件|人物|话题)', r'复盘.{0,10}(事件|人物|话题)',
            r'解读.{0,10}(事件|人物|话题)', r'起底.{0,10}(事件|人物|话题)',
            r'还原.{0,10}(事件|人物|话题)', r'梳理.{0,10}(事件|人物|话题)',
            r'盘点.{0,10}(事件|人物|话题)', r'评价.{0,10}(事件|人物|话题)',
            r'热议',  # 后续热议
        ]
        for pattern in repeat_report_patterns:
            if re.search(pattern, text):
                # 检查是否有实质性的新信息
                # 如果只是"如何评价"、"你怎么看"、"引发热议"而没有新事实，可能是水文
                if not any(kw in text for kw in ['最新', '最新消息', '最新进展', '最新回应', '官方', '首次', '最新披露']):
                    # 检查是否是"如何评价X"的讨论类文章
                    if re.search(r'如何.{0,5}(看待|评价|看待|评论)', text):
                        return True, f"重复讨论类文章: 匹配'{pattern}'"
        
        # 15. 个人感悟/讲故事类文章（只有观点和故事，缺乏客观信息点）
        # 识别作者个人经历、展览观后感、反思类文章
        # 【注意】仅匹配"参观/观展/观后感"等主观体验词，不匹配"博物馆/展览"本身
        # 因为"国家自然博物馆举办霸王龙特展"是新闻，不是个人感悟
        personal_reflection_patterns = [
            r'参观.{0,10}(展览|美术馆|博物馆)', r'观展', r'观后感',
            r'策展', r'策展人', r'反思', r'领悟', r'感受到',
            r'促使.{0,5}反思', r'这让.{0,5}思考', r'引发了.{0,5}思考',
            r'个人.{0,10}(经历|体验|感悟|感受|心得)', r'作者.{0,10}(意识|反思)',
            r'在.{0,10}中.{0,10}(我|本人|作者)', r'通过.{0,10}(参观|游览)',
            r'从.{0,10}展览', r'卢浮宫', r'图案的奇迹',
            r'引导.{0,5}关注', r'叙事框架', r'建设性.{0,5}叙事',
        ]
        for pattern in personal_reflection_patterns:
            if re.search(pattern, text):
                # 【新增】如果是展览/博物馆的客观新闻（有"举办/展出/开幕/特展/开幕"等词），放行
                exhibition_news_keywords = ['举办', '展出', '开幕', '特展', '首展', '开展', '对公众开放', '亮相', '首次']
                if any(kw in text for kw in exhibition_news_keywords):
                    continue
                # 检查是否缺乏客观信息点
                if not any(kw in text for kw in ['数据', '政策', '报告', '统计', '研究', '分析', '官方', '事实', '调查', '发现']):
                    return True, f"个人感悟/讲故事类: 匹配'{pattern}'，缺乏客观信息点"
                # 如果有感悟/反思类词汇但没有实质性内容，也过滤
                if any(kw in text for kw in ['反思', '领悟', '感受', '思考', '启发', '思考过程', '引导', '叙事']):
                    if not any(kw in text for kw in ['数据', '研究', '报告', '官方', '调查', '结论', '政策']):
                        return True, f"个人感悟/讲故事类: 反思类文章缺乏实质信息"
        
        # 16. 经验总结/教育式文章（教导性质，没有实质信息）
        # 识别"关键线索"、"识别特征"、"判断方法"等教导式文章
        teaching_pattern_patterns = [
            r'关键.{0,5}(线索|特征|标志|信号)', r'识别.{0,5}(方法|技巧|方式)',
            r'判断.{0,5}(方法|技巧|方式|标准)', r'判断.{0,5}(依据|要点)',
            r'行为.{0,5}(细节|线索|特征|表现)', r'细节.{0,5}(成为|成为关键)',
            r'观察.{0,5}(判断|识别|发现)', r'无意识.{0,10}(流露|表现|行为)',
            r'非言语.{0,10}(信号|线索|特征)', r'传递.{0,5}(信号|信息)',
            r'容易.{0,5}(被识破|被看穿|被察觉|被识别)', r'识破', r'看穿',
            r'这些.{0,10}共同构成', r'构成了一套',
        ]
        for pattern in teaching_pattern_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏客观数据和研究
                if not any(kw in text for kw in ['数据', '研究', '报告', '统计', '科学', '论文', '实验', '调查', '发现', '结论', '事实']):
                    return True, f"经验总结/教育式文章: 匹配'{pattern}'，缺乏研究数据"
        
        # 17. 地域偏见/容貌焦虑/歧视类文章
        # 识别以地域或先天条件（如容貌）讨论所谓"困境"的文章
        regional_biased_patterns = [
            r'容貌彩票', r'容貌焦虑', r'身体样貌.{0,5}(决定|选择|彩票)',
            r'中了.{0,5}(头彩|彩票|命运)', r'类似.{0,5}(彩票|抽签)',
            r'无法选择.{0,5}(命运|出身|容貌)', r'先天的.{0,5}(命运|容貌|出身)',
            r'地域性.{0,10}(问题|困扰|困境)', r'某地区.{0,10}(问题|困扰|困境)',
            r'普遍.{0,5}(困扰|焦虑|问题|命运)',
        ]
        for pattern in regional_biased_patterns:
            if re.search(pattern, text):
                return True, f"地域偏见/容貌焦虑类: 匹配'{pattern}'"
        
        # 18. 历史总结/回顾类文章
        # 识别回顾历史事件、总结历史经验教训的文章，这些不是当前热点资讯
        history_summary_patterns = [
            r'\d+周年.{0,10}(纪念|回顾|反思|祭)',  # XX周年纪念/回顾
            r'(签订|签署|达成).{0,15}(条约|协定|和约)',  # 历史条约签订
            r'(鸦片战争|甲午战争|八国联军|英法联军|圆明园|马关条约|辛丑条约|南京条约)',
            r'(侵略|掠夺|殖民|不平等条约).{0,10}(回顾|反思|教训|警醒)',
            r'(历史|近代史).{0,10}(教训|启示|反思|经验|铭记)',
            r'(百年|数百年|千年前).{0,15}(历史|事件|教训)',
            r'(洗心革面|不义之财|强盗|掠夺).{0,10}(归还|追索|返还)',
            r'(历史|近代).{0,10}(伤痕|屈辱|耻辱|觉醒)',
            r'(落后.{0,3}挨打|弱国.{0,3}没有外交|制度.{0,5}(落后|革新))',
            r'(帝国|殖民|列强).{0,10}(瓜分|入侵|侵华|侵掠)',
            r'(19|18|17)\d{2}年.{0,15}(战争|条约|签订|签署|割让|赔款)',
        ]
        for pattern in history_summary_patterns:
            if re.search(pattern, text):
                # 排除与当前重大事件相关的历史对比分析（如有最新政策/最新进展则保留）
                if not any(kw in text for kw in ['最新', '今日', '刚刚', '突发', '宣布', '最新进展', '首次回应', '官方通报']):
                    return True, f"历史总结类文章: 匹配'{pattern}'"
        
        # 18.5 单部影视/电影/剧集介绍类文章（不是新闻，没有信息增量）
        # 识别对单部电影、电视剧、综艺节目的介绍/影评/幕后花絮等文章
        # 这类文章只是介绍某部作品的剧情/选角/制作细节，不属于每日资讯
        # 注意：影视行业重大事件（如票房纪录、行业政策、重大投资）不应被过滤
        movie_intro_patterns = [
            r'(导演|主演|编剧).{0,10}(透露|表示|分享|介绍)',  # 导演/主演透露制作细节
            r'(选角|试镜|演员阵容)',  # 选角/阵容介绍
            r'(票房|口碑|评分).{0,5}(扑街|惨淡|崩盘)',  # 票房口碑报道（非行业数据）
            r'(剧情|角色|台词|片尾|彩蛋).{0,5}(解读|分析|揭秘)',  # 剧情/角色解读
            r'(拍摄|取景|杀青|开机|定档|路演|首映)',  # 影视制作/宣发动态
            r'(续集|前传|外传|番外).{0,5}(确认|官宣|宣布)',  # 续集/衍生作品官宣
            r'(宇宙|系列).{0,10}(原点|起源|扩展|布局)',  # 影视宇宙扩展
            r'(特效|CG|动作捕捉|化妆).{0,5}(幕后|制作|技术)',  # 特效制作幕后
        ]
        for pattern in movie_intro_patterns:
            if re.search(pattern, text):
                # 排除影视行业重大事件（票房纪录、行业政策、重大投资/收购等）
                if not any(kw in text for kw in ['票房纪录', '历史纪录', '行业政策', '投资', '收购', '上市',
                                                   '监管', '处罚', '税务', '偷税', ' scandal', '丑闻',
                                                   '破产', '退市', '并购', '重组']):
                    # 排除高质量来源的影视行业深度分析
                    if not is_high_quality:
                        return True, f"单部影视介绍类: 匹配'{pattern}'"
        
        # 18.6 日常医疗健康科普类文章（不是新闻，属于常识/科普）
        # 识别蚊子叮咬、感冒预防、饮食建议等日常健康科普文章
        # 这类文章提供的是通用健康知识，不是当日的新闻资讯
        # 注意：重大公共卫生事件（疫情、疫苗政策、药品审批）不应被过滤
        health_science_patterns = [
            r'(蚊子|蚊虫).{0,10}(叮咬|咬|叮|预防|驱蚊)',  # 蚊虫叮咬科普
            r'(为什么|为啥).{0,15}(总|总是|容易|更).{0,10}(被咬|招蚊子|招虫)',  # "为什么总被蚊子咬"
            r'(止痒|消肿|驱蚊|防蚊)',  # 止痒/驱蚊方法
            r'(O型血|血型).{0,10}(蚊子|叮咬|招)',  # 血型与蚊子
            r'(感冒|流感).{0,10}(预防|治疗|用药|注意)',  # 感冒预防科普
            r'(防晒|防晒霜|SPF).{0,10}(选择|推荐|注意|方法)',  # 防晒科普
            r'(喝水|饮水|补水).{0,10}(多少|方法|建议|最佳)',  # 喝水科普
            r'(睡眠|失眠).{0,10}(改善|方法|建议|技巧)',  # 睡眠改善科普
            r'(久坐|办公).{0,10}(危害|注意|改善|建议)',  # 久坐危害科普
            r'(饮食|饮食搭配|营养).{0,10}(建议|原则|方法|指南)',  # 饮食科普
            r'(中暑|防暑|降温).{0,10}(方法|注意|预防|建议)',  # 防暑科普
        ]
        for pattern in health_science_patterns:
            if re.search(pattern, text):
                # 排除重大公共卫生事件
                if not any(kw in text for kw in ['疫情', '公共卫生', '防控', '疫苗', '药物审批', '医保',
                                                   '政策', '突破', '基因编辑', '世卫', '通报', '召回',
                                                   '食品安全', '中毒', '违规', '查处', '刑拘']):
                    return True, f"日常医疗健康科普类: 匹配'{pattern}'"
        
        # 19. 大学基建/校园日常管理类文章
        # 识别大学修宿舍、装修、采购设备等缺乏新闻价值的校园日常事务
        campus_infrastructure_patterns = [
            r'(大学|学院|学校).{0,15}(修缮|装修|翻新|改造|翻修)',  # 学校装修修缮
            r'(宿舍|公寓|食堂|教学楼).{0,10}(修缮|装修|翻新|改造|翻修|建设)',  # 校舍修缮
            r'(投|拟投|投资|斥资).{0,10}(万|亿)元.{0,15}(修缮|装修|翻新|改造|建设)',  # 投资修缮
            r'(平方米|平米).{0,10}(修缮|装修|翻新|改造)',  # 面积+修缮
            r'(大学|学院).{0,10}(采购|购置|引进).{0,15}(设备|仪器)',  # 学校采购设备
            r'(校园|校区).{0,10}(环境|绿化|道路|管网).{0,10}(提升|改造|建设)',  # 校园环境改造
            r'(大学|学院).{0,15}(扩建|新建).{0,15}(宿舍|食堂|教学楼|体育馆)',  # 校舍扩建
        ]
        for pattern in campus_infrastructure_patterns:
            if re.search(pattern, text):
                # 排除涉及重大安全事故或舆情的
                if not any(kw in text for kw in ['事故', '坍塌', '伤亡', '违法', '违规', '舆情', '曝光', '调查', '丑闻']):
                    return True, f"大学基建/校园日常管理类: 匹配'{pattern}'"
        
        # 20. 师生关系/教育理论分析类文章
        # 识别讨论师生关系、导师角色、教育方法等理论性文章，缺乏新闻价值
        education_theory_patterns = [
            r'(导师|老师|教师).{0,15}(严苛|严格|严厉|温柔|温和)',  # 导师性格讨论
            r'(导师|老师).{0,15}(角色|定位|责任|作用)',  # 导师角色分析
            r'(师生关系|导师制|mentors?ship)',  # 师生关系理论
            r'(改稿|拒稿|论文修改).{0,15}(支持|陪伴|指导)',  # 论文修改类感悟
            r'(科研|学术).{0,10}(引路人|守护者|领路人)',  # 科研引路人
            r'(隐性|隐含|背后).{0,10}(支持|陪伴|关爱)',  # 隐性支持类
            r'(严苛|严格|严厉).{0,10}(背后|之下|外表)',  # 严厉背后
            r'(如何.{0,5}(做|当好|成为)).{0,10}(导师|老师|教师)',  # 如何当导师
            r'(教育|教学).{0,10}(方法|理念|模式|改革|创新).{0,10}(思考|反思|探讨)',  # 教育方法思考
        ]
        for pattern in education_theory_patterns:
            if re.search(pattern, text):
                # 排除涉及重大教育政策或事件的
                if not any(kw in text for kw in ['政策', '教育部', '重大', '违规', '违法', '犯罪', '处罚', '舆情', '热搜', '调查']):
                    return True, f"师生关系/教育理论分析类: 匹配'{pattern}'"
        
        # 21. 产品吹嘘/自夸广告类
        # 识别吹嘘自家产品、打广告的文章，如ESG评级体系更新、吹嘘自己产品
        product_bragging_patterns = [
            r'(首次发布|率先|首发|首发版|率先推出)',  # 强调自己是第一个
            r'(全面覆盖|全覆盖)',  # 夸大覆盖范围
            r'覆盖.{0,5}(股票|债券|多资产)',  # 吹嘘覆盖范围
            r'(超过\d+家.{0,10}(上市公司|发债主体))',  # 用数字夸大
            r'(旨在助力|致力于|帮助.{0,10}(投资者|实现))',  # 目的性吹嘘
            r'(评级体系.{0,5}(更新|V\d\.\d))',  # 吹嘘自家体系更新
        ]
        for pattern in product_bragging_patterns:
            if re.search(pattern, text):
                # 检查是否是自夸广告（强调自家产品、覆盖广、用户多等）
                if any(kw in text for kw in ['评级', '体系', '平台', '产品', '服务', '客户', '投资者', '可持续']):
                    if not any(kw in text for kw in ['政策', '监管', '处罚', '违规', '违法', '舆情', '重大', '调查', '诉讼', '改革', '税务', '征税', '调研', '制度', '法规', '法律', '调整', '深化', '推进', '实施', '落实', '体制', '治理', '试点', '推行']):
                        return True, f"产品吹嘘/自夸广告类: 匹配'{pattern}'"
        
        # 22. 单个股票涨跌/财报信息类
        # 识别只讲单个股票/企业涨跌、财报的文章，这些缺乏宏观价值
        stock_single_patterns = [
            r'(股价|涨幅|跌幅).{0,10}(%|\d+)',  # 股价涨跌信息
            r'(年内|今日|盘中|收盘).{0,10}(涨幅|跌幅|涨|跌)',  # 涨跌描述
            r'领涨|领跌|倒数第',  # 排名描述
            r'(一季|半年|三季|全年).{0,10}(净利|利润|营收|净利润|归母)',  # 财报信息
            r'(同比增长|同比下降|同比实现)',  # 财报增长描述
            r'创.{0,5}(历史新高|新高|新低)',  # 股价新高新低
            r'(大族激光|浦发银行|兴业银行)',  # 具体个股（可能只是涨跌分析）
        ]
        # 检查是否是单纯的个股涨跌/财报文章
        for pattern in stock_single_patterns:
            if re.search(pattern, text):
                # 检查是否只是个股涨跌分析，缺乏宏观/行业价值
                stock_only_indicators = [
                    r'领涨|领跌', r'年内涨幅.{0,3}[-\d]', r'同比实现大增',
                    r'一季度净利', r'归母净利润', r'创历史新高'
                ]
                if any(re.search(ind, text) for ind in stock_only_indicators):
                    # 检查是否缺乏宏观/行业分析
                    if not any(kw in text for kw in ['政策', '行业', '板块', '监管', '宏观', '重大', '处罚', '违规', '违法', '诉讼', '收购', '并购', '投资', '战略', '布局', '研发', '转型', '扩产', '建厂', '出海', '烧钱']):
                        return True, f"单股票涨跌/财报类: 匹配'{pattern}'"
        
        # 23. 主观审美/个人观后感类文章
        # 【修改】只保留真正的主观感受类规则，避免误杀客观分析文章
        # 钛媒体等客观分析类文章不应被过滤
        subjective_aesthetic_patterns = [
            r'让我.{0,10}(感动|流泪|心痛|难过|开心)',  # 作者个人情感反应
            r'(哭得|笑着|哭着)',  # 强烈的个人情感表达
            r'看完.{0,5}后.{0,5}(我想说|我想聊聊|说几句)',  # 观后感开场白
            r'看完.{0,20}(后劲|后遗症|久久)',  # 强调个人情感影响
            r'作为一个.{0,5}(普通人|观众|普通人|吃瓜群众)',  # 强调个人身份
            r'纯属.{0,5}(个人|主观|一时)',  # 强调主观性
            r'我不管.{0,5}(反正|就是|就是觉得)',  # 纯主观表达
        ]
        for pattern in subjective_aesthetic_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏客观数据/政策/分析
                if not any(kw in text for kw in ['数据', '政策', '报告', '统计', '研究', '分析', '官方', '调查', '发现', '结论', '案例', '现象', '影响']):
                    return True, f"主观审美/个人感受类: 匹配'{pattern}'"
        
        # 23.5 产品外观/审丑争议类文章
        # 识别围绕某个产品"丑不丑""好不好看"展开的讨论，本质是主观审美争议，缺乏客观新闻价值
        # 例如："匡威新鞋被群嘲丑""某品牌设计翻车"等
        ugly_product_patterns = [
            r'(丑|好看|难看|怪异).{0,15}(鞋|包|衣服|设计|造型|款式)',
            r'(群嘲|群讽|吐槽|嘲笑|嘲笑).{0,15}(丑|设计|外观|造型)',
            r'(审丑|丑鞋|丑包|丑出圈|丑出天际|丑上热搜)',
            r'(翻车|翻红).{0,15}(设计|造型|外观|审美)',
            r'(设计.{0,5}(翻车|灾难|奇葩|怪异|离谱))',
            r'(被.{0,5}(吐槽|嘲笑|群嘲).{0,15}(丑|设计|外观))',
            r'(收手吧|别.{0,5}了).{0,10}(外面|大家).{0,10}(笑|吐槽|丑)',
            r'(护士鞋|道士鞋|老北京布鞋).{0,15}(像|堪比|被比作)',
        ]
        for pattern in ugly_product_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏重大商业/行业新闻（如收购、退市、重大战略调整等）
                if not any(kw in text for kw in ['收购', '退市', '破产', '清算', '重大战略', '监管处罚', '重大违规', '召回']):
                    return True, f"产品外观/审丑争议类: 匹配'{pattern}'"
        
        # 24. 短剧/微短剧相关内容（归入文化娱乐）
        # 注意：短剧内容应该归入文化娱乐版块，而不是医疗/医药版块
        # 这条规则用于标记，但不直接过滤，而是确保正确分类
        # 短剧内容的分类在AI分类时处理
        
        # 25. 医学实验论文/研究类水文（垃圾科普）
        # 识别具体实验论文、临床试验报告等缺乏重大意义的文章
        medical_research_water_patterns = [
            r'(随机对照试验|RCT)',  # 临床试验
            r'(研究.{0,10}(显示|表明|发现|结果))',  # 研究结果描述
            r'(发表于|发表在|成果发表)',  # 论文发表
            r'(马里兰大学|医学院|研究团队)',  # 具体研究机构
            r'美国.{0,5}家.{0,5}(中心|医院|机构)',  # 多中心研究
            r'(患者.{0,5}随机)',  # 患者随机分组
            r'(切口|手术|感染|对照组|试验组)',  # 临床试验术语
            r'(无显著差异|效果不优于)',  # 研究结论
            r'(TOBRA|妥布霉素|万古霉素)',  # 具体实验名/药物名
            r'《美国医学会杂志》',  # 具体期刊
        ]
        for pattern in medical_research_water_patterns:
            if re.search(pattern, text):
                # 检查是否缺乏重大突破或临床意义
                if not any(kw in text for kw in ['突破', '治愈', '重大', '首次', '批准', '上市', '革命性', '颠覆性']):
                    # 进一步检查是否是实验论文类
                    if any(kw in text for kw in ['随机对照', '对照组', '试验组', '无显著差异', '研究显示']):
                        return True, f"医学实验论文/研究类: 匹配'{pattern}'"
        
        # 26. 医学科普水文（基础医学知识，缺乏新闻价值）
        # 识别具体病症解读、药物使用说明等基础科普
        # 注意：必须先排除明显非医学文章（如国际冲突、经济分析等），避免"核心要素"等通用词误杀
        medical_keywords_check = ['医学', '医疗', '临床', '患者', '药物', '手术', '疾病', '治疗', '诊断', '病历', '症状', '用药', '基因', '病理', '药理']
        is_likely_medical = any(kw in text for kw in medical_keywords_check)
        
        if is_likely_medical:
            medical_science_water_patterns = [
                r'(基因检测|变异致病性|ACMG)',  # 基因检测解读
                r'(NF1|NF2|BRCA)',  # 具体基因
                r'(抗组胺药|苯海拉明|适用.{0,5}(过敏|晕动))',  # 具体药物使用
                r'(适用.{0,10}(疾病|症状|人群))',  # 药物适用症
                r'(禁忌证|注意事项|用法用量)',  # 用药说明
                r'(解读.{0,10}(报告|检测|基因))',  # 报告解读
                r'(四大核心|核心要素)',  # 要素解读类（仅在确认是医学文章时匹配）
                r'精准诊疗',  # 精准医疗
                r'【医学界】',  # 医学界公众号的科普文
            ]
            for pattern in medical_science_water_patterns:
                if re.search(pattern, text):
                    # 检查是否是基础科普水文（缺乏新闻事件）
                    if not any(kw in text for kw in ['政策', '监管', '批准', '上市', '重大', '突破', '违规', '处罚', '舆情', '调查', '召回']):
                        # 检查是否是具体知识科普而非新闻
                        knowledge_indicators = [r'核心要素', r'解读', r'适用.{0,5}(过敏|疾病)', r'禁忌证', r'基因检测']
                        if any(re.search(ind, text) for ind in knowledge_indicators):
                            return True, f"医学科普水文类: 匹配'{pattern}'"
        
        # 26.5 基础科学研究科普类文章（缺乏新闻价值的知识普及）
        # 识别具体研究报告、实验发现等缺乏新闻价值的科普文章
        # 这类文章通常是描述性研究（描述现象/机制），缺乏具体政策、监管、违规等新闻事件
        research_science_water_patterns = [
            r'(荷兰.{0,5}研究|美国.{0,5}研究|英国.{0,5}研究|日本.{0,5}研究|法国.{0,5}研究)',  # 外国研究
            r'(.{0,10}研究.{0,10}发现|.{0,10}研究.{0,10}表明|.{0,10}研究.{0,10}显示)',  # 研究发现描述
            r'(.{0,5}研究团队|.{0,5}科学家|.{0,5}学者)',  # 研究团队/科学家
            r'《.{0,10}期刊》',  # 具体期刊发表
            r'(发表于|发表在|成果发表)',  # 论文发表
            r'(双胞胎|排卵|生育|受孕|生殖)',  # 生殖类研究主题
            r'(假说|科学原理|科学解释)',  # 理论假说类（避免"机制""原理""解释"等通用词误杀金融/经济文章）
        ]
        for pattern in research_science_water_patterns:
            if re.search(pattern, text):
                # 先排除明显的金融/经济/政治/社会新闻文章
                # 这些领域中的"机制""原理""解释"等词不是科学研究含义
                non_research_context = ['股市', '熔断', '涨停', '跌停', '央行', '利率', '汇率', '货币', '金融', '证券', '基金',
                                       '制裁', '关税', '贸易', '外交', '峰会', '协议', '条约', '选举', '议会', '国会',
                                       '战争', '冲突', '军事', '国防', '武器', '军演',
                                       '犯罪', '违法', '监管', '执法', '逮捕', '起诉', '判决',
                                       '裁员', '停工', '降薪', '罢工', '就业', '失业',
                                       'IPO', '上市', '融资', '并购', '收购', '估值', '市值']
                if any(kw in text for kw in non_research_context):
                    continue  # 金融/经济/政治/社会新闻，跳过本条规则
                # 检查是否缺乏具体新闻事件（政策、监管、违规、重大突破等）
                if not any(kw in text for kw in ['政策', '监管', '批准', '上市', '重大', '突破', '违规', '处罚', '舆情', '调查', '召回', '禁售', '违法', '犯罪', '事故', '事件', '裁员', '降薪', '停工', '行业', '市场', '产业', '就业', '失业', '消费']):
                    # 检查是否是基础研究描述（缺乏具体数据支撑或重大发现）
                    if not any(kw in text for kw in ['重大', '突破', '首次', '革命性', '颠覆性', '批准', '上市']):
                        # 进一步检查是否是描述性研究（"可能"、"或许"、"补偿假说"等词汇）
                        # 注意：此处"机制""原理""解释""理论"仍作为辅助判断词，
                        # 但由于前面已排除金融/经济/政治上下文，不会再误杀这些领域的文章
                        if any(kw in text for kw in ['可能', '或许', '补偿', '假说', '机制', '原理', '解释', '理论']):
                            return True, f"基础科学研究科普类: 匹配'{pattern}'，缺乏新闻事件"
        
        # 27. 推文/简报/推广信息类
        # 识别"简报"、"月报"、"榜单"、"评选"等推广类文章
        # 注意：收紧"发布"的匹配，避免误杀正常新闻（如"小红书流量命门"含"发布"但不是简报）
        promotion_report_patterns = [
            r'\d+月.{0,5}简报',  # 月度简报
            r'(优胜|佳作|评出)',  # 评选类
            r'(投稿量|高校参与)',  # 统计报告类（移除"关键词"，太通用）
            r'(红网|青椒评论).{0,10}(简报|统计|分析)',  # 具体平台的统计
            r'发布.{0,5}(简报|月报|年报|季报|周报)',  # 发布简报（收紧：必须紧跟"简报/月报/年报"等）
            r'(特邀|邀请).{0,5}(教授|专家|评委)',  # 专家邀请
        ]
        for pattern in promotion_report_patterns:
            if re.search(pattern, text):
                # 检查是否是推广类文章（排除有实质新闻价值的）
                if not any(kw in text for kw in ['政策', '重大', '违规', '违法', '舆情', '调查', '处罚', '改革', '行业分析', '趋势', '数据', '报告']):
                    return True, f"推文/简报/推广信息类: 匹配'{pattern}'"
        
        # 28. 弱智故事/段子/鸡汤类
        # 识别明显是编造的故事、段子、鸡汤类文章
        stupid_story_patterns = [
            r'(船遇冰山|冰山)',  # 冰山段子
            r'段子.{0,10}(揭示|说明)',  # 用段子说明道理
            r'船.{0,5}(即将撞上|撞上|冰山)',  # 撞冰山场景
            r'(头等舱|二等舱|三等舱|四等舱|五等舱)',  # 船舱等级段子
            r'(认知.{0,5}(以外|差距|差异))',  # 认知差距类
            r'普通人.{0,10}(赚|获得)',  # 普通人赚钱类
            r'(如何赚.{0,5}(认知|认知以外))',  # 赚钱认知类
            r'段子',  # 明确是段子的
        ]
        for pattern in stupid_story_patterns:
            if re.search(pattern, text):
                # 检查是否是用故事/段子讲道理
                if any(kw in text for kw in ['段子', '揭示', '说明', '认知', '普通人', '赚钱']):
                    return True, f"弱智故事/段子类: 匹配'{pattern}'"
        
        # 29. 娱乐八卦/恶俗新闻类
        # 识别娱乐八卦、恶俗新闻、隐私曝光等
        entertainment_gossip_patterns = [
            r'(差评|评论).{0,10}(自曝|暴露)',  # 差评自曝隐私
            r'婚外情',  # 婚外情八卦
            r'(酒店|华住会).{0,10}(差评|自曝)',  # 酒店差评八卦
            r'(金会员|女士).{0,10}(差评|自曝)',  # 会员差评
            r'(吃瓜|围观|全网关注)',  # 吃瓜围观
            r'开房.{0,5}(入住|酒店)',  # 开房八卦
            r'情夫|情妇|小三|出轨',  # 出轨类八卦
        ]
        for pattern in entertainment_gossip_patterns:
            if re.search(pattern, text):
                # 娱乐八卦新闻不需要保留
                if not any(kw in text for kw in ['政策', '监管', '重大', '犯罪', '法律', '调查', '处罚']):
                    return True, f"娱乐八卦/恶俗新闻类: 匹配'{pattern}'"
        
        # 29.5 纯历史内容文章（全篇讲历史事件，无当前新闻价值）
        # 如"特朗普2017年访华签下2534亿美元大单"这种全篇回顾历史事件的文章
        # 注意：如果文章是当前事件分析+引用历史背景，应保留
        pure_history_patterns = [
            r'\d{4}年\d{1,2}月\d{1,2}日.{0,15}(访华|访美|签署|签下|达成)',
            r'(回顾|盘点|重温).{0,10}\d{4}年',
            r'\d{4}年.{0,10}(签下|签署|达成).{0,10}(亿|万)',
            r'(历史上的|当年).{0,15}(访华|访美|签约|合作)',
            r'(曾经|当年|那一年).{0,15}(签下|签署|达成).{0,10}(亿|万|协议|大单)',
        ]
        for pattern in pure_history_patterns:
            if re.search(pattern, text):
                # 排除当前事件+历史背景的文章（有"最新""今日""当前"等词则保留）
                if not any(kw in text for kw in ['最新', '今日', '当前', '今年', '本月', '本周', '刚刚', '宣布', '最新进展', '首次回应', '官方通报', '对比', '启示', '资本', '投资', '分析', '趋势', '风险', '危机', '逃离', '暴跌', '崩盘']):
                    return True, f"纯历史内容文章: 匹配'{pattern}'"

        # 29.6 比赛/竞赛宣传公告类文章
        # 如"全国医学模拟人与健康传感器大赛启动"等比赛招募、报名公告
        competition_announcement_patterns = [
            r'(大赛|竞赛|比赛|挑战赛|创新创业赛).{0,10}(启动|开始|报名|启动报名)',
            r'(大赛|竞赛|比赛).{0,10}(公告|通知|通告|规程)',
            r'(参赛|报名).{0,10}(对象|条件|要求|方式|截止)',
            r'(赛道|赛程|赛制).{0,10}(设置|安排|说明)',
            r'(一等奖|二等奖|三等奖|奖金|奖品).{0,10}(设置|金额|名额)',
            r'(主办|承办|协办).{0,10}(单位|机构|组织)',
            r'(成果转化|激励措施|获奖).{0,10}(方案|措施|办法)',
        ]
        for pattern in competition_announcement_patterns:
            if re.search(pattern, text):
                return True, f"比赛/竞赛宣传公告类: 匹配'{pattern}'"

        # 29.7 地方天气预报类文章（非极端天气）
        # 如"北京16日下午至18日早晨将迎明显降雨"等地方天气预报
        # 但极端天气（地震、火山、龙卷风、海啸等）需保留
        # 【关键原则】所有匹配都必须配合"天气主题词"出现，避免误伤"出行影响""降温"等比喻用法
        # 【强化】地方天气新闻（包括汛期预测、气候预测、防汛通知等）一律过滤，除非是极端灾害
        weather_forecast_patterns = [
            r'(最高温|最低温|最高气温|最低气温).{0,10}(降至|升至|达)',
            r'(降雨|下雨|暴雨|大雨|中雨|小雨|雷阵雨).{0,10}(天气|过程|影响|频繁)',
            # 【修复】"出行/外出"必须与天气词共同出现才算天气预报，避免"无证驾驶""出行影响"等被误判
            r'(外出.{0,5}(备好|添衣|注意).{0,15}(降雨|下雨|降温|大风|寒潮|高温))',
            r'(出行.{0,10}(影响|不利).{0,15}(降雨|下雨|降温|大风|寒潮|高温|降雪|大雪))',
            r'(天气预报|气象台|气象局).{0,10}(发布|预报|预计)',
            r'(降温|升温).{0,10}(度|℃)',
            r'(下周|本周|未来几天).{0,15}(天气|气温|降雨|晴)',
            r'(暴雨|大雨|暴雨预警).{0,10}(蓝色|黄色|橙色|红色)?(预警|信号)',  # 暴雨预警
            r'防汛.{0,10}(应急|响应|通知|启动|准备|形势|发布会)',  # 防汛应急响应/准备/形势
            r'(休市|停课|停运|封闭).{0,10}(暴雨|大雨|防汛|降雨)',  # 因暴雨休市停课
            r'(晴天|多云|阴天).{0,10}(天气|气温|预报)',  # 晴天/多云/阴天需带天气上下文
            r'(汛期|防汛).{0,20}(气候|预测|预判|预计|形势|工作)',  # 汛期气候预测/防汛形势
            r'(已进入|即将进入).{0,5}汛期',  # 进入汛期通知
            r'(降水|雨量).{0,10}(偏多|偏少|较常年)',  # 降水偏多/偏少预测
            r'(高温日数|强降水日数).{0,10}(偏多|预计)',  # 高温日数/强降水日数预测
            r'(雷阵雨|阵雨).{0,10}(频繁|多发)',  # 雷阵雨频繁
            r'(体感|舒适).{0,10}(温度|气温)',  # 体感温度
            r'(备好|常备).{0,10}(雨具|雨衣|雨伞)',  # 备好雨具建议
            r'(防汛办|防汛指挥部).{0,15}(呼吁|提醒|通知)',  # 防汛办呼吁
        ]
        # 【保护名单】包含这些关键词的文章绝对不被天气预报类过滤（投资/监管/重大商业事件）
        # 【重要】保护名单仅在标题（非摘要）包含金融/投资关键词时生效
        # 避免天气预报文章因摘要中顺带提及"上市""投资"等词而被误保护
        weather_filter_safe_keywords = [
            '证监会', '证券', '券商', '基金', '股票', '炒股', '美股', '港股', 'A股',
            '富途', '老虎', '长桥', '理财', '投资', '违法所得',
            '罚款', '罚没', '处罚', '立案', '调查', '稽查', '监管',
            '上市', 'IPO', '退市', '招股书', '财报', '营收', '净利润',
            '并购', '收购', '重组', '破产', '违约', '暴雷',
        ]
        # 仅在标题中包含安全关键词时才保护（避免摘要中顺带提及的无关金融词触发保护）
        title_safe = any(kw in self.title for kw in weather_filter_safe_keywords)
        # 同时检查：标题是否本身就是天气主题
        title_is_weather_topic = any(kw in self.title for kw in ['天气', '气温', '降雨', '雷阵雨', '阵雨', '汛期', '防汛', '暴雨', '气候预测', '高温日', '降水'])
        for pattern in weather_forecast_patterns:
            if re.search(pattern, text):
                # 保留极端天气相关文章
                extreme_weather_keywords = ['地震', '海啸', '火山', '龙卷风', '台风', '飓风', '泥石流', '山体滑坡', '洪灾', '特大暴雨', '极寒', '暴雪', '灾难', '伤亡', '遇难']
                if any(kw in text for kw in extreme_weather_keywords):
                    break  # 极端天气，保留
                # 如果标题本身就是天气主题，直接过滤（无论摘要内容如何）
                if title_is_weather_topic:
                    return True, f"地方天气预报类(标题天气主题): 匹配'{pattern}'"
                # 投资/监管类文章保护（仅在标题含金融关键词时生效）
                if title_safe:
                    break
                return True, f"地方天气预报类: 匹配'{pattern}'"
        # 29.7.1 "回暖"/"降温"/"升温"等词需结合天气语境判断
        # 避免将"行业回暖""经济升温""市场降温"等非天气用法误判为天气预报
        # 关键原则：当文章同时包含天气词和非天气词时，应以文章主旨（标题）判断，
        # 而非仅因出现天气词就过滤。例如"横店短剧大撤退"文中提及降温是比喻，
        # 文章主旨是短剧行业而非天气。
        weather_metaphor_words = ['回暖', '降温', '升温']
        for word in weather_metaphor_words:
            if word in text:
                # 检查是否在天气语境中使用（附近是否有天气相关词）
                weather_context_keywords = ['天气', '气温', '温度', '气象', '预报', '降雨', '降雪', '暴雨', '晴天', '多云', '阴天', '℃', '度', '寒潮', '冷空']
                non_weather_context_keywords = [
                    '行业', '经济', '市场', '楼市', '房价', '消费', '投资', '就业',
                    '票房', '营收', '利润', '需求', '供给', '增长', '复苏', '景气',
                    '短剧', '影视', '拍摄', '剧组', '裁员', '降薪', '停工', '撤退',
                    '饭碗', '挤出', '企业', '公司', '产业', '工厂', '岗位', '失业',
                    '裁员', '缩编', '收缩', '下行', '遇冷', '转冷', '寒冬',
                    '退潮', '降温', '泡沫', '崩盘', '洗牌', '出局', '关门',
                ]
                has_non_weather = any(kw in text for kw in non_weather_context_keywords)
                has_weather = any(kw in text for kw in weather_context_keywords)
                if has_non_weather and not has_weather:
                    continue  # 非天气语境，跳过
                if has_weather and not has_non_weather:
                    return True, f"地方天气预报类: 匹配'{word}'且有天气上下文"
                # 当天气词和非天气词同时存在时，以标题主旨判断文章是否为天气预报
                if has_weather and has_non_weather:
                    title_is_weather = any(kw in self.title for kw in weather_context_keywords)
                    title_is_non_weather = any(kw in self.title for kw in non_weather_context_keywords)
                    # 标题不是天气主题 → 文章主旨不是天气，跳过
                    if not title_is_weather:
                        continue
                    # 标题同时有非天气主题词 → 主旨模糊，保守不过滤
                    if title_is_non_weather:
                        continue
                    # 仅标题纯天气主题时才视为天气预报

        # 29.8 历史科普/历史考证类文章
        # 如"钱穆质疑烽火戏诸侯真实性：西周灭亡或另有隐情"等历史领域科普
        history_science_patterns = [
            r'(质疑|考证|考据|推翻).{0,15}(真实性|历史|记载|说法)',
            r'(或许|可能).{0,10}(另有|并非|不是).{0,10}(隐情|事实|真相)',
            r'(史记|资治通鉴|左传|汉书|三国志).{0,10}(记载|描述|说法)',
            r'(西周|东周|春秋|战国|秦朝|汉朝|唐朝|宋朝|明朝|清朝).{0,10}(灭亡|覆灭|兴衰|更替)',
            r'(烽火戏诸侯|卧薪尝胆|纸上谈兵|围魏救赵)',
            r'(考古.{0,10}发现|出土.{0,10}文物|古墓.{0,10}发现)',
            r'(历史.{0,10}(真相|谎言|谜团|秘密|未解))',
        ]
        for pattern in history_science_patterns:
            if re.search(pattern, text):
                # 排除与当前重大事件相关的历史对比分析
                if not any(kw in text for kw in ['最新', '今日', '当前', '政策', '官方', '首次回应', '对比', '启示', '借鉴']):
                    return True, f"历史科普/考证类: 匹配'{pattern}'"

        # 30. 个人资产配置/赚钱思路/投资策略类（缺乏资讯信息点的个人理财观点）
        # 识别"通缩通胀赚钱"、"资产配置思路"、"XX新思路"等纯个人投资观点文章
        investment_strategy_patterns = [
            r'资产配置.{0,10}(新思路|新方法|新策略)',  # 资产配置新思路
            r'(通缩|通胀).{0,10}(赚钱|获利|收益)',  # 通缩/通胀赚钱
            r'生活在.{0,10}(通缩|通胀).{0,10}赚',  # 生活在通缩赚通胀的钱
            r'(个人|家庭).{0,10}(资产配置|投资策略|理财思路)',  # 个人/家庭资产配置
            r'投资.{0,5}(新思路|新方法|新逻辑)',  # 投资新思路
            r'赚.{0,10}(通胀|通缩).{0,10}(钱|收益)',  # 赚通胀的钱
            r'(渡边太太|日元套利)',  # 渡边太太现象
            r'利用.{0,10}(差异|利差).{0,10}(赚|获利|收益)',  # 利用差异赚钱
        ]
        for pattern in investment_strategy_patterns:
            if re.search(pattern, text):
                # 这类文章只有个人投资思路，没有新闻信息点，直接过滤
                # 除非包含重大政策或市场变动信息
                if not any(kw in text for kw in ['政策', '监管', '央行', '降息', '加息', '重大', '违规', '违法', '调查']):
                    return True, f"个人资产配置/赚钱思路类: 匹配'{pattern}'"
        
        # 31. 一般院校合并/移交/更名类（加强版）
        # 识别具体的院校合并移交事件，如"酒泉卫生学校与工贸中专合并移交酒泉职业技术大学"
        school_merge_patterns = [
            r'(卫生学校|中专|技师学院|职业学校|工业学校|商贸学校).{0,20}(合并|并入|移交|划转)',  # 具体学校合并
            r'(合并|并入|移交|划转).{0,20}(卫生学校|中专|技师学院|职业学校|工业学校|商贸学校)',  # 反向
            r'(职业技术大学|应用技术大学).{0,10}(揭牌|成立|组建)',  # 职业技术大学揭牌
            r'(学校|学院).{0,10}(合并|并入|移交|划转).{0,10}(大学|学院)',  # 学校合并移交
            r'机构编制.{0,10}(通知|批复|调整)',  # 机构编制通知
            r'(编制|人员).{0,10}(划转|移交|调整)',  # 编制划转
        ]
        for pattern in school_merge_patterns:
            if re.search(pattern, text):
                # 一般院校合并/移交属于常规行政操作，直接过滤
                # 除非涉及重大丑闻或安全事件
                if not any(kw in text for kw in ['违规', '丑闻', '犯罪', '舆情', '热搜', '重大事故', '贪污', '腐败', '性侵', '学术不端', '坍塌', '伤亡']):
                    return True, f"一般院校合并/移交类: 匹配'{pattern}'"
        
        # 32. 聚合新闻/快报/简报类（科股快报、新闻快报等）
        # 识别"科股快报"、"新闻快报"、"早报"、"晚报"等聚合式新闻，每条信息只有一两句话
        aggregate_news_patterns = [
            r'科股快报',  # 科股快报
            r'新闻快报',  # 新闻快报
            r'科技快报',  # 科技快报
            r'8点1氪',  # 36氪的聚合快讯栏目
            r'财经早餐',  # 财经早餐
            r'(早报|晚报|晨报|午报).{0,5}[\|｜]',  # 早报/晚报+竖线分隔（聚合新闻特征）
            r'[\|｜].{0,20}(快报|简报|播报)',  # 竖线+快报
            r'(要闻|速递|早知道).{0,5}[\|｜]',  # 要闻速递+竖线
            r'(今日|每日).{0,5}(要闻|速递|资讯|播报).{0,5}[\|｜]',  # 每日要闻+竖线
            r'[\|｜].{2,15}[\|｜].{2,15}[\|｜]',  # 多个竖线分隔（聚合新闻典型特征）
        ]
        for pattern in aggregate_news_patterns:
            if re.search(pattern, text):
                # 聚合新闻每条信息太短，缺乏深度，直接过滤
                # 除非是深度专题报道
                if not any(kw in text for kw in ['深度', '专题', '独家', '调查', '解读']):
                    return True, f"聚合新闻/快报类: 匹配'{pattern}'"
        
        # 33. 城市对比类文章（缺乏新闻价值的话题讨论）
        # 识别"A对标B"、"A是B的缩影"等两个城市或地区进行简单对比的文章
        city_comparison_patterns = [
            r'对标.{0,10}(的|某)',  # 对标某（城市/地区）
            r'是.{0,5}(的|某).{0,5}(缩影|翻版|复制)',  # 是某的缩影/翻版
            r'堪比.{0,5}(的|某)',  # 堪比某（城市）
            r'(安徽|浙江|江苏|广东|四川|湖北|山东|河南).{0,10}的.{0,10}(苏州|深圳|杭州|成都|武汉)',  # 某省/市的苏州/深圳类比
            r'(小苏州|小深圳|小杭州|小成都|小武汉)',  # 小某类比
            r'下一个.{0,5}(苏州|深圳|杭州)',  # 下一个某
            r'逻辑.{0,10}(相似|一致|相同)',  # 逻辑相似
            r'区位.{0,10}(逻辑|优势)',  # 区位逻辑
        ]
        for pattern in city_comparison_patterns:
            if re.search(pattern, text):
                # 城市对比类文章缺乏具体新闻事件，直接过滤
                if not any(kw in text for kw in ['政策', '重大', '签约', '开工', '开通', '突破', '获批', '规划', '建设', '产业转移']):
                    return True, f"城市对比类文章: 匹配'{pattern}'"
        
        # 34. 营养科普/健康建议类文章（缺乏重大新闻价值的健康知识普及）
        # 识别每日摄入XX克XXX可降低XX%风险等健康科普文章
        nutrition_health_patterns = [
            r'每日.{0,10}(摄入|服用|补充).{0,10}(克|毫克|微克)',  # 每日摄入XX克
            r'(降低|减少|增加).{0,10}(%|\d+%)',  # 降低XX%
            r'全因.{0,5}(死亡|风险)',  # 全因死亡风险
            r'心血管.{0,10}(死亡|风险|疾病)',  # 心血管死亡/风险
            r'膳食纤维',  # 膳食纤维相关
            r'可以.{0,10}(降低|减少|预防)',  # 可以降低/预防
            r'几乎.{0,10}(无副作用|无任何副作用|安全)',  # 无副作用
            r'推荐.{0,10}(每日|每天|摄入)',  # 推荐每日摄入
            r'每日.{0,10}(服用|摄入).{0,10}(即可|就能)',  # 每日XX即可
        ]
        for pattern in nutrition_health_patterns:
            if re.search(pattern, text):
                # 营养科普文章缺乏具体新闻事件，直接过滤
                if not any(kw in text for kw in ['政策', '监管', '重大', '违规', '违法', '处罚', '舆情', '调查', '召回', '禁售']):
                    return True, f"营养科普/健康建议类: 匹配'{pattern}'"
        
        # 35. 医学案例/故事类文章（除非涉及重大事件，否则过滤）
        # 识别具体病人的诊疗故事、病历问题等医学案例文章
        medical_case_story_patterns = [
            r'\d{2}岁.{0,10}(男子|女士|先生|患者)',  # XX岁男子/患者
            r'就诊.{0,10}(医院|诊所|门诊)',  # 就诊某医院
            r'医保.{0,10}(凭空|莫名|多出)',  # 医保凭空多出
            r'病历.{0,10}(问题|错误|虚假|伪造)',  # 病历问题
            r'诊断.{0,10}(慢性胃炎|高血压|糖尿病)',  # 具体诊断
            r'医生.{0,10}(已离职|离职)',  # 医生离职
            r'拒保',  # 保险拒保
            r'体检.{0,10}(结果正常|指标正常|无异常)',  # 体检正常
            r'\d{2}岁.{0,10}女医生',  # XX岁女医生案例
            r'每年体检正常.{0,10}确诊',  # 体检正常但确诊晚期
            r'从发现.{0,10}离世',  # 从发现癌症到离世
        ]
        for pattern in medical_case_story_patterns:
            if re.search(pattern, text):
                # 除非涉及重大医疗事件、丑闻、政策，否则过滤
                if not any(kw in text for kw in ['政策', '监管', '重大', '违规', '违法', '处罚', '舆情', '热搜', '丑闻', '调查', '召回', '禁售', '大规模', '群体性']):
                    return True, f"医学案例/故事类: 匹配'{pattern}'"
        
        # 36. 娱乐八卦/争议人物类文章（垃圾八卦信息）
        # 识别明星/运动员/公众人物的八卦、隐私、争议讨论
        celebrity_gossip_patterns = [
            r'他妈.{0,5}(又|出手|干预)',  # 家人干预类八卦
            r'(孙杨|某明星|某运动员).{0,10}(被删|删片段|热搜)',  # 争议人物热搜
            r'凡事.{0,10}找妈妈',  # 找妈妈类吐槽
            r'过度.{0,10}(保护|干预)',  # 过度保护
            r'被逼.{0,10}(学|练)',  # 被逼学/练
            r'口碑.{0,10}(危机|崩塌)',  # 口碑危机
            r'巨婴',  # 巨婴评价
            r'热搜.{0,10}(会过|已过)',  # 热搜会过类感慨
        ]
        for pattern in celebrity_gossip_patterns:
            if re.search(pattern, text):
                # 娱乐八卦/争议人物文章直接过滤
                if not any(kw in text for kw in ['政策', '监管', '重大', '违法', '犯罪', '法律', '调查', '判决', '处罚']):
                    return True, f"娱乐八卦/争议人物类: 匹配'{pattern}'"
        
        # 36.5 非知名院校/地方院校动态过滤（除非涉及重大热点事件）
        # 识别非知名地方院校的日常行政动态，缺乏全国性影响
        local_college_patterns = [
            r'办学条件.{0,20}达标|办学条件.{0,20}达到|办学条件.{0,20}基本要求',  # 地方院校办学条件合格评估
            r'本科教学工作合格评估',  # 地方院校合格评估
        ]
        for pattern in local_college_patterns:
            if re.search(pattern, text):
                # 除非涉及重大热点事件，否则过滤
                if not any(kw in text for kw in ['重大', '舆情', '热搜', '丑闻', '违法', '犯罪', '调查', '判决', '政策']):
                    return True, f"非知名院校/地方院校动态: 匹配'{pattern}'"
        
        # 37. 八卦隐私/暧昧丑闻类文章（绝对不允许保留）
        # 识别领导/官员/公众人物的暧昧信息、隐私八卦
        scandal_gossip_patterns = [
            r'误发.{0,10}(亲爱的|亲密)',  # 误发亲爱的
            r'(副校长|院长|书记|领导).{0,10}(工作群|微信群).{0,10}(发|误发)',  # 领导工作群误发
            r'(暧昧|亲密).{0,10}(短信|消息|聊天)',  # 暧昧短信
            r'(老夫老妻|家人).{0,10}(亲爱的|心疼)',  # 老夫老妻用亲爱的（质疑）
            r'(暗中|偷偷).{0,10}(做|帮忙)',  # 暗中做某事
            r'微信群.{0,10}(手滑|翻车)',  # 微信群手滑
            r'(婚外情|出轨|暧昧关系)',  # 婚外情/出轨
            r'(丑闻|桃色).{0,10}(曝光|爆出|被扒)',  # 丑闻曝光
        ]
        for pattern in scandal_gossip_patterns:
            if re.search(pattern, text):
                # 八卦隐私/暧昧丑闻类文章直接过滤
                return True, f"八卦隐私/暧昧丑闻类: 匹配'{pattern}'"
        
        # 37.5 隐晦暗示/低俗标识类文章
        # 识别公共场所标识设计、隐晦暗示、低俗内容等文章
        subtle_suggestive_patterns = [
            r'(香蕉|西柚|水果).{0,10}(标识|图案|图案|符号|标志)',  # 水果隐晦标识（如影院卫生间香蕉西柚）
            r'标识.{0,10}(隐晦|暗示|性暗示)',  # 标识隐晦暗示
            r'(影院|商场|公共).{0,10}(标识|图案|符号).{0,10}(引发争议|争议)',  # 公共场所标识争议
            r'(男性|女性).{0,5}(标识|标志|图案).{0,10}(水果|食物|水果)',  # 性别标识配食物
            r'抠开.{0,10}(西柚|水果)',  # 抠开西柚等低俗表达
            r'(卫生间|厕所|洗手间).{0,10}(水果|食物|香蕉|西柚)',  # 卫生间水果标识
        ]
        for pattern in subtle_suggestive_patterns:
            if re.search(pattern, text):
                # 隐晦暗示类文章直接过滤
                return True, f"隐晦暗示/低俗标识类: 匹配'{pattern}'"
        
        # 39. 普通学术论文/研究类（非重大发现）
        # 如"地方依恋多维结构影响本地球迷认同与忠诚度"等普通学术研究
        academic_paper_patterns = [
            r'(结构方程模型|SEM|fsQCA|定性比较分析)',
            r'(实证研究|实证分析|问卷调查).{0,15}(影响|相关|中介|调节)',
            r'(中介效应|调节效应|路径分析|因子分析)',
            r'(样本量|信度|效度|拟合度|RMSEA|CFI|TLI|SRMR)',
            r'(地方依恋|粉丝认同|品牌忠诚|消费意愿).{0,10}(研究|影响|相关)',
            r'(基于\d+名|基于\d+个).{0,10}(被试|样本|受访者|参与者)',
            r'(\d{2,3}名).{0,10}(被试|样本|受访者|参与者).{0,10}(研究|调查)',
        ]
        for pattern in academic_paper_patterns:
            if re.search(pattern, text):
                # 普通学术论文直接过滤，除非是重大发现（如Nature/Science级别）
                if not any(kw in text for kw in ['Nature', 'Science', 'Cell', 'Lancet', 'NEJM', '突破', '首次', '革命性', '里程碑', '颠覆']):
                    return True, f"普通学术论文/研究类: 匹配'{pattern}'"
        
        # 40. 非重大学生个人事件类
        # 如"清华大学对一名博士生作出退学处理"等个别学生事件
        student_trivial_patterns = [
            r'(退学处理|予以退学|勒令退学|自动退学)',
            r'(最长修业年限|未毕业且未结业)',
            r'(公告送达|无法直接送达|难于联系)',
            r'(学生申诉处理委员会|申诉)',
            r'(开除学籍|留校察看|记过处分)',
            r'(考试作弊|学术不端).{0,10}(处分|处理|通报)',
        ]
        for pattern in student_trivial_patterns:
            if re.search(pattern, text):
                # 个别学生事件直接过滤，除非涉及重大恶性事件
                if not any(kw in text for kw in ['跳楼', '自杀', '杀人', '投毒', '重大事故', '群体性', '舆情', '热搜', '刑事案件', '死亡', '命案']):
                    return True, f"非重大学生个人事件类: 匹配'{pattern}'"
        
        # 41. 地方政府普通人事任免类
        # 如"湖南省人民政府任免多所高校领导干部"等常规人事调整
        local_personnel_patterns = [
            r'(任免|任命|免去).{0,10}(工作人员|干部|领导|同志)',
            r'(同志).{0,5}(任|免去|免).{0,10}(局长|院长|校长|主任|副)',
            r'(省人民政府|市人民政府|县人民政府).{0,15}(任免|发布任免)',
            r'(省信访局|省地质院|省科学技术事务中心)',
            r'(农业大学|林业科技大学).{0,10}(副校|任免|调整)',
        ]
        for pattern in local_personnel_patterns:
            if re.search(pattern, text):
                # 普通地方人事任免直接过滤，除非涉及高级别（副省级以上）或重大事件
                if not any(kw in text for kw in ['省委', '中央', '国务院', '部长', '省长', '反腐', '落马', '被查', '双开', '违纪', '违法', '犯罪']):
                    return True, f"地方普通人事任免类: 匹配'{pattern}'"
        
        # 42. 微小地方事件类（缺乏全国性新闻价值的地方小事件）
        # 如"救护车驾驶员用救护车捎带化肥"等地方微小事件
        trivial_local_event_patterns = [
            r'(救护车|急救车).{0,15}(捎带|装载|私用|私载|装).{0,10}(化肥|私人物品|菜|货物)',
            r'(乡镇|村级|街道).{0,15}(违规|不当).{0,10}(使用|使用)',
            r'(驾驶员|司机).{0,10}(私用|私载|捎带)',
            r'(一乡镇|某乡镇|某街道).{0,15}(违规|不当)',
        ]
        for pattern in trivial_local_event_patterns:
            if re.search(pattern, text):
                # 微小地方事件直接过滤，除非有全国性影响
                if not any(kw in text for kw in ['全国', '热搜', '舆情', '中央', '国务院', '重大', '群死群伤', '命案']):
                    return True, f"微小地方事件类: 匹配'{pattern}'"
        
        # 28. 股市单日资金流向/龙虎榜数据类（缺乏分析价值）
        # 识别北向资金、主力资金净流入/流出、龙虎榜等纯数据播报类文章
        stock_data_flow_patterns = [
            r'北向资金.{0,10}(成交|净买|净卖|流入|流出)',
            r'(主力资金|资金流向).{0,10}(净流入|净流出|净买|净卖)',
            r'龙虎榜.{0,10}(净买|净卖|居首|上榜)',
            r'(沪股通|深股通).{0,10}(成交|净买|净卖)',
            r'(两市|A股).{0,5}(成交额|缩量|放量).{0,10}(亿元|万亿)',
            r'(主力|机构).{0,10}(净买入|净卖出).{0,10}(居首|排名)',
        ]
        for pattern in stock_data_flow_patterns:
            if re.search(pattern, text):
                # 纯数据播报类缺乏分析价值，除非有行业/政策层面的深度解读
                if not any(kw in text for kw in ['政策', '监管', '行业', '战略', '重大', '改革', '突破', '转型', '制度']):
                    return True, f"股市资金流向/数据播报类: 匹配'{pattern}'"
        
        # 29. 考公/考编/考研备考建议与路线选择类（缺乏新闻价值）
        # 识别给考生的备考建议、路线选择、策略分析等教育指南类文章
        exam_advice_patterns = [
            r'(国考|省考|选调生).{0,15}(差异|选择|区别|路线|解析|攻略|指南)',
            r'(考公|考编|考研).{0,15}(建议|策略|选择|路线|如何选|怎么选|备考)',
            r'(国考|省考|选调生).{0,5}(门槛|要求|条件|专业)',
            r'(行测|申论).{0,10}(题量|难度|备考|策略)',
            r'(应届生|往届生).{0,10}(优先|适合|选择).{0,10}(国考|省考|选调)',
        ]
        for pattern in exam_advice_patterns:
            if re.search(pattern, text):
                # 备考建议类缺乏新闻价值，除非涉及政策变化
                if not any(kw in text for kw in ['改革', '政策变化', '新规', '重大调整', '扩招', '缩招', '首次']):
                    return True, f"考公/考编备考建议类: 匹配'{pattern}'"
        
        # 30. 普通天气景观资讯类（缺乏新闻价值）
        # 识别七彩祥云、彩虹、晚霞、日出等普通天气景观资讯
        # 注意：极端天气灾害（台风、暴雨、地震等）已在其他规则中保留
        weather_landscape_patterns = [
            r'(七彩祥云|彩云|彩虹|晚霞|朝霞|日出景观|火烧云)',
            r'(天空.{0,5}(出现|现).{0,10}(祥云|彩云|彩虹|奇观|美景))',
            r'(市民.{0,10}(拍照|拍摄|记录).{0,10}(祥云|彩云|彩虹|天空|美景))',
            r'(见者.{0,5}(有份|好运|幸运))',
        ]
        for pattern in weather_landscape_patterns:
            if re.search(pattern, text):
                # 普通天气景观直接过滤
                return True, f"普通天气景观资讯类: 匹配'{pattern}'"
        
        # 31. 基于单一企业业绩推断城市经济对比的不客观文章
        # 识别"某企业业绩增长→某城市GDP能否超越另一城市"这种缺乏严谨性的推论
        city_economy_by_company_patterns = [
            r'(能否|能否最终).{0,10}(超越|超过|赶超|反超)',
            r'(GDP.{0,5}(能否|是否).{0,5}(超越|超过|赶超|反超))',
            r'(净利润.{0,10}暴增|营收.{0,10}暴增).{0,30}(超越|超过|赶超)',
            r'(某城市|某省).{0,15}(GDP|经济).{0,10}(超越|超过|赶超|反超)',
            r'(业绩|营收|利润).{0,20}(能否|是否).{0,10}(带动|推动|助力).{0,10}(超越|赶超)',
        ]
        for pattern in city_economy_by_company_patterns:
            if re.search(pattern, text):
                # 检查是否确实是在用单一企业推断城市经济
                if any(kw in text for kw in ['GDP', '经济总量', '超越', '赶超', '反超']) and \
                   any(kw in text for kw in ['营收', '利润', '业绩', '暴增', '净利润']):
                    return True, f"单一企业推断城市经济对比类: 匹配'{pattern}'"
        
        # 32. 纯主观审美判断文章（缺乏客观新闻价值）
        # 识别讨论产品/设计"美丑"的纯主观判断文章，这类文章没有客观事实支撑
        aesthetic_judgment_patterns = [
            r'(丑|好看|美观|审美).{0,5}(出圈|出圈了|火了|刷屏)',
            r'(都在笑你|都在骂|群嘲|吐槽).{0,10}(丑|好看|美)',
            r'(收手吧|别再).{0,10}(丑|设计|外观)',
            r'(设计.{0,5}(奇葩|雷人|丑出|丑到|辣眼睛|不堪入目))',
            r'(丑鞋|丑衣服|丑包|丑到|丑出)',
            r'(审美.{0,5}(崩塌|滑坡|降级|跑偏|灾难))',
            r'(越旧越有态度|有穿上就不易被驯服)',
        ]
        for pattern in aesthetic_judgment_patterns:
            if re.search(pattern, text):
                # 纯审美判断文章直接过滤，除非涉及重大商业事件（如公司出售、重大裁员）
                if not any(kw in text for kw in ['出售', '被收购', '退市', '破产', '重大裁员', '关店潮', '大规模']):
                    return True, f"纯主观审美判断类: 匹配'{pattern}'"
        
        # 33. 主观非资讯类文章（以个人经历/主观感受为主，缺乏客观新闻价值）
        # 如"四川医生自述一人管80张床"等，这类文章以个人叙述、主观感受、情绪表达为主
        # 不是客观的新闻报道，而是个人故事/职业日常/主观吐槽
        subjective_non_news_patterns = [
            r'(自述|自白|独白).{0,10}(崩溃|崩溃瞬间|心酸|崩溃了)',
            r'(一人.{0,5}(管|负责|管理).{0,5}\d+.{0,5}床)',  # 一人管XX床
            r'(疲劳.{0,5}(行医|驾驶|上岗|工作)|行医.*疲劳)',  # 疲劳行医/工作
            r'(崩溃瞬间|破防瞬间)',  # 个人崩溃瞬间
            r'(晒出.{0,10}(截图|聊天记录|排班表))',  # 晒截图
            r'(白\+黑|连轴转).{0,10}(排班|值班|加班|模式)',  # 白+黑连轴转
            r'(下夜班).{0,10}(还需|还要|仍需).{0,10}(查房|开会|处理)',  # 下夜班还要工作
            r'(不是.{0,5}(歌颂|赞美|表彰).{0,5}(而是|而是需要))',  # "需要的不是歌颂而是..."
            r'(值完班|下夜班).{0,10}(毫无|没有).{0,10}(准备|休息)',  # 值完班毫无准备
            r'(标配|日常缩影).{0,10}(疲劳|高强度|超负荷)',  # 成为标配
            r'(极限.{0,5}\d+.{0,5}小时).{0,10}(排班|值班|上班)',  # 极限36小时排班
        ]
        for pattern in subjective_non_news_patterns:
            if re.search(pattern, text):
                # 只有当文章涉及国家级政策/改革/官方正式回应时才保留
                # 仅在结尾顺带提及"制度"不构成保留理由
                # 必须有明确的政策/改革/官方调查等客观新闻要素
                if not any(kw in text for kw in ['国家卫健委', '国务院', '卫健委.*发布', '官方.*回应',
                                                   '改革方案', '立法', '调查组', '通报.*处分',
                                                   '政策.*出台', '制度.*改革']):
                    # 进一步检查：如果"制度/政策/改革"仅出现在文章结尾的呼吁/建议中（而非事实报道），不构成保留理由
                    has_substantial_policy = False
                    if any(kw in text for kw in ['政策', '改革', '制度']):
                        # 检查这些词是否出现在前半部分（事实描述区域），而非仅仅是结尾的呼吁
                        first_half = text[:len(text)//2]
                        if any(kw in first_half for kw in ['政策', '改革方案', '制度改革']):
                            has_substantial_policy = True
                    if not has_substantial_policy:
                        return True, f"主观非资讯类: 匹配'{pattern}'"
        
        # 34. 主观时政评价类（对政治人物/事件的纯主观评价，非专业分析）
        # 如"泽连斯基公开信从五个维度被评为一类文，议题设置能力获赞"
        # 这类文章是对政治人物/事件的主观打分/评价/点赞，缺乏客观事实信息
        # 注意：专业机构的政策分析、国际关系深度解读不应被过滤
        subjective_political_eval_patterns = [
            r'(被评为|获评).{0,15}(一类文|二类文|满分|高分|优秀)',
            r'(议题设置|叙事策略|修辞).{0,10}(能力|技巧|手法).{0,10}(获赞|获评|被赞)',
            r'(从.{1,5}个维度|从.{1,5}个角度).{0,10}(评价|分析|点评|打分)',
            r'(被赞|获赞|赢得).{0,15}(认可|赞誉|好评|点赞)',
            r'(公开信|演讲).{0,15}(被评为|被赞|获赞|获评)',
            r'(政治人物|领导人|总统|总理).{0,10}(表现|发挥).{0,10}(评分|打分|评级)',
        ]
        # 政治人物关键词（用于限定本规则仅针对涉及时政人物的评价文章）
        political_figure_keywords = ['泽连斯基', '普京', '特朗普', '拜登', '马克龙', '朔尔茨', '岸田',
                                     '尹锡悦', '内塔尼亚胡', '金正恩', '莫迪', '鲁比奥', '拉马福萨',
                                     '卢拉', '米莱', '苏纳克', '哈里斯', '万斯']
        has_political_figure = any(kw in text for kw in political_figure_keywords)
        if has_political_figure:
            for pattern in subjective_political_eval_patterns:
                if re.search(pattern, text):
                    # 排除：专业智库/机构的政策分析、国际关系学术解读、事实性报道
                    if not any(kw in text for kw in ['智库', '研究机构', '政策分析', '战略分析',
                                                       '数据', '统计', '调查', '报告', '官方', '通报',
                                                       '制裁', '停火', '协议', '决议', '法案', '军援',
                                                       '突破', '首次', '宣布', '签署']):
                        return True, f"主观时政评价类: 匹配'{pattern}'"
        
        # 35. 演员娱乐八卦类（演员个人表演/标签/形象等八卦内容，非社会重大热点事件）
        # 如"李泽锋凭《主角》摆脱'渣男'标签，自加细节获张嘉益认可"
        # 这类文章关注的是演员的表演标签、形象转变等，缺乏客观新闻价值
        # 除非涉及社会重大热点事件（如#MeToo、行业整顿等）
        actor_gossip_patterns = [
            r'(摆脱|撕掉|洗掉|甩掉).{0,10}(标签|人设|形象|称号)',
            r'(获|得到).{0,10}(认可|好评|称赞|点赞|肯定)',
            r'(自加|添加).{0,10}(细节|设计|动作|桥段)',
            r'(渣男|绿茶|白莲花|霸总).{0,5}(标签|人设|形象)',
            r'(演技|表演).{0,10}(炸裂|封神|突破|蜕变|惊艳)',
            r'(凭|凭借).{0,15}(摆脱|撕掉|洗掉|甩掉)',
            r'(角色|人物).{0,10}(塑造|刻画|演绎).{0,10}(获|得到).{0,5}(认可|好评)',
        ]
        # 演员关键词（用于限定本规则仅针对演员相关文章）
        actor_context_keywords = ['演员', '饰演', '主角', '配角', '男主角', '女主角', '男一号', '女一号',
                                  '影帝', '影后', '视帝', '视后', '老戏骨', '小鲜肉', '流量明星',
                                  '张嘉益', '李泽锋', '演技', '表演', '电视剧', '电影', '综艺']
        has_actor_context = any(kw in text for kw in actor_context_keywords)
        if has_actor_context:
            for pattern in actor_gossip_patterns:
                if re.search(pattern, text):
                    # 排除：社会重大热点事件（如#MeToo、行业整顿、税务违法、犯罪等）
                    if not any(kw in text for kw in ['违法', '犯罪', '偷税', '漏税', '吸毒', '家暴',
                                                       '性侵', '性骚扰', '行业整顿', '封杀', '禁演',
                                                       '政策', '监管', '处罚', '起诉', '判决']):
                        return True, f"演员娱乐八卦类: 匹配'{pattern}'"
        
        # 36. 节假日休市/休市安排类（如沪深北交易所端午节休市安排，无新闻价值）
        holiday_closure_patterns = [
            r'休市.{0,10}(安排|通知|公告)',
            r'(交易所|股市|A股|港股|美股).{0,10}(休市|放假|闭市)',
            r'(端午节|春节|国庆节|中秋节|劳动节|清明节).{0,10}(休市|放假|闭市)',
            r'(节假日|假期).{0,10}(休市|闭市|不开盘)',
        ]
        for pattern in holiday_closure_patterns:
            if re.search(pattern, text):
                return True, f"节假日休市类: 匹配'{pattern}'"
        
        # 37. 高考建议/指南/攻略类（针对高考生的建议、必看、攻略等，非高考新闻本身）
        # 如"2026高考生必看：核心概念与政策全解析"——这是备考建议，不是新闻
        # 注意：高考新闻（如高考作文题、考试时间调整、重大违规事件等）应保留，
        # 仅过滤对高考生的建议/攻略/指南类内容
        gaokao_advice_patterns = [
            r'高考.{0,5}(生|考生).{0,10}(必看|必读|攻略|指南|建议|秘籍|宝典|手册|全解析|必知)',
            r'(必看|必读|攻略|指南).{0,10}高考.{0,5}(生|考生)',
            r'高考.{0,5}(备考|冲刺|提分|复习|答题|志愿填报).{0,10}(技巧|方法|策略|建议|指南|攻略)',
            r'(高考|考生).{0,10}(注意事项|考前提醒|考前必做|考前准备)',
            r'高考.{0,5}(家长|父母).{0,10}(必看|必读|注意|建议|指南|攻略)',
        ]
        for pattern in gaokao_advice_patterns:
            if re.search(pattern, text):
                # 排除：真正的高考新闻（如高考作文题出炉、考试时间调整、重大违规事件等）
                if not any(kw in text for kw in ['作文题', '考试时间', '调整', '违规', '作弊', '泄题', 
                                                   '改革', '政策变化', '录取', '分数线', '成绩', '处罚',
                                                   '舆情', '热搜', '犯罪', '重大']):
                    return True, f"高考建议/指南类: 匹配'{pattern}'"
        
        # 38. 博物馆/场馆闭馆/休整/调整等小新闻（无全国性新闻价值）
        # 如"首都博物馆5月18日起调整闭馆日，每周二闭馆并取消个人停车服务"
        venue_closure_patterns = [
            r'(博物馆|美术馆|图书馆|科技馆|展览馆|纪念馆|文化馆|大剧院|音乐厅).{0,15}(闭馆|休馆|调整闭馆|暂停开放|取消)',
            r'(闭馆|休馆|暂停开放).{0,15}(博物馆|美术馆|图书馆|科技馆|展览馆|纪念馆|文化馆)',
            r'(博物馆|美术馆|图书馆).{0,15}(停车|车位|停车服务).{0,10}(取消|调整|暂停)',
            r'(场馆|展馆).{0,10}(维护|修缮|改造|升级).{0,10}(闭馆|暂停|临时)',
        ]
        for pattern in venue_closure_patterns:
            if re.search(pattern, text):
                # 排除：涉及重大事件（如文物被盗、重大发现、安全事故等）
                if not any(kw in text for kw in ['被盗', '失窃', '重大发现', '安全', '事故', '火灾',
                                                   '丑闻', '争议', '抗议', '犯罪', '处罚']):
                    return True, f"博物馆/场馆闭馆小新闻: 匹配'{pattern}'"
        
        # 39. 教学/指南类文章（教别人做某事，如健身动作教学、烹饪食谱、手工DIY等）
        # 这类文章的核心是"教你怎么做"，而不是报道新闻事件，缺乏客观新闻价值
        # 如"八段锦升级版新增呼吸节奏提示，在家即可跟练全套动作"——这是健身教学，不是新闻
        tutorial_patterns = [
            r'(八段锦|太极拳|五禽戏|易筋经|六字诀).{0,15}(跟练|教学|教程|动作|入门|版|升级版|拆解|详解|示范)',
            r'(瑜伽|普拉提|健身操|广播体操).{0,15}(跟练|教学|教程|动作|入门|示范|拆解)',
            r'(在家|跟着|一起).{0,10}(练|做|学|跟练|跟做).{0,10}(动作|操|式|全套|版)',
            r'(升级版|完整版|全套|详细版).{0,10}(跟练|教学|教程|动作分解)',
            r'(动作|姿势|体式).{0,10}(分解|详解|要点|要领|示范|纠正)',
            r'(烹饪|食谱|做法|菜谱|烘焙).{0,10}(教程|教学|步骤|入门|详解)',
            r'(手工|DIY|编织|折纸).{0,10}(教程|教学|步骤|入门|详解)',
        ]
        for pattern in tutorial_patterns:
            if re.search(pattern, text):
                # 排除：如果文章涉及重大政策/争议/事故等新闻要素则保留
                if not any(kw in text for kw in ['政策', '争议', '事故', '处罚', '违规', '调查',
                                                   '改革', '行业标准', '监管', '质量', '安全', '召回',
                                                   '舆情', '热搜', '犯罪']):
                    return True, f"教学/指南类文章: 匹配'{pattern}'"
        
        return False, ""

def _set_font(paragraph, font_name: str, size):
    """设置段落的字体（辅助函数）"""
    from docx.oxml.ns import qn
    from docx.shared import Pt
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        # 强制清除任何继承的格式
        if run._element.rPr is not None:
            # 确保字号设置生效
            size_elem = run._element.rPr.find(qn('w:sz'))
            if size_elem is None:
                from docx.oxml import OxmlElement
                size_elem = OxmlElement('w:sz')
                size_elem.set(qn('w:val'), str(int(size * 2)))
                run._element.rPr.append(size_elem)
            else:
                size_elem.set(qn('w:val'), str(int(size * 2)))

def _add_internal_hyperlink(doc, text: str, anchor: str, font_name: str = '微软雅黑', font_size: int = 12):
    """
    添加内部超链接（指向文档内书签）
    
    Args:
        doc: Document对象
        text: 超链接显示文本
        anchor: 目标书签名称
        font_name: 字体名称
        font_size: 字体大小
    
    Returns:
        新创建的段落
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    
    # 创建超链接元素
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    
    # 创建run元素
    new_run = OxmlElement('w:r')
    
    # 创建rPr（run属性）
    rPr = OxmlElement('w:rPr')
    
    # 字体设置
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    
    # 字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    
    # 颜色
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0066CC')
    rPr.append(color)
    
    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # 创建文本元素
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    
    # 创建段落并添加超链接
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    para._element.append(hyperlink)
    
    return para

def _add_external_hyperlink(paragraph, url: str, text: str, font_name: str = '微软雅黑', font_size: float = 10.5):
    """
    在段落中添加外部超链接
    
    Args:
        paragraph: 段落对象
        url: 目标URL
        text: 超链接显示文本
        font_name: 字体名称
        font_size: 字体大小（磅）
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 获取段落的XML元素
    para_xml = paragraph._element
    
    # 创建超链接元素
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), '')  # 内部链接用anchor，外部链接用下面的方式
    
    # 使用 relationships 添加外部链接
    # 获取 document part
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink.set(qn('r:id'), r_id)
    
    # 创建 run 元素
    new_run = OxmlElement('w:r')
    
    # 创建 rPr（run属性）
    rPr = OxmlElement('w:rPr')
    
    # 字体设置
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    
    # 字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    
    # 颜色（蓝色）
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # 创建文本元素
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    
    hyperlink.append(new_run)
    
    para_xml.append(hyperlink)

def _add_bookmark(paragraph, bookmark_id: str, bookmark_name: str):
    """
    给段落添加书签
    
    Args:
        paragraph: 段落对象
        bookmark_id: 书签ID
        bookmark_name: 书签名称
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 创建书签开始元素
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # 创建书签结束元素
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    
    # 将书签添加到段落的run之前和之后
    para_element = paragraph._element
    # 在第一个run之前插入开始书签
    first_run = para_element.find(qn('w:r'))
    if first_run is not None:
        first_run.addprevious(bookmark_start)
    else:
        para_element.insert(0, bookmark_start)
    
    # 在最后一个run之后插入结束书签
    last_run = para_element.findall(qn('w:r'))
    if last_run:
        last_run[-1].addnext(bookmark_end)
    else:
        para_element.append(bookmark_end)

class DocumentGenerator:
    """文档生成器"""

    # MECE分类顺序定义（按优先级排序）
    MECE_SORT_ORDER = [
        "1.1", "1.2", "1.3",
        "2.1", "2.2", "2.3", "2.4", "2.5",
        "3.1", "3.2", "3.3", "3.4",
        "4.1", "4.2", "4.3", "4.4",
        "5.1", "5.2", "5.3",
        "6.1", "6.2", "6.3", "6.4",
        "7.1", "7.2", "7.3", "7.4",
        "8.1", "8.2", "8.3", "8.4",
        "9.1", "9.2", "9.3", "9.4",
        "10.1", "10.2", "10.3",
        "11.1", "11.2", "11.3",
    ]

    # 大类名称映射（用于Word显示）
    MECE_MAIN_CATEGORIES = {
        "1": "国际局势与地缘政治",
        "2": "宏观经济与金融",
        "3": "文体娱乐内容",
        "4": "科技产业动态与创新",
        "5": "医疗健康与生命科学",
        "6": "教育发展与人才培养",
        "7": "社会民生与消费",
        "8": "企业与商业",
        "9": "政策法规与监管",
        "10": "能源与资源",
        "11": "社会热点与舆论动态",
    }

    @staticmethod
    def _get_main_category_name(mece_category: str) -> str:
        """从完整分类编号获取大类名称"""
        if '.' in mece_category:
            main_cat = mece_category.split('.')[0]
        else:
            main_cat = mece_category
        return DocumentGenerator.MECE_MAIN_CATEGORIES.get(main_cat, "其他")

    @staticmethod
    def create_document(articles: List[Article], output_path: str, daily_summary: str = "") -> bool:
        """生成Word文档

        Args:
            articles: 文章列表
            output_path: 输出文件路径
            daily_summary: 每日资讯总结段落
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            # 过滤掉广告内容和摘要生成失败的内容
            filtered_articles = [
                a for a in articles 
                if not a.is_advertisement 
                and a.ai_summary 
                and a.ai_summary != "摘要生成失败"
                and a.ai_summary.strip() != ""
            ]

            # 按MECE分类排序文章
            sorted_articles = DocumentGenerator._sort_by_mece_category(filtered_articles)

            # 创建Word文档
            doc = Document()

            # 设置全局字体为微软雅黑5号
            style = doc.styles['Normal']
            font = style.font
            font.name = '微软雅黑'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            font.size = Pt(10.5)

            # 添加标题
            today = datetime.now()
            title = doc.add_heading(f'每日资讯汇总 - {today.strftime("%Y年%m月%d日")}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(16)
                run.font.bold = True  # 大标题加粗

            # 添加总结段落（如果有）
            if daily_summary:
                summary_title_para = doc.add_paragraph()
                summary_title_run = summary_title_para.add_run('【今日导读】')
                summary_title_run.font.name = '微软雅黑'
                summary_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                summary_title_run.font.bold = True
                summary_title_run.font.size = Pt(11)
                summary_title_para.paragraph_format.space_after = Pt(3)

                summary_para = doc.add_paragraph(daily_summary)
                summary_para.paragraph_format.space_after = Pt(15)
                _set_font(summary_para, '微软雅黑', 10.5)

                # 添加分隔线
                sep_para = doc.add_paragraph()
                run = sep_para.add_run("═══════════════════════════════════════════════════════════════════════")
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
                sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sep_para.paragraph_format.space_after = Pt(15)

            # 生成包含分类标签的输出列表
            output_items = []
            current_main_category = None
            article_num = 0  # 总序号（暂时保留备用）
            category_article_num = 0  # 分类内序号，每个分类重置
            
            # 收集所有大类及其在文档中的位置（用于生成目录）
            category_positions = []  # [(category, category_name, position_in_items)]
            
            for article in sorted_articles:
                cat = article.mece_category if article.mece_category else "11.3"
                main_cat = cat.split('.')[0] if '.' in cat else cat
                
                # 如果大分类发生变化，添加分类标题
                if main_cat != current_main_category:
                    current_main_category = main_cat
                    main_cat_name = DocumentGenerator.MECE_MAIN_CATEGORIES.get(main_cat, "其他")
                    # 记录当前位置
                    category_positions.append({
                        'category': main_cat,
                        'category_name': main_cat_name,
                        'position': len(output_items)
                    })
                    output_items.append({
                        'type': 'category',
                        'category': main_cat,
                        'category_name': main_cat_name
                    })
                    category_article_num = 0  # 重置分类内序号
                
                article_num += 1
                category_article_num += 1  # 分类内序号递增
                output_items.append({
                    'type': 'article',
                    'article': article,
                    'num': article_num,
                    'category_num': category_article_num  # 添加分类内序号
                })

            # ========== 生成目录页 ==========
            toc_title = doc.add_heading('目  录', 1)
            toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in toc_title.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(16)
                run.font.bold = True
            
            # 添加大类目录项（带内部超链接）
            chinese_main_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一']
            for idx, cat_info in enumerate(category_positions):
                cat_num_str = chinese_main_numbers[idx] if idx < len(chinese_main_numbers) else str(idx + 1)
                cat_name = cat_info["category_name"]
                anchor = f'cat_{cat_info["category"]}'
                _add_internal_hyperlink(doc, f'{cat_num_str}、{cat_name}', anchor)
            
            # 添加分隔线
            sep_para = doc.add_paragraph()
            run = sep_para.add_run("═══════════════════════════════════════════════════════════════════════")
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep_para.paragraph_format.space_after = Pt(20)
            
            # 添加分页符，开始正文内容
            doc.add_page_break()
            
            # 重新生成output_items的段落，确保分类标题有书签
            # 注意：需要重新遍历，因为之前只是记录了位置
            current_main_category = None
            article_num = 0
            category_article_num = 0
            category_idx = 0  # 用于生成书签ID
            
            # 中文数字序号
            chinese_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                              '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十',
                              '二十一', '二十二', '二十三', '二十四', '二十五', '二十六', '二十七', '二十八', '二十九', '三十',
                              '三十一', '三十二', '三十三', '三十四', '三十五', '三十六', '三十七', '三十八', '三十九', '四十',
                              '四十一', '四十二', '四十三', '四十四', '四十五', '四十六', '四十七', '四十八', '四十九', '五十',
                              '五十一', '五十二', '五十三', '五十四', '五十五', '五十六', '五十七', '五十八', '五十九', '六十',
                              '六十一', '六十二', '六十三', '六十四', '六十五', '六十六', '六十七', '六十八', '六十九', '七十',
                              '七十一', '七十二', '七十三', '七十四', '七十五', '七十六', '七十七', '七十八', '七十九', '八十',
                              '八十一', '八十二', '八十三', '八十四', '八十五', '八十六', '八十七', '八十八', '八十九', '九十',
                              '九十一', '九十二', '九十三', '九十四', '九十五', '九十六', '九十七', '九十八', '九十九', '一百']
            
            for i, item in enumerate(output_items):
                if item['type'] == 'category':
                    # 添加分类标题（只显示大类名称，三号字体16pt）
                    cat_title = doc.add_paragraph()
                    cat_title_run = cat_title.add_run(f"\n【{item['category_name']}】")
                    cat_title_run.font.name = '微软雅黑'
                    cat_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    cat_title_run.font.bold = True
                    cat_title_run.font.size = Pt(16)  # 三号字体
                    cat_title_run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    cat_title.paragraph_format.space_before = Pt(15)
                    cat_title.paragraph_format.space_after = Pt(10)
                    # 添加书签以便目录跳转
                    _add_bookmark(cat_title, category_idx, f'cat_{item["category"]}')
                    category_idx += 1
                else:
                    article = item['article']
                    num = item['category_num']  # 使用分类内序号
                    num_str = chinese_numbers[num-1] if num-1 < len(chinese_numbers) else str(num)

                    # 使用category_tag作为标题显示（AI总结的【】标题）
                    # 【修复】严格验证category_tag，绝不允许回退到原标题
                    display_title = article.category_tag
                    # 验证category_tag有效性：必须包含实质内容（不是空洞占位符）
                    hollow_titles = {'标题', '标题内容', '综合报道', '综合新闻', '新闻摘要', '综合资讯', '资讯', '报道', '综合', '新闻', '热点', '要闻', '事件', '新闻汇总', '综合要闻', '综合消息', '要点', '概要', '摘要', '详情', '整合报道', '专题', '专题报道', '快讯'}
                    if not display_title:
                        # category_tag为空，尝试从ai_summary中提取【】标题
                        tag_match = re.search(r'【([^】]+)】', article.ai_summary or '')
                        if tag_match:
                            raw_tag = tag_match.group(1).strip()
                            if raw_tag not in hollow_titles and len(raw_tag) > 4:
                                display_title = f"【{raw_tag}】"
                                article.category_tag = display_title  # 回填
                            else:
                                display_title = ""
                        if not display_title:
                            # 从摘要首句生成标题
                            summary_text = re.sub(r'【[^】]+】\s*', '', article.ai_summary or '', count=1).strip()
                            first_sentence = re.split(r'[。！？\n]', summary_text)[0]
                            if first_sentence and len(first_sentence) > 4:
                                display_title = f"【{first_sentence[:50]}】"
                                article.category_tag = display_title  # 回填
                            else:
                                # 最终兜底：标记需要人工检查
                                display_title = "【⚠️标题缺失，需人工检查】"
                                article.category_tag = display_title
                    else:
                        # 有category_tag，验证其内容是否为空洞占位符
                        tag_content = display_title.replace('【', '').replace('】', '').strip()
                        if tag_content in hollow_titles or len(tag_content) <= 4:
                            # 空洞占位符，从摘要重新生成
                            summary_text = re.sub(r'【[^】]+】\s*', '', article.ai_summary or '', count=1).strip()
                            first_sentence = re.split(r'[。！？\n]', summary_text)[0]
                            if first_sentence and len(first_sentence) > 4:
                                display_title = f"【{first_sentence[:50]}】"
                                article.category_tag = display_title  # 回填
                            else:
                                display_title = "【⚠️标题缺失，需人工检查】"
                                article.category_tag = display_title

                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(f'{num_str}、{display_title}')
                    title_run.font.name = '微软雅黑'
                    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    title_run.font.bold = True
                    title_run.font.size = Pt(12)
                    title_para.paragraph_format.space_after = Pt(6)

                    # 来源（微软雅黑5号）
                    if getattr(article, 'is_merged', False) and getattr(article, 'merged_sources', None):
                        # 整合文章：显示多个来源
                        source_names = list(set(s['source'] for s in article.merged_sources))
                        source_text = f'         来源：{"、".join(source_names)}（多来源整合）'
                    else:
                        source_text = f'         来源：{article.source_name}'
                    source_para = doc.add_paragraph(source_text)
                    source_para.paragraph_format.space_after = Pt(6)
                    _set_font(source_para, '微软雅黑', 10.5)

                    # 发布时间（微软雅黑5号）
                    time_para = doc.add_paragraph(f'         发布时间：{article.pub_date.strftime("%Y-%m-%d %H:%M:%S")}')
                    time_para.paragraph_format.space_after = Pt(6)
                    _set_font(time_para, '微软雅黑', 10.5)

                    # 内容摘要（标题+内容在同一行，左右对齐）
                    summary_display = article.ai_summary
                    # 对整合文章：去掉AI生成的"来源列表"部分（已在下方单独格式化显示）
                    if getattr(article, 'is_merged', False):
                        summary_display = re.sub(r'\n*来源列表[：:].*$', '', summary_display, flags=re.DOTALL).strip()
                    # 将摘要中的换行替换为空格，确保内容摘要是连续的一整段
                    summary_display = re.sub(r'\n+', ' ', summary_display).strip()
                    # 合并多余空格
                    summary_display = re.sub(r' {2,}', ' ', summary_display)
                    summary_text = f'         内容摘要：{summary_display}'
                    content_para = doc.add_paragraph(summary_text)
                    content_para.paragraph_format.space_after = Pt(6)
                    content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 左右对齐
                    _set_font(content_para, '微软雅黑', 10.5)

                    # 文章链接（微软雅黑5号）
                    if getattr(article, 'is_merged', False):
                        # 整合文章：显示来源列表（加粗标题 + 每个来源配超链接）
                        link_para = doc.add_paragraph()
                        label_run = link_para.add_run('         来源列表：')
                        label_run.font.name = '微软雅黑'
                        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        label_run.font.bold = True
                        label_run.font.size = Pt(10.5)
                        link_para.paragraph_format.space_after = Pt(2)

                        for src_idx, src in enumerate(article.merged_sources):
                            src_para = doc.add_paragraph()
                            src_para.paragraph_format.space_after = Pt(2)
                            src_para.paragraph_format.left_indent = Pt(18)  # 缩进
                            # 先添加非链接部分的文本：• 《
                            prefix_run = src_para.add_run(f"• 《")
                            prefix_run.font.name = '微软雅黑'
                            prefix_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                            prefix_run.font.size = Pt(10.5)
                            # 标题作为超链接（蓝色可点击）
                            if src.get('link'):
                                _add_external_hyperlink(src_para, src['link'], src['title'])
                            else:
                                title_run = src_para.add_run(src['title'])
                                title_run.font.name = '微软雅黑'
                                title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                title_run.font.size = Pt(10.5)
                            # 添加来源名称》（来源）
                            suffix_run = src_para.add_run(f"》（{src['source']}）")
                            suffix_run.font.name = '微软雅黑'
                            suffix_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                            suffix_run.font.size = Pt(10.5)
                        link_para.paragraph_format.space_after = Pt(12)
                    else:
                        link_para = doc.add_paragraph('         文章链接：')
                        _set_font(link_para, '微软雅黑', 10.5)
                        link_run = link_para.add_run(article.link)
                        link_run.font.color.rgb = RGBColor(0, 0, 255)
                        link_para.paragraph_format.space_after = Pt(12)

                    # 分隔线（除了最后一个）
                    if i < len(output_items) - 1:
                        # 检查下一个是否是分类标题
                        next_item = output_items[i + 1]
                        is_last_article = next_item['type'] == 'category'
                        
                        if not is_last_article:
                            sep_para = doc.add_paragraph()
                            run = sep_para.add_run("——————————————————————————————")
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                            run.font.bold = True
                            run.font.size = Pt(9)
                            sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            sep_para.paragraph_format.space_before = Pt(10)
                            sep_para.paragraph_format.space_after = Pt(6)

            # 保存文档
            doc.save(output_path)
            logger.info(f"文档已生成: {output_path}")

            return True

        except Exception as e:
            logger.error(f"生成文档失败: {e}")
            return False

    @staticmethod
    def _sort_by_mece_category(articles: List[Article]) -> List[Article]:
        """
        按MECE分类顺序排列文章
        
        规则：
        1. 按分类编号顺序排列（如1.1 < 1.2 < 1.3 < 2.1...）
        2. 同一大类内，整合文章排在最前（无论子分类）
        3. 同一子分类内非整合文章按发布时间倒序排列
        4. 如果文章没有分类编号，归入11.3（跨界热点）
        """
        def get_sort_key(article: Article) -> tuple:
            category = article.mece_category if article.mece_category else "11.3"
            
            # 处理分类编号，转换为可排序的格式
            # 例如: "1.1" -> (1, 1), "2.3" -> (2, 3), "11" -> (11, 0)
            parts = category.split('.')
            try:
                main_cat = int(parts[0])
                sub_cat = int(parts[1]) if len(parts) > 1 else 0
            except ValueError:
                main_cat = 11
                sub_cat = 3
            
            # 发布时间倒序（最新的在前）
            pub_time = article.pub_date.timestamp() if article.pub_date else 0
            
            # 整合文章排在同大类最前（0=整合文章排前，1=普通文章排后）
            # 注意：is_merged 优先于 sub_cat，确保整合文章总在大类第一个
            is_merged = 0 if getattr(article, 'is_merged', False) else 1
            
            return (main_cat, is_merged, sub_cat, -pub_time)
        
        return sorted(articles, key=get_sort_key)

class WeChatArticleCrawler:
    """微信公众号文章爬虫"""

    def __init__(self, excel_path: str, deepseek_api_key: str):
        """初始化爬虫"""
        self.excel_path = excel_path
        self.articles: List[Article] = []
        self.noise_articles: List[Article] = []  # 被噪音过滤的文章
        self.duplicate_removed: List[Article] = []  # 被去重移除的文章
        self.dedup_details: List[dict] = []  # 去重详情：记录每篇文章被去重的原因
        self.now = datetime.now()  # 使用本地时间
        self.one_day_ago = self.now - timedelta(days=1)  # 1天前（本地时间）
        self.ai_client = DeepSeekClient(deepseek_api_key)
        
        # 缓存相关配置
        self.cache_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "cache")
        os.makedirs(self.cache_dir, exist_ok=True)
        self.cache_file = os.path.join(self.cache_dir, "article_cache.json")
        self.article_cache: dict = self._load_cache()  # {link: article_data}
        self.cache_ttl_days = 7  # 缓存有效期（天）
        
        # 并行爬取配置
        self.max_workers = 8  # 最大并发线程数

    def make_timezone_aware(self, dt: datetime) -> datetime:
        """将datetime对象转换为naive datetime（去除时区信息）"""
        if dt.tzinfo is not None:
            # 如果有时区信息，转换为本地时间并去除时区
            dt = dt.astimezone(None)
        return dt.replace(tzinfo=None)

    def _load_cache(self) -> dict:
        """加载文章缓存"""
        try:
            if os.path.exists(self.cache_file):
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    cache = json.load(f)
                    # 清理过期缓存
                    current_time = datetime.now().timestamp()
                    cleaned_cache = {}
                    for link, data in cache.items():
                        # 如果缓存没有过期时间或者未过期，保留
                        if 'cached_at' not in data or (current_time - data['cached_at']) < (self.cache_ttl_days * 86400):
                            cleaned_cache[link] = data
                    if len(cleaned_cache) < len(cache):
                        logger.info(f"缓存清理：移除{len(cache) - len(cleaned_cache)}条过期缓存")
                    return cleaned_cache
            return {}
        except Exception as e:
            logger.warning(f"缓存加载失败: {e}，将创建新缓存")
            return {}

    def _save_cache(self):
        """保存文章缓存"""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.article_cache, f, ensure_ascii=False, indent=2, default=str)
            logger.info(f"缓存已保存：{len(self.article_cache)}条记录")
        except Exception as e:
            logger.error(f"缓存保存失败: {e}")

    def _get_cache_key(self, article_link: str, title: str) -> str:
        """生成缓存键（使用链接和标题的组合）"""
        import hashlib
        key_str = f"{article_link}|{title[:100]}"
        return hashlib.md5(key_str.encode('utf-8')).hexdigest()

    def _is_article_in_cache(self, article_link: str, title: str, pub_date: datetime) -> bool:
        """检查文章是否在缓存中且未过期"""
        cache_key = self._get_cache_key(article_link, title)
        if cache_key in self.article_cache:
            cached_data = self.article_cache[cache_key]
            # 检查缓存时间
            cached_time = cached_data.get('cached_at', 0)
            if (datetime.now().timestamp() - cached_time) < (self.cache_ttl_days * 86400):
                return True
        return False

    def _add_to_cache(self, article_link: str, title: str, article_data: dict):
        """将文章添加到缓存"""
        cache_key = self._get_cache_key(article_link, title)
        article_data['cached_at'] = datetime.now().timestamp()
        self.article_cache[cache_key] = article_data

    def load_rss_links(self) -> pd.DataFrame:
        """从Excel文件加载RSS链接"""
        try:
            df = pd.read_excel(self.excel_path)
            logger.info(f"成功加载Excel文件，共{len(df)}个公众号")

            required_cols = ['公众号名称', 'RSS链接']
            missing_cols = [col for col in required_cols if col not in df.columns]
            if missing_cols:
                raise ValueError(f"Excel文件缺少必需列: {missing_cols}")

            return df[['公众号名称', 'RSS链接']]

        except Exception as e:
            logger.error(f"加载Excel文件失败: {e}")
            raise

    def _get_rss_headers(self) -> dict:
        """生成随机浏览器请求头，模拟真实浏览器访问"""
        import random
        user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        ]
        return {
            'User-Agent': random.choice(user_agents),
            'Accept': 'application/rss+xml, application/xml, text/xml, application/atom+xml, */*',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Cache-Control': 'max-age=0',
            'Upgrade-Insecure-Requests': '1',
        }

    def _fetch_with_requests(self, rss_url: str, source_name: str, max_retries: int = 3) -> Optional[bytes]:
        """使用requests获取RSS（快速方法）"""
        import random
        
        for attempt in range(max_retries):
            try:
                headers = self._get_rss_headers()
                if attempt > 0:
                    delay = (2 ** attempt) + random.uniform(0, 1)
                    logger.info(f"  {source_name}: requests重试第{attempt + 1}次，等待{delay:.1f}秒...")
                    time.sleep(delay)
                
                response = requests.get(rss_url, headers=headers, timeout=30, allow_redirects=True)
                
                if response.status_code == 403:
                    logger.warning(f"  {source_name}: 请求被拒绝(403)，尝试更换请求头...")
                    time.sleep(2 + random.uniform(0, 2))
                    headers['User-Agent'] = random.choice([
                        'Mozilla/5.0 (iPhone; CPU iPhone OS 17_2 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Mobile/15E148 Safari/604.1',
                        'Mozilla/5.0 (Linux; Android 14; Pixel 7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Mobile Safari/537.36',
                    ])
                    response = requests.get(rss_url, headers=headers, timeout=30, allow_redirects=True)
                
                response.raise_for_status()
                return response.content
                
            except requests.exceptions.Timeout:
                logger.warning(f"  {source_name}: 请求超时(第{attempt + 1}次)")
            except requests.exceptions.ConnectionError as e:
                logger.warning(f"  {source_name}: 连接错误(第{attempt + 1}次): {str(e)[:50]}")
            except requests.exceptions.HTTPError as e:
                if response.status_code == 403:
                    logger.warning(f"  {source_name}: HTTP 403 错误")
                else:
                    logger.warning(f"  {source_name}: HTTP错误(第{attempt + 1}次): {e}")
            except Exception as e:
                logger.warning(f"  {source_name}: 获取失败(第{attempt + 1}次): {e}")
        
        return None

    def _fetch_with_playwright(self, rss_url: str, source_name: str, max_retries: int = 2) -> Optional[bytes]:
        """使用Playwright模拟浏览器获取RSS（绕过403）"""
        import random
        
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            logger.error(f"  {source_name}: Playwright未安装，请运行: pip install playwright && playwright install chromium")
            return None
        
        for attempt in range(max_retries):
            browser = None
            context = None
            try:
                if attempt > 0:
                    delay = 3 + random.uniform(0, 2)
                    logger.info(f"  {source_name}: Playwright重试第{attempt + 1}次，等待{delay:.1f}秒...")
                    time.sleep(delay)
                
                with sync_playwright() as p:
                    # 启动浏览器
                    browser = p.chromium.launch(
                        headless=True,
                        args=['--no-sandbox', '--disable-setuid-sandbox', '--disable-dev-shm-usage']
                    )
                    
                    # 创建上下文（带随机UA）
                    context = browser.new_context(
                        user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                        viewport={'width': 1920, 'height': 1080},
                        locale='zh-CN'
                    )
                    
                    page = context.new_page()
                    
                    # 设置超时
                    page.set_default_timeout(60000)
                    
                    # 访问页面
                    page.goto(rss_url, wait_until='networkidle', timeout=60000)
                    
                    # 获取内容
                    content = page.content()
                    
                    # 关闭
                    page.close()
                    context.close()
                    browser.close()
                    
                    # 返回bytes
                    return content.encode('utf-8') if isinstance(content, str) else content
                    
            except Exception as e:
                logger.warning(f"  {source_name}: Playwright获取失败(第{attempt + 1}次): {str(e)[:80]}")
                # 确保资源被释放
                try:
                    if context:
                        context.close()
                    if browser:
                        browser.close()
                except:
                    pass
        
        logger.error(f"  {source_name}: Playwright达到最大重试次数({max_retries})")
        return None

    def _fetch_with_retry(self, rss_url: str, source_name: str, max_retries: int = 2) -> Optional[bytes]:
        """使用Playwright获取RSS，统一使用playwright避免403错误"""
        logger.info(f"  {source_name}: 使用Playwright模式获取...")
        return self._fetch_with_playwright(rss_url, source_name, max_retries)
        content = self._fetch_with_playwright(rss_url, source_name)
        
        return content

    def _clean_xml_content(self, content: bytes) -> bytes:
        """清理XML内容，处理不规范的XML"""
        try:
            # 尝试解码
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('gbk', errors='ignore')
            
            # 处理常见的XML问题
            # 1. 移除BOM
            if text.startswith('\ufeff'):
                text = text[1:]
            
            # 2. 移除XML声明前的非法字符
            text = text.lstrip()
            
            # 3. 尝试修复不完整的CDATA
            import re
            # 修复未闭合的CDATA
            cdata_pattern = re.compile(r'<!\[CDATA\[([^\]]*)$')
            text = re.sub(cdata_pattern, r'<![CDATA[\1]]>', text)
            
            # 4. 移除XML声明前的多余内容
            xml_start = text.find('<?xml')
            if xml_start > 0:
                text = text[xml_start:]
            
            return text.encode('utf-8')
        except Exception:
            return content

    def _try_parse_xml(self, content: bytes) -> Optional[ET.Element]:
        """尝试解析XML内容，支持多种编码和不规范格式"""
        if content is None:
            return None
            
        # 先尝试直接解析（规范XML）
        try:
            return ET.fromstring(content)
        except ET.ParseError:
            pass
        
        # 尝试清理后解析
        cleaned = self._clean_xml_content(content)
        try:
            return ET.fromstring(cleaned)
        except ET.ParseError:
            pass
        
        # 尝试使用HTMLParser处理混乱的XML
        try:
            from html.parser import HTMLParser
            import io
            
            class XMLFixer(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.result = []
                    self.tags = []
                    
                def handle_starttag(self, tag, attrs):
                    self.tags.append(tag)
                    self.result.append(f'<{tag}>')
                    
                def handle_endtag(self, tag):
                    if self.tags and self.tags[-1] == tag:
                        self.tags.pop()
                        self.result.append(f'</{tag}>')
                        
                def handle_data(self, data):
                    self.result.append(data)
            
            parser = XMLFixer()
            try:
                parser.feed(content.decode('utf-8', errors='ignore'))
                fixed_xml = ''.join(parser.result)
                return ET.fromstring(fixed_xml.encode('utf-8'))
            except Exception:
                pass
        except Exception:
            pass
        
        return None

    def parse_rss_feed(self, rss_url: str, source_name: str) -> List[Article]:
        """解析RSS feed，提取1天内的文章"""
        articles = []
        try:
            # 使用增强的请求头和重试机制获取内容
            content = self._fetch_with_retry(rss_url, source_name)
            if content is None:
                return articles

            root = self._try_parse_xml(content)
            if root is None:
                logger.warning(f"  {source_name}: 无法解析XML内容，跳过此RSS源")
                return articles

            # 查找文章项
            items = []
            for path in ['.//item', './/entry', './/channel/item']:
                items = root.findall(path)
                if items:
                    break

            if not items:
                items = root.findall('.//*[local-name()="item"]')

            logger.info(f"  {source_name}: 找到{len(items)}篇文章")

            for item in items:
                try:
                    article = self._extract_article(item, source_name)
                    if article:
                        # 检查时间是否在1天内（24小时内）
                        pub_date = self.make_timezone_aware(article.pub_date)
                        if pub_date >= self.one_day_ago:
                            # 检查缓存：如果文章已缓存且未过期，跳过AI处理
                            cache_key = self._get_cache_key(article.link, article.title)
                            if cache_key in self.article_cache:
                                cached_data = self.article_cache[cache_key]
                                # 检查缓存是否在有效期内
                                cached_time = cached_data.get('cached_at', 0)
                                if (datetime.now().timestamp() - cached_time) < (self.cache_ttl_days * 86400):
                                    # 使用缓存的AI摘要和判断结果
                                    article.ai_summary = cached_data.get('ai_summary', '')
                                    article.is_advertisement = cached_data.get('is_advertisement', False)
                                    article.rejection_reason = cached_data.get('rejection_reason', '')
                                    article.category_tag = cached_data.get('category_tag', '')
                                    article.mece_category = cached_data.get('mece_category', '')
                                    logger.debug(f"    使用缓存: {article.title[:30]}...")
                                    # 标记为使用缓存
                                    article._from_cache = True
                            articles.append(article)
                        else:
                            logger.debug(f"    文章不在1天内: {article.title[:30]}...")
                except Exception as e:
                    logger.warning(f"    解析文章失败: {e}")
                    continue

        except Exception as e:
            logger.error(f"  处理RSS时发生错误: {e}")

        return articles

    def _extract_article(self, item, source_name: str) -> Optional[Article]:
        """从XML元素提取文章信息"""
        # 提取标题
        title_elem = item.find('title')
        title = title_elem.text if title_elem is not None else "未知标题"
        
        # 标题过长判断（超过50字认为是读取错误）
        if len(title) > 50:
            logger.warning(f"  标题过长跳过: {title[:30]}...")
            return None

        # 提取链接
        link = None
        for elem_name in ['link', '{http://www.w3.org/2005/Atom}link']:
            elem = item.find(elem_name)
            if elem is not None:
                if elem.text:
                    link = elem.text
                elif 'href' in elem.attrib:
                    link = elem.attrib['href']
                if link:
                    break

        if not link:
            link_elem = item.find('guid')
            link = link_elem.text if link_elem is not None else ""

        # 提取发布时间
        pub_date = None
        for date_elem_name in ['pubDate', 'published', 'updated', 'dc:date']:
            date_elem = item.find(date_elem_name)
            if date_elem is not None and date_elem.text:
                pub_date = self._parse_date(date_elem.text)
                if pub_date:
                    break

        if pub_date is None:
            pub_date = self.now

        # 提取内容
        full_content = ""
        for content_elem_name in ['content:encoded', 'description', 'summary', 'content']:
            content_elem = item.find(content_elem_name)
            if content_elem is not None and content_elem.text:
                full_content = self._clean_html(content_elem.text)
                break

        # 创建文章对象
        article = Article(
            source_name=source_name,
            title=title.strip(),
            link=link.strip() if link else "",
            pub_date=pub_date,
            full_content=full_content
        )

        # 暂时只创建文章对象，不进行AI判断和摘要生成
        # 这些操作将在后续统一进行

        return article

    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """解析日期字符串"""
        date_formats = [
            '%a, %d %b %Y %H:%M:%S %z',
            '%a, %d %b %Y %H:%M:%S %Z',
            '%Y-%m-%dT%H:%M:%SZ',
            '%Y-%m-%dT%H:%M:%S.%fZ',
            '%Y-%m-%dT%H:%M:%S%z',
            '%Y-%m-%d %H:%M:%S',
            '%Y/%m/%d %H:%M:%S',
            '%d %b %Y %H:%M:%S',
        ]

        for fmt in date_formats:
            try:
                dt = datetime.strptime(date_str, fmt)
                # 统一转为naive本地时间，避免与datetime.now()做运算时offset-naive/aware冲突
                if dt.tzinfo is not None:
                    dt = dt.astimezone(None).replace(tzinfo=None)
                return dt
            except ValueError:
                continue

        logger.warning(f"无法解析日期: {date_str}")
        return None

    def _clean_html(self, html_text: str) -> str:
        """清理HTML标签"""
        if not html_text:
            return ""

        clean = re.sub(r'<[^>]+>', ' ', html_text)
        clean = re.sub(r'&nbsp;', ' ', clean)
        clean = re.sub(r'&amp;', '&', clean)
        clean = re.sub(r'&lt;', '<', clean)
        clean = re.sub(r'&gt;', '>', clean)
        clean = re.sub(r'\s+', ' ', clean)
        return clean.strip()

    def _deduplicate_articles(self, articles: List[Article]) -> List[Article]:
        """
        去重处理：智能分流策略
        
        - 相同主题文章 > 5篇：使用AI多视角整合，合并成一条信息
        - 相同主题文章 <= 5篇：使用传统去重逻辑，保留最佳一篇
        
        整合后的文章会标记为 is_merged=True，并记录来源列表
        """
        if not articles:
            return articles

        # 只对有效文章进行去重
        valid_articles = [a for a in articles if not a.is_advertisement and a.ai_summary and a.ai_summary != "摘要生成失败"]
        if not valid_articles:
            return articles

        # 按时间排序，最新的在前
        sorted_articles = sorted(valid_articles, key=lambda x: x.pub_date, reverse=True)
        n = len(sorted_articles)
        
        logger.info(f"=" * 60)
        logger.info(f"开始智能去重/整合流水线，共{n}篇文章")
        logger.info(f"=" * 60)

        # ============ Step 1: 文本预处理 ============
        texts = []
        for a in sorted_articles:
            text = (a.category_tag if a.category_tag else a.title) + " " + (a.ai_summary or "")
            texts.append(text)
        logger.info(f"[Step 1] 文本预处理完成，{n}篇文章")

        # ============ Step 2: 生成TF-IDF向量 ============
        embeddings = self.ai_client.compute_tfidf_embeddings(texts)
        if embeddings.size == 0:
            logger.warning("[Step 2] TF-IDF向量为空，跳过去重")
            return sorted_articles
        logger.info(f"[Step 2] TF-IDF向量生成完成，维度: {embeddings.shape}")

        # ============ Step 3: 余弦相似度计算 ============
        sim_matrix = self.ai_client.cosine_similarity_matrix(embeddings)
        
        # 收集所有相似对（分三级）
        high_pairs = []     # >= HIGH_SIM_THRESHOLD: 直接去重
        medium_pairs = []   # >= MEDIUM_SIM_THRESHOLD: 需LLM判断
        low_pairs = []      # >= LOW_SIM_THRESHOLD: 规则辅助
        
        for i in range(n):
            for j in range(i + 1, n):
                sim = float(sim_matrix[i, j])
                if sim >= HIGH_SIM_THRESHOLD:
                    high_pairs.append((i, j, sim))
                elif sim >= MEDIUM_SIM_THRESHOLD:
                    medium_pairs.append((i, j, sim))
                elif sim >= LOW_SIM_THRESHOLD:
                    low_pairs.append((i, j, sim))
        
        logger.info(f"[Step 3] 相似度计算完成: 高相似{len(high_pairs)}对, 中相似{len(medium_pairs)}对, 低相似{len(low_pairs)}对")

        # 【新增】规则驱动的主题预匹配：捕获TF-IDF遗漏的同主题文章对
        # 对于重大外交事件（同一人物+同一外交动作）和重大灾害（同一地点+同一灾害类型），
        # TF-IDF可能因文本差异大而遗漏，但规则匹配可以精准捕获
        rule_matched_pairs = []
        # 预定义知名政治人物列表
        political_figures = ['普京', '特朗普', '拜登', '习近平', '马克龙', '朔尔茨', '岸田',
                            '尹锡悦', '莫迪', '苏纳克', '泽连斯基', '内塔尼亚胡', '金正恩',
                            '拉马福萨', '卢拉', '米莱']
        diplomatic_actions = ['访华', '访美', '访日', '访韩', '访俄', '访欧', '出访', '到访',
                              '国事访问', '正式访问', '抵达北京', '抵京', '抵沪', '抵深',
                              '抵达', '来华', '赴华', '抵达中国']
        disaster_types = ['地震', '海啸', '洪灾', '洪涝', '泥石流', '山体滑坡', '龙卷风', '台风']

        for i in range(n):
            for j in range(i + 1, n):
                a1 = sorted_articles[i]
                a2 = sorted_articles[j]
                t1 = (a1.category_tag or a1.title) + " " + (a1.ai_summary or "")
                t2 = (a2.category_tag or a2.title) + " " + (a2.ai_summary or "")
                title1 = (a1.category_tag or a1.title)
                title2 = (a2.category_tag or a2.title)

                # 【关键修复】跳过聚合文章的规则匹配
                # 聚合文章同时涉及多个不相关领域，如果参与规则匹配，
                # 会把不相关的事件连通到同一个簇中导致错误整合
                def _is_aggregator_title(t):
                    """检测标题是否是聚合快讯类"""
                    aggregator_indicators = ['8点1氪', '早报', '早知道', '快报', '快讯', '财经早餐',
                                             '新闻早参考', '新闻早报', '资讯早报', '今日要闻',
                                             '新闻速递', '资讯速递', '早餐内参', '每日要闻']
                    for ind in aggregator_indicators:
                        if ind in t:
                            return True
                    # 标题用分号分隔3+个独立话题
                    if len(re.split(r'[；;]', t)) >= 3:
                        return True
                    return False
                
                if _is_aggregator_title(title1) or _is_aggregator_title(title2):
                    continue

                # 重大外交事件匹配：同一政治人物 + 同一外交动作
                # 【修复】限制匹配范围：要求两篇文章标题中都包含相同的政治人物和外交动作，
                # 避免仅凭摘要中"提及"某人物就建立匹配
                figures1_title = {f for f in political_figures if f in title1}
                figures2_title = {f for f in political_figures if f in title2}
                has_diplomatic1_title = any(a in title1 for a in diplomatic_actions)
                has_diplomatic2_title = any(a in title2 for a in diplomatic_actions)
                if figures1_title and figures2_title and figures1_title & figures2_title and has_diplomatic1_title and has_diplomatic2_title:
                    key = (min(i, j), max(i, j))
                    rule_matched_pairs.append((key, 0.08, 'diplomatic'))
                
                # 【新增】外交事件补充匹配：同一政治人物+摘要中都有外交动作
                # 解决标题优化后外交动作词丢失导致无法匹配的问题
                # 如："普京抵达北京"和"普京访华"可能因标题优化后表述不同而无法标题级匹配
                if not (figures1_title and figures2_title and figures1_title & figures2_title and has_diplomatic1_title and has_diplomatic2_title):
                    figures1_text = {f for f in political_figures if f in t1}
                    figures2_text = {f for f in political_figures if f in t2}
                    has_diplomatic1_text = any(a in t1 for a in diplomatic_actions)
                    has_diplomatic2_text = any(a in t2 for a in diplomatic_actions)
                    common_figures_text = figures1_text & figures2_text
                    # 要求至少一篇标题中包含该政治人物，避免仅因摘要提及而错误匹配
                    if (common_figures_text and 
                        has_diplomatic1_text and has_diplomatic2_text and
                        (figures1_title & common_figures_text or figures2_title & common_figures_text)):
                        key = (min(i, j), max(i, j))
                        if key not in {k for k, _, _ in rule_matched_pairs}:
                            rule_matched_pairs.append((key, 0.06, 'diplomatic_extended'))

                # 【新增】外交事件弱匹配：标题中有政治人物，摘要中有中国相关地点/活动
                # 解决"普京：学习中文的俄罗斯公民已超10万人"这类标题无外交动作词、
                # 但摘要提到"在北京人民大会堂""中俄教育年"等明确属于该政治人物访华期间活动的文章
                # 此时只要另一篇是同一政治人物的明确外交事件文章，就应建立连接
                if not any(k == (min(i, j), max(i, j)) for k, _, _ in rule_matched_pairs):
                    figures1_title_set = {f for f in political_figures if f in title1}
                    figures2_title_set = {f for f in political_figures if f in title2}
                    common_title_figures = figures1_title_set & figures2_title_set
                    if common_title_figures:
                        # 中国相关地点/活动标识词
                        china_indicators = ['北京', '人民大会堂', '中俄', '中方', '习近平', '钓鱼台',
                                           '天安门', '故宫', '中南海', '中国', '国家主席']
                        # 检查：一篇文章有外交动作词（明确的外交事件），另一篇有中国标识词（外交期间的周边活动）
                        has_diplomatic1 = has_diplomatic1_title or any(a in t1 for a in diplomatic_actions)
                        has_diplomatic2 = has_diplomatic2_title or any(a in t2 for a in diplomatic_actions)
                        has_china1 = any(c in t1 for c in china_indicators)
                        has_china2 = any(c in t2 for c in china_indicators)
                        # 一篇有外交动作+中国标识，另一篇有中国标识+同一政治人物标题 → 弱匹配
                        if ((has_diplomatic1 and has_china1 and has_china2) or
                            (has_diplomatic2 and has_china2 and has_china1)):
                            key = (min(i, j), max(i, j))
                            if key not in {k for k, _, _ in rule_matched_pairs}:
                                rule_matched_pairs.append((key, 0.05, 'diplomatic_china_indicator'))
                                logger.info(f"  [外交弱匹配] 标题1='{title1[:30]}' + 标题2='{title2[:30]}' (共同人物={common_title_figures})")

                # 【新增】摘要级政治人物+外交事件匹配
                # 解决"带了半个内阁，唯独少一人"这类标题既无政治人物也无外交动作词、
                # 但摘要中提到"普京深夜抵京"等明确外交事件的文章
                # 规则：一篇摘要中有政治人物+外交动作，另一篇标题中有同一政治人物 → 建立弱连接
                if not any(k == (min(i, j), max(i, j)) for k, _, _ in rule_matched_pairs):
                    figures1_text = {f for f in political_figures if f in t1}
                    figures2_text = {f for f in political_figures if f in t2}
                    has_diplomatic1_text = any(a in t1 for a in diplomatic_actions)
                    has_diplomatic2_text = any(a in t2 for a in diplomatic_actions)
                    common_figures_text = figures1_text & figures2_text
                    # 一篇标题有政治人物，另一篇摘要有同一政治人物+外交动作 → 弱匹配
                    if common_figures_text:
                        fig_in_title1 = bool(figures1_title & common_figures_text)
                        fig_in_title2 = bool(figures2_title & common_figures_text)
                        # 至少一篇标题中有该人物，另一篇摘要中有该人物+外交动作
                        if ((fig_in_title1 and has_diplomatic2_text and figures2_text & common_figures_text) or
                            (fig_in_title2 and has_diplomatic1_text and figures1_text & common_figures_text)):
                            key = (min(i, j), max(i, j))
                            if key not in {k for k, _, _ in rule_matched_pairs}:
                                rule_matched_pairs.append((key, 0.05, 'diplomatic_summary_figure'))
                                logger.info(f"  [摘要人物匹配] 标题1='{title1[:30]}' + 标题2='{title2[:30]}' (共同人物={common_figures_text})")

                # 重大灾害事件匹配：同一地点+同一灾害类型
                disaster_locs1 = set()
                disaster_locs2 = set()
                for dtype in disaster_types:
                    if dtype in t1:
                        locs1 = re.findall(r'([\u4e00-\u9fa5]{2,6})(?:.{0,10})' + dtype, t1)
                        disaster_locs1.update(locs1)
                    if dtype in t2:
                        locs2 = re.findall(r'([\u4e00-\u9fa5]{2,6})(?:.{0,10})' + dtype, t2)
                        disaster_locs2.update(locs2)
                # 地点匹配：使用子串包含而非精确匹配（"广西"是"广西柳州"的子串）
                if disaster_locs1 and disaster_locs2:
                    location_overlap = False
                    for loc1 in disaster_locs1:
                        for loc2 in disaster_locs2:
                            clean1 = re.sub(r'[发生等级0-9.]', '', loc1)[:4]
                            clean2 = re.sub(r'[发生等级0-9.]', '', loc2)[:4]
                            if clean1 and clean2 and (clean1 in clean2 or clean2 in clean1 or clean1[:2] == clean2[:2]):
                                location_overlap = True
                                break
                        if location_overlap:
                            break
                    if location_overlap:
                        key = (min(i, j), max(i, j))
                        rule_matched_pairs.append((key, 0.08, 'disaster'))

                # 【新增】重大商业事件匹配：同一公司/人物+同一商业事件(IPO/上市/并购等)
                # 解决SpaceX IPO多篇不同角度报道未能去重的问题
                major_business_entities = ['SpaceX', 'OpenAI', 'Anthropic', 'xAI', 'Tesla', 'Apple',
                                          'Google', 'Meta', 'Amazon', 'Microsoft', 'NVIDIA', 'Intel',
                                          'ByteDance', 'TikTok', 'Huawei', 'Samsung', 'Boeing',
                                          '阿里巴巴', '腾讯', '比亚迪', '宁德时代', '字节跳动',
                                          '马斯克', '奥特曼', '黄仁勋', '库克', '扎克伯格', '贝索斯',
                                          '孙正义', '软银', '柠季', '哈根达斯', '通用磨坊',
                                          # 【修复】补充AI模型名/产品名，解决Anthropic Fable 5等模型管制文章未能去重
                                          'Fable', 'Mythos', 'Claude']
                business_events = ['IPO', '上市', '招股书', '募资', '融资', '并购', '收购',
                                   '重组', '破产', '退市', '暴雷', '违约', '被罚', '处罚',
                                   # 【修复】补充AI模型管制/下架相关关键词，解决政府禁止/下架模型类文章未能去重
                                   '下架', '禁令', '禁止', '出口管制', '全球禁', '限制', '管制']
                
                # 从标题+摘要中提取商业实体和事件
                biz_entities1 = {e for e in major_business_entities if e in t1}
                biz_entities2 = {e for e in major_business_entities if e in t2}
                common_biz_entities = biz_entities1 & biz_entities2
                has_biz_event1 = any(e in t1 for e in business_events)
                has_biz_event2 = any(e in t2 for e in business_events)
                
                if common_biz_entities and has_biz_event1 and has_biz_event2:
                    # 至少一个商业实体在两篇标题中都出现，且都有商业事件词
                    common_in_title = {e for e in common_biz_entities if e in title1 and e in title2}
                    if common_in_title:
                        key = (min(i, j), max(i, j))
                        if key not in {k for k, _, _ in rule_matched_pairs}:
                            rule_matched_pairs.append((key, 0.08, 'business_event'))
                            logger.info(f"  [商业事件匹配] 标题1='{title1[:30]}' + 标题2='{title2[:30]}' (共同实体={common_in_title})")
                    elif common_biz_entities:
                        # 实体在摘要中出现，给弱匹配
                        key = (min(i, j), max(i, j))
                        if key not in {k for k, _, _ in rule_matched_pairs}:
                            rule_matched_pairs.append((key, 0.06, 'business_event_extended'))
                            logger.info(f"  [商业事件弱匹配] 标题1='{title1[:30]}' + 标题2='{title2[:30]}' (共同实体={common_biz_entities})")

                # 【新增】国际重大事件主题匹配：同一国际事件关键词组合出现在两篇文章中
                # 解决"美伊和谈"5篇不同角度报道（霍尔木兹海峡开放、鲁比奥表态、伊朗外交部回应等）
                # 未能被TF-IDF或外交人物+动作规则匹配到的问题
                # 引用模块级常量 MAJOR_EVENT_KEYWORD_GROUPS，加上非重大事件的特例组
                international_event_groups = MAJOR_EVENT_KEYWORD_GROUPS + [
                    {
                        'name': '华为韬定律/半导体突破',
                        'keywords': ['韬定律', '何庭波', 'ISCAS', '晶体管密度', '时间缩放'],
                        'min_overlap': 2,
                    },
                    {
                        'name': 'AlphaProof/DeepMind数学',
                        'keywords': ['AlphaProof', 'Erdős', 'DeepMind', '数学难题', 'Lean'],
                        'min_overlap': 2,
                    },
                ]

                for event_group in international_event_groups:
                    event_name = event_group['name']
                    event_keywords = event_group['keywords']
                    min_overlap = event_group.get('min_overlap', 2)
                    is_major_event = event_group.get('is_major_event', False)
                    # 【关键修复】锚定关键词：用于确保匹配的文章确实属于同一具体事件
                    # 避免仅因共享泛化冲突词（如'战争''制裁''导弹'）而错误连通不同事件
                    anchor_keywords = event_group.get('anchor_keywords', None)

                    # 提取两篇文章中命中的关键词
                    hits1 = {kw for kw in event_keywords if kw in t1}
                    hits2 = {kw for kw in event_keywords if kw in t2}
                    overlap = hits1 & hits2

                    # 匹配策略：两种方式任一满足即建立连接
                    # 方式1：关键词重叠 >= min_overlap（两篇文章有相同的关键词）
                    # 方式2：每篇文章都命中 >= min_overlap 个该事件关键词（即使不是同一个词）
                    # 方式2解决"鲁比奥文章说伊朗+海峡+核问题，和谈文章说美伊+海峡+停火"的问题
                    # 【关键修复】无论方式1还是方式2，都必须满足锚定关键词约束：
                    # 两篇文章必须在锚定关键词上有至少1个重叠，确保它们确实属于同一具体事件，
                    # 而非仅因为共享'战争''制裁'等泛化冲突词而被错误连通。
                    matched = False
                    if len(overlap) >= min_overlap:
                        matched = True
                    elif len(hits1) >= min_overlap and len(hits2) >= min_overlap:
                        matched = True
                    
                    # 【关键修复】锚定关键词校验：如果事件组定义了anchor_keywords，
                    # 则要求两篇文章在锚定关键词上至少有1个重叠，才允许建立连接。
                    # 这防止了"乌克兰战争"和"美伊冲突"因共享'战争''导弹'等泛化词而被连通。
                    if matched and anchor_keywords:
                        anchor_hits1 = {kw for kw in anchor_keywords if kw in t1}
                        anchor_hits2 = {kw for kw in anchor_keywords if kw in t2}
                        anchor_overlap = anchor_hits1 & anchor_hits2
                        if not anchor_overlap:
                            matched = False
                            logger.info(
                                f"  [国际事件匹配跳过:{event_name}] 标题1='{title1[:30]}' + 标题2='{title2[:30]}' "
                                f"关键词重叠{len(overlap)}个但无锚定关键词重叠(文章1锚定命中={anchor_hits1}, 文章2锚定命中={anchor_hits2})"
                            )

                    if matched:
                        key = (min(i, j), max(i, j))
                        if key not in {k for k, _, _ in rule_matched_pairs}:
                            # 重大事件用更高相似度(0.10)确保进入连通图；同时携带major_event标记
                            sim_val = 0.10 if is_major_event else 0.08
                            match_tag = f'international_event_{event_name}'
                            if is_major_event:
                                match_tag += ':major_event'
                            rule_matched_pairs.append((key, sim_val, match_tag))
                            logger.info(f"  [国际事件匹配:{event_name}] 标题1='{title1[:30]}' + 标题2='{title2[:30]}' (重叠={overlap}, 文章1命中={hits1}, 文章2命中={hits2})")

                # 【新增】食品安全/产品质量事件匹配：同一事件关键词+同一地点/品牌
                # 解决"泡药杨梅"多篇报道未能去重的问题
                food_safety_terms = ['泡药', '添加剂', '防腐剂', '瘦肉精', '地沟油', '三聚氰胺',
                                     '塑化剂', '农药残留', '重金属超标', '食物中毒', '食品安全',
                                     '假货', '假冒', '掺假', '造假', '以次充好', '问题食品']
                food_products = ['杨梅', '奶粉', '牛奶', '肉类', '鸡蛋', '蔬菜', '水果',
                                 '海鲜', '大米', '食用油', '食品', '白酒', '茶叶', '蜂蜜']
                
                has_food_safety1 = any(t in t1 for t in food_safety_terms)
                has_food_safety2 = any(t in t2 for t in food_safety_terms)
                common_food_products = {p for p in food_products if p in t1 and p in t2}
                
                if has_food_safety1 and has_food_safety2 and common_food_products:
                    key = (min(i, j), max(i, j))
                    if key not in {k for k, _, _ in rule_matched_pairs}:
                        rule_matched_pairs.append((key, 0.08, 'food_safety'))
                        logger.info(f"  [食品安全匹配] 标题1='{title1[:30]}' + 标题2='{title2[:30]}' (共同产品={common_food_products})")

                # 【特例】高考主题强制连通：所有高考相关文章强制归入同一簇，统一整合为一条资讯
                # 高考是年度重大事件，各公众号从不同角度报道（交通管制/噪音管控/考场保障/志愿填报等），
                # TF-IDF难以自动关联，需规则强制连通
                gaokao_core_keywords = ['高考', '高考生', '高考报名', '高考志愿', '高考录取', '高考改革',
                                        '高考保障', '高考考点', '高考考场', '高考噪音', '高考管控', '高考禁噪',
                                        '高考交通', '高考护航', '高考加油', '高考倒计时', '高考时间', '高考科目',
                                        '高考成绩', '高考分数', '高考招生', '高考政策', '高考变化', '高考试卷',
                                        '高考作文', '高考数学', '高考英语', '高考语文', '高考理综', '高考文综',
                                        '高考期间', '高考组织', '高考组织保障', '高考保障工作']
                gaokao_extended_keywords = ['护考', '送考', '考点', '考场', '考生']
                # 排除干扰词：避免把"考研""中考"等误匹配
                gaokao_exclude = ['考研', '中考', '考公', '考编', '自考', '成人高考']

                # 任一标题含排除词 → 跳过（防止考研/中考文章混入）
                if any(ex in title1 for ex in gaokao_exclude) or any(ex in title2 for ex in gaokao_exclude):
                    pass  # 不做高考匹配
                else:
                    # 判断文章是否与高考相关：标题含核心词，或(标题含扩展词 且 正文含核心词)
                    def _is_gaokao_related(title, text):
                        if any(kw in title for kw in gaokao_core_keywords):
                            return True
                        if any(kw in title for kw in gaokao_extended_keywords) and any(kw in text for kw in gaokao_core_keywords):
                            return True
                        return False

                    gaokao_related1 = _is_gaokao_related(title1, t1)
                    gaokao_related2 = _is_gaokao_related(title2, t2)

                    if gaokao_related1 and gaokao_related2:
                        key = (min(i, j), max(i, j))
                        if key not in {k for k, _, _ in rule_matched_pairs}:
                            rule_matched_pairs.append((key, 0.10, 'gaokao'))
                            logger.info(f"  [高考主题匹配] 标题1='{title1[:30]}' + 标题2='{title2[:30]}'")

                # 【新增】世界杯主题强制连通：所有世界杯相关文章强制归入同一簇，统一整合为一条资讯
                # 世界杯是重大体育赛事，各公众号从不同角度报道（赛事结果/转播权/商业合作/球迷文化等），
                # 应整合为一条资讯，分类归入文体娱乐
                world_cup_core_keywords = ['世界杯', 'FIFA世界杯', '2026世界杯', '美加墨世界杯', '世界杯转播',
                                           '世界杯版权', '世界杯赛事', '世界杯小组赛', '世界杯淘汰赛', '世界杯决赛',
                                           '世界杯开幕', '世界杯闭幕', '世界杯举办', '世界杯主办']
                world_cup_extended_keywords = ['FIFA', '美加墨', '足球赛', '球星', '进球', '点球', '越位',
                                               '红牌', '黄牌', '裁判', '国足', '中国队', '出线', '淘汰',
                                               '半决赛', '四强', '八强', '十六强', '小组出线']
                # 排除干扰词：避免把"篮球世界杯""排球世界杯"等误匹配
                world_cup_exclude = ['篮球世界杯', '排球世界杯', '乒乓球世界杯', '羽毛球世界杯', ' Rugby世界杯', '板球世界杯']

                # 任一标题含排除词 → 跳过
                if any(ex in title1 for ex in world_cup_exclude) or any(ex in title2 for ex in world_cup_exclude):
                    pass
                else:
                    def _is_world_cup_related(title, text):
                        if any(kw in title for kw in world_cup_core_keywords):
                            return True
                        if any(kw in title for kw in world_cup_extended_keywords) and any(kw in text for kw in world_cup_core_keywords):
                            return True
                        # 特殊：标题含"足球"+"赛事/比赛/夺冠/出局"等体育事件词
                        if '足球' in title and any(kw in title for kw in ['赛事', '比赛', '夺冠', '出局', '晋级', '半决赛', '决赛']):
                            return True
                        return False

                    wc_related1 = _is_world_cup_related(title1, t1)
                    wc_related2 = _is_world_cup_related(title2, t2)

                    if wc_related1 and wc_related2:
                        key = (min(i, j), max(i, j))
                        if key not in {k for k, _, _ in rule_matched_pairs}:
                            rule_matched_pairs.append((key, 0.10, 'world_cup'))
                            logger.info(f"  [世界杯主题匹配] 标题1='{title1[:30]}' + 标题2='{title2[:30]}'")

        if rule_matched_pairs:
            logger.info(f"[Step 3] 规则预匹配发现{len(rule_matched_pairs)}对同主题文章")

        # 同时使用旧的Jaccard方法作为补充（捕获TF-IDF可能遗漏的短文本匹配）
        articles_for_jaccard = [
            {'title': a.category_tag if a.category_tag else a.title, 'ai_summary': a.ai_summary}
            for a in sorted_articles
        ]
        jaccard_pairs = self.ai_client.find_similar_groups(articles_for_jaccard)
        
        # 精确去重（相同来源+高度相似标题/摘要）
        exact_duplicates = self._find_exact_duplicates(sorted_articles)
        
        # 合并所有相似对来源
        all_pair_map = {}  # (min_i, max_j) -> similarity
        
        for i, j, sim in high_pairs:
            key = (min(i, j), max(i, j))
            all_pair_map[key] = max(all_pair_map.get(key, 0), sim)
        
        for i, j, sim in medium_pairs:
            key = (min(i, j), max(i, j))
            all_pair_map[key] = max(all_pair_map.get(key, 0), sim)
        
        for i, j, sim in low_pairs:
            key = (min(i, j), max(i, j))
            all_pair_map[key] = max(all_pair_map.get(key, 0), sim)
        
        # 【新增】将规则预匹配对加入all_pair_map
        for key, sim, match_type in rule_matched_pairs:
            all_pair_map[key] = max(all_pair_map.get(key, 0), sim)
        
        # 【重要修改】Jaccard匹配不再加入all_pair_map，
        # 因为Jaccard仅基于关键词重叠，容易产生虚假相似度，
        # 导致不相关文章被连通到同一簇中
        # for pair in jaccard_pairs:
        #     if len(pair) >= 2:
        #         key = (min(pair[0], pair[1]), max(pair[0], pair[1]))
        #         all_pair_map.setdefault(key, 0.5)
        
        for pair in exact_duplicates:
            if len(pair) >= 2:
                key = (min(pair[0], pair[1]), max(pair[0], pair[1]))
                all_pair_map[key] = 1.0  # 精确匹配给最高相似度
        
        if not all_pair_map:
            logger.info("[Step 3] 未发现相似文章，无需去重")
            return sorted_articles
        
        logger.info(f"[Step 3] 合并后共{len(all_pair_map)}对相似文章")

        # ============ Step 4: 相似图建模与连通分量聚类 ============
        from collections import defaultdict

        def _build_clusters(pair_map, n):
            """根据给定的相似对构建连通分量"""
            parent = list(range(n))
            def _find(x):
                if parent[x] != x:
                    parent[x] = _find(parent[x])
                return parent[x]
            def _union(x, y):
                px, py = _find(x), _find(y)
                if px != py:
                    parent[px] = py
            for (i, j) in pair_map:
                _union(i, j)
            grp = defaultdict(list)
            for i in range(n):
                grp[_find(i)].append(i)
            return grp, _find

        # 第一轮：只用高+中相似度边 + 精确匹配构建连通图
        # 【重要修改】Jaccard匹配不再以默认0.5加入连通图，
        # 因为Jaccard仅基于关键词重叠，容易把不相关文章（如都提到"联合国"的文章）连通在一起，
        # 导致完全不相关的事件被聚到同一簇中错误整合
        high_medium_pairs = {k: v for k, v in all_pair_map.items() if v >= MEDIUM_SIM_THRESHOLD}
        # 将精确匹配也加入（精确匹配是可靠的去重依据）
        for pair in exact_duplicates:
            if len(pair) >= 2:
                key = (min(pair[0], pair[1]), max(pair[0], pair[1]))
                high_medium_pairs[key] = 1.0
        # 【新增】将规则预匹配对也加入连通图（外交/灾害类同主题文章）
        for key, sim, match_type in rule_matched_pairs:
            high_medium_pairs[key] = max(high_medium_pairs.get(key, 0), sim)
        # Jaccard匹配不加入连通图，仅作为后续LLM判断的参考（已移除）

        groups, find = _build_clusters(high_medium_pairs, n)

        # 只保留有多篇文章的组
        multi_groups = {k: v for k, v in groups.items() if len(v) > 1}
        logger.info(f"[Step 4] 连通分量聚类完成: {len(multi_groups)}个多文章组")

        # ============ Step 5: 智能分流 - 大簇整合 vs 小簇去重 ============
        to_remove = set()  # 最终要移除的文章索引
        merged_articles = []  # 整合生成的文章
        
        # 对每个组进行处理
        groups_for_llm = []  # 需要LLM判断的组
        
        for root, indices in multi_groups.items():
            if len(indices) < 2:
                continue
            
            # 获取组内相似度信息
            group_similarities = {}
            for i in range(len(indices)):
                for j in range(i + 1, len(indices)):
                    key = (min(indices[i], indices[j]), max(indices[i], indices[j]))
                    if key in all_pair_map:
                        group_similarities[key] = all_pair_map[key]
            
            max_sim = max(group_similarities.values()) if group_similarities else 0
            has_exact_dup = any(
                (min(indices[i], indices[j]), max(indices[i], indices[j])) in 
                set((min(p[0], p[1]), max(p[0], p[1])) for p in exact_duplicates)
                for i in range(len(indices))
                for j in range(i + 1, len(indices))
            )
            
            # 【特例】检测该簇是否包含高考规则匹配——高考文章无论数量均走整合路径
            is_gaokao_cluster = False
            for i_idx in range(len(indices)):
                for j_idx in range(i_idx + 1, len(indices)):
                    pair_key = (min(indices[i_idx], indices[j_idx]), max(indices[i_idx], indices[j_idx]))
                    if any(k == pair_key and mt == 'gaokao' for k, _, mt in rule_matched_pairs):
                        is_gaokao_cluster = True
                        break
                if is_gaokao_cluster:
                    break

            # 【特例】检测该簇是否包含世界杯规则匹配——世界杯文章无论数量均走整合路径
            is_world_cup_cluster = False
            for i_idx in range(len(indices)):
                for j_idx in range(i_idx + 1, len(indices)):
                    pair_key = (min(indices[i_idx], indices[j_idx]), max(indices[i_idx], indices[j_idx]))
                    if any(k == pair_key and mt == 'world_cup' for k, _, mt in rule_matched_pairs):
                        is_world_cup_cluster = True
                        break
                if is_world_cup_cluster:
                    break

            # 【特例】检测该簇是否包含重大国际事件规则匹配——重大事件文章无论数量均走整合路径
            # 重大国际事件（如美伊战争/俄乌冲突等）各公众号从不同角度报道（空袭/和谈/制裁/油价影响等），
            # 应整合为一条资讯，避免同一事件碎片化呈现
            # 改用 major_event 标记检测，不再硬编码事件名，更通用
            is_major_event_cluster = False
            for i_idx in range(len(indices)):
                for j_idx in range(i_idx + 1, len(indices)):
                    pair_key = (min(indices[i_idx], indices[j_idx]), max(indices[i_idx], indices[j_idx]))
                    if any(k == pair_key and 'major_event' in str(mt) for k, _, mt in rule_matched_pairs):
                        is_major_event_cluster = True
                        break
                if is_major_event_cluster:
                    break

            # 兼容旧标记：美伊战争（旧格式不带 major_event 后缀的也检测）
            is_us_iran_war_cluster = is_major_event_cluster
            if not is_us_iran_war_cluster:
                for i_idx in range(len(indices)):
                    for j_idx in range(i_idx + 1, len(indices)):
                        pair_key = (min(indices[i_idx], indices[j_idx]), max(indices[i_idx], indices[j_idx]))
                        if any(k == pair_key and '美伊战争' in str(mt) for k, _, mt in rule_matched_pairs):
                            is_us_iran_war_cluster = True
                            break
                    if is_us_iran_war_cluster:
                        break
                if is_us_iran_war_cluster:
                    is_major_event_cluster = True

            # 【关键逻辑】判断是大簇(>=3篇)还是小簇(<3篇)
            # 【修复】把"大簇"阈值从 5 降到 3，让"老虎/富途/长桥被罚"这类3-5篇的同主题
            # 重大事件也走多视角整合，而不是简单去重保留一篇（导致显示原文标题）
            # 【特例】高考/重大事件主题簇无论数量都走整合路径
            if len(indices) >= 3 or is_gaokao_cluster or is_major_event_cluster or is_world_cup_cluster:
                # ====== 大簇：AI多视角整合 ======
                _p(f"[整合] 发现大簇({len(indices)}篇): 准备多视角整合...")
                # 传递簇类型信息，让 _merge_large_cluster 调整实体筛选策略
                _cluster_type = None
                if is_gaokao_cluster:
                    _cluster_type = 'gaokao'
                elif is_world_cup_cluster:
                    _cluster_type = 'world_cup'
                elif is_major_event_cluster:
                    _cluster_type = 'major_event'
                merged_article = self._merge_large_cluster(sorted_articles, indices, cluster_type=_cluster_type)
                if merged_article:
                    merged_articles.append(merged_article)
                    # 标记所有原文章为已移除
                    for idx in indices:
                        to_remove.add(idx)
                        rm_art = sorted_articles[idx]
                        self.dedup_details.append({
                            'removed_title': rm_art.title,
                            'removed_source': rm_art.source_name,
                            'removed_link': rm_art.link,
                            'kept_title': merged_article.category_tag,
                            'kept_source': "多来源整合",
                            'kept_link': "",
                            'reason': f"大簇多视角整合(共{len(indices)}篇)",
                        })
                    _p(f"[整合] 成功整合 {len(indices)} 篇文章为一条信息")
                    logger.info(f"  整合(大簇{len(indices)}篇): 生成'{merged_article.category_tag[:40]}...'")
                else:
                    # 整合失败，降级为传统去重
                    _p(f"[整合] 大簇整合失败，降级为传统去重")
                    best_idx = self._select_best_article(sorted_articles, indices)
                    for idx in indices:
                        if idx != best_idx:
                            to_remove.add(idx)
            else:
                # ====== 小簇：传统去重逻辑 ======
                if has_exact_dup or max_sim >= HIGH_SIM_THRESHOLD:
                    # 高相似度 + 精确匹配：直接去重，用规则打分选最佳
                    best_idx = self._select_best_article(sorted_articles, indices)
                    for idx in indices:
                        if idx != best_idx:
                            to_remove.add(idx)
                            rm_art = sorted_articles[idx]
                            kept_art = sorted_articles[best_idx]
                            self.dedup_details.append({
                                'removed_title': rm_art.title,
                                'removed_source': rm_art.source_name,
                                'removed_link': rm_art.link,
                                'kept_title': kept_art.title,
                                'kept_source': kept_art.source_name,
                                'kept_link': kept_art.link,
                                'reason': f"高相似度直接去重(最高相似度{max_sim:.3f})",
                            })
                            logger.info(f"  去重(高相似): 保留'{kept_art.title[:30]}...'，移除'{rm_art.title[:30]}...'(sim={max_sim:.3f})")
                else:
                    # 中低相似度：需要LLM判断
                    groups_for_llm.append((root, indices))

        logger.info(f"[Step 5] 大簇整合: {len(merged_articles)}组, 高相似去重: 移除{len(to_remove)}篇, 待LLM判断: {len(groups_for_llm)}个组")

        # ============ Step 6: 覆盖性校验（仅对小簇去重） ============
        verified_remove = set()
        for idx in to_remove:
            # 跳过已被整合的文章（它们已经被大簇逻辑处理）
            article = sorted_articles[idx]
            # 查找该文章所在的组
            kept_in_group = []
            for root, indices in multi_groups.items():
                if idx in indices:
                    kept_in_group = [i for i in indices if i not in to_remove]
                    break
            
            if not kept_in_group:
                continue
            
            # 简单覆盖性校验
            removed_text = (article.category_tag or article.title) + " " + (article.ai_summary or "")
            is_covered = False
            for kept_idx in kept_in_group:
                kept_article = sorted_articles[kept_idx]
                kept_text = (kept_article.category_tag or kept_article.title) + " " + (kept_article.ai_summary or "")
                removed_keywords = self._extract_core_entities(removed_text)
                kept_keywords = self._extract_core_entities(kept_text)
                overlap = removed_keywords & kept_keywords
                if len(overlap) >= len(removed_keywords) * 0.5:
                    is_covered = True
                    break
            
            if is_covered:
                verified_remove.add(idx)
            else:
                logger.info(f"  覆盖性校验: '{article.title[:30]}...'的独特信息可能未被覆盖，转为LLM判断")
                for root, indices in multi_groups.items():
                    if idx in indices:
                        groups_for_llm.append((root, [idx] + [i for i in indices if i not in to_remove and i != idx]))
                        break
        
        to_remove = verified_remove
        logger.info(f"[Step 6] 覆盖性校验完成: 确认移除{len(to_remove)}篇")

        # ============ Step 7: LLM仲裁兜底（对小簇中低相似度组） ============
        if groups_for_llm:
            logger.info(f"[Step 7] 开始LLM仲裁，{len(groups_for_llm)}个组待判断（并发{8}线程）")
            
            llm_tasks = []
            for group_id, (root, indices) in enumerate(groups_for_llm):
                active_indices = [i for i in indices if i not in to_remove]
                if len(active_indices) < 2:
                    continue
                
                articles_info = []
                for idx in active_indices:
                    art = sorted_articles[idx]
                    articles_info.append({
                        'title': art.title,
                        'source': art.source_name,
                        'summary': art.ai_summary[:500] if art.ai_summary else ""
                    })
                llm_tasks.append((group_id, articles_info, active_indices))
            
            if llm_tasks:
                llm_results = self.ai_client.batch_dedup_review(
                    [(gid, info) for gid, info, _ in llm_tasks]
                )
                
                for (gid, remove_set), (_, articles_info, active_indices) in zip(llm_results, llm_tasks):
                    for local_idx in remove_set:
                        if local_idx < len(active_indices):
                            original_idx = active_indices[local_idx]
                            to_remove.add(original_idx)
                            rm_art = sorted_articles[original_idx]
                            kept_indices = [active_indices[i] for i in range(len(active_indices)) if i not in remove_set]
                            kept_art = sorted_articles[kept_indices[0]] if kept_indices else sorted_articles[active_indices[0]]
                            self.dedup_details.append({
                                'removed_title': rm_art.title,
                                'removed_source': rm_art.source_name,
                                'removed_link': rm_art.link,
                                'kept_title': kept_art.title,
                                'kept_source': kept_art.source_name,
                                'kept_link': kept_art.link,
                                'reason': "LLM仲裁判断为重复",
                            })
                            logger.info(f"  去重(LLM判断): 保留'{kept_art.title[:30]}...'，移除'{rm_art.title[:30]}...'")
            
            logger.info(f"[Step 7] LLM仲裁完成")

        # ============ Step 8: 吸收与整合文章同主题的独立文章 ============
        # 检查结果列表中是否存在与已整合文章同主题的独立文章，将其吸收进整合文章
        # 【重要修改】提高实体重叠阈值至60%，且必须经AI确认是同一事件才能吸收，
        # 避免不相关文章仅因共享泛化词（如"联合国""万人"）就被错误吸收
        # 【新增】同时检查独立整合文章是否与更大的整合文章同主题，避免子话题整合重复出现
        absorbed_count = 0
        if merged_articles:
            remaining_articles = [article for idx, article in enumerate(sorted_articles) if idx not in to_remove]
            additional_remove = set()

            # 【新增】构建整合文章的来源link集合，用于检测"已被整合来源包含"的文章
            def _get_all_source_links(merged_art):
                """递归获取整合文章的所有原始来源链接"""
                links = set()
                if getattr(merged_art, 'merged_sources', None):
                    for src in merged_art.merged_sources:
                        if 'link' in src and src['link']:
                            links.add(src['link'])
                return links

            for merged_art in merged_articles:
                merged_text = (merged_art.category_tag or merged_art.title) + " " + (merged_art.ai_summary or "")
                merged_keywords = self._extract_core_entities(merged_text)
                merged_source_links = _get_all_source_links(merged_art)

                for idx, article in enumerate(sorted_articles):
                    if idx in to_remove:
                        continue
                    # 【修改】也允许其他整合文章被吸收（解决子话题整合重复出现的问题）
                    # if getattr(article, 'is_merged', False):
                    #     continue

                    # 【新增】如果该文章的link已被整合文章的来源包含，直接吸收（不需要AI确认）
                    if article.link and article.link in merged_source_links:
                        if not getattr(article, 'is_merged', False):
                            merged_art.merged_sources.append({
                                'title': article.title,
                                'source': article.source_name,
                                'link': article.link
                            })
                        elif getattr(article, 'merged_sources', None):
                            # 被吸收的文章本身是整合文章，展开其所有原始来源
                            merged_art.merged_sources.extend(article.merged_sources)
                        additional_remove.add(idx)
                        absorbed_count += 1
                        self.dedup_details.append({
                            'removed_title': article.title,
                            'removed_source': article.source_name,
                            'removed_link': article.link,
                            'kept_title': merged_art.category_tag,
                            'kept_source': "多来源整合",
                            'kept_link': "",
                            'reason': f"被整合文章吸收(来源已被包含在整合文章中)",
                        })
                        logger.info(f"  吸收(来源重复): '{article.title[:30]}...' 的来源已被整合文章 '{merged_art.category_tag[:30]}...' 包含")
                        continue

                    art_text = (article.category_tag or article.title) + " " + (article.ai_summary or "")
                    art_keywords = self._extract_core_entities(art_text)
                    overlap = merged_keywords & art_keywords

                    # 【补充】英文关键实体匹配（如SpaceX、OpenAI等）
                    # 解决中文实体重叠不够但英文实体一致的同主题文章吸收问题
                    english_entities = ['SPACEX', 'OPENAI', 'DEEPSEEK', 'TESLA', 'META', 'GOOGLE', 'AMAZON', 'MICROSOFT', 'NVIDIA', 'ANTHROPIC', 'XAI', 'STARLINK', 'BYTEDANCE', 'TIKTTOK', 'FABLE', 'MYTHOS', 'CLAUDE']
                    merged_eng_entities = set()
                    art_eng_entities = set()
                    for eng in english_entities:
                        if eng in merged_text.upper():
                            merged_eng_entities.add(eng)
                        if eng in art_text.upper():
                            art_eng_entities.add(eng)
                    eng_overlap = merged_eng_entities & art_eng_entities

                    # 核心实体重叠超过60%且至少3个共同实体 → 候选吸收（还需AI确认）
                    # 或：英文关键实体有重叠（至少1个）且中文实体重叠超过30%（至少1个） → 候选吸收
                    # 或：地点+灾害/事故类型匹配（同一地点+同一类事故的文章几乎一定是同一事件）
                    # 【新增】知名政治人物名称匹配：标题中都包含同一知名政治人物 → 候选吸收
                    #   解决"尹锡悦被判30年"(路透午报)未被吸收进"尹锡悦一审获刑30年"(观察者网+澎湃新闻)的问题
                    should_check_ai = False
                    if len(overlap) >= len(art_keywords) * 0.6 and len(overlap) >= 3:
                        should_check_ai = True
                    elif len(eng_overlap) >= 1 and len(overlap) >= max(1, len(art_keywords) * 0.3):
                        should_check_ai = True
                        logger.info(f"  英文实体匹配: 文章'{article.title[:30]}...'与整合文章'{merged_art.category_tag[:30]}...'共享英文实体{list(eng_overlap)}")
                    elif len(overlap) >= 1:
                        # 【新增】知名政治人物名称匹配：如果整合文章和独立文章的标题中都包含同一知名政治人物，
                        # 即使实体重叠不够60%，也应作为候选吸收（仍需AI确认是同一事件）
                        _political_figures = ['普京', '特朗普', '拜登', '习近平', '马克龙', '朔尔茨', '岸田',
                                              '尹锡悦', '莫迪', '苏纳克', '泽连斯基', '内塔尼亚胡', '金正恩',
                                              '拉马福萨', '卢拉', '米莱']
                        merged_title_figures = {f for f in _political_figures if f in (merged_art.category_tag or merged_art.title)}
                        art_title_figures = {f for f in _political_figures if f in (article.category_tag or article.title)}
                        common_figures = merged_title_figures & art_title_figures
                        if common_figures:
                            should_check_ai = True
                            logger.info(f"  政治人物匹配: 文章'{article.title[:30]}'与整合文章'{merged_art.category_tag[:30]}'共享政治人物{list(common_figures)}")
                        
                        if not should_check_ai:
                            # 地点+事故/灾害类型匹配：即使实体重叠很少，
                            # 如果两篇文章都涉及同一地点+同一事故/灾害类型，几乎一定是同一事件
                            # 如：整合文章"山西留神峪煤矿瓦斯爆炸" + 单篇"山西煤矿矿工未佩戴定位卡"
                            disaster_accident_types = ['煤矿', '矿难', '瓦斯', '爆炸', '矿工', '塌方', '坍塌',
                                                       '地震', '洪灾', '火灾', '踩踏', '坠机', '空难', '翻船']
                            location_words = ['山西', '陕西', '河南', '河北', '山东', '四川', '贵州', '云南',
                                              '湖南', '湖北', '安徽', '江西', '广西', '广东', '福建', '浙江',
                                              '江苏', '辽宁', '吉林', '黑龙江', '甘肃', '青海', '宁夏', '新疆',
                                              '西藏', '内蒙古', '北京', '上海', '天津', '重庆']
                            merged_has_disaster = any(dt in merged_text for dt in disaster_accident_types)
                            art_has_disaster = any(dt in art_text for dt in disaster_accident_types)
                            merged_has_location = any(lw in merged_text for lw in location_words)
                            art_has_location = any(lw in art_text for lw in location_words)
                            if merged_has_disaster and art_has_disaster and merged_has_location and art_has_location:
                                # 进一步检查：是否有共同的地点词
                                common_locations = set()
                                for lw in location_words:
                                    if lw in merged_text and lw in art_text:
                                        common_locations.add(lw)
                                if common_locations:
                                    should_check_ai = True
                                    logger.info(f"  地点+事故匹配: 文章'{article.title[:30]}'与整合文章'{merged_art.category_tag[:30]}'共享地点{list(common_locations)}+事故类型")
                    if should_check_ai:
                        # 用AI确认是否真的是同一事件，避免泛化词导致的误吸收
                        try:
                            is_same = self.ai_client.is_same_event(
                                {
                                    'title': merged_art.category_tag or merged_art.title,
                                    'ai_summary': (merged_art.ai_summary or '')[:400],
                                    'source_name': '多来源整合',
                                },
                                {
                                    'title': article.category_tag or article.title,
                                    'ai_summary': (article.ai_summary or '')[:400],
                                    'source_name': article.source_name,
                                }
                            )
                            if not is_same:
                                logger.info(f"  吸收跳过: AI判定'{article.title[:30]}...'与整合文章'{merged_art.category_tag[:30]}...'不是同一事件")
                                continue
                        except Exception as e:
                            logger.warning(f"  吸收AI判断失败，保守跳过: {e}")
                            continue

                        # 吸收该文章到整合文章
                        if getattr(article, 'is_merged', False) and getattr(article, 'merged_sources', None):
                            # 被吸收的文章本身是整合文章，展开其所有原始来源
                            # 但需要过滤掉与当前整合文章主题无关的来源
                            merged_text_lower = (merged_art.category_tag or merged_art.title).lower()
                            for src in article.merged_sources:
                                # 检查来源标题是否与整合文章主题相关
                                src_title_lower = src.get('title', '').lower()
                                # 计算来源标题与整合文章标题的相似度
                                src_relevance = self._calculate_similarity(src_title_lower, merged_text_lower)
                                # 也检查来源标题中是否包含整合文章的核心实体
                                merged_entities = self._extract_core_entities(merged_text_lower)
                                src_entities = self._extract_core_entities(src_title_lower)
                                entity_overlap = merged_entities & src_entities
                                # 来源标题与整合标题相似度>0.1 或 至少1个核心实体重叠 → 相关
                                if src_relevance > 0.1 or len(entity_overlap) >= 1:
                                    merged_art.merged_sources.append(src)
                                else:
                                    logger.info(f"  吸收来源过滤: '{src.get('title', '')[:30]}...' 与整合文章 '{merged_art.category_tag[:30]}...' 主题无关，不吸收此来源")
                        else:
                            merged_art.merged_sources.append({
                                'title': article.title,
                                'source': article.source_name,
                                'link': article.link
                            })
                        additional_remove.add(idx)
                        absorbed_count += 1
                        self.dedup_details.append({
                            'removed_title': article.title,
                            'removed_source': article.source_name,
                            'removed_link': article.link,
                            'kept_title': merged_art.category_tag,
                            'kept_source': "多来源整合",
                            'kept_link': "",
                            'reason': f"被整合文章吸收(AI确认同事件+核心实体重叠{len(overlap)}个)",
                        })
                        logger.info(f"  吸收: '{article.title[:30]}...' 被整合文章 '{merged_art.category_tag[:30]}...' 吸收(AI确认)")

            # 【新增】去除整合文章merged_sources中的重复来源
            # 去重维度：(1) 同一link只保留一条 (2) 同一source下标题高度相似的只保留一条
            for merged_art in merged_articles:
                if getattr(merged_art, 'merged_sources', None):
                    seen_links = set()
                    seen_source_titles = {}  # source -> [标题列表]
                    deduped_sources = []
                    for src in merged_art.merged_sources:
                        src_link = src.get('link', '')
                        src_source = src.get('source', '')
                        src_title = src.get('title', '')
                        
                        # 1. 同一link去重
                        if src_link and src_link in seen_links:
                            continue
                        if src_link:
                            seen_links.add(src_link)
                        
                        # 2. 同一来源下标题高度相似去重
                        if src_source and src_title:
                            if src_source not in seen_source_titles:
                                seen_source_titles[src_source] = []
                            is_dup_title = False
                            for existing_title in seen_source_titles[src_source]:
                                sim = self._calculate_similarity(src_title.lower(), existing_title.lower())
                                if sim > 0.5:
                                    is_dup_title = True
                                    break
                            if is_dup_title:
                                continue
                            seen_source_titles[src_source].append(src_title)
                        
                        deduped_sources.append(src)
                    
                    if len(deduped_sources) < len(merged_art.merged_sources):
                        logger.info(f"  来源去重: 整合文章 '{merged_art.category_tag[:30]}...' 去除{len(merged_art.merged_sources) - len(deduped_sources)}个重复来源")
                    merged_art.merged_sources = deduped_sources

            to_remove.update(additional_remove)
            if absorbed_count > 0:
                logger.info(f"[Step 8] 吸收同主题文章: {absorbed_count}篇被整合文章吸收(AI确认+来源去重)")

        # ============ Step 8.5: 合并同主题的整合文章 ============
        # 如果存在多个整合文章且主题相同（如同一事件的不同整合），将它们合并为一条
        # 【重要修改】提高阈值至60%且需AI确认同一事件，避免不相关整合文章被错误合并
        if len(merged_articles) > 1:
            merged_to_remove = set()  # 需要被合并掉的整合文章索引
            for i in range(len(merged_articles)):
                if i in merged_to_remove:
                    continue
                for j in range(i + 1, len(merged_articles)):
                    if j in merged_to_remove:
                        continue
                    art_i = merged_articles[i]
                    art_j = merged_articles[j]
                    
                    # 提取核心实体比较
                    text_i = (art_i.category_tag or art_i.title) + " " + (art_i.ai_summary or "")
                    text_j = (art_j.category_tag or art_j.title) + " " + (art_j.ai_summary or "")
                    keywords_i = self._extract_core_entities(text_i)
                    keywords_j = self._extract_core_entities(text_j)
                    overlap = keywords_i & keywords_j
                    
                    # 核心实体重叠超过60%且至少3个共同实体 → 候选合并（还需AI确认）
                    # 或：地点+灾害/事故类型匹配（同地点+同类事故的整合文章几乎一定是同一事件）
                    min_keywords = min(len(keywords_i), len(keywords_j))
                    should_merge_check = False
                    if min_keywords > 0 and len(overlap) >= min_keywords * 0.6 and len(overlap) >= 3:
                        should_merge_check = True
                    elif len(overlap) >= 1:
                        # 地点+事故类型匹配
                        disaster_accident_types = ['煤矿', '矿难', '瓦斯', '爆炸', '矿工', '塌方', '坍塌',
                                                   '地震', '洪灾', '火灾', '踩踏', '坠机', '空难', '翻船']
                        location_words = ['山西', '陕西', '河南', '河北', '山东', '四川', '贵州', '云南',
                                          '湖南', '湖北', '安徽', '江西', '广西', '广东', '福建', '浙江',
                                          '江苏', '辽宁', '吉林', '黑龙江', '甘肃', '青海', '宁夏', '新疆',
                                          '西藏', '内蒙古', '北京', '上海', '天津', '重庆']
                        i_has_disaster = any(dt in text_i for dt in disaster_accident_types)
                        j_has_disaster = any(dt in text_j for dt in disaster_accident_types)
                        common_locations = set()
                        for lw in location_words:
                            if lw in text_i and lw in text_j:
                                common_locations.add(lw)
                        if i_has_disaster and j_has_disaster and common_locations:
                            should_merge_check = True
                            logger.info(f"  整合合并地点+事故匹配: '{art_i.category_tag[:30]}...'与'{art_j.category_tag[:30]}...'共享地点{list(common_locations)}+事故类型")
                    
                    if should_merge_check:
                        # 用AI确认是否真的是同一事件
                        try:
                            is_same = self.ai_client.is_same_event(
                                {
                                    'title': art_i.category_tag or art_i.title,
                                    'ai_summary': (art_i.ai_summary or '')[:400],
                                    'source_name': '多来源整合',
                                },
                                {
                                    'title': art_j.category_tag or art_j.title,
                                    'ai_summary': (art_j.ai_summary or '')[:400],
                                    'source_name': '多来源整合',
                                }
                            )
                            if not is_same:
                                logger.info(f"  整合合并跳过: AI判定'{art_j.category_tag[:40]}...'与'{art_i.category_tag[:40]}...'不是同一事件")
                                continue
                        except Exception as e:
                            logger.warning(f"  整合合并AI判断失败，保守跳过: {e}")
                            continue

                        # 将j的来源吸收到i中，但过滤与i主题无关的来源
                        if art_j.merged_sources:
                            merged_i_title = (art_i.category_tag or art_i.title).lower()
                            for src in art_j.merged_sources:
                                src_title_lower = src.get('title', '').lower()
                                src_relevance = self._calculate_similarity(src_title_lower, merged_i_title)
                                merged_entities = self._extract_core_entities(merged_i_title)
                                src_entities = self._extract_core_entities(src_title_lower)
                                entity_overlap = merged_entities & src_entities
                                if src_relevance > 0.1 or len(entity_overlap) >= 1:
                                    art_i.merged_sources.append(src)
                                else:
                                    logger.info(f"  整合合并来源过滤: '{src.get('title', '')[:30]}...' 与 '{art_i.category_tag[:30]}...' 主题无关，不合并此来源")
                        # 合并摘要：将j的摘要追加到i的摘要中
                        if art_j.ai_summary:
                            art_i.ai_summary = art_i.ai_summary.rstrip() + "\n" + art_j.ai_summary.rstrip()
                        # 更新发布时间为两者中较新的
                        if art_j.pub_date and (not art_i.pub_date or art_j.pub_date > art_i.pub_date):
                            art_i.pub_date = art_j.pub_date
                        # 更新标题为更全面的一个
                        if len(art_j.category_tag or '') > len(art_i.category_tag or ''):
                            art_i.category_tag = art_j.category_tag
                            art_i.title = art_j.title
                        
                        merged_to_remove.add(j)
                        logger.info(f"  整合合并: '{art_j.category_tag[:40]}...' 被合并到 '{art_i.category_tag[:40]}...'(AI确认)")
            
            if merged_to_remove:
                merged_articles = [a for idx, a in enumerate(merged_articles) if idx not in merged_to_remove]
                logger.info(f"[Step 8.5] 同主题整合文章合并(AI确认): {len(merged_to_remove)}条被合并，剩余{len(merged_articles)}条整合文章")

        # ============ Step 8.7: 已禁用 ============
        # 原跨分类主题聚合整合逻辑已移除：仅靠实体重叠百分比无法准确判断文章是否同一事件，
        # 导致大量不相关文章被错误整合（如"也门内战"和"三星罢工"因共享"联合国""万人"等泛化词而被合并）。
        # 同事件的跨分类整合由 Step 8.8 (AI事件分组去重) 承担，AI判断更可靠。
        logger.info(f"[Step 8.7] 跨分类主题聚合已禁用（避免不相关文章错误整合）")

        # ============ Step 8.8: AI事件分组去重（补充TF-IDF和实体重叠无法捕获的同事件文章） ============
        # 对去重后剩余的文章，按标题中的关键人物/组织名分组，然后用AI判断是否同一事件
        # 解决如"特朗普访华"相关文章、"世界杯转播权"相关文章等TF-IDF相似度不够但确实是同一事件的问题
        remaining_indices = [idx for idx in range(len(sorted_articles)) if idx not in to_remove]
        
        if len(remaining_indices) > 3:
            # 提取标题和摘要中的关键实体用于预分组
            key_entity_groups = {}  # entity -> [article_indices]
            for idx in remaining_indices:
                art = sorted_articles[idx]
                title_text = (art.category_tag if art.category_tag else art.title)
                summary_text = (art.ai_summary or '')[:600]
                combined_text = title_text + " " + summary_text
                
                # 【关键修复】检测聚合文章（同时涉及多个不相关领域的文章）
                # 聚合文章不应参与实体分组，否则会把不同主题的文章连通在一起
                # 例如"8点1氪"同时包含"普京访华"和"DeepSeek回应"，如果参与分组
                # 会导致普京相关文章和DeepSeek相关文章被错误整合
                is_aggregator_article = False
                aggregator_indicators = ['8点1氪', '早报', '早知道', '快报', '快讯', '财经早餐', 
                                         '新闻早参考', '新闻早报', '资讯早报', '今日要闻',
                                         '新闻速递', '资讯速递', '早餐内参', '每日要闻']
                for indicator in aggregator_indicators:
                    if indicator in title_text:
                        is_aggregator_article = True
                        logger.info(f"[Step 8.8] 跳过聚合文章: '{title_text[:40]}...'")
                        break
                
                # 额外检测：标题中用分号/顿号分隔了多个独立新闻点
                if not is_aggregator_article:
                    # 检查标题是否用"；"或"；"分隔了多个独立话题
                    title_semicolon_parts = re.split(r'[；;]', title_text)
                    if len(title_semicolon_parts) >= 3:
                        is_aggregator_article = True
                        logger.info(f"[Step 8.8] 跳过多话题聚合文章: '{title_text[:40]}...'")
                
                if is_aggregator_article:
                    continue  # 聚合文章不参与实体分组
                
                # 提取关键人物名（2-4字中文 + 政治相关动作）- 仅从标题提取
                # 【修复】限制提取范围：仅从标题提取关键人物+外交动作的组合，
                # 避免从摘要中提取到"提及"而非"聚焦"的实体
                key_persons = re.findall(
                    r'([\u4e00-\u9fa5]{2,4})(?:访华|访美|访日|访韩|访俄|出访|到访|抵达|称赞|邀请|宣布|签署|发布|会见|会谈|通话|回应|表态|会晤|评价|访问)',
                    title_text)  # 仅从标题提取
                
                # 提取关键英文实体（公司名/组织名/人名/赛事名）- 仅从标题提取
                # 【修复】限制提取范围到标题，避免聚合文章摘要中的实体污染分组
                # 【修复】扩充公司名列表，加入SpaceX/Starlink/Anthropic/xAI等
                # 【修复】使用lookaround替代\b，因为\b在中英混合文本中不工作
                key_english = re.findall(r'(?<![a-zA-Z])(Cerebras|OpenAI|FIFA|DeepSeek|Tesla|Apple|Google|Meta|Amazon|Microsoft|NVIDIA|Intel|AMD|IPO|WHO|NBA|GDP|WSBK|MotoGP|F1|UFC|SpaceX|Starlink|Anthropic|xAI|ByteDance|TikTok|Alibaba|Tencent|Huawei|Samsung|Toyota|Volkswagen|Boeing|Airbus|Starship|S-1)(?![a-zA-Z])', title_text, re.IGNORECASE)
                key_english = [e.upper() for e in key_english]
                
                # 提取关键复合事件短语 - 更广泛的事件模式 - 仅从标题提取
                key_events = re.findall(
                    r'([\u4e00-\u9fa5]{2,6}(?:访华|访美|会晤|会谈|转播权|版权|IPO|上市|并购|收购|重组|制裁|冲突|协议|合作|签约|夺冠|坠亡|事故|道歉))',
                    title_text)
                
                # 提取核心事件名词 - 仅从标题提取
                core_event_nouns = re.findall(r'(世界杯|奥运会|亚运会|世博会|进博会|广交会)', title_text)
                core_event_actions = re.findall(r'(转播权|版权|播出权)', title_text)
                
                # 【新增】提取具体事件关键词 - 品牌/地点+事件组合 - 仅从标题提取
                brand_event_keywords = re.findall(r'(张雪机车|华蓥|大摆荡|坠亡|坠亡事件|夺冠)', title_text)
                
                # 【新增】提取中文公司/组织名+商业事件组合 - 仅从标题提取
                # 解决"SpaceX提交IPO"等中文公司名+事件的实体提取问题
                company_event_keywords = re.findall(
                    r'([\u4e00-\u9fa5]{2,6}(?:提交|提交IPO|冲刺IPO|IPO申请|招股书|上市|挂牌|募资|融资|并购|收购|重组|破产|退市|暴雷|违约|被罚|处罚|通报|立案|调查|被查))',
                    title_text)
                
                # 【新增】提取特定知名公司名作为独立实体（不要求后跟事件动词）
                # 解决"富途/老虎/长桥被罚"类文章，公司名本身就足以标识同一事件
                specific_company_names = re.findall(
                    r'(富途|老虎证券|长桥证券|富途控股|老虎证券|长桥|蚂蚁集团|字节跳动|滴滴|恒大|碧桂园|融创|万达|蔚来|理想汽车|小鹏汽车|比亚迪|宁德时代|中芯国际|大疆|商汤|旷视|地平线|百济神州|恒瑞医药|软银|孙正义|柠季|哈根达斯|通用磨坊)',
                    title_text)
                
                # 【新增】提取"跨境+行业关键词"组合（如"跨境券商""跨境炒股"）
                cross_border_keywords = re.findall(
                    r'(跨境[\u4e00-\u9fa5]{2,4})',
                    title_text)
                
                # 【新增】提取影视/文化作品名+事件关键词组合（如"给阿嬷的情书票房""传奇IP版权纠纷"）
                # 解决同一影视作品的多篇报道因TF-IDF相似度不够而未被去重的问题
                movie_event_keywords = re.findall(
                    r'([\u4e00-\u9fa5《》]{2,15})(?:票房|上映|首映|口碑|破亿|破十亿|夺冠|获奖|下架|延期|版权|和解|纠纷|败诉|胜诉)',
                    title_text)
                # 也匹配"《XX》"格式的影视作品名作为独立实体
                movie_name_entities = re.findall(r'《([\u4e00-\u9fa5A-Za-z0-9·\-]{2,20})》', title_text)
                
                # 【新增】提取地点+行业事件关键词组合（如"北京租房""山西煤矿"）
                # 解决地域性民生/安全事件的多篇报道去重问题
                location_event_keywords = re.findall(
                    r'([\u4e00-\u9fa5]{2,4})(?:租房|房源|毕业生.{0,3}租房|矿难|矿工|煤矿|瓦斯|爆炸)',
                    title_text)
                
                # 【新增】提取食品安全/产品质量事件关键词 - 仅从标题提取
                # 解决"泡药杨梅"等食品安全事件的去重问题
                food_safety_keywords = re.findall(
                    r'([\u4e00-\u9fa5]{2,6}(?:泡药|添加剂|防腐剂|瘦肉精|地沟油|三聚氰胺|塑化剂|农药残留|重金属超标|食物中毒|问题食品|食品安全|假货|假冒|掺假|造假|以次充好))',
                    title_text)
                # 也匹配事件本身作为关键词（如"泡药杨梅""问题杨梅"）
                food_safety_event_names = re.findall(
                    r'((?:泡药|问题|毒|假|劣).{0,4}(?:杨梅|奶粉|牛奶|肉类|鸡蛋|蔬菜|水果|海鲜|大米|食用油|食品))',
                    title_text)
                
                # 【新增】提取具体地点+事故组合 - 仅从标题提取
                location_accident = re.findall(r'([\u4e00-\u9fa5]{2,4})(?:坠亡|事故|坍塌|火灾|爆炸|踩踏)', title_text)
                
                # 如果标题提取不到实体，再从摘要中提取（但仅限于标题+摘要组合文本）
                # 【补充】对于整合文章（is_merged=True），如果标题提取不到关键实体，
                # 也从摘要中提取，因为整合文章的标题可能是空洞占位符
                if not (key_persons or key_english or key_events or core_event_nouns or brand_event_keywords or location_accident or company_event_keywords or food_safety_keywords or food_safety_event_names or specific_company_names or cross_border_keywords or movie_event_keywords or movie_name_entities or location_event_keywords):
                    # 标题没有明确实体，尝试从摘要补充
                    key_persons = re.findall(
                        r'([\u4e00-\u9fa5]{2,4})(?:访华|访美|访日|访韩|访俄|出访|到访|抵达|称赞|邀请|宣布|签署|发布|会见|会谈|通话|回应|表态|会晤|评价|访问)',
                        combined_text)
                    key_english = re.findall(r'(?<![a-zA-Z])(Cerebras|OpenAI|FIFA|DeepSeek|Tesla|Apple|Google|Meta|Amazon|Microsoft|NVIDIA|Intel|AMD|IPO|WHO|NBA|GDP|WSBK|MotoGP|F1|UFC|SpaceX|Starlink|Anthropic|xAI|ByteDance|TikTok|Alibaba|Tencent|Huawei|Samsung|Toyota|Volkswagen|Boeing|Airbus|Starship|S-1)(?![a-zA-Z])', combined_text, re.IGNORECASE)
                    key_english = [e.upper() for e in key_english]
                    key_events = re.findall(
                        r'([\u4e00-\u9fa5]{2,6}(?:访华|访美|会晤|会谈|转播权|版权|IPO|上市|并购|收购|重组|制裁|冲突|协议|合作|签约|夺冠|坠亡|事故|道歉))',
                        combined_text)
                    core_event_nouns = re.findall(r'(世界杯|奥运会|亚运会|世博会|进博会|广交会)', combined_text)
                    core_event_actions = re.findall(r'(转播权|版权|播出权)', combined_text)
                    brand_event_keywords = re.findall(r'(张雪机车|华蓥|大摆荡|坠亡|坠亡事件|夺冠)', combined_text)
                    location_accident = re.findall(r'([\u4e00-\u9fa5]{2,4})(?:坠亡|事故|坍塌|火灾|爆炸|踩踏)', combined_text)
                    company_event_keywords = re.findall(
                        r'([\u4e00-\u9fa5]{2,6}(?:提交|提交IPO|冲刺IPO|IPO申请|招股书|上市|挂牌|募资|融资|并购|收购|重组|破产|退市|暴雷|违约|被罚|处罚|通报|立案|调查|被查))',
                        combined_text)
                    food_safety_keywords = re.findall(
                        r'([\u4e00-\u9fa5]{2,6}(?:泡药|添加剂|防腐剂|瘦肉精|地沟油|三聚氰胺|塑化剂|农药残留|重金属超标|食物中毒|问题食品|食品安全|假货|假冒|掺假|造假|以次充好))',
                        combined_text)
                    food_safety_event_names = re.findall(
                        r'((?:泡药|问题|毒|假|劣).{0,4}(?:杨梅|奶粉|牛奶|肉类|鸡蛋|蔬菜|水果|海鲜|大米|食用油|食品))',
                        combined_text)
                    specific_company_names = re.findall(
                        r'(富途|老虎证券|长桥证券|富途控股|长桥|蚂蚁集团|字节跳动|滴滴|恒大|碧桂园|融创|万达|蔚来|理想汽车|小鹏汽车|比亚迪|宁德时代|中芯国际|大疆|商汤|旷视|地平线|百济神州|恒瑞医药|软银|孙正义|柠季|哈根达斯|通用磨坊)',
                        combined_text)
                    cross_border_keywords = re.findall(
                        r'(跨境[\u4e00-\u9fa5]{2,4})',
                        combined_text)
                    movie_event_keywords = re.findall(
                        r'([\u4e00-\u9fa5《》]{2,15})(?:票房|上映|首映|口碑|破亿|破十亿|夺冠|获奖|下架|延期|版权|和解|纠纷|败诉|胜诉)',
                        combined_text)
                    movie_name_entities = re.findall(r'《([\u4e00-\u9fa5A-Za-z0-9·\-]{2,20})》', combined_text)
                    location_event_keywords = re.findall(
                        r'([\u4e00-\u9fa5]{2,4})(?:租房|房源|毕业生.{0,3}租房|矿难|矿工|煤矿|瓦斯|爆炸)',
                        combined_text)
                
                # 组合核心名词和动作为复合实体
                for noun in core_event_nouns:
                    for action in core_event_actions:
                        compound = noun + action
                        key_events.append(compound)
                
                all_key_entities = key_persons + key_english + key_events + core_event_nouns + core_event_actions + brand_event_keywords + location_accident + company_event_keywords + food_safety_keywords + food_safety_event_names + specific_company_names + cross_border_keywords + movie_event_keywords + movie_name_entities + location_event_keywords
                
                for entity in all_key_entities:
                    if entity not in key_entity_groups:
                        key_entity_groups[entity] = []
                    key_entity_groups[entity].append(idx)
                
                # 【重要修复】限制单个文章对同一实体组的贡献：
                # 一篇文章最多只能向一个实体组贡献1次（避免一篇聚合文章把多个不相关实体连到同一个组）
                # 另外，如果一个实体组中超过50%的文章来自同一个来源，很可能是聚合文章的噪声，跳过该组
            
            # 找出包含2篇以上文章的实体组
            multi_article_entities = {e: indices for e, indices in key_entity_groups.items() if len(set(indices)) >= 2}
            
            if multi_article_entities:
                # 对每个实体组，用AI判断是否同一事件
                ai_dedup_groups = []  # 记录被AI确认需要合并的组
                for entity, group_indices in multi_article_entities.items():
                    unique_indices = list(set(group_indices))
                    if len(unique_indices) < 2:
                        continue
                    
                    # 【关键修复】预过滤：只有标题中也包含该实体的文章才参与分组
                    # 避免多篇文章仅因摘要中"提及"同一人物/实体就被错误整合
                    # 例如：特朗普访华、特朗普绿卡政策、美联储加息——虽然都提到"特朗普"，
                    # 但只有前两篇的标题聚焦特朗普，第三篇聚焦美联储
                    entity_filtered_indices = []
                    for idx in unique_indices:
                        art = sorted_articles[idx]
                        title_text = (art.category_tag if art.category_tag else art.title)
                        # 检查实体是否出现在标题中（大小写不敏感）
                        if entity.lower() in title_text.lower():
                            entity_filtered_indices.append(idx)
                        else:
                            logger.info(f"[Step 8.8] 实体'{entity}'预过滤: 文章'{title_text[:40]}...'标题中未包含该实体，排除出分组")
                    
                    if len(entity_filtered_indices) < 2:
                        logger.info(f"[Step 8.8] 实体'{entity}'预过滤后不足2篇，跳过")
                        continue
                    
                    unique_indices = entity_filtered_indices
                    
                    group_articles_info = []
                    for idx in unique_indices:
                        art = sorted_articles[idx]
                        # 【修复】已整合文章也参与AI事件分组，使其能与同主题单篇文章整合
                        group_articles_info.append({
                            'idx': idx,
                            'title': art.title,
                            'source': art.source_name,
                            'summary': (art.ai_summary or '')[:400],
                            'category_tag': art.category_tag or '',
                            'is_merged': getattr(art, 'is_merged', False),
                        })
                    
                    if len(group_articles_info) < 2:
                        continue
                    
                    # 调用AI判断是否同一事件
                    try:
                        is_same = self.ai_client.is_same_event(
                            {
                                'title': group_articles_info[0]['category_tag'] or group_articles_info[0]['title'],
                                'ai_summary': group_articles_info[0]['summary'],
                                'source_name': group_articles_info[0]['source'],
                            },
                            {
                                'title': group_articles_info[1]['category_tag'] or group_articles_info[1]['title'],
                                'ai_summary': group_articles_info[1]['summary'],
                                'source_name': group_articles_info[1]['source'],
                            }
                        )
                        
                        if is_same:
                            # 如果有超过2篇，继续检查其余文章
                            confirmed_indices = [group_articles_info[0]['idx'], group_articles_info[1]['idx']]
                            for k in range(2, len(group_articles_info)):
                                is_same_k = self.ai_client.is_same_event(
                                    {
                                        'title': sorted_articles[confirmed_indices[0]].category_tag or sorted_articles[confirmed_indices[0]].title,
                                        'ai_summary': (sorted_articles[confirmed_indices[0]].ai_summary or '')[:400],
                                        'source_name': sorted_articles[confirmed_indices[0]].source_name,
                                    },
                                    {
                                        'title': group_articles_info[k]['category_tag'] or group_articles_info[k]['title'],
                                        'ai_summary': group_articles_info[k]['summary'],
                                        'source_name': group_articles_info[k]['source'],
                                    }
                                )
                                if is_same_k:
                                    confirmed_indices.append(group_articles_info[k]['idx'])
                            
                            if len(confirmed_indices) >= 2:
                                ai_dedup_groups.append(confirmed_indices)
                                logger.info(f"[Step 8.8] AI事件分组: 实体'{entity}'关联的{len(confirmed_indices)}篇确认为同一事件")
                    except Exception as e:
                        logger.error(f"[Step 8.8] AI事件分组判断失败(实体'{entity}'): {e}")
                
                # 对确认的同一事件组进行整合
                ai_merged_count = 0
                for group_indices in ai_dedup_groups:
                    # 检查是否有文章已经被前面的组处理过
                    already_removed = [idx for idx in group_indices if idx in to_remove]
                    if len(already_removed) == len(group_indices):
                        continue
                    
                    remaining_in_group = [idx for idx in group_indices if idx not in to_remove]
                    if len(remaining_in_group) < 2:
                        continue
                    
                    _p(f"[AI事件去重] 发现同一事件组({len(remaining_in_group)}篇): 准备整合...")
                    merged_article = self._merge_large_cluster(sorted_articles, remaining_in_group)
                    if merged_article:
                        merged_articles.append(merged_article)
                        for idx in remaining_in_group:
                            to_remove.add(idx)
                            rm_art = sorted_articles[idx]
                            self.dedup_details.append({
                                'removed_title': rm_art.title,
                                'removed_source': rm_art.source_name,
                                'removed_link': rm_art.link,
                                'kept_title': merged_article.category_tag,
                                'kept_source': "多来源整合",
                                'kept_link': "",
                                'reason': f"AI事件分组去重(共{len(remaining_in_group)}篇)",
                            })
                        ai_merged_count += 1
                        _p(f"[AI事件去重] 成功整合 {len(remaining_in_group)} 篇文章为一条信息")
                    else:
                        _p(f"[AI事件去重] 整合失败，跳过")
                
                if ai_merged_count > 0:
                    logger.info(f"[Step 8.8] AI事件分组去重: {ai_merged_count}组被整合")

        # ============ Step 9: 输出结果 ============
        result = [article for idx, article in enumerate(sorted_articles) if idx not in to_remove]
        # 添加整合生成的文章
        result.extend(merged_articles)
        # 按发布时间重新排序
        result.sort(key=lambda x: x.pub_date, reverse=True)
        
        logger.info(f"[Step 9] 去重/整合完成！移除{len(to_remove)}篇，整合生成{len(merged_articles)}篇，保留{len(result)}篇")
        logger.info(f"=" * 60)

        return result

    def _determine_merged_category(self, cluster_articles: List[Article], cluster_type: str, latest_article: Article) -> str:
        """
        确定整合文章的MECE分类编号。
        
        对于重大事件簇(major_event)，根据事件类型和簇内文章分类多数派确定分类，
        避免因latest_article被AI误分而导致整个整合文章分类错误。
        对于世界杯簇强制归入3.2体育休闲。其他情况取簇内多数派分类。
        
        Args:
            cluster_articles: 簇内文章列表
            cluster_type: 簇类型
            latest_article: 最新文章（默认回退）
        
        Returns:
            MECE分类编号字符串
        """
        # 世界杯强制归入体育休闲
        if cluster_type == 'world_cup':
            return "3.2"
        
        # 收集簇内所有非广告文章的分类
        categories = [a.mece_category for a in cluster_articles if a.mece_category]
        if not categories:
            return latest_article.mece_category or "11.3"
        
        # 重大事件簇：优先使用1类分类（防止AI把国际事件误分为3/4类导致整合文章分类错误）
        if cluster_type == 'major_event':
            # 检查簇内文章的文本是否匹配重大国际事件关键词组
            all_text = " ".join(
                (a.category_tag or a.title) + " " + (a.ai_summary or "")
                for a in cluster_articles
            )
            for event_group in MAJOR_EVENT_KEYWORD_GROUPS:
                event_keywords = event_group['keywords']
                anchor_keywords = event_group.get('anchor_keywords', event_keywords)
                hits = {kw for kw in event_keywords if kw in all_text}
                anchor_hits = {kw for kw in anchor_keywords if kw in all_text}
                if len(hits) >= 2 and len(anchor_hits) >= 1:
                    # 确认是国际事件簇，优先取1类分类
                    cat_1x = [c for c in categories if c.split('.')[0] == '1']
                    if cat_1x:
                        # 取1类中的多数派
                        from collections import Counter
                        most_common_1x = Counter(cat_1x).most_common(1)[0][0]
                        logger.info(f"  重大事件簇分类决策: 命中事件'{event_group['name']}'，使用1类多数派分类 {most_common_1x}")
                        return most_common_1x
                    else:
                        # 簇内无1类分类，根据文本判断1.1还是1.2
                        conflict_words = ['战争', '冲突', '空袭', '轰炸', '打击', '导弹', '反击', '报复',
                                         '军事', '开战', '进攻', '摧毁', '以军', '防空', '拦截']
                        diplomacy_words = ['和谈', '停火', '停战', '和平', '协议', '谈判', '谅解', '备忘录',
                                         '签署', '签字', '外交', '恢复通行']
                        has_conflict = any(w in all_text for w in conflict_words)
                        has_diplomacy = any(w in all_text for w in diplomacy_words)
                        forced_cat = '1.2' if (has_diplomacy and not has_conflict) else '1.1'
                        logger.info(f"  重大事件簇分类决策: 命中事件'{event_group['name']}'但簇内无1类分类，强制设为{forced_cat}")
                        return forced_cat
        
        # 普通簇：取簇内多数派分类
        from collections import Counter
        if categories:
            most_common = Counter(categories).most_common(1)[0][0]
            return most_common
        
        return latest_article.mece_category or "11.3"

    def _merge_large_cluster(self, articles: List[Article], indices: List[int], cluster_type: str = None) -> Optional[Article]:
        """
        对大簇(>5篇)进行AI多视角整合，合并成一条信息
        
        整合格式：
        - 生成一个综合标题
        - 整合所有文章的关键信息（按时间线/不同视角组织）
        - 记录来源列表
        - 不包含"相关课题"部分
        
        Args:
            articles: 全部文章列表
            indices: 本簇文章在列表中的索引
            cluster_type: 簇类型，'gaokao'=高考主题, 'major_event'=重大事件主题,
                          None=普通簇。重大事件簇在实体预筛选时放宽阈值
        
        Returns:
            整合后的Article对象，失败返回None
        """
        if len(indices) <= 1:
            return None
        
        # 收集簇内所有文章信息
        cluster_articles = [articles[idx] for idx in indices]
        
        # 【问题4修复】对相同来源的文章进行去重：如果多篇来自同一来源且标题相似，只保留一篇
        seen_sources = {}  # source_name -> best_article_index_in_cluster
        deduped_indices = []
        for i, art in enumerate(cluster_articles):
            if art.source_name in seen_sources:
                # 相同来源，检查标题相似度
                prev_idx = seen_sources[art.source_name]
                prev_art = cluster_articles[prev_idx]
                title_sim = self._calculate_similarity(art.title.lower(), prev_art.title.lower())
                if title_sim > 0.6:
                    # 标题高度相似，只保留内容更长的
                    if len(art.ai_summary or '') > len(prev_art.ai_summary or ''):
                        # 替换为更长的版本
                        deduped_indices = [j for j in deduped_indices if j != prev_idx]
                        deduped_indices.append(i)
                        seen_sources[art.source_name] = i
                    logger.info(f"  整合去重: 相同来源'{art.source_name}'标题相似，保留一篇")
                    continue
            seen_sources[art.source_name] = i
            deduped_indices.append(i)
        
        cluster_articles = [cluster_articles[i] for i in deduped_indices]
        
        if len(cluster_articles) <= 1:
            # 去重后只剩1篇，直接返回
            return None
        
        # =================== 【关键修复】语义相关性预筛选 ===================
        # 在送入AI整合前，先用规则剔除主题明显离群的文章，避免AI被强制要求"整合"时
        # 把不相关文章硬拼起来。判定标准：
        # 1) 找出簇内所有文章标题/摘要的"核心实体重叠最多"的中位文章作为"锚点"
        # 2) 计算每篇文章与锚点的核心实体重叠数；重叠少于阈值(普通簇2个,重大事件簇1个)的视为离群，剔除
        # 这样可以彻底防止类似"美国司法部诉中企" + "重庆烤肉店通报"这种完全不同主题被合并
        # 【2026-06修复】重大事件簇(如美伊战争)不同角度报道的实体集差异大，
        # 放宽阈值到1，并用事件关键词补充实体集，避免多角度报道被误剔
        try:
            # 辅助函数：为文章实体集补充命中的重大事件关键词
            # 【关键修复】仅补充锚定关键词（事件特有词），不补充'战争''制裁'等泛化词，
            # 避免不同事件的文章因共享泛化词的实体补充而被误判为同主题
            def _supplement_with_event_keywords(entities_set, text):
                """将文本中命中的重大事件锚定关键词加入实体集，解决多角度报道实体交集过小问题"""
                _hit_kw = {kw for kw in _MAJOR_EVENT_ANCHOR_KEYWORDS if kw in text}
                entities_set.update(_hit_kw)
                return _hit_kw

            anchor_text = (cluster_articles[0].category_tag or cluster_articles[0].title) + " " + (cluster_articles[0].ai_summary or "")[:500]
            anchor_entities = self._extract_core_entities(anchor_text)

            # 【2026-06新增】重大事件簇：将事件关键词也加入锚点实体集
            if cluster_type == 'major_event':
                _anchor_event_kw = _supplement_with_event_keywords(anchor_entities, anchor_text)
                logger.info(f"  重大事件簇：锚点补充事件关键词 {_anchor_event_kw}")

            # 计算每篇与锚点的实体重叠数，找到最大的实体集合作为新锚点
            best_anchor_idx = 0
            best_score = 0
            for i, art in enumerate(cluster_articles):
                art_text = (art.category_tag or art.title) + " " + (art.ai_summary or "")[:500]
                art_entities = self._extract_core_entities(art_text)
                if cluster_type == 'major_event':
                    _supplement_with_event_keywords(art_entities, art_text)
                # 用与其他所有文章的实体重叠总数作为该文章的"中心度"
                centrality = 0
                for j, other in enumerate(cluster_articles):
                    if i == j:
                        continue
                    other_text = (other.category_tag or other.title) + " " + (other.ai_summary or "")[:500]
                    other_entities = self._extract_core_entities(other_text)
                    if cluster_type == 'major_event':
                        _supplement_with_event_keywords(other_entities, other_text)
                    centrality += len(art_entities & other_entities)
                if centrality > best_score:
                    best_score = centrality
                    best_anchor_idx = i

            anchor_text = (cluster_articles[best_anchor_idx].category_tag or cluster_articles[best_anchor_idx].title) + " " + (cluster_articles[best_anchor_idx].ai_summary or "")[:500]
            anchor_entities = self._extract_core_entities(anchor_text)
            if cluster_type == 'major_event':
                _supplement_with_event_keywords(anchor_entities, anchor_text)

            # 【2026-06修复】实体重叠阈值：普通簇 >= 3，重大事件簇 >= 2
            # 原阈值1太低导致完全不相关的文章（如德银加息、莫迪会面、张维为新秩序）
            # 因传递链连通+事件关键词补充实体后仅1个重叠就通过预筛选
            # 重大事件不同角度报道至少共享2个核心实体（如"伊朗"+"美军"）
            # 普通簇从2提高到3，避免仅因共享"美国""AI"等2个泛化实体就通过预筛选
            _entity_overlap_threshold = 2 if cluster_type == 'major_event' else 3

            # 剔除与锚点实体重叠 < 阈值 的离群文章（即"主题与本簇主题不相关"的文章）
            relevant_indices = []
            for i, art in enumerate(cluster_articles):
                if i == best_anchor_idx:
                    relevant_indices.append(i)
                    continue
                art_text = (art.category_tag or art.title) + " " + (art.ai_summary or "")[:500]
                art_entities = self._extract_core_entities(art_text)
                if cluster_type == 'major_event':
                    _supplement_with_event_keywords(art_entities, art_text)
                overlap = anchor_entities & art_entities
                # 实体重叠 >= 阈值 才视为同主题（重大事件簇阈值1，普通簇阈值2）
                if len(overlap) >= _entity_overlap_threshold:
                    relevant_indices.append(i)
                else:
                    logger.warning(
                        f"  整合预筛选剔除离群: '{art.title[:30]}...' 与锚点'{cluster_articles[best_anchor_idx].title[:30]}...' "
                        f"实体重叠仅{len(overlap)}个 (overlap={list(overlap)[:5]})"
                    )

            # =================== 【关键修复】标题级核心实体相关性检查 ===================
            # 对所有簇类型（包括major_event），额外检查每篇文章的标题是否包含锚点标题的核心实体。
            # 解决SpaceX IPO簇中混入"SK海力士IPO""AI版次贷危机""美加墨角力"等标题不含SpaceX/马斯克
            # 但因摘要中共享泛化金融词汇而通过实体预筛选的问题。
            # 【2026-06修复】major_event簇也启用标题检查，但使用更宽松的标准：
            # major_event允许锚点特定实体集为空（完全靠泛化实体匹配），其他簇则严格要求特定实体
            if cluster_type not in ('gaokao',) and len(relevant_indices) >= 2:
                anchor_title_text = cluster_articles[best_anchor_idx].category_tag or cluster_articles[best_anchor_idx].title
                anchor_title_entities = self._extract_core_entities(anchor_title_text)
                # 泛化实体：这些实体出现在很多不同主题的文章中，不能作为相关性依据
                generic_title_entities = {'IPO', '上市', '融资', '并购', '收购', '重组', '破产', '暴雷',
                                          '万亿', '美元', '破亿', '夺冠', '冠军', '开幕', '闭幕',
                                          '报道', '回应', '发布', '宣布', '最新', '突破', '首次', '创纪录',
                                          '纳斯达克', '纽交所', '港交所', '首富', '富豪', '亿万', '史上最大',
                                          '创纪录', '创历史', '刷新纪录', '市值', '股价', '市值破', '营收',
                                          '纪录', '历史', '史上', '全球', '全球最大', '全球首', '万亿美元',
                                          # 国名/地区名：出现在大量不同主题文章中，不能作为相关性依据
                                          '美国', '中国', '俄罗斯', '日本', '韩国', '英国', '法国', '德国',
                                          '印度', '伊朗', '以色列', '乌克兰', '巴西', '澳大利亚',
                                          # 泛化话题词：AI、科技等话题词出现在大量不同事件文章中
                                          'AI', '人工智能', '大模型', '科技', '军事', '经济', '政治',
                                          '安全', '监管', '管制', '限制', '禁令', '禁止', '制裁',
                                          # 泛化事件/状态词
                                          '危机', '改革', '政策', '战略', '竞争', '合作', '冲突',
                                          '战争', '挑战', '机遇', '趋势', '未来', '发展'}
                anchor_specific_entities = anchor_title_entities - generic_title_entities

                # 【2026-06修复】major_event簇锚点标题可能全是泛化实体（如"美伊冲突最新进展"），
                # 导致anchor_specific_entities为空。此时用摘要级实体重叠代替标题级检查
                if not anchor_specific_entities and cluster_type == 'major_event':
                    # major_event簇锚点无特定实体：用摘要级非泛化实体重叠做二次过滤
                    title_filtered = []
                    for i in relevant_indices:
                        art_text = (cluster_articles[i].category_tag or cluster_articles[i].title) + " " + (cluster_articles[i].ai_summary or "")[:500]
                        art_entities = self._extract_core_entities(art_text)
                        # 补充事件关键词
                        if cluster_type == 'major_event':
                            _supplement_with_event_keywords(art_entities, art_text)
                        overlap_with_anchor = anchor_entities & art_entities
                        non_generic_overlap = overlap_with_anchor - generic_title_entities
                        # major_event簇：至少2个非泛化实体重叠才保留（避免仅因共享泛化实体通过）
                        if len(non_generic_overlap) >= 2:
                            title_filtered.append(i)
                        else:
                            logger.warning(
                                f"  整合预筛选(major_event摘要级): '{cluster_articles[i].title[:30]}...' "
                                f"非泛化实体重叠仅{len(non_generic_overlap)}个({list(non_generic_overlap)[:5]})，排除"
                            )
                    if len(title_filtered) >= 2:
                        relevant_indices = title_filtered
                elif anchor_specific_entities:
                    title_filtered = []
                    for i in relevant_indices:
                        art_title_text = cluster_articles[i].category_tag or cluster_articles[i].title
                        art_title_entities = self._extract_core_entities(art_title_text)
                        # 标题中包含锚点的特定实体 → 保留
                        if anchor_specific_entities & art_title_entities:
                            title_filtered.append(i)
                        else:
                            # 标题不含锚点特定实体，但实体总重叠很高(>=5)且至少含1个非泛化实体 → 可能是高度相关的，保留
                            art_text = (cluster_articles[i].category_tag or cluster_articles[i].title) + " " + (cluster_articles[i].ai_summary or "")[:500]
                            art_entities = self._extract_core_entities(art_text)
                            overlap_with_anchor = anchor_entities & art_entities
                            # 【修复】从重叠实体中排除泛化实体，要求至少1个非泛化实体重叠
                            non_generic_overlap = overlap_with_anchor - generic_title_entities
                            if len(overlap_with_anchor) >= 5 and len(non_generic_overlap) >= 1:
                                title_filtered.append(i)
                                logger.info(f"  标题无特定实体但总重叠高(>={len(overlap_with_anchor)}): '{cluster_articles[i].title[:30]}...' 保留")
                            else:
                                logger.warning(
                                    f"  整合预筛选标题去离群: '{cluster_articles[i].title[:30]}' "
                                    f"标题不含锚点核心实体{list(anchor_specific_entities)[:5]}，且总实体重叠仅{len(overlap_with_anchor)}个(非泛化{len(non_generic_overlap)}个)，排除"
                                )
                    if len(title_filtered) >= 2:
                        relevant_indices = title_filtered
                    # 如果标题过滤后不足2篇，保留原有relevant_indices（保守策略）

            if len(relevant_indices) < 2:
                logger.warning(f"  整合预筛选: 剔除离群文章后剩余<2篇，放弃整合")
                return None

            # 进一步用AI抽样校验：取锚点+最远的1-2篇做 is_same_event 校验，
            # 如果AI判定不是同一事件，则放弃整合本簇（避免规则筛选漏网）
            if len(relevant_indices) >= 2:
                anchor_art = cluster_articles[best_anchor_idx]
                # 找出与锚点实体重叠最少（但 >= 2）的那篇做AI校验
                weakest_idx = None
                weakest_overlap = 999
                for i in relevant_indices:
                    if i == best_anchor_idx:
                        continue
                    art = cluster_articles[i]
                    art_text = (art.category_tag or art.title) + " " + (art.ai_summary or "")[:500]
                    art_entities = self._extract_core_entities(art_text)
                    ovl = len(anchor_entities & art_entities)
                    if ovl < weakest_overlap:
                        weakest_overlap = ovl
                        weakest_idx = i
                # 【2026-06修复】不再跳过major_event簇的AI校验
                # 原逻辑跳过AI校验会导致完全不相关的文章（如德银加息、莫迪会面、张维为新秩序）
                # 因传递链连通后被强制整合。现在改为：major_event簇也做AI校验，
                # 但使用更宽松的判断标准——只要"涉及同一重大事件的不同角度"就算同一事件
                _skip_ai_verify = cluster_type in ('gaokao', 'world_cup')
                if weakest_idx is not None and weakest_overlap <= 4 and not _skip_ai_verify:
                    try:
                        is_same = self.ai_client.is_same_event(
                            {
                                'title': anchor_art.category_tag or anchor_art.title,
                                'ai_summary': (anchor_art.ai_summary or '')[:400],
                                'source_name': anchor_art.source_name,
                            },
                            {
                                'title': cluster_articles[weakest_idx].category_tag or cluster_articles[weakest_idx].title,
                                'ai_summary': (cluster_articles[weakest_idx].ai_summary or '')[:400],
                                'source_name': cluster_articles[weakest_idx].source_name,
                            }
                        )
                        if not is_same:
                            logger.warning(
                                f"  整合预筛选: AI校验最弱关联文章不是同一事件，放弃整合本簇 "
                                f"('{anchor_art.title[:30]}...' vs '{cluster_articles[weakest_idx].title[:30]}...')"
                            )
                            return None
                    except Exception as _e:
                        logger.warning(f"  整合预筛选: AI校验异常，保守跳过本簇: {_e}")
                        return None

            cluster_articles = [cluster_articles[i] for i in relevant_indices]
        except Exception as _e:
            logger.warning(f"  整合预筛选异常（保守继续）：{_e}")

        if len(cluster_articles) < 2:
            return None

        # 构建整合prompt
        articles_text = []
        for i, art in enumerate(cluster_articles):
            articles_text.append(
                f"文章{i+1}:\n"
                f"标题：{art.title}\n"
                f"来源：{art.source_name}\n"
                f"摘要：{art.ai_summary[:600] if art.ai_summary else ''}\n"
            )
        articles_str = "\n\n".join(articles_text)
        
        system_prompt = "你是一位资深新闻编辑，擅长将多篇报道同一具体事件的文章整合为一条完整、信息丰富的资讯。"
        
        # 重大事件簇使用更宽松的整合提示，允许不同角度（军事/外交/经济）的报道整合
        _major_event_hint = ""
        if cluster_type == 'world_cup':
            _major_event_hint = """
**【世界杯赛事多角度整合说明】**：本组文章属于同一重大体育赛事的不同角度报道。同一赛事往往涵盖多个维度：
- **赛事维度**：比赛结果、进球、点球、淘汰赛等赛事进程
- **商业维度**：转播权、赞助商、版权、广告等商业合作
- **社会维度**：球迷文化、主办国影响、安保等社会影响
- **地缘维度**：美加墨合办、国际关系等政治关联
这些维度虽然看似不同，但都是同一赛事的不同侧面，**必须整合为一条资讯**，按维度或时间线组织内容。
"""
        elif cluster_type == 'major_event':
            _major_event_hint = """
**【重大事件多角度整合说明】**：本组文章属于同一重大国际事件的不同角度报道。同一重大事件往往涵盖多个维度：
- **军事维度**：空袭、打击、导弹、反击等军事行动
- **外交维度**：和谈、谈判、停火、制裁等外交进展
- **经济维度**：石油、海峡封锁、航运等经济影响
- **各方声明**：不同国家/组织的表态和回应
这些维度虽然看似不同，但都是同一重大事件的不同侧面，**必须整合为一条资讯**，按时间线或维度组织内容。不要因为"军事打击"与"和谈"表面不同就判定为不同事件而拒绝整合。
"""
        
        user_prompt = f"""请将以下{len(cluster_articles)}篇文章整合为一条完整的资讯。
{_major_event_hint}
**【极其重要的前提检查】**：在整合之前，你必须先检查这些文章是否确实在报道同一个具体事件。如果这些文章涉及多个完全不相关的事件（如一篇讲体育赛事、另一篇讲外交政治），则**绝对不能整合在一起**，你应该在输出开头标注【整合失败：文章主题不相关】，然后分别列出各篇文章的核心事件。只有当所有文章都在报道同一个具体事件的不同角度/进展时，才进行整合。

**【严禁跨事件整合】**：
- 绝对不能将不同事件的文章拼凑整合，即使它们都属于"社会争议"或"舆论事件"的泛化分类
- 例如：一篇讲"莫奈真迹被误判AI画作"、另一篇讲"盲道摆拍造假"——虽然都涉及"信任危机"，但这是两个完全不同的具体事件，绝对不能整合
- 例如：一篇讲"中美元首会晤"、另一篇讲"学术打假事件"——虽然可能出现在同一篇聚合文章中，但这是两个完全不同的具体事件，绝对不能整合
- **同一事件的判断标准：核心新闻事实完全一致**（同一个人、同一件事、同一个新闻事实），而不是仅因为共享某个抽象主题

{articles_str}

【整合要求】：
1. **标题**：生成一个简洁有力的综合标题，格式为【具体标题内容】，聚焦核心事件。标题必须是具体的事件描述，例如【特朗普访华：中美确立建设性战略稳定关系】或【外交部回应中美元首会晤多项议题】，绝对不要输出【标题】或【综合报道】等空洞占位符。**禁止使用问号结尾的反问句，必须用陈述句直接陈述核心事实。禁止使用比喻和修辞手法，用平实语言描述事实。标题信息量要高，总结全文主旨。**
2. **内容组织**：
   - 按照事件发展的时间线或不同视角（如前因后果、各方反应、不同角度细节）组织内容
   - 确保每篇文章的独特信息都被保留，不要遗漏任何重要细节
   - 如果不同文章提供了同一事件的不同细节（如：一篇讲经贸成果、一篇讲人物会谈、一篇讲后续影响），都要整合进去
   - 信息密度要高，去除重复表达
   - **绝对不能将不相关的事件拼凑在一起**——如果文章A讲的是"张雪机车夺冠"，文章B讲的是"特朗普访华"，这两个是完全不同的事件，不能整合
3. **格式要求**：
   - 首行是【具体标题】，标题必须包含事件的核心关键词
   - 正文是一段完整的叙述（4-6句、总字数不超过100字，严格限制），涵盖所有文章最核心的信息，正文内不要换行
   - 正文后空一行，加上"来源列表："，列出所有来源文章的标题（用《》包裹，用顿号分隔），不要在正文下方再单独列出每篇来源文章的详情
   - **绝对不要出现"相关课题"或类似标签**
4. **信息完整性**：
   - 保留所有关键数据、人名、时间、地点
   - 保留不同文章提供的不同视角和细节
   - 确保整合后的内容比单篇文章更全面
5. **字数严格限制**：
   - 正文内容（不含标题和来源列表）用 4-6 句短句，严格不超过 100 字
   - 只保留最关键信息，不要废话和冗余表述，宁可少写也不要超
   - 信息密度要高，用最精炼的语言整合多篇文章的核心内容

请直接输出整合后的内容，不要有任何解释说明。"""

        try:
            response = self.ai_client._call_api_with_system(system_prompt, user_prompt)
            merged_content = response.strip()
            
            # 检查AI是否判定文章主题不相关，拒绝整合
            if '整合失败' in merged_content and '主题不相关' in merged_content:
                logger.warning(f"AI判定文章主题不相关，拒绝整合: {[articles[idx].title[:30] for idx in indices]}")
                return None
            
            # 解析整合内容
            # 提取标题
            title_match = re.search(r'【([^】]+)】', merged_content)
            # 空洞占位符标题列表（扩充版，覆盖更多AI可能返回的占位符）
            hollow_titles = {'标题', '综合报道', '综合新闻', '新闻摘要', '综合资讯', '资讯', '报道', '综合', '新闻', '热点', '要闻', '事件', '新闻汇总', '综合要闻', '综合消息', '要点', '概要', '摘要', '详情', '详情请看', '整合报道', '专题', '专题报道', '快讯'}
            if title_match:
                raw_title = title_match.group(1).strip()
                # 如果AI返回的是空洞占位符标题，则从摘要内容中提取关键词生成标题
                if raw_title in hollow_titles or len(raw_title) <= 4:
                    # 从摘要中提取内容生成标题
                    summary_for_title = re.sub(r'【[^】]+】\s*', '', merged_content, count=1).strip()
                    # 去掉来源列表部分
                    summary_for_title = re.sub(r'\n*来源列表[：:].*$', '', summary_for_title, flags=re.DOTALL).strip()
                    # 取第一句话作为标题
                    first_sentence = re.split(r'[。！？\n]', summary_for_title)[0]
                    raw_title = first_sentence[:40] + ('…' if len(first_sentence) > 40 else '')
                    # 如果提取的标题仍然为空或空洞，使用第一篇文章的标题
                    if not raw_title.strip() or raw_title.strip() in hollow_titles:
                        raw_title = cluster_articles[0].title[:40]
                merged_title = f"【{raw_title}】"
                # 移除标题部分，剩余为摘要
                merged_summary = re.sub(r'【[^】]+】\s*', '', merged_content, count=1)
            else:
                # 如果没有【】标题，用第一行作为标题
                lines = merged_content.split('\n')
                raw_title = lines[0].strip()[:50]
                if raw_title in hollow_titles or len(raw_title) <= 4:
                    summary_for_title = '\n'.join(lines[1:]).strip()
                    first_sentence = re.split(r'[。！？\n]', summary_for_title)[0]
                    raw_title = first_sentence[:40] + ('…' if len(first_sentence) > 40 else '')
                    if (not raw_title.strip() or raw_title.strip() in hollow_titles) and cluster_articles:
                        raw_title = cluster_articles[0].title[:40]
                merged_title = f"【{raw_title}】"
                merged_summary = '\n'.join(lines[1:]).strip()
            
            # 提取来源列表（如果AI生成了的话）
            source_list = []
            for art in cluster_articles:
                if getattr(art, 'is_merged', False) and getattr(art, 'merged_sources', None):
                    for src in art.merged_sources:
                        source_list.append(f"《{src['title']}》")
                else:
                    source_list.append(f"《{art.title}》")
            
            # 如果摘要中没有来源列表，追加一个
            if "来源列表" not in merged_summary:
                source_str = "、".join(source_list)
                merged_summary = f"{merged_summary}\n\n来源列表：{source_str}"
            
            # 【字数检查】整合文章摘要正文严格不超过100字（4-6句）
            # 提取正文部分（去掉标题和来源列表）
            content_for_check = merged_summary
            if "来源列表" in content_for_check:
                content_for_check = re.sub(r'\n*来源列表[：:].*$', '', content_for_check, flags=re.DOTALL).strip()
            content_for_check = re.sub(r'【[^】]+】\s*', '', content_for_check, count=1).strip()
            content_chars = len(content_for_check.replace(' ', '').replace('\n', ''))
            
            if content_chars > 100:
                logger.info(f"整合摘要字数超限: '{merged_title[:30]}...', 当前{content_chars}字，开始精简...")
                # 调用AI精简
                condensed = self._condense_merged_summary(merged_content, 100)
                if condensed:
                    # 重新解析精简后的内容
                    title_match2 = re.search(r'【([^】]+)】', condensed)
                    if title_match2:
                        raw_title2 = title_match2.group(1).strip()
                        # 精简后的标题也可能是空洞占位符，需要检测
                        if raw_title2 not in hollow_titles and len(raw_title2) > 4:
                            merged_title = f"【{raw_title2}】"
                        else:
                            # 精简后标题为空洞占位符，保留之前的标题
                            logger.warning(f"  精简后标题为空洞占位符'{raw_title2}'，保留原标题")
                        merged_summary = re.sub(r'【[^】]+】\s*', '', condensed, count=1).strip()
                    else:
                        merged_summary = condensed
                    
                    # 重新追加来源列表
                    source_str = "、".join(source_list)
                    merged_summary = f"{merged_summary}\n\n来源列表：{source_str}"
                    
                    logger.info(f"整合摘要精简完成: '{merged_title[:30]}...'")

            
            # 【关键修复】整合后交叉主题检测
            # 使用AI判断整合结果是否包含多个不相关事件
            merged_summary_text = merged_summary
            if "来源列表" in merged_summary_text:
                merged_summary_text = re.sub(r'\n*来源列表[：:].*$', '', merged_summary_text, flags=re.DOTALL).strip()
            
            # 先做快速规则层检测：识别多个不相关的具体事件关键词
            # 定义不同具体事件的关键词（比领域更细粒度）
            specific_event_indicators = {
                '绿卡/移民政策': ['绿卡', '移民', '签证', '原籍国', '离境申请', '公民身份'],
                '普京/俄罗斯访华': ['普京', '专机', '俄方', '俄罗斯', '大礼包'],
                '美联储/沃什': ['美联储', '沃什', '加息', '降息', '利率决议'],
                '中东/美伊': ['美伊', '伊朗', '中东', '协议草案', '以伊', '霍尔木兹', '德黑兰', '伊核'],
                '特朗普个人': ['特朗普', '白宫', '取消出席', '婚礼'],
                '航天/神舟': ['神舟', '航天员', '发射', '乘组', '登月'],
                '跨境券商/富途老虎': ['富途', '老虎证券', '跨境', '券商', '非法展业', '罚没'],
                '食品安全': ['泡药', '添加剂', '防腐剂', '食品安全', '食物中毒', '假货', '掺假'],
                'AI科技': ['DeepSeek', 'AI', '大模型', 'ChatGPT', 'OpenAI', 'Gemini', '人工智能'],
                '地震灾害': ['地震', '洪灾', '暴雨', '坍塌', '踩踏', '火灾', '伤亡'],
                '俄乌冲突': ['俄乌', '乌克兰', '泽连斯基', '克里米亚', '顿巴斯'],
                '以巴/以黎冲突': ['以色列', '真主党', '加沙', '哈马斯', '内塔尼亚胡'],
                '日本/日韩': ['日本', '日韩', '日媒', '石脑油', '岸田', '日菲'],
                '朝鲜半岛': ['朝鲜', '朝核', '金正恩', '半岛'],
            }
            events_in_summary = set()
            for event_name, keywords in specific_event_indicators.items():
                if any(kw in merged_summary_text for kw in keywords):
                    events_in_summary.add(event_name)
            
            # 天然共现的事件组合（同一事件天然涉及两个子话题）
            natural_event_pairs = [
                {'特朗普个人', '绿卡/移民政策'},    # 特朗普政府出台移民政策
                {'特朗普个人', '中东/美伊'},       # 特朗普政府的中东政策
                {'特朗普个人', '美联储/沃什'},      # 特朗普提名美联储主席
                {'普京/俄罗斯访华', '特朗普个人'},  # 大国关系互动
                {'跨境券商/富途老虎', '美联储/沃什'},  # 金融监管联动
            ]
            
            if len(events_in_summary) >= 3:
                # 3个以上不相关的具体事件 → 几乎确定是过度整合
                # 检查是否有天然共现
                is_natural = any(
                    pair.issubset(events_in_summary) for pair in natural_event_pairs
                )
                if not is_natural:
                    logger.warning(f"整合后交叉主题检测: 整合摘要涵盖{len(events_in_summary)}个不相关事件{list(events_in_summary)}，拒绝整合")
                    _p(f"[整合] 交叉主题检测: 整合结果涵盖{len(events_in_summary)}个不相关事件{list(events_in_summary)}，拒绝整合")
                    return None
            elif len(events_in_summary) == 2:
                # 2个具体事件，检查是否天然共现
                is_natural = any(
                    pair.issubset(events_in_summary) for pair in natural_event_pairs
                )
                if not is_natural:
                    # 2个不相关事件：用AI做最终判断
                    try:
                        ai_cross_check = self.ai_client._call_api_with_system(
                            "你是一位资深新闻编辑。请判断以下整合后的资讯是否错误地将多个不相关的事件拼凑在一起。",
                            f"以下是一条整合后的资讯摘要：\n\n{merged_summary_text[:800]}\n\n"
                            f"检测到该摘要同时涉及以下事件关键词：{list(events_in_summary)}\n\n"
                            f"请判断：这条资讯是否把2个及以上完全不相关的具体新闻事件错误地拼凑在一起了？\n"
                            f"判断标准：如果这些事件之间没有直接的因果/影响/同一事件关系，就是错误整合。\n"
                            f"请严格按以下JSON格式回复：\n"
                            f'{{\"is_over_merged\": true/false, \"reason\": \"一句话说明\"}}'
                        )
                        import json as _json
                        check_result = _json.loads(ai_cross_check.strip())
                        if check_result.get('is_over_merged', False):
                            logger.warning(f"整合后AI交叉主题检测: AI判定过度整合({check_result.get('reason', '')})，拒绝整合")
                            _p(f"[整合] AI交叉主题检测: 判定过度整合，拒绝整合")
                            return None
                        else:
                            logger.info(f"整合后AI交叉主题检测: AI判定非过度整合({check_result.get('reason', '')})，允许整合")
                    except Exception as e:
                        logger.warning(f"整合后AI交叉主题检测失败: {e}，保守拒绝整合")
                        return None

            # 【关键兜底】如果整合后的标题仍然为空洞占位符，强制根据摘要内容重新生成
            # 提取当前标题内容（去掉【】）
            current_title_content = merged_title.replace('【', '').replace('】', '').strip()
            if current_title_content in hollow_titles or len(current_title_content) <= 4:
                # 尝试从摘要中提取第一句作为标题
                summary_text_for_title = merged_summary
                if "来源列表" in summary_text_for_title:
                    summary_text_for_title = re.sub(r'\n*来源列表[：:].*$', '', summary_text_for_title, flags=re.DOTALL).strip()
                first_sentence = re.split(r'[。！？\n]', summary_text_for_title)[0]
                if first_sentence and len(first_sentence) > 4 and first_sentence.strip() not in hollow_titles:
                    raw_fallback = first_sentence[:50]
                    merged_title = f"【{raw_fallback}】"
                    merged_article_for_tag = merged_title
                    logger.info(f"  整合标题兜底: 从摘要首句生成标题 '{raw_fallback[:30]}...'")
                else:
                    # 最终兜底：使用第一篇文章的标题
                    if cluster_articles:
                        fallback_title = cluster_articles[0].title[:50]
                        merged_title = f"【{fallback_title}】"
                        logger.warning(f"  整合标题兜底: 使用第一篇文章标题 '{fallback_title[:30]}...'")
            
            # 创建整合后的Article对象
            # 使用簇内最新文章的发布时间
            latest_article = max(cluster_articles, key=lambda x: x.pub_date)
            
            merged_article = Article(
                source_name="多来源整合",
                title=merged_title,
                link="",  # 整合文章无单一链接
                pub_date=latest_article.pub_date,
                full_content=merged_summary,
                ai_summary=merged_summary,
                category_tag=merged_title,
                mece_category=self._determine_merged_category(cluster_articles, cluster_type, latest_article),
            )
            # 标记为整合文章
            merged_article.is_merged = True
            merged_article.merged_sources = []
            seen_source_titles = set()  # 【修复】去重：同名文章只保留一个
            for art in cluster_articles:
                if getattr(art, 'is_merged', False) and getattr(art, 'merged_sources', None):
                    # 文章本身是整合文章，展开其所有原始来源
                    for src in art.merged_sources:
                        src_title_key = src.get('title', '').strip()
                        # 同名文章去重：如果标题完全一致，只保留第一个
                        if src_title_key not in seen_source_titles:
                            merged_article.merged_sources.append(src)
                            seen_source_titles.add(src_title_key)
                        else:
                            logger.info(f"  整合来源去重: 跳过重复标题'{src_title_key[:30]}...'")
                else:
                    src_title_key = art.title.strip()
                    if src_title_key not in seen_source_titles:
                        merged_article.merged_sources.append({
                            'title': art.title, 'source': art.source_name, 'link': art.link
                        })
                        seen_source_titles.add(src_title_key)
                    else:
                        logger.info(f"  整合来源去重: 跳过重复标题'{art.title[:30]}...'")
            
            logger.info(f"多视角整合完成: '{merged_title}' ({len(cluster_articles)}篇)")
            return merged_article
            
        except Exception as e:
            logger.error(f"大簇整合失败: {e}")
            return None

    def _condense_merged_summary(self, original_content: str, max_length: int = 100) -> Optional[str]:
        """
        对整合文章的超长摘要进行精简（4-6句、≤100字）
        
        Args:
            original_content: 原始整合内容（含标题和来源列表）
            max_length: 正文最大字数
            
        Returns:
            精简后的内容，失败返回None
        """
        system_prompt = "你是一位资深新闻编辑，擅长将冗长的资讯精简为极简、信息密度极高的内容。"
        
        user_prompt = f"""请将以下整合资讯的正文内容精简到不超过100字，用 4-6 句短句表达。

【极其重要】
- 正文用 4 到 6 个短句，总字数严格不超过 100 字，只保留最关键的信息
- 去除冗余表述和重复信息，宁可少写也不要超
- 保留核心事实、关键数据、重要结论
- 信息密度要高，不要废话
- **严禁包含人物百科式背景信息**（出生日期、毕业院校、职业履历等与核心事件无关的信息）
- **严禁将不同事件拼凑在一起**，如果原文包含多个不相关事件，只保留最核心的事件

原始内容：
{original_content[:6000]}

【格式要求】
1. 保持【标题】格式开头
2. 正文用 4-6 句、严格不超过 100 字
3. 不要包含来源列表（会在代码中单独追加）
4. 直接输出精简后的内容，不要有任何解释说明"""

        try:
            response = self.ai_client._call_api_with_system(system_prompt, user_prompt)
            condensed = response.strip()
            condensed = condensed.replace('\n', ' ').replace('\r', ' ')
            condensed = re.sub(r'\*\*(.*?)\*\*', r'\1', condensed)
            return condensed
        except Exception as e:
            logger.error(f"整合摘要精简失败: {e}")
            return None

    def _select_best_article(self, articles: List[Article], indices: List[int]) -> int:
        """
        从一组相似文章中选择最佳代表文（信息最全面的）
        
        打分规则：
        - 内容长度权重
        - 来源权威度权重
        - 摘要完整度权重
        - 时间新鲜度权重
        
        Args:
            articles: 全部文章列表
            indices: 候选文章的索引列表
        
        Returns:
            最佳文章的索引
        """
        best_idx = indices[0]
        best_score = -1
        
        for idx in indices:
            art = articles[idx]
            score = 0
            
            # 内容长度（0-30分）
            content_len = len(art.full_content or "")
            if content_len > 3000:
                score += 30
            elif content_len > 1500:
                score += 25
            elif content_len > 800:
                score += 20
            elif content_len > 400:
                score += 15
            else:
                score += 10
            
            # 摘要完整度（0-25分）
            summary_len = len(art.ai_summary or "")
            if summary_len > 300:
                score += 25
            elif summary_len > 200:
                score += 20
            elif summary_len > 100:
                score += 15
            else:
                score += 10
            
            # 来源权威度（0-25分）
            authoritative_sources = {
                '澎湃新闻', '财新', '新华社', '人民日报', '观察者网', '第一财经',
                '财经杂志', '华尔街见闻', '路透午报', '东方财富网', '36氪',
                '量子位', '机器之心', '钛媒体', '半月谈', '新财富杂志',
            }
            if art.source_name in authoritative_sources:
                score += 25
            else:
                score += 15
            
            # 时间新鲜度（0-20分）- 越新越好
            if art.pub_date:
                pub_dt = self.make_timezone_aware(art.pub_date)
                hours_ago = (datetime.now() - pub_dt).total_seconds() / 3600
                if hours_ago < 6:
                    score += 20
                elif hours_ago < 12:
                    score += 18
                elif hours_ago < 18:
                    score += 15
                else:
                    score += 10
            else:
                score += 10
            
            if score > best_score:
                best_score = score
                best_idx = idx
        
        return best_idx

    def _extract_core_entities(self, text: str) -> set:
        """
        提取文本中的核心实体（用于覆盖性校验）
        
        Args:
            text: 文本
        
        Returns:
            核心实体集合
        """
        if not text:
            return set()
        
        entities = set()
        
        # 提取人名（2-4字中文 + 事件关键词）
        person_event = re.findall(r'([\u4e00-\u9fa5]{2,4})(?:访华|访美|访日|访韩|抵达|认罪|被捕|被查|被诉|判刑|起诉|称赞|邀请|宣布|签署|发布|会见|会谈|通话|回应|表态|访俄|访欧|出访|到访)', text)
        entities.update(person_event)
        
        # 提取复合事件短语（人名+动作，作为一个整体实体，更精准地标识事件）
        compound_event_patterns = re.findall(
            r'([\u4e00-\u9fa5]{2,4}(?:访华|访美|访日|访韩|访俄|出访|到访|抵达北京|抵达上海|签署|宣布|回应|表态|称赞|邀请|会见|会谈|通话))', text)
        entities.update(compound_event_patterns)
        
        # 提取地名/国名 + 动作
        place_event = re.findall(r'([\u4e00-\u9fa5]{2,4})(?:冲突|战争|制裁|协议|合作|磋商|会谈)', text)
        entities.update(place_event)
        
        # 提取品牌/公司名
        brands = re.findall(r'(?:苹果|谷歌|微软|英伟达|Meta|特斯拉|亚马逊|百度|阿里|腾讯|字节|快手|华为|小米|OPPO|宇树|爱奇艺|Cerebras|OpenAI|FIFA|DeepSeek|高通|波音|联想|英特尔|AMD|英伟达|阿里云|AWS|Anthropic|Claude|Fable|Mistral|Perplexity|Midjourney|Stability|Cohere|Gemini|ChatGPT|Sora|GPT)', text)
        entities.update(brands)
        
        # 提取具体数字 + 单位
        numbers_with_unit = re.findall(r'(\d+(?:\.\d+)?(?:亿|万|千|百|块|元|美元|点|分|％|%))', text)
        entities.update(numbers_with_unit)
        
        # 提取英文缩写（2-6字母大写，如FIFA、AI、GDP、NBA、WHO）
        english_acronyms = re.findall(r'\b([A-Z]{2,6})\b', text)
        entities.update(english_acronyms)
        
        # 提取2-3字关键中文专有名词（高频名词，用于主题识别）
        # 模式：匹配2-3字中文词 + 后面的助词/标点，提取前面的名词
        short_noun_patterns = re.findall(
            r'([\u4e00-\u9fa5]{2,3})(?:的|了|是|与|和|在|将|被|把|向|对|由|为|至|从|因|按|经|比|及|等|之|'
            r'，|。|：|；|！|？|、|）|】|\"|\'|\s|$)', text)
        # 过滤通用词和停用词
        stop_words = {
            '事件', '问题', '情况', '原因', '影响', '表示', '认为', '目前',
            '记者', '消息', '报道', '公开', '相关', '如何', '对于', '关于',
            '通过', '进行', '已经', '可能', '需要', '我们', '你们', '他们',
            '社会', '全国', '各地', '网络', '网友', '引发', '关注', '成为',
            '一些', '很多', '不少', '部分', '多数', '少数', '基本', '主要',
            '这是', '这是', '什么', '怎么', '哪有', '如何', '为何', '但是',
            '因为', '所以', '如果', '虽然', '而且', '或者', '以及', '不过',
            '然而', '此外', '另外', '同时', '之后', '之前', '以来', '之间',
            '其中', '其他', '这个', '那个', '这些', '那些', '这种', '那种',
            '没有', '不是', '可以', '应该', '必须', '已经', '正在', '将要',
            '还要', '只是', '只有', '就是', '还是', '也有', '又有', '先有',
            '还要', '更要', '虽已', '也已', '早已', '曾有', '现有', '现有',
            '一年', '两年', '三年', '四年', '五年', '一次', '两次', '三次',
            '包括', '属于', '位于', '来自', '得到', '受到', '实现', '完成',
            '开始', '继续', '保持', '超过', '接近', '达到', '突破', '创下',
            '一个', '一种', '一项', '一些', '一行', '一起', '一定', '一致',
            '国内', '国外', '海外', '世界', '全球', '国际', '地区', '地方',
            '当地', '本身', '总体', '整体', '长期', '短期', '近期', '前期',
            '过去', '未来', '今年', '去年', '明年', '本月', '上月', '下月',
            '今天', '昨天', '明天', '目前', '当前', '现在', '以后', '此前',
            '此时', '届时', '以来', '期间', '以来', '以内', '以上', '以下',
            '之后', '之前', '之中', '之内', '之外', '之上', '之下', '之间',
        }
        # 只保留出现2次以上的2-3字词（提高精准度）
        from collections import Counter
        short_noun_counts = Counter(short_noun_patterns)
        for noun, count in short_noun_counts.items():
            if noun not in stop_words and len(noun) >= 2 and count >= 1:
                entities.add(noun)
        
        # 提取4字以上专有名词
        long_entities = re.findall(r'([\u4e00-\u9fa5]{4,8})(?:的|是|在|了|和|与|，|。|！|？|\s|$)', text)
        # 过滤通用词
        common_words = {
            '事件', '问题', '情况', '原因', '影响', '表示', '认为', '目前',
            '记者', '消息', '报道', '公开', '相关', '如何', '对于', '关于',
            '通过', '进行', '已经', '可能', '需要', '我们', '你们', '他们',
            '社会', '全国', '各地', '网络', '网友', '引发', '关注', '成为',
            '一些', '很多', '不少', '部分', '多数', '少数', '基本', '主要',
        }
        entities.update(e for e in long_entities if e not in common_words)
        
        return entities

    def _find_exact_duplicates(self, articles: List[Article]) -> List[Tuple[int, int]]:
        """检测完全重复的文章（相同来源+相似标题或摘要，或不同来源但事件核心相同）"""
        pairs = []
        n = len(articles)
        
        for i in range(n):
            for j in range(i + 1, n):
                article1 = articles[i]
                article2 = articles[j]
                
                # 相同来源
                if article1.source_name == article2.source_name:
                    # 标题完全相同或高度相似（>80%相似度）
                    title1 = article1.title.lower()
                    title2 = article2.title.lower()
                    title_similarity = self._calculate_similarity(title1, title2)
                    
                    if title_similarity > 0.8:
                        pairs.append((i, j))
                        logger.debug(f"精确重复: 文章{i}和文章{j}标题高度相似({title_similarity:.2f})")
                        continue
                    
                    # 摘要完全相同（直接复制）
                    summary1 = article1.ai_summary[:200] if article1.ai_summary else ""
                    summary2 = article2.ai_summary[:200] if article2.ai_summary else ""
                    if summary1 and summary2:
                        summary_similarity = self._calculate_similarity(summary1, summary2)
                        if summary_similarity > 0.9:
                            pairs.append((i, j))
                            logger.debug(f"精确重复: 文章{i}和文章{j}摘要高度相似({summary_similarity:.2f})")
                
                # 【新增】不同来源但标题高度相似（>85%相似度）- 也标记为精确重复
                # 例如同一事件的不同公众号转载
                else:
                    title1 = article1.title.lower()
                    title2 = article2.title.lower()
                    title_similarity = self._calculate_similarity(title1, title2)
                    if title_similarity > 0.85:
                        pairs.append((i, j))
                        logger.debug(f"跨源精确重复: 文章{i}({article1.source_name})和文章{j}({article2.source_name})标题高度相似({title_similarity:.2f})")
                        continue
                    
                    # 摘要高度相似
                    summary1 = article1.ai_summary[:300] if article1.ai_summary else ""
                    summary2 = article2.ai_summary[:300] if article2.ai_summary else ""
                    if summary1 and summary2:
                        summary_similarity = self._calculate_similarity(summary1, summary2)
                        if summary_similarity > 0.85:
                            pairs.append((i, j))
                            logger.debug(f"跨源精确重复: 文章{i}和文章{j}摘要高度相似({summary_similarity:.2f})")
        
        return pairs

    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """计算两个文本的相似度（简单Jaccard）"""
        if not text1 or not text2:
            return 0.0
        
        # Jaccard相似度
        set1 = set(text1)
        set2 = set(text2)
        intersection = len(set1 & set2)
        union = len(set1 | set2)
        
        if union == 0:
            return 0.0
        return intersection / union

    def crawl_all_feeds(self):
        """爬取所有RSS feed（多线程并行）"""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        logger.info(f"第一步：开始爬取并筛选1天内的文章（从{self.one_day_ago.strftime('%Y-%m-%d %H:%M:%S')}到现在）...")
        _p("\n==== 第一步 / 共六步：爬取 RSS 源（多线程并行）====")

        df = self.load_rss_links()
        total_sources = len(df)
        _p(f"[步骤1] 待爬取公众号数: {total_sources}，并发线程数: {self.max_workers}")

        # 准备RSS源列表
        rss_tasks = []
        for idx, row in df.iterrows():
            source_name = row['公众号名称']
            rss_url = row['RSS链接']
            if pd.isna(rss_url) or not str(rss_url).startswith('http'):
                logger.warning(f"{source_name}: RSS链接无效，跳过")
                continue
            rss_tasks.append((source_name, rss_url))

        # 多线程并行爬取
        all_articles = []
        completed_count = 0
        lock = __import__('threading').Lock()
        
        def fetch_single_feed(task):
            source_name, rss_url = task
            try:
                articles = self.parse_rss_feed(rss_url, source_name)
                return source_name, articles
            except Exception as e:
                logger.error(f"  {source_name}: 爬取失败: {e}")
                return source_name, []
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_task = {executor.submit(fetch_single_feed, task): task for task in rss_tasks}
            
            for future in as_completed(future_to_task):
                completed_count += 1
                source_name, articles = future.result()
                
                if articles:
                    with lock:
                        all_articles.extend(articles)
                        logger.info(f"  {source_name}: 获取{len(articles)}篇文章")
                        _p(f"[步骤1][{completed_count}/{len(rss_tasks)}] {source_name} -> 获取 {len(articles)} 篇（累计 {len(all_articles)} 篇）")
                else:
                    _p(f"[步骤1][{completed_count}/{len(rss_tasks)}] {source_name} -> 未获取到 1 天内的新文章")

        self.articles = all_articles
        logger.info(f"第一步完成！共获取{len(all_articles)}篇文章")
        _p(f"[步骤1] 完成 ✓ 获取 {len(all_articles)} 篇 1 天内文章")

        # 第一步(补)：注入「网站资讯源」（监控源3：AIbase / AI HOT / TopHub / NewsNow /
        # TechURLs / Buzzing / AI HubToday 等中文 AI 媒体+社区站）。抓取方法照搬
        # ai-news-radar，仅保留今日(24h内)条目；注入后与公众号文章走同一条 AI 过滤/分类/
        # 去重/限长摘要管线，因此现有的去广告等内容筛选依然生效。
        if os.environ.get("ENABLE_WEB_SOURCES", "1") != "0":
            try:
                from web_sources import collect_web_items
                _p("\n==== 第一步(补)：采集网站资讯源（监控源3）====")
                web_items = collect_web_items(self.now, self.one_day_ago)
                _added = 0
                for it in web_items:
                    self.articles.append(Article(
                        source_name=it["source"],
                        title=it["title"],
                        link=it["url"],
                        pub_date=it["published_at"],
                        full_content=it.get("content") or it["title"],
                    ))
                    _added += 1
                logger.info(f"网站资讯源注入 {_added} 篇今日文章")
                _p(f"[步骤1补] 网站源注入 {_added} 篇今日资讯（累计 {len(self.articles)} 篇）")
            except Exception as e:
                logger.error(f"网站资讯源采集失败（不影响公众号流程）: {e}")
                _p(f"[步骤1补] 网站源采集失败，已跳过：{e}")

        # 【关键修复】基于链接的去重：不同公众号可能转载同一文章（链接相同），
        # 导致同一篇文章在结果中出现两次。此处按链接去重，保留第一篇（通常来自原发公众号）。
        if self.articles:
            seen_links = {}
            deduped_articles = []
            link_dup_count = 0
            for art in self.articles:
                if art.link and art.link in seen_links:
                    link_dup_count += 1
                    logger.debug(f"  链接级去重: '{art.title[:30]}...' 与 '{seen_links[art.link][:30]}...' 链接相同，跳过")
                    continue
                if art.link:
                    seen_links[art.link] = art.title
                deduped_articles.append(art)
            if link_dup_count > 0:
                logger.info(f"链接级去重: 移除{link_dup_count}篇重复文章（相同链接不同来源）")
                _p(f"[步骤1] 链接级去重: 移除 {link_dup_count} 篇重复文章")
            self.articles = deduped_articles

        # 按发布时间排序
        if self.articles:
            self.articles.sort(key=lambda x: x.pub_date, reverse=True)

        valid_articles = 0
        ad_articles = 0

        # 第二步：AI合并判断广告+MECE分类（8线程并发，跳过已缓存的文章）
        logger.info("第二步：AI合并判断广告+MECE分类（8线程并发）...")
        _p(f"\n==== 第二步 / 共六步：AI 判断广告 + 分类（8线程并发）====")
        total_step2 = 0
        cached_count = 0
        
        # 分离需要判断的文章和已缓存的文章
        articles_to_judge = []
        for i, article in enumerate(self.articles):
            if getattr(article, '_from_cache', False):
                cached_count += 1
                continue
            articles_to_judge.append((i, article))
        
        total_step2 = len(articles_to_judge)
        _p(f"[步骤2] 需要判断: {total_step2}篇，缓存: {cached_count}篇")
        
        if articles_to_judge:
            # 批量并发判断广告+分类（合并为一次API调用）
            judge_data = [
                {'title': art.title, 'content': art.full_content}
                for _, art in articles_to_judge
            ]
            ad_mece_results = self.ai_client.batch_classify_ad_and_mece(judge_data)
            
            ad_articles = 0
            for (orig_idx, article), result in zip(articles_to_judge, ad_mece_results):
                article.is_advertisement = result['is_ad']
                article.rejection_reason = result.get('reason', '') if result['is_ad'] else ""
                # 非广告文章直接设置MECE分类
                if not result['is_ad'] and result.get('category'):
                    article.mece_category = result['category']
                if result['is_ad']:
                    ad_articles += 1
                logger.info(f"  AI判断: {article.title[:30]}... -> {'广告' if result['is_ad'] else '有效(' + result.get('category', '?') + ')'}")
            
            _p(f"[步骤2] 完成 共 {total_step2} 篇判断，广告 {ad_articles} 篇，{cached_count} 篇来自缓存")
        else:
            _p(f"[步骤2] 无新文章需要判断，{cached_count} 篇来自缓存")

        # 【关键修复】MECE分类后处理：对匹配重大国际事件但被AI误分到非1类的文章进行纠正
        # 解决美伊停战谅解备忘录等国际外交事件被误分为文体娱乐(3类)或科技产业(4类)的问题
        # 【2026-06修复】大幅收紧阈值：原阈值(2关键词+1锚定词)太低，
        # 导致仅因摘要中提及"制裁""石油""战争"等泛化词就被错误纠正到1类，
        # 把本该属于2类(经济/金融)、4类(科技)、8类(企业)的文章错误归入国际局势
        _mece_fixed_count = 0
        for article in self.articles:
            if article.is_advertisement or not article.mece_category:
                continue
            # 检查文章是否匹配重大国际事件关键词组
            article_text = (article.category_tag or article.title) + " " + (article.ai_summary or article.full_content or "")
            for event_group in MAJOR_EVENT_KEYWORD_GROUPS:
                event_keywords = event_group['keywords']
                anchor_keywords = event_group.get('anchor_keywords', event_keywords)
                # 文章命中该事件的关键词数量
                hits = {kw for kw in event_keywords if kw in article_text}
                anchor_hits = {kw for kw in anchor_keywords if kw in article_text}
                # 【2026-06修复】大幅收紧阈值：
                # 1) 关键词命中数从2提高到4（2个泛化词如"石油""制裁"不足以确定是国际事件）
                # 2) 锚定关键词命中数从1提高到2（仅1个锚定词可能只是提及）
                # 3) 额外约束：文章标题中必须至少包含1个锚定关键词（确保核心主题是国际事件，而非仅摘要中提及）
                title_anchor_hits = {kw for kw in anchor_keywords if kw in (article.category_tag or article.title)}
                if len(hits) >= 4 and len(anchor_hits) >= 2 and len(title_anchor_hits) >= 1:
                    current_main = article.mece_category.split('.')[0] if '.' in article.mece_category else article.mece_category
                    # 如果文章被分到非1类，强制纠正为1.1或1.2
                    if current_main != '1':
                        old_cat = article.mece_category
                        # 根据事件内容判断1.1(冲突)还是1.2(外交)
                        conflict_words = ['战争', '冲突', '空袭', '轰炸', '打击', '导弹', '反击', '报复',
                                         '军事', '开战', '进攻', '摧毁', '以军', '防空', '拦截']
                        diplomacy_words = ['和谈', '停火', '停战', '和平', '协议', '谈判', '谅解', '备忘录',
                                         '签署', '签字', '外交', '恢复通行']
                        has_conflict = any(w in article_text for w in conflict_words)
                        has_diplomacy = any(w in article_text for w in diplomacy_words)
                        if has_diplomacy and not has_conflict:
                            article.mece_category = '1.2'
                        else:
                            article.mece_category = '1.1'
                        _mece_fixed_count += 1
                        logger.info(
                            f"  [MECE纠正] 文章'{article.title[:40]}...' 从{old_cat}纠正为{article.mece_category} "
                            f"(命中事件'{event_group['name']}'，关键词命中={hits})"
                        )
                    break  # 匹配到一个事件组即可，无需继续检查
        if _mece_fixed_count > 0:
            _p(f"[MECE纠正] 共纠正 {_mece_fixed_count} 篇重大国际事件文章的分类（误分→1类国际局势）")

        # 第三步：AI生成400-500字摘要 + 高标题（并行处理，标题已在摘要生成中一并产出）
        logger.info("第三步：AI生成400-500字摘要+高标题（并行处理）...")
        _step3_total = sum(1 for a in self.articles if not a.is_advertisement)
        _step3_cached = sum(1 for a in self.articles if not a.is_advertisement and getattr(a, '_from_cache', False))
        _p(f"\n==== 第三步 / 共六步：AI 生成摘要 + 高标题（并行处理）===")
        _step3_done = 0

        # 并行生成摘要（已合并标题优化，无需再调用batch_optimize_titles）
        articles_to_summarize = [
            a for a in self.articles
            if not a.is_advertisement and not getattr(a, '_from_cache', False) and a.full_content
        ]
        
        if articles_to_summarize:
            from concurrent.futures import ThreadPoolExecutor, as_completed
            
            def _gen_summary(article):
                try:
                    summary, category_tag = self.ai_client.generate_summary_with_retry(
                        article.title, article.full_content
                    )
                    return article, summary, category_tag, None
                except Exception as e:
                    logger.error(f"摘要生成失败: {article.title[:30]}... 错误: {e}")
                    return article, "摘要生成失败", "", str(e)
            
            with ThreadPoolExecutor(max_workers=8) as executor:
                futures = {executor.submit(_gen_summary, a): a for a in articles_to_summarize}
                completed = 0
                for future in as_completed(futures):
                    article, summary, category_tag, error = future.result()
                    article.ai_summary = summary
                    if category_tag:
                        article.category_tag = category_tag
                    completed += 1
                    if error:
                        logger.warning(f"[{completed}/{len(articles_to_summarize)}] {article.title[:30]}... 失败")
                    else:
                        logger.info(f"[{completed}/{len(articles_to_summarize)}] {article.title[:30]}... -> {len(summary)}字, 标题={category_tag}")
        
        _step3_done = len(articles_to_summarize)
        valid_articles = _step3_total

        logger.info(f"AI摘要生成完成！为{valid_articles}篇有效文章生成摘要（{_step3_cached}篇来自缓存）")
        _p(f"[步骤3] 完成 共生成 {_step3_done} 篇摘要（{_step3_cached} 篇来自缓存）")

        # 高标题后处理：从摘要中移除【】标签部分，并兜底未获得标题的文章
        import re

        INVALID_TITLE_PLACEHOLDERS = {
            "标题", "某某事件如何怎样", "某某事件", "某某", "xxx", "XXX",
            "新闻标题", "文章标题", "无标题",
        }

        def _is_invalid_tag(t: str) -> bool:
            t = (t or "").strip()
            return (not t) or len(t) < 4 or t in INVALID_TITLE_PLACEHOLDERS

        for article in self.articles:
            if article.is_advertisement or not article.ai_summary:
                continue
            # 从ai_summary中移除【】标题部分，只保留正文
            article.ai_summary = re.sub(r'【[^】]+】\s*', '', article.ai_summary)

            # 兜底：如果没有从摘要生成中获得高标题，根据内容/摘要生成而非直接用原文标题
            _need_fallback = False
            if not getattr(article, 'category_tag', None) or not article.category_tag:
                _need_fallback = True
            else:
                _inner = article.category_tag[1:-1] if article.category_tag.startswith("【") and article.category_tag.endswith("】") else article.category_tag
                if _is_invalid_tag(_inner):
                    _need_fallback = True
                else:
                    # ============ 【关键修复】检测AI生成的标题是否过于接近原文标题 ============
                    # 如果AI偷懒直接复用了原文标题（常见于带比喻/网络用语的标题），强制重新生成
                    # 判定：标题去除标点后的字符 Jaccard 相似度 > 0.6 视为复用
                    def _norm(s):
                        return re.sub(r'[【】！？，。、；：""''《》（）()\s—\-…·]', '', s or '')
                    inner_norm = _norm(_inner)
                    title_norm = _norm(article.title or '')
                    if inner_norm and title_norm:
                        s1, s2 = set(inner_norm), set(title_norm)
                        sim = len(s1 & s2) / max(1, len(s1 | s2))
                        # 同时检测是否包含原文标题中的"比喻/悬念/反问"等不良元素
                        bad_indicators = [
                            '？', '?', '！！', '——', '…', '弃子', '一把火', '凉了',
                            '都回来', '走火入魔', '黄蓉', '巧妇', '冰火', '炸锅', '失控',
                            '太离谱', '天塌了', '栽了', '崩了', '瓜了', '薄如蝉翼',
                        ]
                        has_bad = any(b in (article.title or '') for b in bad_indicators) or \
                                  any(b in _inner for b in bad_indicators)
                        # 高相似 或 有不良比喻元素 → 强制走 optimize_title 重新生成
                        if sim > 0.6 or has_bad:
                            try:
                                logger.info(
                                    f"标题强制重生(sim={sim:.2f}, has_bad={has_bad}): "
                                    f"原标题='{(article.title or '')[:30]}' AI标题='{_inner[:30]}'"
                                )
                                _new_title = self.ai_client.optimize_title(
                                    article.title or "",
                                    article.ai_summary or article.full_content or "",
                                    force=True
                                )
                                if _new_title:
                                    article.category_tag = _new_title
                                    logger.info(f"  -> 重生后='{_new_title}'")
                            except Exception as _e:
                                logger.warning(f"  标题强制重生失败，保留原AI标题: {_e}")
            if _need_fallback:
                # 优先根据摘要内容调用AI生成标题
                _fallback_title = None
                try:
                    _content_for_title = article.ai_summary or article.full_content or ""
                    if _content_for_title:
                        _fallback_title = self.ai_client.optimize_title(
                            article.title or "",
                            _content_for_title,
                            force=True  # 兜底场景强制根据内容生成，不沿用原文标题
                        )
                        if _fallback_title:
                            logger.info(f"兜底标题生成: 原标题='{(article.title or '')[:30]}' -> AI生成='{_fallback_title}'")
                except Exception as e:
                    logger.warning(f"兜底标题AI生成失败，将用摘要首句: {e}")
                # 如果AI生成也失败了，用摘要首句截取作为标题（比原文标题更贴合内容）
                if not _fallback_title:
                    first_sentence = re.split(r'[。！？；]', article.ai_summary or "")[0]
                    if len(first_sentence) >= 10:
                        _fallback_title = f"【{first_sentence[:50]}】"
                    else:
                        _fallback_title = f"【{(article.title or '').strip()[:40]}】"
                article.category_tag = _fallback_title

        # 第四步：生成每日总结（在所有处理完成后）
        logger.info("第四步：生成每日资讯总结...")
        _p(f"\n==== 第四步 / 共六步：生成每日总结 ====")
        valid_articles_list = [a for a in self.articles if not a.is_advertisement and a.ai_summary]
        if valid_articles_list:
            _p(f"[步骤5] 基于 {len(valid_articles_list)} 篇有效文章生成每日总结...")
            articles_for_summary = [
                {'title': a.title, 'ai_summary': a.ai_summary}
                for a in valid_articles_list
            ]
            self.daily_summary = self.ai_client.generate_daily_summary(articles_for_summary)
            logger.info(f"每日总结生成完成: {len(self.daily_summary)}字")
            _p(f"[步骤5] 完成 每日总结 {len(self.daily_summary)} 字")
        else:
            self.daily_summary = "今日无资讯更新。"
            logger.info("无有效文章，跳过总结生成")
            _p("[步骤5] 无有效文章，跳过总结生成")

        # 第五步：规则过滤（过滤鸡汤文、医学案例、综合式文章、地方新闻等）
        logger.info("第五步：规则过滤噪音内容...")
        _p(f"\n==== 第五步 / 共六步：规则过滤噪音内容 ====")
        noise_count = 0
        for article in self.articles:
            if article.is_advertisement:
                continue  # 已经是广告的跳过
            is_noise, reason = article.is_noise_content()
            if not is_noise:
                continue

            # 【AI 复核】对仅凭关键词命中而被判为"主题杂糅/聚合快讯/综合集成"的文章
            # 进行 DeepSeek 二次判断：AI 认为不是杂糅就撤销过滤，避免误杀跨领域深度分析。
            needs_ai_review = any(tag in reason for tag in (
                "内容杂糅类", "聚合快讯类", "综合集成式文章"
            ))
            if needs_ai_review:
                try:
                    _p(f"[步骤5] 规则疑似噪音，AI 复核中: {article.title[:40]} ({reason})")
                    is_mess, ai_reason = self.ai_client.is_topic_mess(
                        article.title,
                        article.full_content or article.ai_summary or "",
                        rule_reason=reason,
                    )
                    if not is_mess:
                        # AI 判定不是杂糅：保留文章，不进入噪音列表
                        logger.info(
                            f"  噪音过滤撤销(AI复核): {article.title[:30]}... "
                            f"规则理由='{reason}'，AI认为='{ai_reason}'"
                        )
                        _p(f"   └─ AI 复核：保留（{ai_reason}）")
                        continue
                    else:
                        # AI 也认为杂糅，追加 AI 理由
                        reason = f"{reason} | AI复核确认：{ai_reason}"
                except Exception as _e:
                    logger.error(f"  AI 杂糅复核异常，保留规则判定: {_e}")
                    _p(f"   └─ AI 复核异常（按规则过滤）: {_e}")

            article.is_advertisement = True
            article.rejection_reason = reason
            self.noise_articles.append(article)  # 保存噪音文章
            noise_count += 1
            logger.info(f"  噪音过滤: {article.title[:30]}... -> {reason}")
            _p(f"[步骤5] 噪音过滤: {article.title[:40]} -> {reason}")
        logger.info(f"第五步完成！规则过滤出{noise_count}篇噪音文章")
        _p(f"[步骤5] 完成 过滤 {noise_count} 篇噪音文章")

        # 第六步：去重处理（AI判断同一事件，保留发布时间更近的）
        logger.info("第六步：去重处理，合并同一事件的报道...")
        _p(f"\n==== 第六步 / 共六步：去重（同一事件合并）====")
        original_articles = self.articles.copy()  # 保存原始列表用于对比
        deduped_articles = self._deduplicate_articles(self.articles)
        removed_count = len(self.articles) - len(deduped_articles)
        
        # 记录被去重移除的文章（仅记录非广告/非噪音的有效文章）
        deduped_links = {a.link for a in deduped_articles}
        for article in original_articles:
            if article.link not in deduped_links:
                # 只有非广告、非噪音的文章才算"去重移除"
                # 广告和噪音文章已经在之前的步骤中被标记，不应该出现在去重移除列表
                if not article.is_advertisement:
                    self.duplicate_removed.append(article)
                # 如果是广告/噪音文章被去重逻辑也排除了，不需要额外记录
                # 因为它们已经在噪音过滤列表中了
        
        logger.info(f"去重完成！移除{removed_count}篇重复文章，保留{len(deduped_articles)}篇")
        _p(f"[步骤7] 完成 去重移除 {removed_count} 篇，最终保留 {len(deduped_articles)} 篇")
        self.articles = deduped_articles

def _clean_summaries(articles):
    """
    清理所有文章摘要中的违规内容：
    1. 去掉发布元信息（如"文章由XX发布，作者XX，编辑XX"）
    2. 去掉"原文未进一步说明""后续处理措施未提及"等表述
    3. 确保摘要首句不重复文章标题
    """
    import re
    
    # 发布元信息清理模式
    meta_patterns = [
        r'[，。]?\s*文章由[^。]+发布[，，][^。]+',
        r'[，。]?\s*作者[：:][^，。]+[，，]编辑[：:][^，。]+',
        r'[，。]?\s*编辑[：:][^，。]+[，，]作者[：:][^，。]+',
        r'[，。]?\s*来源[：:][^，。]+编辑[：:][^，。]+',
        r'[，。]?\s*本文由[^。]+发布',
        r'[，。]?\s*转载自[^，。]+',
    ]
    
    # "原文未说明"类表述清理模式
    no_info_patterns = [
        r'[，。]?\s*但原文未[进一步]{0,3}(说明|提及|公布|披露|给出|提供)[^。]*。',
        r'[，。]?\s*原文未[进一步]{0,3}(说明|提及|公布|披露|给出|提供)[^。]*。',
        r'[，。]?\s*(具体|详细|后续)(处理|处罚|结果|措施|进展)(未|暂未|尚无)[^。]*。',
        r'[，。]?\s*(未|暂未|尚无)(公布|披露|说明|提及|给出)(具体|详细|后续)[^。]*。',
    ]
    
    cleaned_count = 0
    for article in articles:
        if article.is_advertisement or not article.ai_summary:
            continue
        
        original = article.ai_summary
        cleaned = original
        
        # 清理发布元信息
        for pattern in meta_patterns:
            cleaned = re.sub(pattern, '', cleaned)
        
        # 清理"原文未说明"类表述
        for pattern in no_info_patterns:
            cleaned = re.sub(pattern, '', cleaned)
        
        # 清理多余空格和连续标点
        cleaned = re.sub(r'\s{2,}', ' ', cleaned)
        cleaned = re.sub(r'[，。]{2,}', '，', cleaned)
        cleaned = cleaned.strip()
        
        if cleaned != original:
            article.ai_summary = cleaned
            cleaned_count += 1
    
    if cleaned_count > 0:
        logger.info(f"摘要后处理清理: 清理了{cleaned_count}篇文章的摘要")

def _convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """
    将docx文件转换为pdf文件
    
    支持多种转换方式（按优先级尝试）：
    1. docx2pdf Python库（macOS 调用 Microsoft Word，Windows 调用 COM）— 格式与Word原文完全一致
    2. LibreOffice (soffice) 命令行（完整保留格式）
    3. fpdf2 纯Python库（run 级精确渲染，保留超链接/粗体/字号/颜色）
    """
    import subprocess
    import platform
    
    # 方式1: 使用docx2pdf库（优先：格式与Word"另存为PDF"完全一致）
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            logger.info(f"PDF生成成功(docx2pdf): {pdf_path}")
            return
    except ImportError:
        logger.info("docx2pdf未安装，尝试其他方式")
    except Exception as e:
        logger.warning(f"docx2pdf转换失败: {e}")
    
    # 方式2: 使用LibreOffice命令行转换
    libreoffice_names = ['soffice', 'libreoffice']
    soffice_path = None
    for name in libreoffice_names:
        try:
            result = subprocess.run([name, '--version'], capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                soffice_path = name
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    if soffice_path:
        output_dir = os.path.dirname(pdf_path)
        try:
            cmd = [soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            if result.returncode == 0 and os.path.exists(pdf_path):
                logger.info(f"PDF生成成功(LibreOffice): {pdf_path}")
                return
            else:
                logger.warning(f"LibreOffice转换失败: {result.stderr}")
        except (subprocess.TimeoutExpired, Exception) as e:
            logger.warning(f"LibreOffice转换异常: {e}")
    
    # 方式3: 使用fpdf2纯Python库（run 级精确渲染）
    try:
        _convert_docx_to_pdf_via_fpdf2(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            logger.info(f"PDF生成成功(fpdf2): {pdf_path}")
            return
    except ImportError:
        logger.warning("fpdf2未安装，尝试其他方式")
    except Exception as e:
        logger.warning(f"fpdf2转换失败: {e}")
    
    # 所有方式都失败
    raise RuntimeError(f"无法将docx转换为pdf：docx2pdf、LibreOffice和fpdf2均不可用。请安装Microsoft Word后运行 pip install docx2pdf，或安装LibreOffice，或运行 pip install fpdf2")


def _find_chinese_font() -> tuple:
    """
    查找系统中支持中文的字体文件路径。
    返回 (regular_path, bold_path) 元组；粗体可能和常规体相同。
    """
    import platform
    system = platform.system()

    if system == 'Darwin':  # macOS
        # 优先用 STHeiti Light + Medium（黑体族，粗细搭配最佳）
        candidates_regular = [
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/System/Library/Fonts/Hiragino Sans GB.ttc',
            '/System/Library/Fonts/Supplemental/Songti.ttc',
            '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
            '/Library/Fonts/Arial Unicode.ttf',
        ]
        candidates_bold = [
            '/System/Library/Fonts/STHeiti Medium.ttc',
            '/System/Library/Fonts/Hiragino Sans GB.ttc',
            '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
        ]
    elif system == 'Windows':
        base = os.path.join(os.environ.get('WINDIR', r'C:\Windows'), 'Fonts')
        candidates_regular = [
            os.path.join(base, 'msyh.ttc'),
            os.path.join(base, 'simhei.ttf'),
            os.path.join(base, 'simsun.ttc'),
        ]
        candidates_bold = [
            os.path.join(base, 'msyhbd.ttc'),
            os.path.join(base, 'simhei.ttf'),
        ]
    else:  # Linux
        candidates_regular = [
            '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
            '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
        ]
        candidates_bold = candidates_regular[:]

    regular = bold = ""
    for p in candidates_regular:
        if os.path.exists(p):
            regular = p
            break
    for p in candidates_bold:
        if os.path.exists(p):
            bold = p
            break
    if not bold and regular:
        bold = regular
    return regular, bold


def _emu_to_pt(emu_val) -> float:
    """将 Word 的 EMU (English Metric Unit) 值转为 pt（磅）。
    Word 中 w:sz 的值是半磅单位，即 val=24 表示 12pt。"""
    if emu_val is None:
        return 0
    return emu_val / 2.0


def _extract_para_text_with_runs(para):
    """
    从段落中提取所有文本片段（含超链接），返回列表：
    [(text, is_bold, font_size_pt, color_hex), ...]
    font_size_pt 为 0 表示使用默认字号。
    """
    from docx.oxml.ns import qn
    segments = []

    def _parse_rpr(rpr, default_bold, default_size, default_color):
        """从 w:rPr 中解析粗体 / 字号 / 颜色"""
        bold = default_bold
        size = default_size
        color = default_color
        if rpr is not None:
            b_elem = rpr.find(qn('w:b'))
            if b_elem is not None:
                bold = True
            sz_elem = rpr.find(qn('w:sz'))
            if sz_elem is not None:
                try:
                    size = _emu_to_pt(int(sz_elem.get(qn('w:val'))))
                except (ValueError, TypeError):
                    pass
            color_elem = rpr.find(qn('w:color'))
            if color_elem is not None:
                c = color_elem.get(qn('w:val'))
                if c:
                    color = c
        return bold, size, color

    for child in para._element:
        if child.tag == qn('w:r'):
            # 普通 run
            rpr = child.find(qn('w:rPr'))
            bold, size, color = _parse_rpr(rpr, False, 0, '000000')
            t_elem = child.find(qn('w:t'))
            if t_elem is not None and t_elem.text:
                segments.append((t_elem.text, bold, size, color))
        elif child.tag == qn('w:hyperlink'):
            # 超链接 run
            for r in child.findall(qn('w:r')):
                rpr = r.find(qn('w:rPr'))
                # 超链接默认蓝色
                bold, size, color = _parse_rpr(rpr, False, 0, '0000FF')
                t_elem = r.find(qn('w:t'))
                if t_elem is not None and t_elem.text:
                    segments.append((t_elem.text, bold, size, color))

    return segments


def _convert_docx_to_pdf_via_fpdf2(docx_path: str, pdf_path: str):
    """使用fpdf2读取docx内容并生成PDF（纯Python，无需外部依赖）
    
    改进：按 run 级别精确渲染，保留超链接文本、蓝色标记、粗体/字号差异。
    """
    from fpdf import FPDF
    from docx import Document as DocxDocument

    # 【降噪】抑制 fontTools / fpdf2 内部的大量 INFO 日志（只保留 WARNING 以上）
    for noisy_logger_name in ('fontTools', 'fontTools.subset', 'fontTools.ttLib', 'fpdf', 'fpdf.font'):
        logging.getLogger(noisy_logger_name).setLevel(logging.WARNING)

    # 查找中文字体（常规 + 粗体）
    font_regular, font_bold = _find_chinese_font()
    if not font_regular:
        raise RuntimeError("未找到系统中支持中文的字体文件")

    # 读取docx文档
    doc = DocxDocument(docx_path)

    # 创建PDF
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # 注册字体
    font_name = 'zhfont'
    pdf.add_font(font_name, '', font_regular)
    if font_bold != font_regular:
        pdf.add_font(font_name, 'B', font_bold)
    else:
        pdf.add_font(font_name, 'B', font_bold)

    page_width = pdf.w - pdf.l_margin - pdf.r_margin

    # 默认字号
    DEFAULT_SIZE = 10.5
    # 行高系数（行高 = 字号 * LINE_HEIGHT_FACTOR）
    LINE_HEIGHT_FACTOR = 1.5

    def _render_segments(segments, default_size=DEFAULT_SIZE):
        """将一个段落的 segments 用 write 逐段渲染（保留 run 级格式）"""
        for text, is_bold, size_pt, color_hex in segments:
            sz = size_pt if size_pt > 0 else default_size
            style = 'B' if is_bold else ''
            pdf.set_font(font_name, style, sz)
            # 设置颜色
            try:
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
                pdf.set_text_color(r, g, b)
            except (ValueError, IndexError):
                pdf.set_text_color(0, 0, 0)
            line_h = sz * LINE_HEIGHT_FACTOR * 0.3528  # pt → mm
            pdf.write(line_h, text)
        # 段落结束换行
        pdf.ln()

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ''

        # 空行 → 小间距
        if not text:
            pdf.ln(2)
            continue

        # 提取 run 级数据
        segments = _extract_para_text_with_runs(para)

        # 推断段落默认字号
        if style_name == 'Title':
            para_default_size = 16
        elif 'Heading 1' in style_name:
            para_default_size = 18
        elif 'Heading 2' in style_name:
            para_default_size = 15
        elif 'Heading 3' in style_name:
            para_default_size = 13
        elif 'Heading 4' in style_name:
            para_default_size = 12
        else:
            # 从段落的 run 中探测字号
            sizes_in_para = [s for _, _, s, _ in segments if s > 0]
            if sizes_in_para:
                # 取众数（最常见的字号）
                from collections import Counter
                size_counter = Counter(sizes_in_para)
                para_default_size = size_counter.most_common(1)[0][0]
            else:
                para_default_size = DEFAULT_SIZE

        # 分隔线
        if not style_name.startswith('Heading') and style_name != 'Title':
            if all(c in '-=—─═━' for c in text.replace(' ', '')):
                pdf.set_font(font_name, '', 8)
                pdf.set_text_color(128, 128, 128)
                pdf.multi_cell(w=page_width, h=3, text='─' * 70)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(2)
                continue

        # 判断段落类型以设置段前/段后间距
        is_category_title = text.startswith('【') and text.endswith('】')
        is_article_title = any(text.startswith(f'{cn}、') for cn in
                               '一二三四五六七八九十百千') and '【' in text
        is_toc_line = any(text.startswith(f'{cn}、') for cn in
                          '一二三四五六七八九十') and not '【' in text and len(text) < 30

        # 段前间距
        if style_name == 'Title':
            pass  # 标题不需要段前间距
        elif is_category_title:
            pdf.ln(4)
        elif is_article_title:
            pdf.ln(2)

        # 渲染段落内容
        _render_segments(segments, para_default_size)

        # 段后间距
        if style_name == 'Title':
            pdf.ln(6)
        elif 'Heading 1' in style_name:
            pdf.ln(4)
        elif is_category_title:
            pdf.ln(2)
        elif is_article_title:
            pdf.ln(1)

    # 处理表格（如果有的话）
    for table in doc.tables:
        pdf.ln(3)
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.replace('|', '').strip():
                pdf.set_font(font_name, '', 9)
                pdf.set_text_color(0, 0, 0)
                pdf.multi_cell(w=page_width, h=4, text=row_text)

    pdf.output(pdf_path)


def _auto_fix_duplicates(crawler, suggestions: str) -> bool:
    """
    根据API验证结果的重复建议，自动在文章列表中合并重复文章
    
    Args:
        crawler: WeChatArticleCrawler实例
        suggestions: API返回的重复建议文本
    
    Returns:
        是否成功修复了重复
    """
    import re
    
    # 构建非广告文章的有序列表（与文档中的序号对应）
    valid_articles = [a for a in crawler.articles if not a.is_advertisement and a.ai_summary and a.ai_summary != "摘要生成失败"]
    
    # 按组分解析建议，格式如"第1条和第2条都涉及ZZZ；第3条和第5条都涉及YYY"
    # 先按分号/句号拆分成独立的重复组
    group_texts = re.split(r'[；;。\n]', suggestions)
    
    merge_groups = []  # 每个元素是一组需要合并的文章编号列表
    for group_text in group_texts:
        # 从每组中提取"第X条"的编号
        refs_in_group = re.findall(r'第(\d+)条', group_text)
        if len(refs_in_group) >= 2:
            # 转为0-based索引
            group_indices = []
            for ref_str in refs_in_group:
                ref_num = int(ref_str)
                if 1 <= ref_num <= len(valid_articles):
                    group_indices.append(ref_num - 1)
            group_indices = sorted(set(group_indices))
            if len(group_indices) >= 2:
                merge_groups.append(group_indices)
    
    if not merge_groups:
        logger.info("自动修复：无法从建议中解析文章编号，跳过")
        return False
    
    # 合并重叠的组（如组A包含[1,2]，组B包含[2,3]，则合并为[1,2,3]）
    merged_groups = []
    for group in merge_groups:
        group_set = set(group)
        merged_into_existing = False
        for i, existing in enumerate(merged_groups):
            if group_set & set(existing):  # 有重叠
                merged_groups[i] = sorted(set(existing) | group_set)
                merged_into_existing = True
                break
        if not merged_into_existing:
            merged_groups.append(group)
    
    any_fixed = False
    for group_indices in merged_groups:
        articles_to_merge = [valid_articles[i] for i in group_indices if i < len(valid_articles)]
        
        if len(articles_to_merge) < 2:
            continue
        
        # 【重要新增】合并前必须AI确认这些文章确实是同一事件
        # 避免AI验证文本解析错误导致不相关文章被合并
        if len(articles_to_merge) >= 2:
            # ============ 【关键修复】实体重叠预筛选 ============
            # 在调AI前先用规则筛一遍：如果文章对之间核心实体重叠为0，直接判定不是同一事件
            # 这样可以彻底防止AI偶尔判错把"美国司法部诉中企"+"重庆烤肉店"等完全不相关文章合并
            # 【2026-06修复】加入重大事件关键词补充：同一重大事件（如美伊战争）的不同角度报道
            # 实体交集可能很小，但都命中同一事件关键词列表，此时放宽阈值到1
            base_art = articles_to_merge[0]
            base_text = (base_art.category_tag or base_art.title) + " " + (base_art.ai_summary or "")[:400]
            base_entities = crawler._extract_core_entities(base_text)
            # 检查是否命中重大事件关键词
            base_event_kw = {kw for kw in _MAJOR_EVENT_ALL_KEYWORDS if kw in base_text}
            is_major_event_group = len(base_event_kw) >= 2
            if is_major_event_group:
                base_entities.update(base_event_kw)
            entity_check_passed = True
            zero_overlap_pairs = []  # 记录实体重叠为0的文章对，后续用AI兜底
            for k in range(1, len(articles_to_merge)):
                k_art = articles_to_merge[k]
                k_text = (k_art.category_tag or k_art.title) + " " + (k_art.ai_summary or "")[:400]
                k_entities = crawler._extract_core_entities(k_text)
                k_event_kw = {kw for kw in _MAJOR_EVENT_ALL_KEYWORDS if kw in k_text}
                # 如果双方都命中了重大事件关键词，也加入实体集
                if len(k_event_kw) >= 2:
                    k_entities.update(k_event_kw)
                    is_major_event_group = True
                overlap = base_entities & k_entities
                # 【修复】重大事件组放宽阈值到1，普通组也降为1
                # 因为自动修复来自文档验证AI的建议，AI已确认相关性，阈值应更宽松
                # 实体重叠为0时仍需AI兜底判定，不能直接拒绝
                _auto_fix_threshold = 1
                if len(overlap) < _auto_fix_threshold:
                    zero_overlap_pairs.append((0, k))
                    logger.info(
                        f"自动修复实体重叠不足: '{k_art.title[:30]}...'与'{base_art.title[:30]}...' "
                        f"核心实体重叠仅{len(overlap)}个 ({list(overlap)[:3]})，将用AI兜底判定"
                    )
            # 【注意】实体重叠不足的对不会直接跳过，而是由下方AI is_same_event判定
            # AI确认同一事件即可通过，避免实体提取遗漏（如Anthropic/Fable未被提取）导致误拒

            # 对组内文章两两检查是否同一事件（取第一篇与后续每篇比较）
            all_same_event = True
            for k in range(1, len(articles_to_merge)):
                try:
                    is_same = crawler.ai_client.is_same_event(
                        {
                            'title': base_art.category_tag if hasattr(base_art, 'category_tag') and base_art.category_tag else base_art.title,
                            'ai_summary': (base_art.ai_summary or '')[:400],
                            'source_name': base_art.source_name,
                        },
                        {
                            'title': articles_to_merge[k].category_tag if hasattr(articles_to_merge[k], 'category_tag') and articles_to_merge[k].category_tag else articles_to_merge[k].title,
                            'ai_summary': (articles_to_merge[k].ai_summary or '')[:400],
                            'source_name': articles_to_merge[k].source_name,
                        }
                    )
                    if not is_same:
                        all_same_event = False
                        logger.info(f"自动修复跳过: AI判定'{articles_to_merge[k].title[:30]}...'与'{base_art.title[:30]}...'不是同一事件，拒绝合并本组")
                        break
                except Exception as e:
                    logger.warning(f"自动修复AI判断失败，保守跳过本组: {e}")
                    all_same_event = False
                    break
            
            if not all_same_event:
                continue
        
        logger.info(f"自动修复：准备合并{len(articles_to_merge)}篇重复文章(AI已确认同一事件)...")
        
        # 使用AI整合这些文章
        article_links = {a.link for a in articles_to_merge}
        
        # 构建整合数据
        articles_text = []
        for i, art in enumerate(articles_to_merge):
            articles_text.append(
                f"文章{i+1}:\n"
                f"标题：{art.title}\n"
                f"来源：{art.source_name}\n"
                f"摘要：{art.ai_summary[:600] if art.ai_summary else ''}\n"
            )
        articles_str = "\n\n".join(articles_text)
        
        system_prompt = "你是一位资深新闻编辑，擅长将多篇报道同一具体事件的文章整合为一条完整、信息丰富的资讯。"
        
        user_prompt = f"""请将以下{len(articles_to_merge)}篇文章整合为一条完整的资讯。

**【极其重要的前提检查】**：在整合之前，你必须先检查这些文章是否确实在报道同一个具体事件。如果这些文章涉及多个完全不相关的事件，则**绝对不能整合在一起**，你应该在输出开头标注【整合失败：文章主题不相关】。只有当所有文章都在报道同一个具体事件的不同角度/进展时，才进行整合。

{articles_str}

【整合要求】：
1. **标题**：生成一个简洁有力的综合标题，格式为【具体标题内容】，聚焦核心事件。绝对不要输出【标题】或【综合报道】等空洞占位符。
2. **内容组织**：
   - 按照事件发展的时间线或不同视角组织内容
   - 确保每篇文章的独特信息都被保留，不要遗漏任何重要细节
   - 信息密度要高，去除重复表达
   - **绝对不能将不相关的事件拼凑在一起**
3. **格式要求**：
   - 首行是【具体标题】，标题必须包含事件的核心关键词
   - 正文是一段完整的叙述（4-6句、总字数不超过100字，严格限制），涵盖所有文章最核心的信息，正文内不要换行
   - 正文后空一行，加上"来源列表："，列出所有来源文章的标题（用《》包裹，用顿号分隔）
4. **字数严格限制**：
   - 正文内容（不含标题和来源列表）用 4-6 句短句，严格不超过 100 字

请直接输出整合后的内容，不要有任何解释说明。"""
        
        try:
            response = crawler.ai_client._call_api_with_system(system_prompt, user_prompt)
            merged_content = response.strip()
            
            # 检查AI是否判定文章主题不相关，拒绝整合
            if '整合失败' in merged_content and '主题不相关' in merged_content:
                logger.warning(f"自动修复：AI判定文章主题不相关，拒绝整合本组")
                continue
            
            # 解析整合内容
            title_match = re.search(r'【([^】]+)】', merged_content)
            hollow_titles = {'标题', '综合报道', '综合新闻', '新闻摘要', '综合资讯', '资讯', '报道', '综合', '新闻', '热点', '要闻', '事件'}
            if title_match:
                raw_title = title_match.group(1).strip()
                if raw_title in hollow_titles or len(raw_title) <= 2:
                    summary_for_title = re.sub(r'【[^】]+】\s*', '', merged_content, count=1).strip()
                    summary_for_title = re.sub(r'\n*来源列表[：:].*$', '', summary_for_title, flags=re.DOTALL).strip()
                    first_sentence = re.split(r'[。！？\n]', summary_for_title)[0]
                    raw_title = first_sentence[:40] + ('…' if len(first_sentence) > 40 else '')
                    if not raw_title.strip() and articles_to_merge:
                        raw_title = articles_to_merge[0].title[:40]
                merged_title = f"【{raw_title}】"
                merged_summary = re.sub(r'【[^】]+】\s*', '', merged_content, count=1)
            else:
                lines = merged_content.split('\n')
                raw_title = lines[0].strip()[:50]
                if raw_title in hollow_titles or len(raw_title) <= 2:
                    summary_for_title = '\n'.join(lines[1:]).strip()
                    first_sentence = re.split(r'[。！？\n]', summary_for_title)[0]
                    raw_title = first_sentence[:40] + ('…' if len(first_sentence) > 40 else '')
                    if not raw_title.strip() and articles_to_merge:
                        raw_title = articles_to_merge[0].title[:40]
                merged_title = f"【{raw_title}】"
                merged_summary = '\n'.join(lines[1:]).strip()
            
            # 提取来源列表
            source_list = []
            source_match = re.search(r'来源列表[：:](.+)', merged_summary, re.DOTALL)
            if source_match:
                source_text = source_match.group(1)
                source_titles = re.findall(r'《([^》]+)》', source_text)
                source_list = source_titles
                merged_summary = re.sub(r'\n*来源列表[：:].*$', '', merged_summary, flags=re.DOTALL).strip()
            
            # 清理摘要
            merged_summary = merged_summary.replace('\n', ' ').replace('\r', ' ').strip()
            
            # 字数检查，超100字硬截断
            if len(merged_summary) > 100:
                truncated = merged_summary[:100]
                last_punct = max(truncated.rfind('。'), truncated.rfind('！'), truncated.rfind('？'), truncated.rfind('；'))
                if last_punct > 60:
                    truncated = truncated[:last_punct+1]
                merged_summary = truncated
            
            # 创建整合后的Article
            base_article = articles_to_merge[0]
            
            # 收集所有来源信息
            merged_sources = []
            seen_links = set()
            seen_titles = set()  # 【修复】同名文章去重
            for art in articles_to_merge:
                if getattr(art, 'is_merged', False) and getattr(art, 'merged_sources', None):
                    # 文章本身是整合文章，展开其所有原始来源
                    for src in art.merged_sources:
                        src_link = src.get('link', '')
                        src_title_key = src.get('title', '').strip()
                        if src_link not in seen_links and src_title_key not in seen_titles:
                            merged_sources.append(src)
                            seen_links.add(src_link)
                            seen_titles.add(src_title_key)
                else:
                    if art.link not in seen_links and art.title.strip() not in seen_titles:
                        merged_sources.append({
                            'title': art.title,
                            'source': art.source_name,
                            'link': art.link,
                        })
                        seen_links.add(art.link)
                        seen_titles.add(art.title.strip())
            
            # 创建新的整合文章
            # 【修复】如果合并组中已有整合文章（is_merged=True），优先继承其 category_tag/title，
            # 避免合并后丢失之前精心生成的整合标题，回退到原文标题
            existing_merged = next((a for a in articles_to_merge if getattr(a, 'is_merged', False)), None)
            if existing_merged and existing_merged.category_tag and "⚠️" not in existing_merged.category_tag:
                # 用现有整合文章作为基础
                merged_article = Article(
                    title=existing_merged.title,
                    link=existing_merged.link,
                    source_name="多来源整合",
                    pub_date=max((a.pub_date for a in articles_to_merge if a.pub_date), default=base_article.pub_date),
                    full_content=existing_merged.full_content or base_article.full_content,
                )
                # category_tag 优先用本次AI重新生成的 merged_title，但若新生成的是空洞占位则保留原整合标题
                _new_tag_inner = (merged_title or '').replace('【', '').replace('】', '').strip()
                if _new_tag_inner and _new_tag_inner not in hollow_titles and len(_new_tag_inner) > 4:
                    merged_article.category_tag = merged_title
                else:
                    merged_article.category_tag = existing_merged.category_tag
                    logger.info(f"自动修复: 保留原整合标题 '{existing_merged.category_tag}' 而非新生成的空洞标题")
            else:
                merged_article = Article(
                    title=base_article.title,
                    link=base_article.link,
                    source_name=base_article.source_name,
                    pub_date=max((a.pub_date for a in articles_to_merge if a.pub_date), default=base_article.pub_date),
                    full_content=base_article.full_content,
                )
                merged_article.category_tag = merged_title
            merged_article.ai_summary = merged_summary
            merged_article.mece_category = base_article.mece_category
            merged_article.is_merged = True
            merged_article.merged_sources = merged_sources
            
            # 从crawler.articles中移除被合并的文章，添加整合文章
            # 【修复】用 id() 标识被合并的文章对象，避免 link="" 的多个整合文章被误删
            ids_to_remove = {id(a) for a in articles_to_merge}
            new_articles = []
            replaced = False
            for art in crawler.articles:
                if id(art) in ids_to_remove:
                    if not replaced:
                        new_articles.append(merged_article)
                        replaced = True
                else:
                    new_articles.append(art)
            
            crawler.articles = new_articles
            logger.info(f"自动修复：成功合并{len(articles_to_merge)}篇重复文章为1篇整合文章")
            _p(f"[自动修复] 成功合并{len(articles_to_merge)}篇重复文章为1篇整合文章")
            any_fixed = True
            
        except Exception as e:
            logger.error(f"自动修复失败: {e}")
            _p(f"[自动修复] 合并失败: {e}")
            continue
    
    return any_fixed


def main(custom_excel_path: str = None):
    """主函数"""
    try:
        # DeepSeek API 密钥：从环境变量读取（由 run.py 通过 auth 模块注入）
        # 直接调用 main()（不经 run.py）时，会触发交互式授权流程作为兜底
        DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "").strip()
        if not DEEPSEEK_API_KEY:
            try:
                # 兜底：直跑 src/main.py 时也走授权流程
                sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
                from auth import get_api_key  # type: ignore
                DEEPSEEK_API_KEY = get_api_key() or ""
            except Exception as _e:
                logger.error(f"加载授权模块失败: {_e}")
            if not DEEPSEEK_API_KEY:
                print("[错误] 未获取到 DeepSeek API 密钥，请通过 run.py 启动并完成授权。")
                return False

        # 文件路径
        current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        parent_dir = os.path.dirname(current_dir)
        # 如果指定了自定义Excel文件则使用，否则使用默认的
        if custom_excel_path:
            if os.path.isabs(custom_excel_path):
                excel_path = custom_excel_path
            else:
                excel_path = os.path.join(current_dir, custom_excel_path)
        else:
            excel_path = os.path.join(parent_dir, "公众号监测源.xlsx")
        # 输出到上级目录的每日资讯文件夹
        output_dir = os.path.join(parent_dir, "每日资讯")
        config_dir = os.path.join(current_dir, "config")
        # 筛选记录目录
        filter_record_dir = os.path.join(current_dir, "筛选记录")

        # 确保目录存在
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(config_dir, exist_ok=True)
        os.makedirs(filter_record_dir, exist_ok=True)

        print("=" * 70)
        print("每日资讯Skill - AI版本（DeepSeek）")
        print(f"开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("=" * 70)

        # 检查Excel文件
        if not os.path.exists(excel_path):
            logger.error(f"Excel文件不存在: {excel_path}")
            print(f"错误: 请确保文件存在: {excel_path}")
            return False

        # 创建爬虫实例
        crawler = WeChatArticleCrawler(excel_path, DEEPSEEK_API_KEY)

        # 爬取所有RSS feed
        crawler.crawl_all_feeds()

        # 生成输出文件名
        today = datetime.now()
        date_str = today.strftime('%Y-%m-%d')
        
        # 【修改】输出到月度文件夹下的按日期命名的子文件夹，如"2026_05/每日资讯_2026-05-25/"
        month_folder = today.strftime('%Y_%m')  # 如 "2026_05"
        output_month_dir = os.path.join(output_dir, month_folder)
        os.makedirs(output_month_dir, exist_ok=True)
        folder_name = f"每日资讯_{date_str}"
        output_sub_dir = os.path.join(output_month_dir, folder_name)
        os.makedirs(output_sub_dir, exist_ok=True)
        
        filename = f"每日资讯_{date_str}.docx"
        output_path = os.path.join(output_sub_dir, filename)

        # 生成文档（传递每日总结）
        _p(f"\n==== 生成 Word 文档 ====")
        _p(f"输出路径: {output_path}")
        
        # 在生成文档前，清理所有摘要中的违规内容
        _clean_summaries(crawler.articles)
        
        document_generator = DocumentGenerator()
        success = document_generator.create_document(crawler.articles, output_path, crawler.daily_summary)
        _p(f"Word 文档生成: {'成功' if success else '失败'}")
        
        # 【新增】生成PDF版本
        if success:
            pdf_filename = f"每日资讯_{date_str}.pdf"
            pdf_output_path = os.path.join(output_sub_dir, pdf_filename)
            try:
                _convert_docx_to_pdf(output_path, pdf_output_path)
                _p(f"PDF 文档生成: 成功")
            except Exception as pdf_err:
                _p(f"PDF 文档生成: 失败({pdf_err})，跳过PDF生成")
                logger.warning(f"PDF生成失败: {pdf_err}")
        
        # ========== 第八步：DeepSeek API验证文档重复 + 自动修复（最多3轮） ==========
        if success:
            _p(f"\n==== 第八步 / 共八步：DeepSeek API验证文档重复 ====")
            _p(f"[步骤8] 调用DeepSeek API检查文档中是否还存在重复内容...")
            
            max_verify_rounds = 3
            for verify_round in range(1, max_verify_rounds + 1):
                # 构建文档文本用于验证（包含标题和摘要）
                # 【重要修复】使用顺序编号（1, 2, 3...）而非原始列表索引，
                # 确保 _auto_fix_duplicates 中的 valid_articles 索引与编号一致
                verify_text_parts = []
                seq_num = 0
                for article in crawler.articles:
                    if article.is_advertisement:
                        continue
                    if not article.ai_summary or article.ai_summary == "摘要生成失败":
                        continue
                    seq_num += 1
                    title = article.category_tag if article.category_tag else article.title
                    summary = article.ai_summary[:400] if article.ai_summary else ""
                    verify_text_parts.append(f"{seq_num}. {title}\n{summary}\n")
                
                verify_text = "\n".join(verify_text_parts)
                
                try:
                    has_duplicates, duplicate_groups, suggestions = crawler.ai_client.verify_document_duplicates(verify_text)
                    
                    if has_duplicates:
                        _p(f"[步骤8] 第{verify_round}轮验证：⚠️ 发现重复内容！")
                        _p(f"[步骤8] 建议: {suggestions}")
                        logger.warning(f"文档验证第{verify_round}轮发现重复: {suggestions}")
                        
                        # 自动修复：调用AI解析重复组，然后在文章列表中合并
                        fixed = _auto_fix_duplicates(crawler, suggestions)
                        
                        if fixed:
                            _p(f"[步骤8] 第{verify_round}轮：已自动修复重复内容，重新生成文档...")
                            # 重新生成文档
                            success = document_generator.create_document(crawler.articles, output_path, crawler.daily_summary)
                            if success:
                                _p(f"[步骤8] 文档重新生成成功，进入第{verify_round+1}轮验证...")
                                continue
                            else:
                                _p(f"[步骤8] 文档重新生成失败，停止验证")
                                break
                        else:
                            _p(f"[步骤8] 第{verify_round}轮：无法自动修复，记录问题")
                            crawler.doc_verify_result = {
                                'has_duplicates': True,
                                'duplicate_groups': duplicate_groups,
                                'suggestions': suggestions,
                            }
                            break
                    else:
                        _p(f"[步骤8] 第{verify_round}轮验证：✓ 文档无重复内容")
                        crawler.doc_verify_result = {
                            'has_duplicates': False,
                            'duplicate_groups': [],
                            'suggestions': "",
                        }
                        break
                except Exception as e:
                    _p(f"[步骤8] 第{verify_round}轮验证过程出错: {e}")
                    logger.error(f"文档重复验证第{verify_round}轮出错: {e}")
                    crawler.doc_verify_result = {
                        'has_duplicates': False,
                        'duplicate_groups': [],
                        'suggestions': f"验证出错: {e}",
                    }
                    break
            
            if verify_round >= max_verify_rounds:
                _p(f"[步骤8] 已达到最大验证轮数({max_verify_rounds})，停止验证")
                crawler.doc_verify_result = {
                    'has_duplicates': True,
                    'duplicate_groups': [],
                    'suggestions': f"经过{max_verify_rounds}轮验证仍有重复",
                }
        
        # ========== 保存缓存 ==========
        if success:
            _p(f"\n==== 保存文章缓存 ====")
            _cache_count = 0
            for article in crawler.articles:
                if not article.is_advertisement and article.ai_summary and article.ai_summary != "摘要生成失败":
                    # 准备缓存数据
                    article_data = {
                        'title': article.title,
                        'link': article.link,
                        'source_name': article.source_name,
                        'ai_summary': article.ai_summary,
                        'is_advertisement': article.is_advertisement,
                        'rejection_reason': article.rejection_reason,
                        'category_tag': article.category_tag,
                        'mece_category': article.mece_category,
                        'pub_date': article.pub_date.isoformat() if article.pub_date else None,
                    }
                    crawler._add_to_cache(article.link, article.title, article_data)
                    _cache_count += 1
            crawler._save_cache()
            _p(f"[缓存] 保存 {_cache_count} 篇文章到缓存")
        
        # ========== 生成筛选记录 ==========
        _p(f"\n==== 生成筛选记录 ====")
        # 获取噪音过滤和去重移除的文章信息
        noise_articles = crawler.noise_articles
        duplicate_articles = crawler.duplicate_removed
        
        # 生成筛选记录Markdown文件
        record_filename = f"筛选记录_{today.strftime('%Y-%m-%d')}.md"
        record_path = os.path.join(filter_record_dir, record_filename)
        
        with open(record_path, 'w', encoding='utf-8') as f:
            f.write(f"# 每日资讯筛选记录\n\n")
            f.write(f"**生成时间**: {today.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"## 统计信息\n\n")
            total_original = len(crawler.articles) + len(noise_articles) + len(duplicate_articles)
            f.write(f"- 总爬取文章数: {total_original}篇\n")
            f.write(f"- 噪音过滤: {len(noise_articles)}篇\n")
            f.write(f"- 去重移除: {len(duplicate_articles)}篇\n")
            f.write(f"- 最终收录: {len(crawler.articles)}篇\n\n")
            
            # 噪音过滤文章列表
            f.write(f"## 噪音过滤文章列表\n\n")
            if noise_articles:
                for i, article in enumerate(noise_articles, 1):
                    f.write(f"### {i}. {article.title}\n\n")
                    f.write(f"- **来源**: {article.source_name}\n")
                    f.write(f"- **过滤原因**: {article.rejection_reason}\n")
                    f.write(f"- **链接**: {article.link}\n\n")
            else:
                f.write("（无）\n\n")
            
            # 去重移除文章列表（含详细去重原因）
            f.write(f"## 去重移除文章列表\n\n")
            dedup_details = crawler.dedup_details
            if duplicate_articles:
                for i, article in enumerate(duplicate_articles, 1):
                    f.write(f"### {i}. {article.title}\n\n")
                    f.write(f"- **来源**: {article.source_name}\n")
                    f.write(f"- **发布时间**: {article.pub_date.strftime('%Y-%m-%d %H:%M')}\n")
                    f.write(f"- **链接**: {article.link}\n")
                    # 查找该文章的去重详情
                    detail = next((d for d in dedup_details if d['removed_link'] == article.link), None)
                    if detail:
                        f.write(f"- **去重原因**: {detail['reason']}\n")
                        f.write(f"- **保留的文章**: [{detail['kept_title']}]({detail['kept_link']}) (来源: {detail['kept_source']})\n")
                    f.write(f"\n")
            else:
                f.write("（无）\n\n")
            
            # 附录：分类统计
            f.write(f"## 附录：分类统计\n\n")
            category_stats = {}
            for article in crawler.articles:
                cat = article.mece_category if article.mece_category else "11.3"
                main_cat = cat.split('.')[0]
                cat_name = DocumentGenerator.MECE_MAIN_CATEGORIES.get(main_cat, "其他")
                category_stats[cat_name] = category_stats.get(cat_name, 0) + 1
            
            for cat_name, count in sorted(category_stats.items(), key=lambda x: list(DocumentGenerator.MECE_MAIN_CATEGORIES.values()).index(x[0]) if x[0] in DocumentGenerator.MECE_MAIN_CATEGORIES.values() else 999):
                f.write(f"- **{cat_name}**: {count}篇\n")
            
            # 附录：文档重复验证结果
            f.write(f"\n## 附录：文档重复验证结果\n\n")
            doc_verify = getattr(crawler, 'doc_verify_result', None)
            if doc_verify:
                if doc_verify.get('has_duplicates', False):
                    f.write(f"- **验证结果**: ⚠️ 发现重复内容\n")
                    duplicate_groups = doc_verify.get('duplicate_groups', [])
                    suggestions = doc_verify.get('suggestions', '')
                    if duplicate_groups:
                        f.write(f"\n**重复组列表**（共{len(duplicate_groups)}组）：\n\n")
                        for idx, group in enumerate(duplicate_groups, 1):
                            f.write(f"{idx}. {group}\n")
                    if suggestions:
                        f.write(f"\n**建议**：\n\n")
                        # 尝试将建议按编号拆分为多行
                        suggestion_items = re.split(r'(?=\d+\.\s)', suggestions)
                        suggestion_items = [s.strip() for s in suggestion_items if s.strip()]
                        for idx, item in enumerate(suggestion_items, 1):
                            f.write(f"{idx}. {item}\n")
                else:
                    f.write(f"- **验证结果**: ✓ 文档无重复内容\n")
            else:
                f.write(f"- **验证结果**: 未进行验证\n")
        
        logger.info(f"筛选记录已生成: {record_path}")

        if success:
            print("\n" + "=" * 70)
            print("[完成] 任务执行成功！")
            print("=" * 70)

            # 统计信息
            valid_articles = [a for a in crawler.articles if not a.is_advertisement]

            print(f"\n统计信息:")
            total_original = len(crawler.articles) + len(noise_articles) + len(duplicate_articles)
            print(f"- 总爬取: {total_original}篇")
            print(f"- 噪音过滤: {len(noise_articles)}篇")
            print(f"- 去重移除: {len(duplicate_articles)}篇")
            print(f"- 最终收录: {len(crawler.articles)}篇")
            
            # 显示整合文章数量
            merged_count = sum(1 for a in crawler.articles if getattr(a, 'is_merged', False))
            if merged_count > 0:
                print(f"- 多视角整合: {merged_count}条")

            # 显示验证结果
            doc_verify = getattr(crawler, 'doc_verify_result', None)
            if doc_verify:
                if doc_verify.get('has_duplicates', False):
                    print(f"- 文档验证: ⚠️ 发现重复，建议查看筛选记录")
                else:
                    print(f"- 文档验证: ✓ 无重复")

            print(f"\n输出文档: {output_path}")
            print(f"筛选记录: {record_path}")

        else:
            print("\n[错误] 文档生成失败")
            return False

        print("\n" + "=" * 70)

    except Exception as e:
        logger.error(f"程序执行出错: {e}", exc_info=True)
        print(f"\n[错误] 程序执行出错: {e}")
        import traceback
        traceback.print_exc()
        return False

    return True

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)
